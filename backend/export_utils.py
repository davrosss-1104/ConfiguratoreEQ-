"""
export_utils.py — Generazione DOCX e XLSX per Preventivi e Ordini
Posizionare in: backend/export_utils.py

Dipendenze: pip install python-docx openpyxl cairosvg --break-system-packages
"""
import io
import os
from datetime import datetime
from pathlib import Path


# ═══════════════════════════════════════════════════════
# PRODUCT ICON LOADER — SVG → PNG per DOCX
# ═══════════════════════════════════════════════════════

# Percorsi dove cercare le icone SVG del prodotto.
# Il backend prova ciascun percorso in ordine (relativi a CWD,
# poi assoluti comuni) finché trova il file.
ICON_SEARCH_PATHS = [
    "frontend/public/icons",         # dev: CWD = root progetto
    "../frontend/public/icons",      # dev: CWD = backend/
    "public/icons",                  # build / static
    "static/icons",
    "icons",
    "../public/icons",
]


def load_product_icon_png(icon_name: str, size: int = 120) -> bytes | None:
    """
    Cerca l'icona prodotto come PNG (preferito) o SVG convertito.
    Ritorna bytes PNG oppure None se non trovata.
    """
    if not icon_name:
        return None

    # 1. Cerca PNG diretto (pre-convertito) — normalizza DPI
    png_path = _find_icon_file(icon_name, ".png")
    if png_path:
        try:
            from PIL import Image
            img = Image.open(png_path)
            # Forza DPI uniforme per evitare distorsioni in Word
            buf = io.BytesIO()
            img.save(buf, format="PNG", dpi=(150, 150))
            buf.seek(0)
            print(f"[DEBUG ICON] {icon_name}: {img.size[0]}x{img.size[1]} dpi={img.info.get('dpi', 'N/A')}")
            return buf.read()
        except Exception as e:
            print(f"WARN: Errore caricamento PNG ({icon_name}): {e}")

    # 2. Fallback: SVG con cairosvg (se disponibile)
    svg_path = _find_icon_file(icon_name, ".svg")
    if svg_path:
        try:
            import cairosvg
            return cairosvg.svg2png(url=str(svg_path), output_width=size, output_height=size)
        except ImportError:
            pass
        except Exception:
            pass

    print(f"WARN: Icona non trovata ({icon_name}) — salva un .png nella cartella icons/")
    return None


def _find_icon_file(icon_name: str, ext: str) -> Path | None:
    """Cerca un file icona con l'estensione specificata"""
    filename = f"{icon_name}{ext}"
    for base in ICON_SEARCH_PATHS:
        candidate = Path(base) / filename
        if candidate.exists():
            return candidate
    for base_dir in [os.path.expanduser("~")]:
        for rel in ICON_SEARCH_PATHS:
            candidate = Path(base_dir) / rel / filename
            if candidate.exists():
                return candidate
    return None


def _find_icon_svg(icon_name: str) -> Path | None:
    """Cerca il file SVG dell'icona in tutti i percorsi configurati"""
    filename = f"{icon_name}.svg"
    # Percorsi relativi a CWD
    for base in ICON_SEARCH_PATHS:
        candidate = Path(base) / filename
        if candidate.exists():
            return candidate
    # Percorsi assoluti comuni
    for base_dir in ["/app", "/opt/app", os.path.expanduser("~")]:
        for rel in ICON_SEARCH_PATHS:
            candidate = Path(base_dir) / rel / filename
            if candidate.exists():
                return candidate
    return None


# ═══════════════════════════════════════════════════════
# HELPER: Product Icon Header per DOCX
# ═══════════════════════════════════════════════════════
def _add_product_icon_header(doc, product_icon_png: bytes = None,
                              product_name: str = "", product_category: str = ""):
    """
    Aggiunge header con icona prodotto + titolo ELETTROQUADRI + nome prodotto.
    Layout: tabella senza bordi [Icona | Intestazione]
    """
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn

    num_cols = 2 if product_icon_png else 1
    t = doc.add_table(rows=1, cols=num_cols)

    # Rimuovi bordi dalla tabella
    tbl = t._tbl
    tbl_pr = tbl.tblPr if tbl.tblPr is not None else tbl.makeelement(qn('w:tblPr'), {})
    borders = tbl_pr.makeelement(qn('w:tblBorders'), {})
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        el = borders.makeelement(qn(f'w:{edge}'), {
            qn('w:val'): 'none', qn('w:sz'): '0', qn('w:space'): '0', qn('w:color'): 'auto'
        })
        borders.append(el)
    tbl_pr.append(borders)
    if tbl.tblPr is None:
        tbl.insert(0, tbl_pr)

    col_idx = 0

    # Colonna icona
    if product_icon_png:
        cell_icon = t.cell(0, 0)
        p_icon = cell_icon.paragraphs[0]
        p_icon.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_icon = p_icon.add_run()
        try:
            run_icon.add_picture(io.BytesIO(product_icon_png), width=Cm(2.5))
        except Exception as e:
            print(f"WARN: Errore inserimento icona nel DOCX: {e}")

        # Larghezza colonna icona fissa
        tc_pr = cell_icon._element.get_or_add_tcPr()
        tc_w = tc_pr.makeelement(qn('w:tcW'), {qn('w:w'): '1800', qn('w:type'): 'dxa'})
        tc_pr.append(tc_w)
        col_idx = 1

    # Colonna titolo
    cell_title = t.cell(0, col_idx)

    # ELETTROQUADRI S.r.l.
    p_company = cell_title.paragraphs[0]
    p_company.alignment = WD_ALIGN_PARAGRAPH.LEFT if product_icon_png else WD_ALIGN_PARAGRAPH.CENTER
    r = p_company.add_run("ELETTROQUADRI S.r.l.")
    r.font.size = Pt(14)
    r.font.bold = True
    r.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
    r.font.name = 'Arial'

    # Nome prodotto
    if product_name:
        p_prod = cell_title.add_paragraph()
        p_prod.alignment = WD_ALIGN_PARAGRAPH.LEFT if product_icon_png else WD_ALIGN_PARAGRAPH.CENTER
        r_prod = p_prod.add_run(product_name.upper())
        r_prod.font.size = Pt(11)
        r_prod.font.bold = True
        r_prod.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
        r_prod.font.name = 'Arial'

    # Badge categoria
    if product_category:
        p_cat = cell_title.add_paragraph()
        p_cat.alignment = WD_ALIGN_PARAGRAPH.LEFT if product_icon_png else WD_ALIGN_PARAGRAPH.CENTER
        cat_color = RGBColor(0x16, 0xA3, 0x4A) if product_category == "RISE" else RGBColor(0xD9, 0x77, 0x06)
        r_cat = p_cat.add_run(f"● {product_category}")
        r_cat.font.size = Pt(9)
        r_cat.font.bold = True
        r_cat.font.color.rgb = cat_color
        r_cat.font.name = 'Arial'


def _set_repeat_header_row(table):
    """Imposta la prima riga della tabella come intestazione ripetuta su ogni pagina"""
    from docx.oxml.ns import qn
    try:
        tr = table.rows[0]._tr
        trPr = tr.get_or_add_trPr()
        trPr.append(trPr.makeelement(qn('w:tblHeader'), {}))
    except Exception:
        pass


# ═══════════════════════════════════════════════════════
# HELPER: formattazione
# ═══════════════════════════════════════════════════════
def fmt_euro(val):
    """Formatta valore in euro"""
    if val is None:
        return "€ 0,00"
    return f"€ {val:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")


def safe_str(val, default=""):
    """Stringa sicura, mai None"""
    return str(val) if val is not None else default


def safe_float(val, default=0.0):
    try:
        return float(val)
    except (TypeError, ValueError):
        return default


# ═══════════════════════════════════════════════════════
# DOCX: PREVENTIVO
# ═══════════════════════════════════════════════════════
def genera_docx_preventivo(preventivo, dati_commessa, dati_principali, normative, argano, materiali,
                           cliente=None, product_icon_png=None, product_name="", product_category=""):
    """
    Genera un documento DOCX professionale per il preventivo.
    Ritorna un io.BytesIO pronto per StreamingResponse.
    """
    from docx import Document
    from docx.shared import Inches, Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT

    doc = Document()

    # Stili
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(10)

    # ── INTESTAZIONE ──
    if product_icon_png or product_name:
        _add_product_icon_header(doc, product_icon_png, product_name, product_category)
    else:
        h = doc.add_heading('ELETTROQUADRI S.r.l.', level=1)
        h.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in h.runs:
            run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run('Sede Legale: Via Puccini, 1 — Tel. 0332 470049\n').font.size = Pt(8)
    p.add_run('info@elettroquadri.net — www.elettroquadri.net').font.size = Pt(8)

    doc.add_paragraph('')  # spacer

    # ── DATI OFFERTA ──
    numero = getattr(preventivo, 'numero', None) or getattr(preventivo, 'numero_preventivo', '') or ''
    rev = getattr(preventivo, 'revisione_corrente', 0) or 0
    if rev > 0:
        numero = f"{numero} REV.{rev}"
    stato = getattr(preventivo, 'stato', None) or getattr(preventivo, 'status', 'bozza') or 'bozza'
    customer = getattr(preventivo, 'customer_name', '') or ''
    if cliente:
        customer = getattr(cliente, 'ragione_sociale', customer) or customer

    t = doc.add_table(rows=4, cols=4)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = 'Table Grid'
    cells = [
        ("Offerta N°", safe_str(numero), "Data", safe_str(getattr(dati_commessa, 'data_offerta', '') if dati_commessa else '')),
        ("Stato", safe_str(stato).upper(), "Riferimento", safe_str(getattr(dati_commessa, 'riferimento_cliente', '') if dati_commessa else '')),
        ("Cliente", customer, "Quantità", safe_str(getattr(dati_commessa, 'quantita', '') if dati_commessa else '')),
        ("Pagamento", safe_str(getattr(dati_commessa, 'pagamento', '') if dati_commessa else ''), "Trasporto", safe_str(getattr(dati_commessa, 'trasporto', '') if dati_commessa else '')),
    ]
    for i, (l1, v1, l2, v2) in enumerate(cells):
        t.cell(i, 0).text = l1
        t.cell(i, 1).text = v1
        t.cell(i, 2).text = l2
        t.cell(i, 3).text = v2
        for j in [0, 2]:
            for run in t.cell(i, j).paragraphs[0].runs:
                run.font.bold = True
                run.font.size = Pt(9)
            for run in t.cell(i, j + 1).paragraphs[0].runs:
                run.font.size = Pt(9)

    doc.add_paragraph('')

    # ── DATI PRINCIPALI ──
    if dati_principali:
        doc.add_heading('Dati Principali Impianto', level=2)
        dp_fields = [
            ("Tipo impianto", getattr(dati_principali, 'tipo_impianto', '')),
            ("N° Fermate", getattr(dati_principali, 'numero_fermate', '')),
            ("N° Servizi", getattr(dati_principali, 'numero_servizi', '')),
            ("Velocità (m/s)", getattr(dati_principali, 'velocita', '')),
            ("Corsa (m)", getattr(dati_principali, 'corsa', '')),
            ("Tipo Trazione", getattr(dati_principali, 'tipo_trazione', '')),
            ("Forza Motrice", getattr(dati_principali, 'forza_motrice', '')),
            ("Tensione Manovra", getattr(dati_principali, 'tensione_manovra', '')),
            ("Tensione Freno", getattr(dati_principali, 'tensione_freno', '')),
        ]
        t2 = doc.add_table(rows=len(dp_fields), cols=2)
        t2.style = 'Table Grid'
        for i, (label, val) in enumerate(dp_fields):
            t2.cell(i, 0).text = label
            t2.cell(i, 1).text = safe_str(val)
            for run in t2.cell(i, 0).paragraphs[0].runs:
                run.font.bold = True
                run.font.size = Pt(9)

    # ── ARGANO ──
    if argano:
        doc.add_heading('Argano / Motore', level=2)
        arg_fields = [
            ("Trazione", getattr(argano, 'trazione', '')),
            ("Potenza Motore (kW)", getattr(argano, 'potenza_motore_kw', '')),
            ("Corrente Nominale (A)", getattr(argano, 'corrente_nom_motore_amp', '')),
            ("Tipo VVVF", getattr(argano, 'tipo_vvvf', '')),
        ]
        t3 = doc.add_table(rows=len(arg_fields), cols=2)
        t3.style = 'Table Grid'
        for i, (label, val) in enumerate(arg_fields):
            t3.cell(i, 0).text = label
            t3.cell(i, 1).text = safe_str(val)
            for run in t3.cell(i, 0).paragraphs[0].runs:
                run.font.bold = True
                run.font.size = Pt(9)

    # ── NORMATIVE ──
    if normative:
        doc.add_heading('Normative', level=2)
        norms_attive = []
        for n in ['en_81_1', 'en_81_20', 'en_81_21', 'en_81_28', 'en_81_70',
                   'en_81_72', 'en_81_73', 'a3_95_16', 'dm236_legge13']:
            if getattr(normative, n, False):
                norms_attive.append(n.replace('_', ' ').upper())
        if norms_attive:
            doc.add_paragraph(', '.join(norms_attive))

    # ── MATERIALI ──
    doc.add_heading('Distinta Materiali', level=2)

    if materiali:
        col_widths = [Cm(3), Cm(7), Cm(2), Cm(2.5), Cm(3)]
        t4 = doc.add_table(rows=1, cols=5)
        t4.style = 'Table Grid'
        t4.alignment = WD_TABLE_ALIGNMENT.CENTER

        # Header
        headers = ['Codice', 'Descrizione', 'Qtà', 'Prezzo Unit.', 'Totale']
        for i, h_text in enumerate(headers):
            cell = t4.rows[0].cells[i]
            cell.text = h_text
            for run in cell.paragraphs[0].runs:
                run.font.bold = True
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            from docx.oxml.ns import qn
            shading = cell._element.get_or_add_tcPr()
            shading_elm = shading.makeelement(qn('w:shd'), {
                qn('w:fill'): '333333', qn('w:val'): 'clear'
            })
            shading.append(shading_elm)

        _set_repeat_header_row(t4)

        # Righe materiali
        totale_mat = 0.0
        for m in materiali:
            row = t4.add_row().cells
            row[0].text = safe_str(m.codice)
            desc = safe_str(m.descrizione)
            if m.aggiunto_da_regola:
                regola_id = safe_str(getattr(m, 'regola_id', ''))
                if regola_id.startswith('TEMPLATE_BASE'):
                    desc += ' [Template]'
                else:
                    desc += ' [Auto]'
            row[1].text = desc
            row[2].text = safe_str(m.quantita)
            row[3].text = fmt_euro(m.prezzo_unitario)
            row[4].text = fmt_euro(m.prezzo_totale)
            totale_mat += safe_float(m.prezzo_totale)
            for cell in row:
                for run in cell.paragraphs[0].runs:
                    run.font.size = Pt(9)

        # Riga totale
        row_tot = t4.add_row().cells
        row_tot[0].text = ''
        row_tot[1].text = ''
        row_tot[2].text = ''
        row_tot[3].text = 'TOTALE'
        row_tot[4].text = fmt_euro(totale_mat)
        for run in row_tot[3].paragraphs[0].runs:
            run.font.bold = True
        for run in row_tot[4].paragraphs[0].runs:
            run.font.bold = True

    # ── TOTALI ──
    doc.add_paragraph('')
    totale = safe_float(getattr(preventivo, 'totale_materiali', None) or getattr(preventivo, 'total_price', 0))
    netto = safe_float(getattr(preventivo, 'totale_netto', None) or getattr(preventivo, 'total_price_finale', 0))
    sconto = safe_float(getattr(preventivo, 'sconto_percentuale', 0))

    p_tot = doc.add_paragraph()
    p_tot.add_run(f'Totale Materiali: {fmt_euro(totale)}\n').font.size = Pt(11)
    if sconto > 0:
        p_tot.add_run(f'Sconto: {sconto}%\n').font.size = Pt(11)
    if netto > 0 and netto != totale:
        run_netto = p_tot.add_run(f'Totale Netto: {fmt_euro(netto)}')
        run_netto.font.size = Pt(12)
        run_netto.font.bold = True

    # ── NOTE ──
    note = getattr(preventivo, 'note_interne', '')
    if note:
        doc.add_heading('Note', level=2)
        doc.add_paragraph(safe_str(note))

    # ── FOOTER ──
    doc.add_paragraph('')
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run(f'Documento generato il {datetime.now().strftime("%d/%m/%Y %H:%M")}').font.size = Pt(8)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


# ═══════════════════════════════════════════════════════
# DOCX: PREVENTIVO v2 — Data-driven da /dati-documento
# ═══════════════════════════════════════════════════════
def genera_docx_preventivo_v2(preventivo_info: dict, dati_documento: dict, materiali: list,
                              product_icon_png: bytes = None, product_name: str = "",
                              product_category: str = ""):
    """
    Genera documento DOCX leggendo la struttura dinamica da /dati-documento.
    
    Args:
        preventivo_info: {"numero": str, "customer": str, "status": str,
                          "totale": float, "sconto": float, "netto": float, "note": str}
        dati_documento:  output di GET /preventivi/{id}/dati-documento
        materiali:       lista ORM Materiale (per tabella materiali dettagliata)
        product_icon_png: bytes PNG dell'icona prodotto (opzionale)
        product_name:    nome display del prodotto (opzionale)
        product_category: categoria "RISE" o "HOME" (opzionale)
    """
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn

    doc = Document()

    # Stili base
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(10)

    sezioni = dati_documento.get("sezioni", [])

    for sez in sezioni:
        codice = sez.get("codice", "")
        titolo = sez.get("titolo", "")
        tipo = sez.get("tipo", "tabella")
        mostra_titolo = sez.get("mostra_titolo", True)
        nota = sez.get("nota")
        stile = sez.get("stile")
        campi = sez.get("campi", [])
        tipo_speciale = sez.get("_tipo_speciale")

        # ── INTESTAZIONE ──
        if tipo == "intestazione" or codice == "intestazione":
            if product_icon_png or product_name:
                _add_product_icon_header(doc, product_icon_png, product_name, product_category)
            else:
                h = doc.add_heading('ELETTROQUADRI S.r.l.', level=1)
                h.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in h.runs:
                    run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)

            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run('Sede Legale: Via Puccini, 1 — Tel. 0332 470049\n').font.size = Pt(8)
            p.add_run('info@elettroquadri.net — www.elettroquadri.net').font.size = Pt(8)

            # Info base preventivo
            doc.add_paragraph('')
            numero = preventivo_info.get("numero", "")
            customer = preventivo_info.get("customer", "")
            status = preventivo_info.get("status", "bozza")

            t = doc.add_table(rows=2, cols=4)
            t.alignment = WD_TABLE_ALIGNMENT.CENTER
            t.style = 'Table Grid'
            revisione = preventivo_info.get("revisione", 0)
            numero_display = safe_str(numero)
            if revisione:
                numero_display += f" — REV.{revisione}"
            rows_data = [
                ("Offerta N°", numero_display, "Stato", safe_str(status).upper()),
                ("Cliente", safe_str(customer), "", ""),
            ]
            for i, (l1, v1, l2, v2) in enumerate(rows_data):
                t.cell(i, 0).text = l1
                t.cell(i, 1).text = v1
                t.cell(i, 2).text = l2
                t.cell(i, 3).text = v2
                for j in [0, 2]:
                    for run in t.cell(i, j).paragraphs[0].runs:
                        run.font.bold = True
                        run.font.size = Pt(9)
                    for run in t.cell(i, j + 1).paragraphs[0].runs:
                        run.font.size = Pt(9)
            doc.add_paragraph('')
            continue

        # ── MATERIALI ──
        if tipo_speciale == "materiali" or codice == "materiali":
            if mostra_titolo:
                doc.add_heading(titolo or 'Distinta Materiali', level=2)

            if materiali:
                col_widths = [Cm(3), Cm(7), Cm(2), Cm(2.5), Cm(3)]
                t_mat = doc.add_table(rows=1, cols=5)
                t_mat.style = 'Table Grid'
                t_mat.alignment = WD_TABLE_ALIGNMENT.CENTER

                headers = ['Codice', 'Descrizione', 'Qtà', 'Prezzo Unit.', 'Totale']
                for i, h_text in enumerate(headers):
                    cell = t_mat.rows[0].cells[i]
                    cell.text = h_text
                    for run in cell.paragraphs[0].runs:
                        run.font.bold = True
                        run.font.size = Pt(9)
                        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                    shading = cell._element.get_or_add_tcPr()
                    shading_elm = shading.makeelement(qn('w:shd'), {
                        qn('w:fill'): '333333', qn('w:val'): 'clear'
                    })
                    shading.append(shading_elm)

                _set_repeat_header_row(t_mat)

                totale_mat = 0.0
                for m in materiali:
                    row = t_mat.add_row().cells
                    row[0].text = safe_str(getattr(m, 'codice', ''))
                    desc = safe_str(getattr(m, 'descrizione', ''))
                    if getattr(m, 'aggiunto_da_regola', False):
                        regola_id = safe_str(getattr(m, 'regola_id', ''))
                        desc += ' [Template]' if regola_id.startswith('TEMPLATE_BASE') else ' [Auto]'
                    row[1].text = desc
                    row[2].text = safe_str(getattr(m, 'quantita', ''))
                    row[3].text = fmt_euro(getattr(m, 'prezzo_unitario', 0))
                    prezzo_tot = safe_float(getattr(m, 'prezzo_totale', 0))
                    row[4].text = fmt_euro(prezzo_tot)
                    totale_mat += prezzo_tot
                    for cell in row:
                        for run in cell.paragraphs[0].runs:
                            run.font.size = Pt(9)

                # Riga totale
                row_tot = t_mat.add_row().cells
                row_tot[3].text = 'TOTALE'
                row_tot[4].text = fmt_euro(totale_mat)
                for run in row_tot[3].paragraphs[0].runs:
                    run.font.bold = True
                for run in row_tot[4].paragraphs[0].runs:
                    run.font.bold = True

            # Totali con sconti
            doc.add_paragraph('')
            totale = safe_float(preventivo_info.get("totale", 0))
            netto = safe_float(preventivo_info.get("netto", 0))
            sconto = safe_float(preventivo_info.get("sconto", 0))

            p_tot = doc.add_paragraph()
            p_tot.add_run(f'Totale Materiali: {fmt_euro(totale)}\n').font.size = Pt(11)
            if sconto > 0:
                p_tot.add_run(f'Sconto: {sconto}%\n').font.size = Pt(11)
            if netto > 0 and netto != totale:
                run_netto = p_tot.add_run(f'Totale Netto: {fmt_euro(netto)}')
                run_netto.font.size = Pt(12)
                run_netto.font.bold = True
            continue

        # ── TESTO LIBERO (note, condizioni) ──
        if tipo in ("testo_libero", "testo"):
            if mostra_titolo:
                doc.add_heading(titolo, level=2)
            # I campi di testo libero hanno valore nel primo campo
            for campo in campi:
                val = campo.get("valore", "")
                if val:
                    doc.add_paragraph(safe_str(val))
            # Nota default della sezione
            if nota:
                doc.add_paragraph(safe_str(nota))
            continue

        # ── VALORI STANDARD ──
        if codice == "valori_standard":
            if mostra_titolo:
                doc.add_heading(titolo, level=2)
            if nota:
                p_nota = doc.add_paragraph()
                run_nota = p_nota.add_run(safe_str(nota))
                run_nota.font.size = Pt(8)
                run_nota.font.italic = True
                run_nota.font.color.rgb = RGBColor(0x99, 0x66, 0x00)

            if campi:
                t_std = doc.add_table(rows=1, cols=3)
                t_std.style = 'Table Grid'
                for i, h_text in enumerate(['Campo', 'Valore', 'Sezione']):
                    cell = t_std.rows[0].cells[i]
                    cell.text = h_text
                    for run in cell.paragraphs[0].runs:
                        run.font.bold = True
                        run.font.size = Pt(9)
                _set_repeat_header_row(t_std)
                for campo in campi:
                    row = t_std.add_row().cells
                    row[0].text = safe_str(campo.get("etichetta", ""))
                    valore = safe_str(campo.get("valore", ""))
                    um = campo.get("unita_misura")
                    if um:
                        valore += f" {um}"
                    row[1].text = valore
                    row[2].text = safe_str(campo.get("sezione_config", ""))
                    for cell in row:
                        for run in cell.paragraphs[0].runs:
                            run.font.size = Pt(9)
            continue

        # ── TABELLA GENERICA (dati_commessa, specifiche_tecniche, configurazione, ecc.) ──
        if campi:
            if mostra_titolo:
                doc.add_heading(titolo, level=2)

            t_sez = doc.add_table(rows=len(campi), cols=2)
            t_sez.style = 'Table Grid'
            has_defaults = False
            for i, campo in enumerate(campi):
                etichetta = safe_str(campo.get("etichetta", ""))
                valore = safe_str(campo.get("valore", ""))
                um = campo.get("unita_misura")
                if um:
                    valore += f" {um}"
                is_def = campo.get("is_default", False)

                t_sez.cell(i, 0).text = etichetta
                t_sez.cell(i, 1).text = valore
                for run in t_sez.cell(i, 0).paragraphs[0].runs:
                    run.font.bold = True
                    run.font.size = Pt(9)
                    if is_def:
                        run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
                for run in t_sez.cell(i, 1).paragraphs[0].runs:
                    run.font.size = Pt(9)
                    if is_def:
                        run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
                        run.font.italic = True

                if is_def:
                    has_defaults = True

            # Nota a piè di sezione se ci sono valori di default
            if has_defaults:
                p_leg = doc.add_paragraph()
                run_leg = p_leg.add_run("I valori in grigio corsivo non sono stati modificati rispetto alla configurazione standard.")
                run_leg.font.size = Pt(7)
                run_leg.font.italic = True
                run_leg.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    # ── NOTE PREVENTIVO ──
    note = preventivo_info.get("note", "")
    if note:
        doc.add_heading('Note', level=2)
        doc.add_paragraph(safe_str(note))

    # ── FOOTER ──
    doc.add_paragraph('')
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run(f'Documento generato il {datetime.now().strftime("%d/%m/%Y %H:%M")}').font.size = Pt(8)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


# ═══════════════════════════════════════════════════════
# XLSX: PREVENTIVO
# ═══════════════════════════════════════════════════════
def genera_xlsx_preventivo(preventivo, dati_commessa, dati_principali, normative, argano, materiali, cliente=None):
    """
    Genera un file XLSX con fogli: Riepilogo, Materiali, Dati Tecnici.
    Ritorna un io.BytesIO pronto per StreamingResponse.
    """
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side

    wb = Workbook()
    header_font = Font(name='Arial', bold=True, size=10, color='FFFFFF')
    header_fill = PatternFill('solid', fgColor='333333')
    label_font = Font(name='Arial', bold=True, size=10)
    value_font = Font(name='Arial', size=10)
    euro_fmt = '#,##0.00 €'
    thin_border = Border(
        left=Side(style='thin'), right=Side(style='thin'),
        top=Side(style='thin'), bottom=Side(style='thin')
    )

    numero = getattr(preventivo, 'numero', None) or getattr(preventivo, 'numero_preventivo', '') or ''
    rev = getattr(preventivo, 'revisione_corrente', 0) or 0
    if rev > 0:
        numero = f"{numero} REV.{rev}"
    customer = getattr(preventivo, 'customer_name', '') or ''
    if cliente:
        customer = getattr(cliente, 'ragione_sociale', customer) or customer

    # ── FOGLIO 1: RIEPILOGO ──
    ws = wb.active
    ws.title = "Riepilogo"
    ws.sheet_properties.tabColor = "CC0000"

    ws.merge_cells('A1:D1')
    ws['A1'] = 'ELETTROQUADRI S.r.l. — Preventivo'
    ws['A1'].font = Font(name='Arial', bold=True, size=14, color='CC0000')

    info = [
        ('Numero Preventivo', safe_str(numero)),
        ('Cliente', customer),
        ('Stato', safe_str(getattr(preventivo, 'stato', '') or getattr(preventivo, 'status', 'bozza'))),
        ('Data Offerta', safe_str(getattr(dati_commessa, 'data_offerta', '') if dati_commessa else '')),
        ('Riferimento', safe_str(getattr(dati_commessa, 'riferimento_cliente', '') if dati_commessa else '')),
        ('Pagamento', safe_str(getattr(dati_commessa, 'pagamento', '') if dati_commessa else '')),
        ('Trasporto', safe_str(getattr(dati_commessa, 'trasporto', '') if dati_commessa else '')),
    ]
    for i, (label, val) in enumerate(info, start=3):
        ws.cell(row=i, column=1, value=label).font = label_font
        ws.cell(row=i, column=2, value=val).font = value_font

    row = len(info) + 4
    totale = safe_float(getattr(preventivo, 'totale_materiali', None) or getattr(preventivo, 'total_price', 0))
    sconto = safe_float(getattr(preventivo, 'sconto_percentuale', 0))
    netto = safe_float(getattr(preventivo, 'totale_netto', None) or getattr(preventivo, 'total_price_finale', 0))

    ws.cell(row=row, column=1, value='Totale Materiali').font = Font(name='Arial', bold=True, size=12)
    c = ws.cell(row=row, column=2, value=totale)
    c.number_format = euro_fmt
    c.font = Font(name='Arial', bold=True, size=12)
    if sconto > 0:
        row += 1
        ws.cell(row=row, column=1, value='Sconto').font = label_font
        ws.cell(row=row, column=2, value=f'{sconto}%').font = value_font
    if netto > 0 and netto != totale:
        row += 1
        ws.cell(row=row, column=1, value='Totale Netto').font = Font(name='Arial', bold=True, size=12, color='CC0000')
        c = ws.cell(row=row, column=2, value=netto)
        c.number_format = euro_fmt
        c.font = Font(name='Arial', bold=True, size=12, color='CC0000')

    ws.column_dimensions['A'].width = 25
    ws.column_dimensions['B'].width = 30

    # ── FOGLIO 2: MATERIALI ──
    ws2 = wb.create_sheet("Materiali")
    ws2.sheet_properties.tabColor = "006600"

    headers = ['Codice', 'Descrizione', 'Categoria', 'Qtà', 'Prezzo Unit.', 'Totale', 'Origine']
    for i, h in enumerate(headers, start=1):
        c = ws2.cell(row=1, column=i, value=h)
        c.font = header_font
        c.fill = header_fill
        c.alignment = Alignment(horizontal='center')
        c.border = thin_border

    for idx, m in enumerate(materiali, start=2):
        regola_id = safe_str(getattr(m, 'regola_id', ''))
        if m.aggiunto_da_regola and regola_id.startswith('TEMPLATE_BASE'):
            origine = 'Template'
        elif m.aggiunto_da_regola:
            origine = 'Auto (Regola)'
        else:
            origine = 'Manuale'

        vals = [
            safe_str(m.codice), safe_str(m.descrizione), safe_str(m.categoria),
            m.quantita, m.prezzo_unitario, m.prezzo_totale, origine
        ]
        for j, v in enumerate(vals, start=1):
            c = ws2.cell(row=idx, column=j, value=v)
            c.font = value_font
            c.border = thin_border
            if j in (5, 6):
                c.number_format = euro_fmt

        # Colora righe automatiche
        if m.aggiunto_da_regola:
            fill_color = 'DBEAFE' if regola_id.startswith('TEMPLATE_BASE') else 'DCFCE7'
            for j in range(1, 8):
                ws2.cell(row=idx, column=j).fill = PatternFill('solid', fgColor=fill_color)

    # Riga totale
    last_row = len(materiali) + 2
    ws2.cell(row=last_row, column=5, value='TOTALE').font = Font(name='Arial', bold=True, size=11)
    totale_formula = f'=SUM(F2:F{last_row - 1})'
    c = ws2.cell(row=last_row, column=6, value=totale_formula)
    c.font = Font(name='Arial', bold=True, size=11)
    c.number_format = euro_fmt

    for col, w in [('A', 18), ('B', 40), ('C', 15), ('D', 8), ('E', 14), ('F', 14), ('G', 14)]:
        ws2.column_dimensions[col].width = w

    # ── FOGLIO 3: DATI TECNICI ──
    ws3 = wb.create_sheet("Dati Tecnici")
    ws3.sheet_properties.tabColor = "0066CC"

    row = 1
    ws3.cell(row=row, column=1, value='DATI PRINCIPALI').font = Font(name='Arial', bold=True, size=12, color='0066CC')
    row += 1
    if dati_principali:
        for label, attr in [
            ('Tipo impianto', 'tipo_impianto'), ('N° Fermate', 'numero_fermate'),
            ('N° Servizi', 'numero_servizi'), ('Velocità (m/s)', 'velocita'),
            ('Corsa (m)', 'corsa'), ('Tipo Trazione', 'tipo_trazione'),
            ('Forza Motrice', 'forza_motrice'), ('Luce', 'luce'),
            ('Tensione Manovra', 'tensione_manovra'), ('Tensione Freno', 'tensione_freno'),
        ]:
            ws3.cell(row=row, column=1, value=label).font = label_font
            ws3.cell(row=row, column=2, value=safe_str(getattr(dati_principali, attr, ''))).font = value_font
            row += 1

    row += 1
    ws3.cell(row=row, column=1, value='ARGANO / MOTORE').font = Font(name='Arial', bold=True, size=12, color='0066CC')
    row += 1
    if argano:
        for label, attr in [
            ('Trazione', 'trazione'), ('Potenza Motore (kW)', 'potenza_motore_kw'),
            ('Corrente Nominale (A)', 'corrente_nom_motore_amp'), ('Tipo VVVF', 'tipo_vvvf'),
        ]:
            ws3.cell(row=row, column=1, value=label).font = label_font
            ws3.cell(row=row, column=2, value=safe_str(getattr(argano, attr, ''))).font = value_font
            row += 1

    row += 1
    ws3.cell(row=row, column=1, value='NORMATIVE').font = Font(name='Arial', bold=True, size=12, color='0066CC')
    row += 1
    if normative:
        for n in ['en_81_1', 'en_81_20', 'en_81_21', 'en_81_28', 'en_81_70',
                   'en_81_72', 'en_81_73', 'a3_95_16', 'dm236_legge13']:
            val = getattr(normative, n, False)
            ws3.cell(row=row, column=1, value=n.replace('_', ' ').upper()).font = value_font
            ws3.cell(row=row, column=2, value='✓' if val else '').font = value_font
            if val:
                ws3.cell(row=row, column=2).fill = PatternFill('solid', fgColor='DCFCE7')
            row += 1

    ws3.column_dimensions['A'].width = 25
    ws3.column_dimensions['B'].width = 30

    buf = io.BytesIO()
    wb.save(buf)
    buf.seek(0)
    return buf


# ═══════════════════════════════════════════════════════
# DOCX: ORDINE
# ═══════════════════════════════════════════════════════
def genera_docx_ordine(ordine_data, materiali, esplosi=None, lista_acquisti=None):
    """
    Genera un documento DOCX per l'ordine con BOM e lista acquisti.
    ordine_data: dict dal DB, materiali: lista Materiale ORM, esplosi: lista dict, lista_acquisti: dict fornitori
    """
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT

    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(10)

    # ── INTESTAZIONE ──
    h = doc.add_heading('ELETTROQUADRI S.r.l.', level=1)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in h.runs:
        run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)

    doc.add_heading(f'Ordine {ordine_data.get("numero_ordine", "")}', level=2)

    # ── INFO ORDINE ──
    t = doc.add_table(rows=5, cols=4)
    t.style = 'Table Grid'
    info_rows = [
        ('N° Ordine', safe_str(ordine_data.get('numero_ordine')), 'Stato', safe_str(ordine_data.get('stato', '')).upper()),
        ('Cliente', safe_str(ordine_data.get('cliente', '')), 'Tipo Impianto', safe_str(ordine_data.get('tipo_impianto', ''))),
        ('Totale Materiali', fmt_euro(ordine_data.get('totale_materiali', 0)), 'Totale Netto', fmt_euro(ordine_data.get('totale_netto', 0))),
        ('Lead Time', f"{ordine_data.get('lead_time_giorni', 15)} giorni", 'Consegna Prevista', safe_str(ordine_data.get('data_consegna_prevista', '')[:10] if ordine_data.get('data_consegna_prevista') else '')),
        ('Data Creazione', safe_str(ordine_data.get('created_at', '')[:10] if ordine_data.get('created_at') else ''), '', ''),
    ]
    for i, (l1, v1, l2, v2) in enumerate(info_rows):
        t.cell(i, 0).text = l1
        t.cell(i, 1).text = v1
        t.cell(i, 2).text = l2
        t.cell(i, 3).text = v2
        for j in [0, 2]:
            for run in t.cell(i, j).paragraphs[0].runs:
                run.font.bold = True
                run.font.size = Pt(9)

    # ── MATERIALI PREVENTIVO ──
    if materiali:
        doc.add_heading('Materiali Preventivo', level=2)
        t2 = doc.add_table(rows=1, cols=5)
        t2.style = 'Table Grid'
        for i, h_text in enumerate(['Codice', 'Descrizione', 'Qtà', 'Prezzo Unit.', 'Totale']):
            t2.rows[0].cells[i].text = h_text
            for run in t2.rows[0].cells[i].paragraphs[0].runs:
                run.font.bold = True
                run.font.size = Pt(9)
        _set_repeat_header_row(t2)

        for m in materiali:
            row = t2.add_row().cells
            row[0].text = safe_str(m.codice)
            row[1].text = safe_str(m.descrizione)
            row[2].text = safe_str(m.quantita)
            row[3].text = fmt_euro(m.prezzo_unitario)
            row[4].text = fmt_euro(m.prezzo_totale)

    # ── BOM ESPLOSA ──
    if esplosi:
        doc.add_heading('Distinta Base Esplosa', level=2)
        doc.add_paragraph(f'Totale componenti: {len(esplosi)}')

        # Raggruppa per tipo
        by_tipo = {}
        for e in esplosi:
            tipo = e.get('tipo', 'ALTRO')
            if tipo not in by_tipo:
                by_tipo[tipo] = []
            by_tipo[tipo].append(e)

        for tipo, items in by_tipo.items():
            doc.add_heading(f'{tipo} ({len(items)} componenti)', level=3)
            t3 = doc.add_table(rows=1, cols=5)
            t3.style = 'Table Grid'
            for i, h_text in enumerate(['Codice', 'Descrizione', 'Qtà', 'UM', 'Costo Tot.']):
                t3.rows[0].cells[i].text = h_text
                for run in t3.rows[0].cells[i].paragraphs[0].runs:
                    run.font.bold = True
                    run.font.size = Pt(9)
            _set_repeat_header_row(t3)

            for e in items:
                row = t3.add_row().cells
                row[0].text = safe_str(e.get('codice'))
                row[1].text = safe_str(e.get('descrizione'))
                row[2].text = safe_str(e.get('quantita'))
                row[3].text = safe_str(e.get('unita_misura', 'pz'))
                row[4].text = fmt_euro(e.get('costo_totale', 0))

    # ── LISTA ACQUISTI ──
    if lista_acquisti and lista_acquisti.get('fornitori'):
        doc.add_heading('Lista Acquisti per Fornitore', level=2)
        for fornitore, data in lista_acquisti['fornitori'].items():
            doc.add_heading(f'{fornitore} — {fmt_euro(data["totale"])}', level=3)
            t4 = doc.add_table(rows=1, cols=5)
            t4.style = 'Table Grid'
            for i, h_text in enumerate(['Codice', 'Descrizione', 'Qtà', 'Costo Unit.', 'Costo Tot.']):
                t4.rows[0].cells[i].text = h_text
                for run in t4.rows[0].cells[i].paragraphs[0].runs:
                    run.font.bold = True
                    run.font.size = Pt(9)
            _set_repeat_header_row(t4)
            for art in data['articoli']:
                row = t4.add_row().cells
                row[0].text = safe_str(art.get('codice'))
                row[1].text = safe_str(art.get('descrizione'))
                row[2].text = safe_str(art.get('quantita'))
                row[3].text = fmt_euro(art.get('costo_unitario', 0))
                row[4].text = fmt_euro(art.get('costo_totale', 0))

    # ── FOOTER ──
    doc.add_paragraph('')
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run(f'Documento generato il {datetime.now().strftime("%d/%m/%Y %H:%M")}').font.size = Pt(8)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


# ═══════════════════════════════════════════════════════
# XLSX: ORDINE
# ═══════════════════════════════════════════════════════
def genera_xlsx_ordine(ordine_data, materiali, esplosi=None, lista_acquisti=None):
    """
    Genera file XLSX ordine con fogli: Riepilogo, Materiali, BOM Esplosa, Lista Acquisti.
    """
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side

    wb = Workbook()
    hdr_font = Font(name='Arial', bold=True, size=10, color='FFFFFF')
    hdr_fill = PatternFill('solid', fgColor='333333')
    lbl_font = Font(name='Arial', bold=True, size=10)
    val_font = Font(name='Arial', size=10)
    euro_fmt = '#,##0.00 €'
    thin = Border(left=Side(style='thin'), right=Side(style='thin'),
                  top=Side(style='thin'), bottom=Side(style='thin'))

    # ── FOGLIO 1: RIEPILOGO ──
    ws = wb.active
    ws.title = "Riepilogo Ordine"
    ws.sheet_properties.tabColor = "CC0000"

    ws.merge_cells('A1:C1')
    ws['A1'] = f'Ordine {ordine_data.get("numero_ordine", "")}'
    ws['A1'].font = Font(name='Arial', bold=True, size=14, color='CC0000')

    info = [
        ('N° Ordine', safe_str(ordine_data.get('numero_ordine'))),
        ('Cliente', safe_str(ordine_data.get('cliente', ''))),
        ('Stato', safe_str(ordine_data.get('stato', '')).upper()),
        ('Tipo Impianto', safe_str(ordine_data.get('tipo_impianto', ''))),
        ('Totale Materiali', safe_float(ordine_data.get('totale_materiali', 0))),
        ('Totale Netto', safe_float(ordine_data.get('totale_netto', 0))),
        ('Lead Time', f"{ordine_data.get('lead_time_giorni', 15)} giorni"),
        ('Consegna Prevista', safe_str(ordine_data.get('data_consegna_prevista', '')[:10] if ordine_data.get('data_consegna_prevista') else '')),
    ]
    for i, (label, val) in enumerate(info, start=3):
        ws.cell(row=i, column=1, value=label).font = lbl_font
        c = ws.cell(row=i, column=2, value=val)
        c.font = val_font
        if isinstance(val, float):
            c.number_format = euro_fmt
    ws.column_dimensions['A'].width = 22
    ws.column_dimensions['B'].width = 30

    # ── FOGLIO 2: MATERIALI ──
    ws2 = wb.create_sheet("Materiali")
    ws2.sheet_properties.tabColor = "006600"
    mat_headers = ['Codice', 'Descrizione', 'Qtà', 'Prezzo Unit.', 'Totale']
    for i, h in enumerate(mat_headers, 1):
        c = ws2.cell(row=1, column=i, value=h)
        c.font = hdr_font
        c.fill = hdr_fill
        c.border = thin
    for idx, m in enumerate(materiali, 2):
        vals = [safe_str(m.codice), safe_str(m.descrizione), m.quantita, m.prezzo_unitario, m.prezzo_totale]
        for j, v in enumerate(vals, 1):
            c = ws2.cell(row=idx, column=j, value=v)
            c.font = val_font
            c.border = thin
            if j in (4, 5):
                c.number_format = euro_fmt
    for col, w in [('A', 18), ('B', 40), ('C', 8), ('D', 14), ('E', 14)]:
        ws2.column_dimensions[col].width = w

    # ── FOGLIO 3: BOM ESPLOSA ──
    if esplosi:
        ws3 = wb.create_sheet("BOM Esplosa")
        ws3.sheet_properties.tabColor = "0066CC"
        bom_headers = ['Codice', 'Descrizione', 'Tipo', 'Categoria', 'Qtà', 'UM', 'Costo Unit.', 'Costo Tot.']
        for i, h in enumerate(bom_headers, 1):
            c = ws3.cell(row=1, column=i, value=h)
            c.font = hdr_font
            c.fill = hdr_fill
            c.border = thin

        tipo_colors = {
            'MASTER': 'DBEAFE',
            'SEMILAVORATO': 'FEF3C7',
            'ACQUISTO': 'DCFCE7',
        }
        for idx, e in enumerate(esplosi, 2):
            tipo = e.get('tipo', '')
            vals = [
                safe_str(e.get('codice')), safe_str(e.get('descrizione')), tipo,
                safe_str(e.get('categoria', '')), e.get('quantita', 0),
                safe_str(e.get('unita_misura', 'pz')),
                e.get('costo_unitario', 0), e.get('costo_totale', 0)
            ]
            fill_color = tipo_colors.get(tipo, '')
            for j, v in enumerate(vals, 1):
                c = ws3.cell(row=idx, column=j, value=v)
                c.font = val_font
                c.border = thin
                if j in (7, 8):
                    c.number_format = euro_fmt
                if fill_color:
                    c.fill = PatternFill('solid', fgColor=fill_color)

        last = len(esplosi) + 2
        ws3.cell(row=last, column=7, value='TOTALE').font = lbl_font
        c = ws3.cell(row=last, column=8, value=f'=SUM(H2:H{last - 1})')
        c.font = Font(name='Arial', bold=True, size=11)
        c.number_format = euro_fmt

        for col, w in [('A', 20), ('B', 40), ('C', 14), ('D', 14), ('E', 8), ('F', 8), ('G', 14), ('H', 14)]:
            ws3.column_dimensions[col].width = w

    # ── FOGLIO 4: LISTA ACQUISTI ──
    if lista_acquisti and lista_acquisti.get('fornitori'):
        ws4 = wb.create_sheet("Lista Acquisti")
        ws4.sheet_properties.tabColor = "FF6600"

        acq_headers = ['Fornitore', 'Codice', 'Descrizione', 'Qtà', 'UM', 'Costo Unit.', 'Costo Tot.', 'Cod. Fornitore', 'Lead Time (gg)']
        for i, h in enumerate(acq_headers, 1):
            c = ws4.cell(row=1, column=i, value=h)
            c.font = hdr_font
            c.fill = hdr_fill
            c.border = thin

        row = 2
        for fornitore, data in lista_acquisti['fornitori'].items():
            for art in data['articoli']:
                vals = [
                    fornitore, safe_str(art.get('codice')), safe_str(art.get('descrizione')),
                    art.get('quantita', 0), safe_str(art.get('unita', 'pz')),
                    art.get('costo_unitario', 0), art.get('costo_totale', 0),
                    safe_str(art.get('codice_fornitore', '')), art.get('lead_time_giorni')
                ]
                for j, v in enumerate(vals, 1):
                    c = ws4.cell(row=row, column=j, value=v)
                    c.font = val_font
                    c.border = thin
                    if j in (6, 7):
                        c.number_format = euro_fmt
                row += 1

        for col, w in [('A', 20), ('B', 18), ('C', 35), ('D', 8), ('E', 8), ('F', 14), ('G', 14), ('H', 16), ('I', 12)]:
            ws4.column_dimensions[col].width = w

    buf = io.BytesIO()
    wb.save(buf)
    buf.seek(0)
    return buf
