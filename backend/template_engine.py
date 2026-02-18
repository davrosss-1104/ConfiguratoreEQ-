"""
template_engine.py - Motore template documenti per Preventivi e Ordini
Genera DOCX a partire da una configurazione JSON salvata nel DB.
Sezioni e campi dinamici letti da sezioni_configuratore + campi_configuratore.

Posizionare in: backend/template_engine.py (stesso livello di main.py)
Dipendenze: pip install python-docx --break-system-packages
"""
import io
import json
import base64
from datetime import datetime
from typing import Optional, Dict, Any, List

from sqlalchemy import Column, Integer, String, Boolean, DateTime, Text, LargeBinary, JSON, text
from sqlalchemy.sql import func
from database import Base


# ═══════════════════════════════════════════════════════
# MODEL
# ═══════════════════════════════════════════════════════

class DocumentTemplate(Base):
    __tablename__ = "document_templates"
    id = Column(Integer, primary_key=True, index=True)
    nome = Column(String(200), nullable=False)
    tipo = Column(String(50), nullable=False)        # "preventivo" | "ordine"
    descrizione = Column(Text)
    config = Column(JSON, nullable=False)
    logo_data = Column(LargeBinary)
    logo_filename = Column(String(200))
    logo_mime = Column(String(50))
    attivo = Column(Boolean, default=True)
    is_default = Column(Boolean, default=False)
    created_at = Column(DateTime(timezone=True), server_default=func.now())
    updated_at = Column(DateTime(timezone=True), onupdate=func.now())


# ═══════════════════════════════════════════════════════
# SEZIONI STATICHE (strutturali, non dal DB)
# ═══════════════════════════════════════════════════════

STATIC_SECTIONS = {
    "intestazione": {
        "label": "Intestazione",
        "icon": "Building2",
        "type": "header",
        "fields": [
            {"key": "company_name", "label": "Ragione Sociale Azienda", "default": "ELETTROQUADRI S.r.l."},
            {"key": "company_address", "label": "Indirizzo Azienda", "default": "Via Puccini, 1"},
            {"key": "company_phone", "label": "Telefono Azienda", "default": "Tel. 0332 470049"},
            {"key": "company_email", "label": "Email Azienda", "default": "info@elettroquadri.net"},
            {"key": "company_web", "label": "Sito Web", "default": "www.elettroquadri.net"},
        ]
    },
    "materiali": {
        "label": "Tabella Materiali",
        "icon": "Table",
        "type": "materials_table",
        "columns": [
            {"key": "codice", "label": "Codice", "width_pct": 15},
            {"key": "descrizione", "label": "Descrizione", "width_pct": 40},
            {"key": "categoria", "label": "Categoria", "width_pct": 10},
            {"key": "quantita", "label": "Qta", "width_pct": 8},
            {"key": "prezzo_unitario", "label": "Prezzo Unit.", "width_pct": 13, "format": "euro"},
            {"key": "prezzo_totale", "label": "Totale", "width_pct": 14, "format": "euro"},
        ]
    },
    "riepilogo_prezzi": {
        "label": "Riepilogo Prezzi",
        "icon": "Calculator",
        "type": "price_summary",
        "fields": [
            {"key": "totale_materiali", "label": "Totale Materiali", "source": "preventivo.totale_materiali", "format": "euro"},
            {"key": "totale_manodopera", "label": "Manodopera", "source": "preventivo.totale_manodopera", "format": "euro"},
            {"key": "totale_trasporto", "label": "Trasporto", "source": "preventivo.totale_trasporto", "format": "euro"},
            {"key": "sconto_percentuale", "label": "Sconto %", "source": "preventivo.sconto_percentuale", "format": "percent"},
            {"key": "totale_netto", "label": "Totale Netto", "source": "preventivo.totale_netto", "format": "euro"},
            {"key": "totale_iva", "label": "IVA", "source": "preventivo.totale_iva", "format": "euro"},
            {"key": "totale_lordo", "label": "Totale Lordo", "source": "preventivo.totale_lordo", "format": "euro"},
        ]
    },
    "note": {
        "label": "Note",
        "icon": "MessageSquare",
        "type": "text_block",
        "fields": [
            {"key": "note_cliente", "label": "Note per il Cliente", "source": "preventivo.note_cliente"},
            {"key": "note_interne", "label": "Note Interne", "source": "preventivo.note_interne"},
        ]
    },
    "footer": {
        "label": "Pie di Pagina",
        "icon": "AlignBottom",
        "type": "footer",
        "fields": [
            {"key": "data_generazione", "label": "Data Generazione", "computed": True},
            {"key": "numero_pagina", "label": "Numero Pagina", "computed": True},
            {"key": "testo_libero_footer", "label": "Testo Personalizzato", "editable": True},
        ]
    },
}

# Icona di default per sezioni dal DB in base al codice
_ICON_HINTS = {
    "dati_commessa": "FileText", "dati_offerta": "FileText",
    "dati_principali": "Settings", "dati_cliente": "User", "cliente": "User",
    "argano": "Cog", "normative": "Shield", "porte": "DoorOpen",
    "disposizione_vano": "Layout", "disposizione": "Layout",
}


# ═══════════════════════════════════════════════════════
# BUILDER DINAMICO — Legge sezioni/campi dal DB
# ═══════════════════════════════════════════════════════

def get_available_fields_from_db(db) -> Dict[str, Any]:
    """
    Costruisce il dizionario AVAILABLE_FIELDS unendo:
    - sezioni statiche (intestazione, materiali, prezzi, note, footer)
    - sezioni dinamiche da sezioni_configuratore + campi_configuratore
    """
    result = dict(STATIC_SECTIONS)

    try:
        # Leggi sezioni dal DB
        sezioni_rows = db.execute(text("""
            SELECT id, codice, etichetta, descrizione, icona, ordine, attivo
            FROM sezioni_configuratore
            WHERE attivo = 1
            ORDER BY ordine
        """)).fetchall()
    except Exception:
        # Tabella non esiste o errore: ritorna solo le statiche
        return result

    for srow in sezioni_rows:
        s_id, s_codice, s_etichetta, s_desc, s_icona, s_ordine, s_attivo = srow

        # Non sovrascrivere sezioni statiche strutturali
        if s_codice in ("materiali", "riepilogo_prezzi", "footer", "intestazione"):
            continue

        # Leggi campi di questa sezione
        try:
            campi_rows = db.execute(text("""
                SELECT codice, etichetta, tipo, unita_misura
                FROM campi_configuratore
                WHERE sezione = :sez AND attivo = 1
                ORDER BY ordine
            """), {"sez": s_codice}).fetchall()
        except Exception:
            campi_rows = []

        fields = []
        for crow in campi_rows:
            c_codice, c_etichetta, c_tipo, c_unita = crow
            field = {
                "key": c_codice,
                "label": c_etichetta or c_codice,
                # source dinamica: il resolver cerca in valori_configurazione
                "source": f"dinamico.{c_codice}",
                "tipo_campo": c_tipo,  # testo, numero, booleano, dropdown, data
            }
            if c_unita:
                field["unita_misura"] = c_unita
            if c_tipo == "numero":
                field["format"] = "number"
            elif c_tipo == "booleano":
                field["format"] = "boolean"
            fields.append(field)

        # Determina icona
        icon = s_icona or _ICON_HINTS.get(s_codice, "FileText")

        # Tipo sezione: se ha campi booleani potrebbe essere normative-like
        bool_count = sum(1 for f in fields if f.get("tipo_campo") == "booleano")
        if bool_count > 0 and bool_count == len(fields):
            sec_type = "normative_list"
        else:
            sec_type = "key_value"

        result[s_codice] = {
            "label": s_etichetta or s_codice,
            "icon": icon,
            "type": sec_type,
            "fields": fields,
            "db_ordine": s_ordine,  # per ordinamento default
            "dynamic": True,        # flag per distinguerle
        }

    return result


def get_default_template_config_from_db(db, tipo="preventivo") -> Dict[str, Any]:
    """
    Costruisce config default usando le sezioni dal DB.
    Ordine: intestazione, [sezioni dinamiche ordinate], materiali, riepilogo, note, footer
    """
    available = get_available_fields_from_db(db)

    # Separa statiche e dinamiche
    static_start = ["intestazione"]
    static_end_prev = ["materiali", "riepilogo_prezzi", "note", "footer"]
    static_end_ord = ["materiali", "riepilogo_prezzi", "footer"]
    static_end = static_end_prev if tipo == "preventivo" else static_end_ord

    # Sezioni dinamiche ordinate per db_ordine
    dynamic_ids = [
        k for k, v in available.items()
        if v.get("dynamic") and k not in static_start and k not in static_end
    ]
    dynamic_ids.sort(key=lambda k: available[k].get("db_ordine", 999))

    section_order = static_start + dynamic_ids + static_end

    sections = []
    for order, section_id in enumerate(section_order):
        if section_id not in available:
            continue
        meta = available[section_id]
        section = {
            "id": section_id,
            "type": meta["type"],
            "title": meta["label"],
            "enabled": True,
            "order": order,
        }
        if "fields" in meta:
            section["fields"] = [
                {"key": f["key"], "label": f["label"], "enabled": True}
                for f in meta["fields"]
            ]
        if "columns" in meta:
            section["columns"] = [dict(c) for c in meta["columns"]]
            section["show_totals"] = True
        sections.append(section)

    return {
        "page": {"size": "A4", "orientation": "portrait",
                 "margins": {"top": 2.0, "bottom": 2.0, "left": 2.5, "right": 2.0}},
        "style": {"font": "Arial", "font_size": 10, "heading_color": "#CC0000",
                  "table_header_bg": "#333333", "table_header_text": "#FFFFFF",
                  "table_alt_row_bg": "#F9F9F9"},
        "sections": sections
    }


# ═══════════════════════════════════════════════════════
# DATA RESOLVER — Ora legge anche valori_configurazione
# ═══════════════════════════════════════════════════════

def _resolve_value(source, data_context):
    """
    Risolve un valore. Supporta:
    - 'preventivo.numero'         -> ORM object / dict
    - 'cliente.ragione_sociale'   -> ORM object / dict
    - 'dinamico.codice_campo'     -> cerca in valori_configurazione
    """
    parts = source.split(".", 1)
    if len(parts) != 2:
        return ""
    entity, field = parts

    if entity == "dinamico":
        # Cerca nei valori dinamici caricati nel contesto
        dyn = data_context.get("_valori_dinamici", {})
        return dyn.get(field, "")

    obj = data_context.get(entity)
    if obj is None:
        return ""
    if isinstance(obj, dict):
        return obj.get(field, "")
    return getattr(obj, field, "")


def _build_data_context(preventivo, dati_commessa, dati_principali,
                         normative, argano, materiali, cliente,
                         valori_dinamici=None):
    """Costruisce contesto dati con supporto valori dinamici"""
    ctx = {
        "preventivo": preventivo,
        "dati_commessa": dati_commessa,
        "dati_principali": dati_principali,
        "normative": normative,
        "argano": argano,
        "materiali": materiali,
        "cliente": cliente,
    }
    if valori_dinamici:
        ctx["_valori_dinamici"] = valori_dinamici
    return ctx


def load_valori_dinamici(db, preventivo_id: int) -> Dict[str, str]:
    """Carica tutti i valori da valori_configurazione per un preventivo"""
    try:
        result = db.execute(
            text("SELECT codice_campo, valore FROM valori_configurazione WHERE preventivo_id = :pid"),
            {"pid": preventivo_id}
        )
        return {row[0]: row[1] for row in result.fetchall() if row[0] and row[1] is not None}
    except Exception:
        return {}


# ═══════════════════════════════════════════════════════
# FORMATTAZIONE
# ═══════════════════════════════════════════════════════

def _fmt_euro(val):
    if val is None: return "€ 0,00"
    try:
        v = float(val)
        return f"€ {v:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")
    except (TypeError, ValueError):
        return "€ 0,00"

def _fmt_percent(val):
    if val is None: return "0%"
    try: return f"{float(val):.1f}%"
    except (TypeError, ValueError): return "0%"

def _safe_str(val, default=""):
    if val is None: return default
    if isinstance(val, bool): return "Si" if val else "No"
    return str(val)

def _format_value(val, fmt=None):
    if fmt == "euro": return _fmt_euro(val)
    elif fmt == "percent": return _fmt_percent(val)
    elif fmt == "boolean":
        if isinstance(val, str):
            return "Si" if val.lower() in ("1", "true", "si", "yes") else "No"
        return "Si" if val else "No"
    elif fmt == "number":
        try: return str(float(val)) if val else ""
        except: return _safe_str(val)
    return _safe_str(val)


# ═══════════════════════════════════════════════════════
# DOCX GENERATOR
# ═══════════════════════════════════════════════════════

def genera_docx_da_template(template_config, preventivo, dati_commessa, dati_principali,
                             normative, argano, materiali, cliente,
                             logo_data=None, logo_mime=None,
                             valori_dinamici=None, available_fields=None):
    """
    Genera un DOCX professionale basato sulla configurazione del template.
    
    Args:
        available_fields: dizionario AVAILABLE_FIELDS (statico+dinamico).
                          Se None, usa solo STATIC_SECTIONS (no risoluzione source dinamiche).
        valori_dinamici:  dict {codice_campo: valore} da valori_configurazione.
    """
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn

    doc = Document()
    ctx = _build_data_context(
        preventivo, dati_commessa, dati_principali,
        normative, argano, materiali, cliente,
        valori_dinamici=valori_dinamici
    )

    af = available_fields or STATIC_SECTIONS
    config = template_config
    style_cfg = config.get("style", {})
    page_cfg = config.get("page", {})

    font_name = style_cfg.get("font", "Arial")
    font_size = style_cfg.get("font_size", 10)
    hc_hex = style_cfg.get("heading_color", "#CC0000").lstrip("#")
    th_bg = style_cfg.get("table_header_bg", "#333333").lstrip("#")
    th_txt = style_cfg.get("table_header_text", "#FFFFFF").lstrip("#")

    normal_style = doc.styles['Normal']
    normal_style.font.name = font_name
    normal_style.font.size = Pt(font_size)

    margins = page_cfg.get("margins", {})
    sec0 = doc.sections[0]
    sec0.top_margin = Cm(margins.get("top", 2.0))
    sec0.bottom_margin = Cm(margins.get("bottom", 2.0))
    sec0.left_margin = Cm(margins.get("left", 2.5))
    sec0.right_margin = Cm(margins.get("right", 2.0))

    def _hc():
        return RGBColor(int(hc_hex[0:2], 16), int(hc_hex[2:4], 16), int(hc_hex[4:6], 16))

    def _set_shading(cell, color):
        sh = cell._element.get_or_add_tcPr()
        el = sh.makeelement(qn('w:shd'), {qn('w:fill'): color, qn('w:val'): 'clear'})
        sh.append(el)

    # Lookup globale dei field per key (da available_fields)
    _fl = {}
    for sid, meta in af.items():
        for f in meta.get("fields", []):
            _fl[f["key"]] = f

    sorted_secs = sorted(config.get("sections", []), key=lambda s: s.get("order", 0))

    for sec in sorted_secs:
        if not sec.get("enabled", True):
            continue
        st = sec.get("type", "")
        title = sec.get("title", "")
        fields = sec.get("fields", [])

        # ─────── HEADER ───────
        if st == "header":
            if logo_data:
                try:
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = p.add_run()
                    run.add_picture(io.BytesIO(logo_data), width=Cm(5))
                except Exception: pass

            ef = {f["key"]: f for f in fields if f.get("enabled", True)}
            if "company_name" in ef:
                cn = ef["company_name"].get("value") or _fl.get("company_name", {}).get("default", "")
                if cn:
                    h = doc.add_heading(cn, level=1)
                    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in h.runs: run.font.color.rgb = _hc()

            parts = []
            for fk in ["company_address", "company_phone", "company_email", "company_web"]:
                if fk in ef:
                    v = ef[fk].get("value") or _fl.get(fk, {}).get("default", "")
                    if v: parts.append(v)
            if parts:
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.add_run(" - ".join(parts)).font.size = Pt(8)
            doc.add_paragraph('')

        # ─────── KEY_VALUE ───────
        elif st == "key_value":
            efl = [f for f in fields if f.get("enabled", True)]
            if not efl: continue
            rows = []
            for f in efl:
                lookup = _fl.get(f["key"], {})
                src = lookup.get("source", "")
                fmt = lookup.get("format")
                val = _resolve_value(src, ctx) if src else ""
                rows.append((f.get("label", f["key"]), _format_value(val, fmt)))
            if not any(v for _, v in rows): continue

            doc.add_heading(title, level=2)
            if len(rows) <= 6:
                t = doc.add_table(rows=len(rows), cols=2)
                t.style = 'Table Grid'
                for i, (l, v) in enumerate(rows):
                    t.cell(i, 0).text = l
                    t.cell(i, 1).text = v
                    for r in t.cell(i, 0).paragraphs[0].runs: r.font.bold = True; r.font.size = Pt(font_size - 1)
                    for r in t.cell(i, 1).paragraphs[0].runs: r.font.size = Pt(font_size - 1)
            else:
                half = (len(rows) + 1) // 2
                t = doc.add_table(rows=half, cols=4)
                t.style = 'Table Grid'
                for i in range(half):
                    l1, v1 = rows[i]
                    t.cell(i, 0).text = l1; t.cell(i, 1).text = v1
                    if i + half < len(rows):
                        l2, v2 = rows[i + half]
                        t.cell(i, 2).text = l2; t.cell(i, 3).text = v2
                    for j in [0, 2]:
                        for r in t.cell(i, j).paragraphs[0].runs: r.font.bold = True; r.font.size = Pt(font_size - 1)
                    for j in [1, 3]:
                        for r in t.cell(i, j).paragraphs[0].runs: r.font.size = Pt(font_size - 1)
            doc.add_paragraph('')

        # ─────── NORMATIVE ───────
        elif st == "normative_list":
            efl = [f for f in fields if f.get("enabled", True)]
            norms = []
            for f in efl:
                lookup = _fl.get(f["key"], {})
                src = lookup.get("source", "")
                if src:
                    val = _resolve_value(src, ctx)
                    if val and val not in (False, "False", "0", "", None, "No", "no"):
                        norms.append(f.get("label", f["key"]))
            if norms:
                doc.add_heading(title, level=2)
                doc.add_paragraph(", ".join(norms))
                doc.add_paragraph('')

        # ─────── TABELLA MATERIALI ───────
        elif st == "materials_table":
            ml = ctx.get("materiali") or []
            if not ml: continue
            doc.add_heading(title, level=2)
            cols = sec.get("columns", STATIC_SECTIONS["materiali"]["columns"])
            ecols = [c for c in cols if c.get("enabled", True)]
            if not ecols: ecols = cols

            t = doc.add_table(rows=1, cols=len(ecols))
            t.style = 'Table Grid'
            t.alignment = WD_TABLE_ALIGNMENT.CENTER
            for i, col in enumerate(ecols):
                cell = t.rows[0].cells[i]
                cell.text = col["label"]
                for r in cell.paragraphs[0].runs:
                    r.font.bold = True; r.font.size = Pt(font_size - 1)
                    r.font.color.rgb = RGBColor(int(th_txt[0:2],16), int(th_txt[2:4],16), int(th_txt[4:6],16))
                _set_shading(cell, th_bg)

            totale = 0.0
            for m in ml:
                row = t.add_row().cells
                for i, col in enumerate(ecols):
                    val = m.get(col["key"], "") if isinstance(m, dict) else getattr(m, col["key"], "")
                    row[i].text = _format_value(val, col.get("format"))
                    for r in row[i].paragraphs[0].runs: r.font.size = Pt(font_size - 1)
                pt = m.get("prezzo_totale", 0) if isinstance(m, dict) else getattr(m, "prezzo_totale", 0)
                try: totale += float(pt or 0)
                except: pass

            if sec.get("show_totals", True):
                tr = t.add_row().cells
                tr[max(0, len(ecols)-2)].text = "TOTALE"
                tr[len(ecols)-1].text = _fmt_euro(totale)
                for c in tr:
                    for r in c.paragraphs[0].runs: r.font.bold = True; r.font.size = Pt(font_size)
            doc.add_paragraph('')

        # ─────── RIEPILOGO PREZZI ───────
        elif st == "price_summary":
            efl = [f for f in fields if f.get("enabled", True)]
            if not efl: continue
            doc.add_heading(title, level=2)
            t = doc.add_table(rows=len(efl), cols=2)
            t.style = 'Table Grid'
            for i, f in enumerate(efl):
                lookup = _fl.get(f["key"], {})
                src = lookup.get("source", ""); fmt = lookup.get("format")
                val = _resolve_value(src, ctx) if src else ""
                t.cell(i, 0).text = f.get("label", f["key"])
                t.cell(i, 1).text = _format_value(val, fmt)
                is_tot = f["key"] in ("totale_netto", "totale_lordo")
                for r in t.cell(i, 0).paragraphs[0].runs: r.font.bold = True; r.font.size = Pt(font_size + (2 if is_tot else 0))
                for r in t.cell(i, 1).paragraphs[0].runs: r.font.bold = is_tot; r.font.size = Pt(font_size + (2 if is_tot else 0))
            doc.add_paragraph('')

        # ─────── NOTE ───────
        elif st == "text_block":
            efl = [f for f in fields if f.get("enabled", True)]
            has = False
            for f in efl:
                lookup = _fl.get(f["key"], {})
                src = lookup.get("source", "")
                if src and _resolve_value(src, ctx): has = True; break
            if has:
                doc.add_heading(title, level=2)
                for f in efl:
                    lookup = _fl.get(f["key"], {})
                    src = lookup.get("source", "")
                    if src:
                        val = _resolve_value(src, ctx)
                        if val:
                            p = doc.add_paragraph()
                            p.add_run(f"{f.get('label', '')}: ").font.bold = True
                            p.add_run(_safe_str(val))
                doc.add_paragraph('')

        # ─────── FOOTER ───────
        elif st == "footer":
            doc.add_paragraph('')
            fp = doc.add_paragraph()
            fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            efl = [f for f in fields if f.get("enabled", True)]
            parts = []
            for f in efl:
                if f["key"] == "data_generazione":
                    parts.append(f"Documento generato il {datetime.now().strftime('%d/%m/%Y %H:%M')}")
                elif f["key"] == "testo_libero_footer":
                    cv = f.get("value", "")
                    if cv: parts.append(cv)
            text_out = " - ".join(parts) if parts else f"Documento generato il {datetime.now().strftime('%d/%m/%Y %H:%M')}"
            fp.add_run(text_out).font.size = Pt(8)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf
