import sys; sys.stdout.reconfigure(line_buffering=True)
from fastapi import FastAPI, Depends, HTTPException, Query, Header, UploadFile, File, Form, Request
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import StreamingResponse
from sqlalchemy.orm import Session
from sqlalchemy import text, or_
from typing import List, Optional
from datetime import datetime, timedelta
import json
import os
import io
import math
import re


import base64
from template_engine import (
    DocumentTemplate, STATIC_SECTIONS,
    get_available_fields_from_db, get_default_template_config_from_db,
    genera_docx_da_template, load_valori_dinamici, load_defaults_info
)

from database import engine, SessionLocal
from models import (
    Base, Preventivo, DatiCommessa, DatiPrincipali, Normative,
    DisposizioneVano, Porte, Materiale, ProductTemplate,
    Utente, GruppoUtenti, PermessoGruppo, Ruolo, PermessoRuolo,
    Articolo, CategoriaArticoli, Cliente,
    RigaRicambio, Argano, ParametriSistema, BomStruttura
)
from schemas import (
    PreventivoCreate, PreventivoUpdate, Preventivo as PreventivoSchema,
    DatiCommessaCreate, DatiCommessaUpdate, DatiCommessa as DatiCommessaSchema,
    DatiPrincipaliCreate, DatiPrincipaliUpdate, DatiPrincipali as DatiPrincipaliSchema,
    NormativeCreate, NormativeUpdate, Normative as NormativeSchema,
    DisposizioneVanoCreate, DisposizioneVanoUpdate, DisposizioneVano as DisposizioneVanoSchema,
    PorteCreate, PorteUpdate, Porte as PorteSchema,
    MaterialeCreate, MaterialeUpdate, Materiale as MaterialeSchema,
    ProductTemplateCreate, ProductTemplateUpdate, ProductTemplate as ProductTemplateSchema
)
from auth import (
    authenticate_user, create_access_token, get_password_hash,
    get_user_from_token, create_default_admin,
    get_user_permissions,
    ACCESS_TOKEN_EXPIRE_MINUTES
)
from fastapi.responses import JSONResponse
from fastapi.encoders import jsonable_encoder
import json as json_module
from excel_data_loader import ExcelDataLoader
from excel_import import ExcelImporter
from api_permessi import router as permessi_router

from moduli_attivabili import router as moduli_router, richiedi_modulo
from fatturazione_api import router as fatturazione_router
from ordini_stato import router as ordini_stato_router
from starlette.responses import StreamingResponse
    

# Custom JSON encoder per SQLAlchemy objects
class SQLAlchemyEncoder(json_module.JSONEncoder):
    def default(self, obj):
        if hasattr(obj, '__table__'):
            row = {}
            for col in obj.__table__.columns:
                val = getattr(obj, col.name, None)
                if hasattr(val, 'isoformat'):
                    val = val.isoformat()
                row[col.name] = val
            return row
        if hasattr(obj, 'isoformat'):
            return obj.isoformat()
        return super().default(obj)

def _orm_to_dict(obj):
    """Converte un oggetto ORM in dict serializzabile"""
    if obj is None:
        return None
    if isinstance(obj, dict):
        return obj
    if isinstance(obj, list):
        return [_orm_to_dict(item) for item in obj]
    if hasattr(obj, '__table__'):
        row = {}
        for col in obj.__table__.columns:
            val = getattr(obj, col.name, None)
            if hasattr(val, 'isoformat'):
                val = val.isoformat()
            row[col.name] = val
        return row
    return obj

# ─────────────────────────────────────────────────────────────────
# DEFAULT VALUE EXPRESSION RESOLVER
# Sintassi nel campo valore_default di campi_configuratore:
#   {{TODAY}}                  → data odierna in formato gg/mm/aaaa
#   {{YEAR}}                   → anno corrente (4 cifre)
#   {{MONTH}}                  → mese corrente (2 cifre, es. 03)
#   {{campo:codice.campo}}     → valore di un altro campo del preventivo
# Variante con fallback:
#   {{TODAY|01/01/2000}}       → data odierna, o 01/01/2000 se non disponibile
#   {{campo:dati_commessa.numero_offerta|N/A}}
# ─────────────────────────────────────────────────────────────────
def _resolve_default_expr(expr: str, preventivo_id: int, db, context: dict = None) -> str:
    if not expr or not expr.startswith('{{'):
        return expr
    inner = expr.strip()
    if inner.startswith('{{') and inner.endswith('}}'):
        inner = inner[2:-2]
    else:
        return expr
    fallback = ''
    if '|' in inner:
        inner, fallback = inner.split('|', 1)
    inner = inner.strip()
    from datetime import date as _date
    today = _date.today()
    if inner == 'TODAY':
        return today.strftime('%d/%m/%Y')
    elif inner == 'YEAR':
        return str(today.year)
    elif inner == 'MONTH':
        return today.strftime('%m')
    elif inner.startswith('campo:'):
        campo_codice = inner[6:].strip()
        # Prima cerca nel contesto in-memoria (stessa sessione, non ancora committato)
        if context and campo_codice in context:
            val = context[campo_codice]
            if val is not None and str(val).strip() != '':
                return str(val)
        try:
            row = db.execute(
                text("SELECT valore FROM valori_configurazione WHERE preventivo_id=:pid AND codice_campo=:cod LIMIT 1"),
                {"pid": preventivo_id, "cod": campo_codice}
            ).fetchone()
            if row and row[0] is not None and str(row[0]).strip() != '':
                return str(row[0])
        except Exception:
            pass
        # Fallback: campi ORM hardcoded (dati_commessa, disposizione_vano)
        ORM_TABLE_MAP = {
            'dati_commessa': 'dati_commessa',
            'disposizione_vano': 'disposizione_vano',
        }
        parts = campo_codice.split('.', 1)
        if len(parts) == 2:
            tabella, colonna = parts
            if tabella in ORM_TABLE_MAP:
                # Whitelist colonne per sicurezza
                COLONNE_CONSENTITE = {
                    'dati_commessa': {'numero_offerta', 'data_offerta', 'riferimento_cliente',
                                      'quantita', 'consegna_richiesta', 'prezzo_unitario',
                                      'pagamento', 'trasporto', 'destinazione'},
                    'disposizione_vano': {'altezza_vano', 'piano_piu_alto', 'piano_piu_basso',
                                          'posizione_quadro_lato', 'posizione_quadro_piano'},
                }
                if colonna in COLONNE_CONSENTITE.get(tabella, set()):
                    try:
                        orm_row = db.execute(
                            text(f"SELECT {colonna} FROM {tabella} WHERE preventivo_id=:pid LIMIT 1"),
                            {"pid": preventivo_id}
                        ).fetchone()
                        if orm_row and orm_row[0] is not None and str(orm_row[0]).strip() != '':
                            return str(orm_row[0])
                    except Exception:
                        pass
        return fallback
    return fallback or expr

# Crea tabelle
Base.metadata.create_all(bind=engine)

# FastAPI app
app = FastAPI(title="Configuratore Elettroquadri API", version="0.11.0")

# ==========================================
# CORS CONFIGURATION
# ==========================================
app.add_middleware(
    CORSMiddleware,
    allow_origins=[
        "http://localhost:5173",
        "http://localhost:3000",
        "http://127.0.0.1:5173",
        "http://127.0.0.1:3000"
    ],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# ==========================================
# ROUTER: Gruppi, Ruoli, Permessi
# ==========================================
app.include_router(permessi_router)
app.include_router(moduli_router)
fatturazione_router.dependencies = [Depends(richiedi_modulo("fatturazione"))]
app.include_router(fatturazione_router)
app.include_router(ordini_stato_router)

# ==========================================
# DEPENDENCY
# ==========================================
def get_db():
    db = SessionLocal()
    try:
        yield db
    finally:
        db.close()

# ==========================================
# DEPENDENCY: Current User (opzionale e obbligatoria)
# ==========================================
def _extract_user_from_request(
    authorization: Optional[str] = Header(None),
    db: Session = Depends(get_db)
) -> Optional["Utente"]:
    """Estrae l'utente dal token Bearer. Ritorna None se non autenticato."""
    if not authorization or not authorization.startswith("Bearer "):
        return None
    token = authorization[7:]
    return get_user_from_token(token, db)

def get_current_user_required(
    authorization: Optional[str] = Header(None),
    db: Session = Depends(get_db)
) -> "Utente":
    """Dependency: utente obbligatorio. Alza 401 se non autenticato."""
    user = _extract_user_from_request(authorization, db)
    if not user:
        raise HTTPException(status_code=401, detail="Autenticazione richiesta")
    return user

def user_can_see_all_preventivi(user, db: Session) -> bool:
    """True se l'utente può vedere tutti i preventivi (interni/admin)"""
    if user.is_admin:
        return True
    permessi = get_user_permissions(user, db)
    return "preventivi.view_all" in permessi

def filter_preventivi_by_user(query, user, db: Session):
    """Filtra la query preventivi in base all'utente.
    - Utente interno (con view_all o senza cliente_id): vede tutto
    - Utente cliente (con cliente_id, senza view_all): solo i suoi
    """
    if user_can_see_all_preventivi(user, db):
        return query
    if user.cliente_id:
        return query.filter(Preventivo.cliente_id == user.cliente_id)
    # Utente senza cliente_id e senza view_all: vede solo quelli creati da lui
    return query.filter(Preventivo.created_by == user.id)

# ==========================================
# HELPER: Accesso sicuro ai campi Preventivo
# Nomi reali da models.py: stato, numero, tipo, totale_materiali, totale_netto
# ==========================================
def _prev_get(preventivo, *names, default=None):
    """Cerca il primo attributo esistente tra i nomi forniti"""
    for name in names:
        val = getattr(preventivo, name, None)
        if val is not None:
            return val
    return default

def _prev_set(preventivo, value, *names):
    """Setta il primo attributo esistente tra i nomi forniti"""
    for name in names:
        if hasattr(preventivo, name):
            setattr(preventivo, name, value)
            return True
    setattr(preventivo, names[0], value)
    return False

# Shortcuts - nomi REALI prima, fallback dopo
def _prev_stato(p):      return _prev_get(p, 'stato', 'status', default='bozza')
def _prev_numero(p):     return _prev_get(p, 'numero', 'numero_preventivo', default='')
def _prev_tipo(p):       return _prev_get(p, 'tipo', 'tipo_preventivo', default='')
def _prev_totale(p):     return _prev_get(p, 'totale_materiali', 'total_price', default=0)
def _prev_netto(p):      return _prev_get(p, 'totale_netto', 'total_price_finale', default=0)
def _prev_categoria(p):  return _prev_get(p, 'categoria', default='')

def safe_float(val, default=0.0):
    try:
        return float(val) if val is not None else default
    except (TypeError, ValueError):
        return default


# ==========================================
# RULE ENGINE - VERSIONE MIGLIORATA
# ==========================================
def load_rules():
    """Carica le regole dai file JSON nella directory rules/"""
    rules = []
    rules_dir = "./rules"
    
    if not os.path.exists(rules_dir):
        print("Ã¢Å¡Â Ã¯Â¸Â Directory rules/ non trovata")
        return rules
    
    for filename in os.listdir(rules_dir):
        if filename.endswith(".json"):
            try:
                with open(os.path.join(rules_dir, filename), "r", encoding="utf-8") as f:
                    rule = json.load(f)
                    rules.append(rule)
                    print(f"Ã¢Å“â€¦ Regola caricata: {filename}")
            except Exception as e:
                print(f"Ã¢ÂÅ’ Errore caricamento {filename}: {e}")
    
    print(f"Ã°Å¸â€œâ€¹ Caricate {len(rules)} regole")
    return rules

def evaluate_condition(condition: dict, config_data: dict) -> bool:
    """Valuta una singola condizione con gestione robusta di NULL/None e tipi misti"""
    if isinstance(condition, str):
        return False
    
    field = condition.get("field")
    operator = condition.get("operator")
    expected_value = condition.get("value")
    
    if field not in config_data:
        return False
    
    config_value = config_data.get(field)
    
    if config_value is None:
        return False
    
    if isinstance(config_value, str) and config_value.strip() == "":
        return False
    
    result = False
    
    if operator == "equals":
        # Confronto robusto: prova prima diretto, poi stringa vs stringa
        result = config_value == expected_value
        if not result:
            result = str(config_value).strip().lower() == str(expected_value).strip().lower()
    elif operator == "not_equals":
        result = str(config_value).strip().lower() != str(expected_value).strip().lower()
    elif operator == "contains":
        result = str(expected_value).lower() in str(config_value).lower() if config_value else False
    elif operator in ("greater_than", "gt"):
        try:
            result = float(config_value) > float(expected_value)
        except (ValueError, TypeError):
            result = False
    elif operator in ("greater_equal", "gte", "greater_than_or_equal"):
        try:
            result = float(config_value) >= float(expected_value)
        except (ValueError, TypeError):
            result = False
    elif operator in ("less_than", "lt"):
        try:
            result = float(config_value) < float(expected_value)
        except (ValueError, TypeError):
            result = False
    elif operator in ("less_equal", "lte", "less_than_or_equal"):
        try:
            result = float(config_value) <= float(expected_value)
        except (ValueError, TypeError):
            result = False
    elif operator == "in":
        if isinstance(expected_value, list):
            cv_lower = str(config_value).strip().lower()
            result = cv_lower in [str(v).strip().lower() for v in expected_value]
        else:
            result = False
    elif operator == "not_in":
        if isinstance(expected_value, list):
            cv_lower = str(config_value).strip().lower()
            result = cv_lower not in [str(v).strip().lower() for v in expected_value]
        else:
            result = True
    elif operator == "is_true":
        result = str(config_value).strip().lower() in ("true", "1", "si", "yes")
    elif operator == "is_false":
        result = str(config_value).strip().lower() in ("false", "0", "no", "")
    
    return result

def build_config_context(preventivo_id: int, db: Session) -> dict:
    """
    Costruisce il contesto COMPLETO per la valutazione regole,
    leggendo da TUTTE le fonti dati del preventivo.
    
    Fonti (in ordine di priorita, le ultime sovrascrivono):
      0. defaults da campi_configuratore (base, priorità minima)
      1. metadati preventivo (tipo, categoria)
      2. dati_principali (tabella ORM)
      3. normative (tabella ORM)
      4. argano (query diretta)
      5. porte (query diretta)
      6. disposizione_vano (query diretta)
      7. valori_configurazione (tabella chiave/valore per campi dinamici)
    """
    preventivo = db.query(Preventivo).filter(Preventivo.id == preventivo_id).first()
    if not preventivo:
        return {}
    
    config_data = {}
    
    # 0. DEFAULTS da campi_configuratore (priorità più bassa, base per tutte le sezioni)
    #    Garantisce che i valori default siano disponibili anche per sezioni non ancora visitate
    try:
        defaults_rows = db.execute(
            text("SELECT codice, valore_default FROM campi_configuratore WHERE attivo=1 AND valore_default IS NOT NULL AND valore_default != '-'")
        ).fetchall()
        for codice, valore_default in defaults_rows:
            if codice and valore_default:
                config_data[codice] = valore_default
    except Exception:
        pass
    
    # 1. Metadati preventivo
    config_data["preventivo_tipo"] = _prev_tipo(preventivo)
    config_data["preventivo_categoria"] = _prev_categoria(preventivo)
    
    # 1b. Info template prodotto (per regole TEMPLATE_BASE)
    try:
        tmpl_id = getattr(preventivo, 'template_id', None)
        if tmpl_id:
            tmpl = db.query(ProductTemplate).filter(ProductTemplate.id == tmpl_id).first()
            if tmpl:
                config_data["template_id"] = tmpl.id
                config_data["template_categoria"] = tmpl.categoria
                config_data["template_sottocategoria"] = tmpl.sottocategoria
                config_data["template_nome"] = tmpl.nome_display
    except Exception:
        pass
    
    # 2. Dati principali (query diretta, NO relationship)
    try:
        dp = db.query(DatiPrincipali).filter(DatiPrincipali.preventivo_id == preventivo_id).first()
        if dp:
            config_data.update({
                "tipo_impianto": dp.tipo_impianto,
                "nuovo_impianto": dp.nuovo_impianto,
                "numero_fermate": dp.numero_fermate,
                "numero_servizi": dp.numero_servizi,
                "velocita": dp.velocita,
                "corsa": dp.corsa,
                "con_locale_macchina": dp.con_locale_macchina,
                "posizione_locale_macchina": dp.posizione_locale_macchina,
                "tipo_trazione": dp.tipo_trazione,
                "forza_motrice": dp.forza_motrice,
                "luce": dp.luce,
                "tensione_manovra": dp.tensione_manovra,
                "tensione_freno": dp.tensione_freno,
            })
    except Exception:
        pass
    
    # 3. Normative (query diretta, NO relationship)
    try:
        norm = db.query(Normative).filter(Normative.preventivo_id == preventivo_id).first()
        if norm:
            config_data.update({
                "en_81_1": norm.en_81_1,
                "en_81_20": norm.en_81_20,
                "en_81_21": norm.en_81_21,
                "en_81_28": norm.en_81_28,
                "en_81_70": norm.en_81_70,
                "en_81_72": norm.en_81_72,
                "en_81_73": norm.en_81_73,
                "a3_95_16": norm.a3_95_16,
                "dm236_legge13": norm.dm236_legge13,
                "emendamento_a3": norm.emendamento_a3,
                "uni_10411_1": norm.uni_10411_1,
            })
    except Exception:
        pass
    
    # 4. Argano (non ha relazione ORM, query diretta)
    try:
        argano = db.query(Argano).filter(Argano.preventivo_id == preventivo_id).first()
        if argano:
            config_data.update({
                "trazione": argano.trazione,
                "potenza_motore_kw": argano.potenza_motore_kw,
                "corrente_nom_motore_amp": argano.corrente_nom_motore_amp,
                "tipo_vvvf": argano.tipo_vvvf,
                "vvvf_nel_vano": argano.vvvf_nel_vano,
                "freno_tensione": argano.freno_tensione,
                "ventilazione_forzata": argano.ventilazione_forzata,
                "tipo_teleruttore": argano.tipo_teleruttore,
            })
    except Exception:
        pass
    
    # 5. Porte (query diretta, leggi tutti gli attributi)
    try:
        porte = db.query(Porte).filter(Porte.preventivo_id == preventivo_id).first()
        if porte:
            skip = {'id', 'preventivo_id', 'preventivo', 'metadata', 'registry',
                    '_sa_instance_state', '_sa_class_manager'}
            for col in porte.__table__.columns:
                if col.name not in ('id', 'preventivo_id'):
                    val = getattr(porte, col.name, None)
                    if val is not None:
                        config_data[col.name] = val
    except Exception:
        pass
    
    # 6. Disposizione vano (query diretta)
    try:
        dv = db.query(DisposizioneVano).filter(
            DisposizioneVano.preventivo_id == preventivo_id
        ).first()
        if dv:
            for col in dv.__table__.columns:
                if col.name not in ('id', 'preventivo_id'):
                    val = getattr(dv, col.name, None)
                    if val is not None:
                        config_data[col.name] = val
    except Exception:
        pass
    
    # 7. Valori configurazione (tabella chiave/valore, massima priorita)
    try:
        result = db.execute(
            text("SELECT codice_campo, valore FROM valori_configurazione WHERE preventivo_id = :pid"),
            {"pid": preventivo_id}
        )
        for row in result.fetchall():
            campo, valore = row[0], row[1]
            if campo and valore is not None:
                config_data[campo] = valore
    except Exception:
        pass  # Tabella potrebbe non esistere
    
    # Pulizia: rimuovi None
    config_data = {k: v for k, v in config_data.items() if v is not None}
    
    return config_data


def evaluate_rules(preventivo_id: int, db: Session):
    """
    Valuta le regole JSON per un preventivo.
    
    1. Costruisce contesto da TUTTE le sezioni (ORM + valori_configurazione + JSON)
    2. Carica regole da file JSON (./rules/)
    3. Per ogni regola: valuta condizioni â†’ se tutte OK â†’ aggiunge materiali
    4. Rimuove materiali orfani (regole non piu attive)
    
    Returns: dict con risultato valutazione
    """
    db.expire_all()
    
    preventivo = db.query(Preventivo).filter(Preventivo.id == preventivo_id).first()
    if not preventivo:
        return {"error": "Preventivo non trovato"}
    
    # 1. Costruisci contesto completo
    config_data = build_config_context(preventivo_id, db)
    
    # 2. Carica regole da file JSON
    rules = load_rules()
    
    if not rules:
        return {
            "status": "warning",
            "message": "Nessuna regola trovata in ./rules/",
            "materiali_aggiunti": 0,
            "materiali_rimossi": 0,
            "regole_attive": [],
            "context_keys": list(config_data.keys())
        }
    
    # 3. Valuta regole
    active_rules = set()
    materials_to_add = []
    
    for rule in rules:
        if not rule.get("enabled", True):
            continue
        
        rule_id = rule.get("id", "unknown")
        conditions = rule.get("conditions", [])
        
        # Valuta tutte le condizioni (AND implicito)
        all_met = True
        for condition in conditions:
            if not evaluate_condition(condition, config_data):
                all_met = False
                break
        
        if all_met and conditions:
            active_rules.add(rule_id)
            for material in rule.get("materials", []):
                materials_to_add.append({
                    "rule_id": rule_id,
                    "codice": material.get("codice"),
                    "descrizione": material.get("descrizione"),
                    "quantita": material.get("quantita", 1),
                    "prezzo_unitario": material.get("prezzo_unitario", 0.0),
                    "categoria": material.get("categoria", "Materiale Automatico"),
                    "note": material.get("note", ""),
                })
    
    # 4. Lookup prezzi da tabella articoli per materiali senza prezzo
    codici_senza_prezzo = set(
        m["codice"] for m in materials_to_add
        if not m.get("prezzo_unitario") and m.get("codice")
    )
    prezzi_articoli = {}
    if codici_senza_prezzo:
        articoli_trovati = db.query(Articolo).filter(
            Articolo.codice.in_(codici_senza_prezzo)
        ).all()
        for art in articoli_trovati:
            costo_base = art.costo_fisso or 0
            ricarico = art.ricarico_percentuale or 0
            prezzo = costo_base * (1 + ricarico / 100) if ricarico else costo_base
            prezzi_articoli[art.codice] = round(prezzo, 4)
            
        # Fallback: cerca anche in articoli_bom (tabella legacy)
        codici_mancanti = codici_senza_prezzo - set(prezzi_articoli.keys())
        if codici_mancanti:
            try:
                placeholders = ",".join(f":c{i}" for i in range(len(codici_mancanti)))
                params = {f"c{i}": c for i, c in enumerate(codici_mancanti)}
                result = db.execute(text(
                    f"SELECT codice, costo_fisso, costo_variabile FROM articoli_bom WHERE codice IN ({placeholders})"
                ), params)
                for row in result.fetchall():
                    costo = (row[1] or 0) + (row[2] or 0)
                    if costo > 0:
                        prezzi_articoli[row[0]] = round(costo, 4)
            except Exception:
                pass  # tabella articoli_bom potrebbe non esistere

    # Aggiorna prezzi nei materiali da aggiungere
    for mat_data in materials_to_add:
        if not mat_data.get("prezzo_unitario") and mat_data["codice"] in prezzi_articoli:
            mat_data["prezzo_unitario"] = prezzi_articoli[mat_data["codice"]]

    # 5. Rimozione orfani: materiali da regole non piu attive
    materiali_rimossi = 0
    existing_auto = db.query(Materiale).filter(
        Materiale.preventivo_id == preventivo_id,
        Materiale.aggiunto_da_regola == True
    ).all()
    
    for mat in existing_auto:
        if mat.regola_id and mat.regola_id not in active_rules:
            db.delete(mat)
            materiali_rimossi += 1
    
    # 6. Aggiungi nuovi materiali (evita duplicati) + aggiorna prezzi mancanti
    materiali_aggiunti = 0
    materiali_prezzo_aggiornato = 0
    for mat_data in materials_to_add:
        existing = db.query(Materiale).filter(
            Materiale.preventivo_id == preventivo_id,
            Materiale.codice == mat_data["codice"],
            Materiale.regola_id == mat_data["rule_id"]
        ).first()
        
        if existing:
            # Aggiorna prezzo se era 0 e ora abbiamo un prezzo valido
            if (not existing.prezzo_unitario or existing.prezzo_unitario == 0) and mat_data["prezzo_unitario"]:
                existing.prezzo_unitario = mat_data["prezzo_unitario"]
                existing.prezzo_totale = existing.quantita * mat_data["prezzo_unitario"]
                materiali_prezzo_aggiornato += 1
        else:
            qta = mat_data["quantita"]
            prezzo = mat_data["prezzo_unitario"]
            nuovo = Materiale(
                preventivo_id=preventivo_id,
                codice=mat_data["codice"],
                descrizione=mat_data["descrizione"],
                quantita=qta,
                prezzo_unitario=prezzo,
                prezzo_totale=qta * prezzo,
                categoria=mat_data["categoria"],
                aggiunto_da_regola=True,
                regola_id=mat_data["rule_id"],
                note=mat_data.get("note", ""),
            )
            db.add(nuovo)
            materiali_aggiunti += 1
    
    # 7. Ricalcola totale preventivo
    if materiali_aggiunti > 0 or materiali_rimossi > 0 or materiali_prezzo_aggiornato > 0:
        db.commit()
        tutti_materiali = db.query(Materiale).filter(
            Materiale.preventivo_id == preventivo_id
        ).all()
        totale_calc = sum(m.prezzo_totale or 0 for m in tutti_materiali)
        _prev_set(preventivo, totale_calc, "totale_materiali", "total_price")
        db.commit()
    
    return {
        "status": "ok",
        "materiali_aggiunti": materiali_aggiunti,
        "materiali_rimossi": materiali_rimossi,
        "materiali_prezzo_aggiornato": materiali_prezzo_aggiornato,
        "regole_attive": list(active_rules),
        "regole_totali": len(rules),
        "prezzi_da_articoli": len(prezzi_articoli),
        "codici_senza_prezzo": list(codici_senza_prezzo - set(prezzi_articoli.keys())),
        "context_keys": list(config_data.keys()),
    }

# ============================================================
# AGGIUNGERE IN main.py dopo l'endpoint evaluate_rules
# ============================================================

@app.post("/preventivi/{preventivo_id}/test-rules")
def test_rules_endpoint(preventivo_id: int, body: dict = None, db: Session = Depends(get_db)):
    """Test regole senza side-effects. Restituisce report diagnostico completo."""
    from rule_engine import RuleEngine
    preventivo = db.query(Preventivo).get(preventivo_id)
    if not preventivo:
        raise HTTPException(404, "Preventivo non trovato")
    engine = RuleEngine(db)
    override_ctx = body.get("override_context") if body else None
    report = engine.test_rules(preventivo, override_context=override_ctx)
    return report


@app.get("/preventivi/{preventivo_id}/rule-context")
def get_rule_context(preventivo_id: int, db: Session = Depends(get_db)):
    """Restituisce il context completo che il rule engine vede per questo preventivo."""
    from rule_engine import RuleEngine
    preventivo = db.query(Preventivo).get(preventivo_id)
    if not preventivo:
        raise HTTPException(404, "Preventivo non trovato")
    engine = RuleEngine(db)
    return engine.get_context_debug(preventivo)


def safe_evaluate_rules(preventivo_id: int, db: Session):
    """
    Wrapper che non crasha mai - logga errori ma non blocca il chiamante.
    Usa il nuovo RuleEngine che supporta actions (lookup_table, ecc.)
    e placeholder nei materiali.
    """
    try:
        from rule_engine import RuleEngine
        preventivo = db.query(Preventivo).filter(Preventivo.id == preventivo_id).first()
        if not preventivo:
            return {"error": "Preventivo non trovato"}
        
        engine = RuleEngine(db)
        result = engine.evaluate_rules(preventivo)
        db.commit()
        return result
    except Exception as e:
        import traceback
        print(f"[RULE ENGINE ERROR] preventivo {preventivo_id}: {e}")
        traceback.print_exc()
        db.rollback()
        return {"error": str(e)}


def touch_preventivo(preventivo_id: int, db: Session):
    """Aggiorna updated_at del preventivo. Necessario per il dirty-check delle revisioni."""
    try:
        db.execute(
            text("UPDATE preventivi SET updated_at = :now WHERE id = :pid"),
            {"now": datetime.now().isoformat(), "pid": preventivo_id}
        )
        db.commit()
    except Exception as e:
        print(f"[WARN] touch_preventivo({preventivo_id}): {e}")


# ============================================================
# ORDINI E BOM - FUNZIONI HELPER
# ============================================================

def calcola_lead_time(preventivo_id: int, db: Session) -> int:
    conn = db.get_bind().raw_connection()
    cursor = conn.cursor()
    materiali = db.query(Materiale).filter(Materiale.preventivo_id == preventivo_id).all()
    if not materiali:
        return 15
    max_lt = 0
    for mat in materiali:
        cursor.execute("SELECT lead_time_produzione, lead_time_acquisto, tipo FROM articoli_bom WHERE codice = ?", (mat.codice,))
        row = cursor.fetchone()
        if row:
            lt = row[0] if row[2] in ("MASTER", "SEMILAVORATO") else row[1]
            max_lt = max(max_lt, lt or 0)
    return max_lt + 5 + 3


def esplodi_bom_ricorsiva(cursor, codice_padre, qta_padre, contesto, ordine_id,
                           livello=0, percorso="", risultati=None):
    if risultati is None:
        risultati = []
    if livello > 10:
        return risultati

    cursor.execute(
        "SELECT bs.figlio_codice, bs.quantita, bs.formula_quantita, bs.condizione_esistenza, "
        "ab.descrizione, ab.tipo, ab.categoria, ab.costo_fisso, ab.costo_variabile, "
        "ab.unita_misura, ab.lead_time_produzione, ab.lead_time_acquisto, "
        "ab.parametro1_nome, ab.parametro2_nome, ab.parametro3_nome "
        "FROM bom_struttura bs "
        "JOIN articoli_bom ab ON ab.codice = bs.figlio_codice "
        "WHERE bs.padre_codice = ? ORDER BY bs.posizione",
        (codice_padre,)
    )
    figli = cursor.fetchall()

    if not figli:
        # LEAF
        cursor.execute(
            "SELECT descrizione, tipo, categoria, costo_fisso, costo_variabile, "
            "unita_misura, lead_time_produzione, lead_time_acquisto, "
            "parametro1_nome, parametro2_nome, parametro3_nome "
            "FROM articoli_bom WHERE codice = ?",
            (codice_padre,)
        )
        art = cursor.fetchone()
        if art:
            costo_unit = (art[3] or 0) + (art[4] or 0)
            param_vals = {}
            for i, pname in enumerate(art[8:11]):
                if pname:
                    val = contesto.get(pname, contesto.get(pname.upper(), ""))
                    param_vals[f"parametro{i+1}_nome"] = pname
                    param_vals[f"parametro{i+1}_valore"] = str(val) if val else ""
            risultati.append({
                "ordine_id": ordine_id, "codice": codice_padre,
                "descrizione": art[0], "tipo": art[1], "categoria": art[2],
                "quantita": qta_padre, "unita_misura": art[5] or "PZ",
                "costo_unitario": costo_unit, "costo_totale": costo_unit * qta_padre,
                "padre_codice": percorso.split(" > ")[-2] if " > " in percorso else None,
                "livello_esplosione": livello, "percorso": percorso,
                "lead_time_giorni": art[6] if art[1] != "ACQUISTO" else art[7],
                **param_vals
            })
        return risultati

    for figlio in figli:
        (f_cod, f_qta, f_formula, f_cond, f_desc, f_tipo, f_cat,
         f_cf, f_cv, f_um, f_lt_p, f_lt_a, f_p1, f_p2, f_p3) = figlio

        if f_cond and not _valuta_condizione_bom(f_cond, contesto):
            continue

        qta = f_qta
        if f_formula:
            qta_calc = _calcola_formula(f_formula, contesto)
            if qta_calc is not None:
                qta = qta_calc

        qta_tot = qta * qta_padre
        nuovo_percorso = f"{percorso} > {f_cod}" if percorso else f_cod

        if f_tipo == "ACQUISTO":
            costo_unit = (f_cf or 0) + (f_cv or 0)
            param_vals = {}
            for i, pname in enumerate([f_p1, f_p2, f_p3]):
                if pname:
                    val = contesto.get(pname, contesto.get(pname.upper(), ""))
                    param_vals[f"parametro{i+1}_nome"] = pname
                    param_vals[f"parametro{i+1}_valore"] = str(val) if val else ""
            risultati.append({
                "ordine_id": ordine_id, "codice": f_cod,
                "descrizione": f_desc, "tipo": f_tipo, "categoria": f_cat,
                "quantita": qta_tot, "unita_misura": f_um or "PZ",
                "costo_unitario": costo_unit, "costo_totale": costo_unit * qta_tot,
                "padre_codice": codice_padre,
                "livello_esplosione": livello + 1, "percorso": nuovo_percorso,
                "lead_time_giorni": f_lt_a or 0, **param_vals
            })
        else:
            esplodi_bom_ricorsiva(cursor, f_cod, qta_tot, contesto, ordine_id,
                                  livello + 1, nuovo_percorso, risultati)
    return risultati


def _valuta_condizione_bom(condizione, contesto):
    try:
        expr = condizione
        for match in re.finditer(r'\[(\w+)\]', condizione):
            campo = match.group(1)
            valore = contesto.get(campo, contesto.get(campo.upper(), ""))
            if isinstance(valore, (int, float)):
                expr = expr.replace(f"[{campo}]", str(valore))
            else:
                expr = expr.replace(f"[{campo}]", f"'{valore}'" if valore else "''")
        return bool(eval(expr))
    except:
        return True


def _calcola_formula(formula, contesto):
    try:
        expr = formula
        for match in re.finditer(r'\[(\w+)\]', formula):
            campo = match.group(1)
            valore = contesto.get(campo, contesto.get(campo.upper(), 0))
            try:
                valore = float(valore)
            except (ValueError, TypeError):
                valore = 0
            expr = expr.replace(f"[{campo}]", str(valore))
        return max(1, math.ceil(eval(expr)))
    except:
        return None



# ==========================================
# AUTH
# ==========================================
@app.post("/auth/login")
def login(data: dict, db: Session = Depends(get_db)):
    """Login utente, restituisce JWT token"""
    username = data.get("username", "")
    password = data.get("password", "")
    
    user = authenticate_user(db, username, password)
    if not user:
        raise HTTPException(status_code=401, detail="Credenziali non valide")
    
    # Aggiorna last_login
    user.last_login = datetime.now()
    db.commit()
    
    # Carica permessi e info ruolo/gruppo
    permessi = get_user_permissions(user, db)
    
    gruppo_nome = None
    if user.gruppo_id:
        gruppo = db.query(GruppoUtenti).filter(GruppoUtenti.id == user.gruppo_id).first()
        gruppo_nome = gruppo.nome if gruppo else None
    
    ruolo_nome = None
    ruolo_codice = None
    if user.ruolo_id:
        ruolo = db.query(Ruolo).filter(Ruolo.id == user.ruolo_id).first()
        if ruolo:
            ruolo_nome = ruolo.nome
            ruolo_codice = ruolo.codice
    
    access_token = create_access_token(data={"sub": user.username})
    return {
        "access_token": access_token,
        "token_type": "bearer",
        "user": {
            "id": user.id,
            "username": user.username,
            "nome": user.nome,
            "cognome": user.cognome,
            "email": user.email,
            "is_admin": user.is_admin,
            "gruppo_id": user.gruppo_id,
            "gruppo_nome": gruppo_nome,
            "ruolo_id": user.ruolo_id,
            "ruolo_nome": ruolo_nome,
            "ruolo_codice": ruolo_codice,
            "permessi": permessi,
            "cliente_id": user.cliente_id,
        }
    }

@app.get("/auth/me")
def get_current_user_info(
    authorization: Optional[str] = Header(None),
    db: Session = Depends(get_db)
):
    """Ottieni info utente corrente dal token"""
    token = None
    if authorization and authorization.startswith("Bearer "):
        token = authorization[7:]
    
    if not token:
        raise HTTPException(status_code=401, detail="Token non fornito")
    
    user = get_user_from_token(token, db)
    if not user:
        raise HTTPException(status_code=401, detail="Token non valido")
    
    permessi = get_user_permissions(user, db)
    
    gruppo_nome = None
    if user.gruppo_id:
        gruppo = db.query(GruppoUtenti).filter(GruppoUtenti.id == user.gruppo_id).first()
        gruppo_nome = gruppo.nome if gruppo else None
    
    ruolo_nome = None
    ruolo_codice = None
    if user.ruolo_id:
        ruolo = db.query(Ruolo).filter(Ruolo.id == user.ruolo_id).first()
        if ruolo:
            ruolo_nome = ruolo.nome
            ruolo_codice = ruolo.codice
    
    return {
        "id": user.id,
        "username": user.username,
        "nome": user.nome,
        "cognome": user.cognome,
        "email": user.email,
        "is_admin": user.is_admin,
        "is_active": user.is_active,
        "gruppo_id": user.gruppo_id,
        "gruppo_nome": gruppo_nome,
        "ruolo_id": user.ruolo_id,
        "ruolo_nome": ruolo_nome,
        "ruolo_codice": ruolo_codice,
        "permessi": permessi,
        "cliente_id": user.cliente_id,
    }


# ==========================================
# PREVENTIVI
# ==========================================
@app.get("/preventivi")
def get_preventivi(
    authorization: Optional[str] = Header(None),
    db: Session = Depends(get_db)
):
    """Ottieni preventivi ordinati per data. Filtra per cliente se utente è cliente."""
    user = _extract_user_from_request(authorization, db)
    
    query = db.query(Preventivo)
    if user:
        query = filter_preventivi_by_user(query, user, db)
    # Se nessun token (backward compat), mostra tutto
    
    preventivi = query.order_by(Preventivo.created_at.desc()).all()
    results = []
    for p in preventivi:
        row = {}
        for col in p.__table__.columns:
            val = getattr(p, col.name, None)
            if hasattr(val, 'isoformat'):
                val = val.isoformat()
            row[col.name] = val
        # Risolvi nome cliente da cliente_id
        if p.cliente_id and not row.get('customer_name'):
            cl = db.query(Cliente).filter(Cliente.id == p.cliente_id).first()
            if cl:
                row['customer_name'] = cl.ragione_sociale
        results.append(row)
    return results


@app.get("/preventivi/search")
def search_preventivi(
    q: str = None,
    status: str = None,
    categoria: str = None,
    data_da: str = None,
    data_a: str = None,
    limit: int = 50,
    authorization: Optional[str] = Header(None),
    db: Session = Depends(get_db)
):
    """Ricerca preventivi con filtri multipli. Filtra per cliente se utente è cliente."""
    query = db.query(Preventivo).outerjoin(
        Cliente, Preventivo.cliente_id == Cliente.id
    )

    # Filtro visibilità per utente
    user = _extract_user_from_request(authorization, db)
    if user:
        if not user_can_see_all_preventivi(user, db):
            if user.cliente_id:
                query = query.filter(Preventivo.cliente_id == user.cliente_id)
            else:
                query = query.filter(Preventivo.created_by == user.id)

    if q and q.strip():
        pattern = f"%{q.strip()}%"
        query = query.filter(
            or_(
                Preventivo.numero_preventivo.ilike(pattern),
                Preventivo.customer_name.ilike(pattern),
                Cliente.ragione_sociale.ilike(pattern),
                Cliente.codice.ilike(pattern),
            )
        )

    if status:
        query = query.filter(Preventivo.status == status)

    if categoria:
        query = query.filter(Preventivo.categoria == categoria.upper())

    if data_da:
        try:
            from datetime import datetime as dt
            d = dt.strptime(data_da, "%Y-%m-%d")
            query = query.filter(Preventivo.created_at >= d)
        except ValueError:
            pass

    if data_a:
        try:
            from datetime import datetime as dt
            d = dt.strptime(data_a, "%Y-%m-%d")
            # Include tutto il giorno
            d = d.replace(hour=23, minute=59, second=59)
            query = query.filter(Preventivo.created_at <= d)
        except ValueError:
            pass

    preventivi = query.order_by(Preventivo.created_at.desc()).limit(limit).all()

    results = []
    for p in preventivi:
        row = {}
        for col in p.__table__.columns:
            row[col.name] = getattr(p, col.name, None)
        # Aggiungi ragione_sociale cliente se disponibile
        if p.cliente_id:
            cl = db.query(Cliente).filter(Cliente.id == p.cliente_id).first()
            if cl:
                row['cliente_ragione_sociale'] = cl.ragione_sociale
                row['cliente_codice'] = cl.codice
        results.append(row)

    return results

@app.get("/preventivi/{preventivo_id}")
def get_preventivo(
    preventivo_id: int,
    authorization: Optional[str] = Header(None),
    db: Session = Depends(get_db)
):
    """Ottieni un preventivo specifico (con controllo accesso cliente)"""
    preventivo = db.query(Preventivo).filter(Preventivo.id == preventivo_id).first()
    if not preventivo:
        raise HTTPException(status_code=404, detail="Preventivo non trovato")
    
    # Controllo accesso
    user = _extract_user_from_request(authorization, db)
    if user and not user_can_see_all_preventivi(user, db):
        if user.cliente_id and preventivo.cliente_id != user.cliente_id:
            raise HTTPException(status_code=403, detail="Accesso negato a questo preventivo")
        elif not user.cliente_id and preventivo.created_by != user.id:
            raise HTTPException(status_code=403, detail="Accesso negato a questo preventivo")
    
    result = {}
    for col in preventivo.__table__.columns:
        result[col.name] = getattr(preventivo, col.name, None)
    return result

@app.post("/preventivi", status_code=201)
def create_preventivo(
    data: PreventivoCreate,
    authorization: Optional[str] = Header(None),
    db: Session = Depends(get_db)
):
    """Crea nuovo preventivo con generazione automatica numero preventivo."""
    # Genera numero preventivo automatico se non fornito
    if not data.numero_preventivo:
        current_year = datetime.now().year
        last_preventivo = (
            db.query(Preventivo)
            .filter(Preventivo.numero_preventivo.like(f"{current_year}/%"))
            .order_by(Preventivo.id.desc())
            .first()
        )
        
        if last_preventivo and "/" in (_prev_numero(last_preventivo)):
            try:
                last_number = int((_prev_numero(last_preventivo)).split("/")[1])
                new_number = last_number + 1
            except (ValueError, IndexError):
                new_number = 1
        else:
            new_number = 1
        
        numero_preventivo = f"{current_year}/{new_number:04d}"
    else:
        numero_preventivo = data.numero_preventivo
    
    # Carica template se specificato
    template_data_parsed = None
    template = None
    if data.template_id:
        template = db.query(ProductTemplate).filter(ProductTemplate.id == data.template_id).first()
        if template and template.template_data:
            try:
                template_data_parsed = json.loads(template.template_data)
            except json.JSONDecodeError:
                pass
    
    # Crea preventivo
    preventivo_dict = data.dict(exclude_unset=True)
    preventivo_dict['numero_preventivo'] = numero_preventivo
    template_id = preventivo_dict.pop('template_id', None)
    
    if template_id:
        preventivo_dict['template_id'] = template_id
        if template:
            preventivo_dict['categoria'] = template.categoria
    
    if 'created_at' not in preventivo_dict:
        preventivo_dict['created_at'] = datetime.now()
    if 'updated_at' not in preventivo_dict:
        preventivo_dict['updated_at'] = datetime.now()
    
    # Auto-set created_by e cliente_id dall'utente loggato
    user = _extract_user_from_request(authorization, db)
    if user:
        preventivo_dict['created_by'] = user.id
        if user.cliente_id and 'cliente_id' not in preventivo_dict:
            preventivo_dict['cliente_id'] = user.cliente_id

    preventivo = Preventivo(**preventivo_dict)
    db.add(preventivo)
    db.commit()
    db.refresh(preventivo)
    
    # Crea record ORM vuoti per le sezioni dedicate (backward compat)
    db.add_all([
        DatiCommessa(preventivo_id=preventivo.id, data_offerta=datetime.now().strftime('%d/%m/%Y')),
        DatiPrincipali(preventivo_id=preventivo.id),
        Normative(preventivo_id=preventivo.id),
        DisposizioneVano(preventivo_id=preventivo.id),
        Porte(preventivo_id=preventivo.id),
        Argano(preventivo_id=preventivo.id),
    ])
    db.commit()
    
    # Scrivi valori template in valori_configurazione (fonte primaria)
    if template_data_parsed:
        skip_keys = {"materiali", "field_config"}
        for sez_codice, campi in template_data_parsed.items():
            if sez_codice in skip_keys or not isinstance(campi, dict):
                continue
            for codice_campo, valore in campi.items():
                # Salta valori che sono dict/list (sotto-config, non valori reali)
                if isinstance(valore, (dict, list)):
                    continue
                if valore is not None and str(valore).strip() != "":
                    db.execute(text("""
                        INSERT INTO valori_configurazione
                            (preventivo_id, sezione, codice_campo, valore, is_default)
                        VALUES (:pid, :sez, :campo, :val, 1)
                    """), {
                        "pid": preventivo.id,
                        "sez": sez_codice,
                        "campo": codice_campo,
                        "val": str(valore),
                    })
        db.commit()
    
    # Aggiungi materiali da template se presenti
    if template_data_parsed and "materiali" in template_data_parsed:
        for mat in template_data_parsed["materiali"]:
            quantita = mat.get("quantita", 1)
            prezzo_unitario = mat.get("prezzo_unitario", 0.0)
            materiale = Materiale(
                preventivo_id=preventivo.id,
                codice=mat.get("codice", ""),
                descrizione=mat.get("descrizione", ""),
                quantita=quantita,
                prezzo_unitario=prezzo_unitario,
                prezzo_totale=quantita * prezzo_unitario,
                categoria=mat.get("categoria", "Template"),
                aggiunto_da_regola=False,
                note=f"Da template: {template.nome_display}" if template else None
            )
            db.add(materiale)
        
        db.commit()
        all_materials = db.query(Materiale).filter(Materiale.preventivo_id == preventivo.id).all()
        totale_calc = sum(m.prezzo_totale for m in all_materials)
        _prev_set(preventivo, totale_calc, "totale_materiali", "total_price")
        db.commit()
        db.refresh(preventivo)
    
    # Valuta regole (incluse TEMPLATE_BASE) per aggiungere articoli base automatici
    safe_evaluate_rules(preventivo.id, db)
    db.refresh(preventivo)

    # Inizializza defaults e applica field_config dal template
    try:
        campi_default = db.execute(text("""
            SELECT codice, sezione, valore_default
            FROM campi_configuratore
            WHERE attivo = 1 AND valore_default IS NOT NULL AND valore_default != ''
        """)).fetchall()

        existing_campos = db.execute(text("""
            SELECT codice_campo FROM valori_configurazione WHERE preventivo_id = :pid
        """), {"pid": preventivo.id}).fetchall()
        existing_set = {r[0] for r in existing_campos}

        # Ordina per ordine così {{campo:X}} trova X già nel context
        campi_default_ordinati = sorted(campi_default, key=lambda r: r[0])
        try:
            campi_ordine = {r[0]: r[1] for r in db.execute(text(
                "SELECT codice, ordine FROM campi_configuratore WHERE attivo=1"
            )).fetchall()}
            campi_default_ordinati = sorted(campi_default, key=lambda r: campi_ordine.get(r[0], 9999))
        except Exception:
            pass

        context_init: dict = {}
        for codice, sezione, valore_default in campi_default_ordinati:
            if codice not in existing_set:
                val = _resolve_default_expr(str(valore_default), preventivo.id, db, context=context_init)
                db.execute(text("""
                    INSERT INTO valori_configurazione
                        (preventivo_id, sezione, codice_campo, valore, is_default)
                    VALUES (:pid, :sez, :campo, :val, 1)
                """), {"pid": preventivo.id, "sez": sezione,
                       "campo": codice, "val": val})
                context_init[codice] = val
        db.commit()

        if data.template_id:
            _apply_template_field_config(db, preventivo.id, data.template_id)
            db.commit()
    except Exception as e:
        logger.warning(f"Errore inizializzazione defaults/field_config: {e}")
    
    return _orm_to_dict(preventivo)

@app.put("/preventivi/{preventivo_id}")
def update_preventivo(preventivo_id: int, data: PreventivoUpdate, db: Session = Depends(get_db)):
    """Aggiorna preventivo"""
    preventivo = db.query(Preventivo).filter(Preventivo.id == preventivo_id).first()
    if not preventivo:
        raise HTTPException(status_code=404, detail="Preventivo non trovato")
    
    update_dict = data.dict(exclude_unset=True)
    update_dict['updated_at'] = datetime.now()
    
    for key, value in update_dict.items():
        setattr(preventivo, key, value)
    
    db.commit()
    db.refresh(preventivo)
    return _orm_to_dict(preventivo)

@app.delete("/preventivi/{preventivo_id}")
def delete_preventivo(preventivo_id: int, db: Session = Depends(get_db)):
    """Elimina preventivo"""
    preventivo = db.query(Preventivo).filter(Preventivo.id == preventivo_id).first()
    if not preventivo:
        raise HTTPException(status_code=404, detail="Preventivo non trovato")
    db.delete(preventivo)
    db.commit()
    return {"message": "Preventivo eliminato"}

# ==========================================
# DATI COMMESSA
# ==========================================
@app.get("/preventivi/{preventivo_id}/dati-commessa")
def get_dati_commessa(preventivo_id: int, db: Session = Depends(get_db)):
    dati = db.query(DatiCommessa).filter(DatiCommessa.preventivo_id == preventivo_id).first()
    if not dati:
        dati = DatiCommessa(preventivo_id=preventivo_id)
        db.add(dati)
        db.commit()
        db.refresh(dati)
    # Ritorna dict manuale per evitare errori schema
    result = {}
    for col in dati.__table__.columns:
        result[col.name] = getattr(dati, col.name, None)
    # Alias per frontend
    if "consegna_richiesta" in result:
        result["data_consegna_richiesta"] = result["consegna_richiesta"]
    return result

@app.put("/preventivi/{preventivo_id}/dati-commessa")
def update_dati_commessa(preventivo_id: int, data: DatiCommessaUpdate, db: Session = Depends(get_db)):
    try:
        dati = db.query(DatiCommessa).filter(DatiCommessa.preventivo_id == preventivo_id).first()
        if not dati:
            dati = DatiCommessa(preventivo_id=preventivo_id)
            db.add(dati)
        update_data = data.dict(exclude_unset=True)
        print(f"[DATI COMMESSA] PUT preventivo_id={preventivo_id}, fields={list(update_data.keys())}")
        # Salva campi dati_commessa
        cliente_id_value = update_data.pop('cliente_id', None)
        # Mapping nomi frontend -> nomi ORM
        field_aliases = {
            "data_consegna_richiesta": "consegna_richiesta",
        }
        for key, value in update_data.items():
            orm_key = field_aliases.get(key, key)
            if hasattr(dati, orm_key):
                setattr(dati, orm_key, value)
            else:
                print(f"[DATI COMMESSA] WARNING: campo '{key}' (orm: '{orm_key}') non esiste nel modello DatiCommessa")
        
        # Se presente cliente_id, aggiorna direttamente il preventivo
        if cliente_id_value is not None:
            preventivo = db.query(Preventivo).filter(Preventivo.id == preventivo_id).first()
            if preventivo:
                preventivo.cliente_id = cliente_id_value
                print(f"[DATI COMMESSA] Aggiornato cliente_id={cliente_id_value} sul preventivo {preventivo_id}")
        
        db.commit()
        db.refresh(dati)
        safe_evaluate_rules(preventivo_id, db)
        touch_preventivo(preventivo_id, db)
        # Ritorna dict manuale per evitare errori schema
        result = {}
        for col in dati.__table__.columns:
            result[col.name] = getattr(dati, col.name, None)
        return result
    except Exception as e:
        import traceback
        traceback.print_exc()
        db.rollback()
        raise HTTPException(status_code=500, detail=f"Errore salvataggio dati commessa: {str(e)}")
    
# ============================================================
# FILE: main.py â€” DELETE CLIENTE (aggiungere dopo PUT /clienti/{id})
# ============================================================

@app.delete("/clienti/{cliente_id}")
def delete_cliente(cliente_id: int, db: Session = Depends(get_db)):
    """Disattiva un cliente (soft delete)"""
    cliente = db.query(Cliente).filter(Cliente.id == cliente_id).first()
    if not cliente:
        raise HTTPException(404, "Cliente non trovato")
    cliente.is_active = False
    db.commit()
    return {"status": "ok"}

# ==========================================
# DATI PRINCIPALI
# ==========================================
@app.get("/preventivi/{preventivo_id}/dati-principali")
def get_dati_principali(preventivo_id: int, db: Session = Depends(get_db)):
    dati = db.query(DatiPrincipali).filter(DatiPrincipali.preventivo_id == preventivo_id).first()
    if not dati:
        dati = DatiPrincipali(preventivo_id=preventivo_id)
        db.add(dati)
        db.commit()
        db.refresh(dati)
    return _orm_to_dict(dati)

@app.put("/preventivi/{preventivo_id}/dati-principali")
def update_dati_principali(preventivo_id: int, data: DatiPrincipaliUpdate, db: Session = Depends(get_db)):
    dati = db.query(DatiPrincipali).filter(DatiPrincipali.preventivo_id == preventivo_id).first()
    if not dati:
        dati = DatiPrincipali(preventivo_id=preventivo_id)
        db.add(dati)
    update_data = data.dict(exclude_unset=True)
    print(f"[DATI_PRINCIPALI] PUT preventivo_id={preventivo_id}, fields={list(update_data.keys())}")
    for key, value in update_data.items():
        if hasattr(dati, key):
            setattr(dati, key, value)
        else:
            print(f"[DATI_PRINCIPALI] WARNING: campo '{key}' non esiste nel modello")
    db.commit()
    db.refresh(dati)
    rule_result = safe_evaluate_rules(preventivo_id, db)
    touch_preventivo(preventivo_id, db)
    result = _orm_to_dict(dati)
    if rule_result and isinstance(rule_result, dict):
        result["materials_added"] = rule_result.get("materiali_aggiunti", 0)
        result["materials_removed"] = rule_result.get("materiali_rimossi", 0)
        result["active_rules"] = rule_result.get("regole_attive", [])
    return result

# ==========================================
# NORMATIVE
# ==========================================
@app.get("/preventivi/{preventivo_id}/normative")
def get_normative(preventivo_id: int, db: Session = Depends(get_db)):
    normative = db.query(Normative).filter(Normative.preventivo_id == preventivo_id).first()
    if not normative:
        normative = Normative(preventivo_id=preventivo_id)
        db.add(normative)
        db.commit()
        db.refresh(normative)
    return _orm_to_dict(normative)

@app.put("/preventivi/{preventivo_id}/normative")
def update_normative(preventivo_id: int, data: NormativeUpdate, db: Session = Depends(get_db)):
    normative = db.query(Normative).filter(Normative.preventivo_id == preventivo_id).first()
    if not normative:
        normative = Normative(preventivo_id=preventivo_id)
        db.add(normative)
    update_data = data.dict(exclude_unset=True)
    print(f"[NORMATIVE] PUT preventivo_id={preventivo_id}, fields={update_data}")
    for key, value in update_data.items():
        if hasattr(normative, key):
            setattr(normative, key, value)
        else:
            print(f"[NORMATIVE] WARNING: campo '{key}' non esiste nel modello")
    db.commit()
    db.refresh(normative)
    rule_result = safe_evaluate_rules(preventivo_id, db)
    touch_preventivo(preventivo_id, db)
    result = _orm_to_dict(normative)
    # Aggiungi info regole alla risposta
    if rule_result and isinstance(rule_result, dict):
        result["materials_added"] = rule_result.get("materiali_aggiunti", 0)
        result["materials_removed"] = rule_result.get("materiali_rimossi", 0)
        result["active_rules"] = rule_result.get("regole_attive", [])
    return result

# ==========================================
# DISPOSIZIONE VANO
# ==========================================
@app.get("/preventivi/{preventivo_id}/disposizione-vano")
def get_disposizione_vano(preventivo_id: int, db: Session = Depends(get_db)):
    disposizione = db.query(DisposizioneVano).filter(DisposizioneVano.preventivo_id == preventivo_id).first()
    if not disposizione:
        disposizione = DisposizioneVano(preventivo_id=preventivo_id)
        db.add(disposizione)
        db.commit()
        db.refresh(disposizione)
    return _orm_to_dict(disposizione)

@app.put("/preventivi/{preventivo_id}/disposizione-vano")
def update_disposizione_vano(preventivo_id: int, data: DisposizioneVanoUpdate, db: Session = Depends(get_db)):
    disposizione = db.query(DisposizioneVano).filter(DisposizioneVano.preventivo_id == preventivo_id).first()
    if not disposizione:
        disposizione = DisposizioneVano(preventivo_id=preventivo_id)
        db.add(disposizione)
    
    data_dict = data.dict(exclude_unset=True)
    cleaned_data = {}
    for key, value in data_dict.items():
        if value == '' or value == '{}' or value == '[]':
            cleaned_data[key] = None
        else:
            cleaned_data[key] = value
    
    for key, value in cleaned_data.items():
        if hasattr(disposizione, key):
            setattr(disposizione, key, value)
        else:
            print(f"[DISPOSIZIONE_VANO] WARNING: campo '{key}' non esiste nel modello")
    
    db.commit()
    db.refresh(disposizione)
    touch_preventivo(preventivo_id, db)
    return _orm_to_dict(disposizione)

@app.delete("/preventivi/{preventivo_id}/disposizione-vano")
def delete_disposizione_vano(preventivo_id: int, db: Session = Depends(get_db)):
    disposizione = db.query(DisposizioneVano).filter(DisposizioneVano.preventivo_id == preventivo_id).first()
    if disposizione:
        db.delete(disposizione)
        db.commit()
    return {"message": "Disposizione vano eliminata"}

# ==========================================
# PORTE
# ==========================================
@app.get("/preventivi/{preventivo_id}/porte")
def get_porte(preventivo_id: int, db: Session = Depends(get_db)):
    porte = db.query(Porte).filter(Porte.preventivo_id == preventivo_id).first()
    if not porte:
        porte = Porte(preventivo_id=preventivo_id)
        db.add(porte)
        db.commit()
        db.refresh(porte)
    return _orm_to_dict(porte)

@app.put("/preventivi/{preventivo_id}/porte")
def update_porte(preventivo_id: int, data: PorteUpdate, db: Session = Depends(get_db)):
    porte = db.query(Porte).filter(Porte.preventivo_id == preventivo_id).first()
    if not porte:
        porte = Porte(preventivo_id=preventivo_id)
        db.add(porte)
    for key, value in data.dict(exclude_unset=True).items():
        if hasattr(porte, key):
            setattr(porte, key, value)
        else:
            print(f"[PORTE] WARNING: campo '{key}' non esiste nel modello")
    db.commit()
    db.refresh(porte)
    safe_evaluate_rules(preventivo_id, db)
    touch_preventivo(preventivo_id, db)
    return _orm_to_dict(porte)

# ==========================================
# ARGANO
# ==========================================
@app.get("/preventivi/{preventivo_id}/argano")
def get_argano(preventivo_id: int, db: Session = Depends(get_db)):
    """Ottieni dati argano, crea se non esiste"""
    argano = db.query(Argano).filter(Argano.preventivo_id == preventivo_id).first()
    if not argano:
        argano = Argano(preventivo_id=preventivo_id)
        db.add(argano)
        db.commit()
        db.refresh(argano)
    return {
        "id": argano.id,
        "preventivo_id": argano.preventivo_id,
        "trazione": argano.trazione,
        "potenza_motore_kw": argano.potenza_motore_kw,
        "corrente_nom_motore_amp": argano.corrente_nom_motore_amp,
        "tipo_vvvf": argano.tipo_vvvf,
        "vvvf_nel_vano": argano.vvvf_nel_vano,
        "freno_tensione": argano.freno_tensione,
        "ventilazione_forzata": argano.ventilazione_forzata,
        "tipo_teleruttore": argano.tipo_teleruttore
    }

@app.put("/preventivi/{preventivo_id}/argano")
def update_argano(preventivo_id: int, data: dict, db: Session = Depends(get_db)):
    """Aggiorna dati argano"""
    argano = db.query(Argano).filter(Argano.preventivo_id == preventivo_id).first()
    if not argano:
        argano = Argano(preventivo_id=preventivo_id)
        db.add(argano)
    
    for key in ["trazione", "potenza_motore_kw", "corrente_nom_motore_amp",
                "tipo_vvvf", "vvvf_nel_vano", "freno_tensione",
                "ventilazione_forzata", "tipo_teleruttore"]:
        if key in data:
            setattr(argano, key, data[key])
    
    db.commit()
    db.refresh(argano)

    safe_evaluate_rules(preventivo_id, db)    # <-- AGGIUNGI QUESTA RIGA
    touch_preventivo(preventivo_id, db)

    return {
        "id": argano.id,
        "preventivo_id": argano.preventivo_id,
        "trazione": argano.trazione,
        "potenza_motore_kw": argano.potenza_motore_kw,
        "corrente_nom_motore_amp": argano.corrente_nom_motore_amp,
        "tipo_vvvf": argano.tipo_vvvf,
        "vvvf_nel_vano": argano.vvvf_nel_vano,
        "freno_tensione": argano.freno_tensione,
        "ventilazione_forzata": argano.ventilazione_forzata,
        "tipo_teleruttore": argano.tipo_teleruttore
    }


# ==========================================
# MATERIALI
# ==========================================
@app.get("/preventivi/{preventivo_id}/materiali")
def get_materiali(preventivo_id: int, db: Session = Depends(get_db)):
    materiali = db.query(Materiale).filter(Materiale.preventivo_id == preventivo_id).all()
    result = []
    for m in materiali:
        d = _orm_to_dict(m)
        art = db.execute(
            text("SELECT id FROM articoli WHERE codice=:codice"),
            {"codice": m.codice}
        ).fetchone()
        d["articolo_non_trovato"] = art is None
        result.append(d)
    return result

@app.post("/preventivi/{preventivo_id}/materiali")
def create_materiale(preventivo_id: int, data: MaterialeCreate, db: Session = Depends(get_db)):
    prezzo_totale = data.quantita * data.prezzo_unitario
    materiale = Materiale(
        preventivo_id=preventivo_id,
        codice=data.codice, descrizione=data.descrizione,
        quantita=data.quantita, prezzo_unitario=data.prezzo_unitario,
        prezzo_totale=prezzo_totale, categoria=data.categoria,
        aggiunto_da_regola=data.aggiunto_da_regola,
        regola_id=data.regola_id, note=data.note
    )
    db.add(materiale)
    db.commit()
    db.refresh(materiale)
    
    preventivo = db.query(Preventivo).filter(Preventivo.id == preventivo_id).first()
    if preventivo:
        all_materials = db.query(Materiale).filter(Materiale.preventivo_id == preventivo_id).all()
        totale_calc = sum(m.prezzo_totale for m in all_materials)
        _prev_set(preventivo, totale_calc, "totale_materiali", "total_price")
        db.commit()
    touch_preventivo(preventivo_id, db)
    return _orm_to_dict(materiale)

@app.put("/preventivi/{preventivo_id}/materiali/{materiale_id}")
def update_materiale(preventivo_id: int, materiale_id: int, data: MaterialeUpdate, db: Session = Depends(get_db)):
    materiale = db.query(Materiale).filter(Materiale.id == materiale_id, Materiale.preventivo_id == preventivo_id).first()
    if not materiale:
        raise HTTPException(status_code=404, detail="Materiale non trovato")
    
    update_data = data.dict(exclude_unset=True)
    for key, value in update_data.items():
        setattr(materiale, key, value)
    
    if 'quantita' in update_data or 'prezzo_unitario' in update_data:
        materiale.prezzo_totale = materiale.quantita * materiale.prezzo_unitario
    
    db.commit()
    db.refresh(materiale)
    
    preventivo = db.query(Preventivo).filter(Preventivo.id == preventivo_id).first()
    if preventivo:
        all_materials = db.query(Materiale).filter(Materiale.preventivo_id == preventivo_id).all()
        totale_calc = sum(m.prezzo_totale for m in all_materials)
        _prev_set(preventivo, totale_calc, "totale_materiali", "total_price")
        db.commit()
    touch_preventivo(preventivo_id, db)
    return _orm_to_dict(materiale)

@app.delete("/preventivi/{preventivo_id}/materiali/{materiale_id}")
def delete_materiale(preventivo_id: int, materiale_id: int, db: Session = Depends(get_db)):
    materiale = db.query(Materiale).filter(Materiale.id == materiale_id, Materiale.preventivo_id == preventivo_id).first()
    if not materiale:
        raise HTTPException(status_code=404, detail="Materiale non trovato")
    db.delete(materiale)
    db.commit()
    
    preventivo = db.query(Preventivo).filter(Preventivo.id == preventivo_id).first()
    if preventivo:
        all_materials = db.query(Materiale).filter(Materiale.preventivo_id == preventivo_id).all()
        totale_calc = sum(m.prezzo_totale for m in all_materials)
        _prev_set(preventivo, totale_calc, "totale_materiali", "total_price")
        db.commit()
    touch_preventivo(preventivo_id, db)
    return {"message": "Materiale eliminato"}

# ==========================================
# RIGHE RICAMBIO
# ==========================================
@app.get("/preventivi/{preventivo_id}/righe-ricambio")
def get_righe_ricambio(preventivo_id: int, db: Session = Depends(get_db)):
    righe = db.query(RigaRicambio).filter(
        RigaRicambio.preventivo_id == preventivo_id
    ).order_by(RigaRicambio.ordine).all()
    return [{
        "id": r.id, "preventivo_id": r.preventivo_id,
        "articolo_id": r.articolo_id, "codice": r.codice,
        "descrizione": r.descrizione, "tipo_articolo": r.tipo_articolo,
        "quantita": r.quantita,
        "parametro_1": r.parametro_1, "unita_param_1": r.unita_param_1,
        "desc_param_1": r.desc_param_1, "costo_var_1": r.costo_var_1,
        "parametro_2": r.parametro_2, "unita_param_2": r.unita_param_2,
        "desc_param_2": r.desc_param_2, "costo_var_2": r.costo_var_2,
        "parametro_3": r.parametro_3, "unita_param_3": r.unita_param_3,
        "desc_param_3": r.desc_param_3, "costo_var_3": r.costo_var_3,
        "parametro_4": r.parametro_4, "unita_param_4": r.unita_param_4,
        "desc_param_4": r.desc_param_4, "costo_var_4": r.costo_var_4,
        "costo_fisso": r.costo_fisso, "costo_base_unitario": r.costo_base_unitario,
        "ricarico_percentuale": r.ricarico_percentuale,
        "prezzo_listino_unitario": r.prezzo_listino_unitario,
        "sconto_cliente": r.sconto_cliente,
        "prezzo_cliente_unitario": r.prezzo_cliente_unitario,
        "prezzo_totale_listino": r.prezzo_totale_listino,
        "prezzo_totale_cliente": r.prezzo_totale_cliente,
        "ordine": r.ordine, "note": r.note
    } for r in righe]

@app.post("/preventivi/{preventivo_id}/righe-ricambio")
def create_riga_ricambio(preventivo_id: int, data: dict, db: Session = Depends(get_db)):
    riga = RigaRicambio(preventivo_id=preventivo_id)
    for key in [
        "articolo_id", "codice", "descrizione", "tipo_articolo", "quantita",
        "parametro_1", "unita_param_1", "desc_param_1", "costo_var_1",
        "parametro_2", "unita_param_2", "desc_param_2", "costo_var_2",
        "parametro_3", "unita_param_3", "desc_param_3", "costo_var_3",
        "parametro_4", "unita_param_4", "desc_param_4", "costo_var_4",
        "costo_fisso", "costo_base_unitario", "ricarico_percentuale",
        "prezzo_listino_unitario", "sconto_cliente", "prezzo_cliente_unitario",
        "prezzo_totale_listino", "prezzo_totale_cliente", "ordine", "note"
    ]:
        if key in data:
            setattr(riga, key, data[key])
    
    db.add(riga)
    db.commit()
    db.refresh(riga)
    touch_preventivo(preventivo_id, db)
    return {"id": riga.id, "status": "ok"}

@app.put("/preventivi/{preventivo_id}/righe-ricambio/{riga_id}")
def update_riga_ricambio(preventivo_id: int, riga_id: int, data: dict, db: Session = Depends(get_db)):
    riga = db.query(RigaRicambio).filter(
        RigaRicambio.id == riga_id,
        RigaRicambio.preventivo_id == preventivo_id
    ).first()
    if not riga:
        raise HTTPException(status_code=404, detail="Riga ricambio non trovata")
    
    for key in [
        "articolo_id", "codice", "descrizione", "tipo_articolo", "quantita",
        "parametro_1", "unita_param_1", "desc_param_1", "costo_var_1",
        "parametro_2", "unita_param_2", "desc_param_2", "costo_var_2",
        "parametro_3", "unita_param_3", "desc_param_3", "costo_var_3",
        "parametro_4", "unita_param_4", "desc_param_4", "costo_var_4",
        "costo_fisso", "costo_base_unitario", "ricarico_percentuale",
        "prezzo_listino_unitario", "sconto_cliente", "prezzo_cliente_unitario",
        "prezzo_totale_listino", "prezzo_totale_cliente", "ordine", "note"
    ]:
        if key in data:
            setattr(riga, key, data[key])
    
    db.commit()
    db.refresh(riga)
    touch_preventivo(preventivo_id, db)
    return {"id": riga.id, "status": "ok"}

@app.delete("/preventivi/{preventivo_id}/righe-ricambio/{riga_id}")
def delete_riga_ricambio(preventivo_id: int, riga_id: int, db: Session = Depends(get_db)):
    riga = db.query(RigaRicambio).filter(
        RigaRicambio.id == riga_id,
        RigaRicambio.preventivo_id == preventivo_id
    ).first()
    if not riga:
        raise HTTPException(status_code=404, detail="Riga ricambio non trovata")
    db.delete(riga)
    db.commit()
    touch_preventivo(preventivo_id, db)
    return {"status": "ok"}

# ==========================================
# SCONTO ADMIN
# ==========================================
@app.put("/preventivi/{preventivo_id}/sconto-admin")
def update_sconto_admin(preventivo_id: int, data: dict, db: Session = Depends(get_db)):
    """Aggiorna sconto extra admin per un preventivo"""
    preventivo = db.query(Preventivo).filter(Preventivo.id == preventivo_id).first()
    if not preventivo:
        raise HTTPException(status_code=404, detail="Preventivo non trovato")
    
    if "sconto_extra_admin" in data:
        preventivo.sconto_extra_admin = data["sconto_extra_admin"]
    
    # Ricalcola prezzo finale
    sconto_totale = (preventivo.sconto_cliente or 0) + (preventivo.sconto_extra_admin or 0)
    if _prev_totale(preventivo):
        totale_netto_calc = _prev_totale(preventivo) * (1 - sconto_totale / 100)
        _prev_set(preventivo, totale_netto_calc, 'totale_netto', 'total_price_finale')
    
    preventivo.updated_at = datetime.now()
    db.commit()
    db.refresh(preventivo)
    return {
        "sconto_extra_admin": preventivo.sconto_extra_admin,
        "total_price_finale": _prev_netto(preventivo)
    }

# ==========================================
# EXPORT PREVENTIVO
# ==========================================
@app.get("/preventivi/{preventivo_id}/export/{formato}")
def export_preventivo(preventivo_id: int, formato: str, request: Request, db: Session = Depends(get_db)):
    with open("debug_export.txt", "a", encoding="utf-8") as dbg:
        dbg.write(f"EXPORT CHIAMATO preventivo_id={preventivo_id} formato={formato}\n")
    preventivo = db.query(Preventivo).filter(Preventivo.id == preventivo_id).first()
    if not preventivo:
        raise HTTPException(status_code=404, detail="Preventivo non trovato")
    
    dati_commessa = db.query(DatiCommessa).filter(DatiCommessa.preventivo_id == preventivo_id).first()
    dati_principali = db.query(DatiPrincipali).filter(DatiPrincipali.preventivo_id == preventivo_id).first()
    normative_data = db.query(Normative).filter(Normative.preventivo_id == preventivo_id).first()
    materiali = db.query(Materiale).filter(Materiale.preventivo_id == preventivo_id).all()
    argano_data = None
    try:
        argano_data = db.query(Argano).filter(Argano.preventivo_id == preventivo_id).first()
    except Exception:
        pass
    cliente = None
    try:
        cid = getattr(preventivo, 'cliente_id', None)
        if cid:
            cliente = db.query(Cliente).filter(Cliente.id == cid).first()
    except Exception:
        pass
    
    # ── Carica icona prodotto dal template associato ──
    product_icon_png = None
    product_name = ""
    product_category = ""
    try:
        tid = getattr(preventivo, 'template_id', None)
        if tid:
            prod_tpl = db.query(ProductTemplate).filter(ProductTemplate.id == tid).first()
            if prod_tpl:
                product_name = getattr(prod_tpl, 'nome_display', '') or ''
                product_category = getattr(prod_tpl, 'categoria', '') or ''
                icon_name = getattr(prod_tpl, 'icona', None)
                if icon_name:
                    from export_utils import load_product_icon_png
                    product_icon_png = load_product_icon_png(icon_name, size=160)
    except Exception as e:
        print(f"WARN: Impossibile caricare icona prodotto: {e}")
    
    fmt = formato.lower()
    filename_base = f"preventivo_{_prev_numero(preventivo).replace('/', '_')}"
    
    if fmt == "json":
        return {
            "preventivo": {
                "id": preventivo.id,
                "numero_preventivo": _prev_numero(preventivo),
                "status": _prev_stato(preventivo),
                "total_price": _prev_totale(preventivo),
                "customer_name": getattr(preventivo, 'customer_name', ''),
                "created_at": str(preventivo.created_at) if preventivo.created_at else None,
            },
            "dati_commessa": {k: v for k, v in (dati_commessa.__dict__.items() if dati_commessa else {}) if not k.startswith('_')},
            "dati_principali": {k: v for k, v in (dati_principali.__dict__.items() if dati_principali else {}) if not k.startswith('_')},
            "materiali": [{
                "codice": m.codice, "descrizione": m.descrizione,
                "quantita": m.quantita, "prezzo_unitario": m.prezzo_unitario,
                "prezzo_totale": m.prezzo_totale, "categoria": m.categoria
            } for m in materiali]
        }
    
    elif fmt in ("docx", "pdf"):
        print(f"[DEBUG EXPORT] Cerco template... doc_template_id={request.query_params.get('template_id')}")
        # --- Template personalizzato ---
        doc_template_id = request.query_params.get("template_id")
        doc_template = None
        if doc_template_id:
            doc_template = db.query(DocumentTemplate).filter(
                DocumentTemplate.id == int(doc_template_id)
            ).first()
        else:
            doc_template = db.query(DocumentTemplate).filter(
                DocumentTemplate.tipo == "preventivo",
                DocumentTemplate.is_default == True,
                DocumentTemplate.attivo == True
            ).first()

        if doc_template and doc_template.config:
            print(f"[DEBUG EXPORT] Template trovato: id={doc_template.id} nome={doc_template.nome}")
            try:
                available = get_available_fields_from_db(db)
                valori_din = load_valori_dinamici(db, preventivo_id)
                def_info = load_defaults_info(db, preventivo_id)
                buf = genera_docx_da_template(
                    template_config=doc_template.config,
                    preventivo=preventivo,
                    dati_commessa=dati_commessa,
                    dati_principali=dati_principali,
                    normative=normative_data,
                    argano=argano_data,
                    materiali=materiali,
                    cliente=cliente,
                    logo_data=doc_template.logo_data,
                    logo_mime=doc_template.logo_mime,
                    valori_dinamici=valori_din,
                    available_fields=available,
                    defaults_info=def_info,
                    product_icon_png=product_icon_png,
                    product_name=product_name,
                    product_category=product_category,
                )
                return StreamingResponse(buf,
                    media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    headers={"Content-Disposition": f'attachment; filename="{filename_base}.docx"'})
            except Exception as e:
                print(f"Errore template personalizzato, fallback: {e}")

        # --- Fallback: generazione standard ---
        print(f"[DEBUG EXPORT] FALLBACK - nessun template, uso genera_docx_preventivo_v2")
        try:
            dati_doc = get_dati_documento_preventivo(preventivo_id, db)
            preventivo_info = {
                "numero": _prev_numero(preventivo),
                "customer": getattr(preventivo, 'customer_name', '') or '',
                "status": _prev_stato(preventivo),
                "totale": safe_float(_prev_totale(preventivo)),
                "sconto": safe_float(getattr(preventivo, 'sconto_cliente', 0)) + safe_float(getattr(preventivo, 'sconto_extra_admin', 0)),
                "netto": safe_float(_prev_netto(preventivo)),
                "note": getattr(preventivo, 'note', '') or '',
                "revisione": getattr(preventivo, 'revisione_corrente', 0) or 0,
            }
            if cliente:
                preventivo_info["customer"] = getattr(cliente, 'ragione_sociale', preventivo_info["customer"]) or preventivo_info["customer"]
            from export_utils import genera_docx_preventivo_v2
            buf = genera_docx_preventivo_v2(preventivo_info, dati_doc, materiali,
                                             product_icon_png=product_icon_png,
                                             product_name=product_name,
                                             product_category=product_category)
            return StreamingResponse(buf,
                media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                headers={"Content-Disposition": f'attachment; filename="{filename_base}.docx"'})
        except ImportError as e:
            raise HTTPException(status_code=500, detail=f"Dipendenza mancante: {e}. Installa: pip install python-docx")
    
    elif fmt in ("xlsx", "excel", "xls"):
        try:
            from export_utils import genera_xlsx_preventivo
            buf = genera_xlsx_preventivo(preventivo, dati_commessa, dati_principali, normative_data, argano_data, materiali, cliente)
            return StreamingResponse(buf,
                media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                headers={"Content-Disposition": f'attachment; filename="{filename_base}.xlsx"'})
        except ImportError as e:
            raise HTTPException(status_code=500, detail=f"Dipendenza mancante: {e}. Installa: pip install openpyxl")
    
    raise HTTPException(status_code=400, detail=f"Formato '{formato}' non supportato. Usa: json, docx, xlsx")


@app.get("/ordini/{ordine_id}/export/{formato}")
def export_ordine(ordine_id: int, formato: str, db: Session = Depends(get_db)):
    """Esporta ordine in formato JSON, DOCX o XLSX (include BOM esplosa e lista acquisti)"""
    conn = db.get_bind().raw_connection()
    cursor = conn.cursor()
    
    cursor.execute("SELECT * FROM ordini WHERE id = ?", (ordine_id,))
    r = cursor.fetchone()
    if not r:
        raise HTTPException(status_code=404, detail="Ordine non trovato")
    cols = [desc[0] for desc in cursor.description]
    ordine_data = dict(zip(cols, r))
    
    if ordine_data.get('cliente_id'):
        try:
            cl = db.query(Cliente).filter(Cliente.id == ordine_data['cliente_id']).first()
            ordine_data['cliente'] = cl.ragione_sociale if cl else ''
        except Exception:
            ordine_data['cliente'] = ''
    
    materiali = []
    if ordine_data.get('preventivo_id'):
        materiali = db.query(Materiale).filter(
            Materiale.preventivo_id == ordine_data['preventivo_id']
        ).order_by(Materiale.ordine).all()
    
    esplosi = []
    try:
        cursor.execute("SELECT * FROM esplosi WHERE ordine_id = ? ORDER BY tipo, categoria, codice", (ordine_id,))
        ecols = [d[0] for d in cursor.description]
        esplosi = [dict(zip(ecols, row)) for row in cursor.fetchall()]
    except Exception:
        pass
    
    lista_acquisti = None
    try:
        cursor.execute(
            "SELECT e.codice, e.descrizione, e.quantita, e.unita_misura, e.costo_unitario, e.costo_totale, "
            "ab.fornitore, ab.codice_fornitore, ab.lead_time_acquisto "
            "FROM esplosi e LEFT JOIN articoli_bom ab ON ab.codice = e.codice "
            "WHERE e.ordine_id = ? AND e.tipo = 'ACQUISTO' ORDER BY ab.fornitore, e.codice",
            (ordine_id,)
        )
        fornitori = {}
        for row in cursor.fetchall():
            f = row[6] or "Non assegnato"
            if f not in fornitori:
                fornitori[f] = {"articoli": [], "totale": 0}
            fornitori[f]["articoli"].append({
                "codice": row[0], "descrizione": row[1], "quantita": row[2],
                "unita": row[3], "costo_unitario": row[4], "costo_totale": row[5],
                "codice_fornitore": row[7], "lead_time_giorni": row[8]
            })
            fornitori[f]["totale"] += row[5] or 0
        if fornitori:
            lista_acquisti = {"fornitori": {n: {"totale": round(d["totale"], 2), "articoli": d["articoli"]} for n, d in fornitori.items()}}
    except Exception:
        pass
    
    fmt = formato.lower()
    num_ordine = ordine_data.get('numero_ordine', str(ordine_id)).replace('/', '_')
    
    if fmt == "json":
        return {
            "ordine": ordine_data,
            "materiali": [{"codice": m.codice, "descrizione": m.descrizione,
                "quantita": m.quantita, "prezzo_unitario": m.prezzo_unitario,
                "prezzo_totale": m.prezzo_totale} for m in materiali],
            "esplosi": esplosi,
            "lista_acquisti": lista_acquisti
        }
    
    elif fmt in ("docx", "pdf"):
        try:
            from export_utils import genera_docx_ordine
            buf = genera_docx_ordine(ordine_data, materiali, esplosi, lista_acquisti)
            return StreamingResponse(buf,
                media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                headers={"Content-Disposition": f'attachment; filename="ordine_{num_ordine}.docx"'})
        except ImportError as e:
            raise HTTPException(status_code=500, detail=f"Dipendenza mancante: {e}. Installa: pip install python-docx")
    
    elif fmt in ("xlsx", "excel", "xls"):
        try:
            from export_utils import genera_xlsx_ordine
            buf = genera_xlsx_ordine(ordine_data, materiali, esplosi, lista_acquisti)
            return StreamingResponse(buf,
                media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                headers={"Content-Disposition": f'attachment; filename="ordine_{num_ordine}.xlsx"'})
        except ImportError as e:
            raise HTTPException(status_code=500, detail=f"Dipendenza mancante: {e}. Installa: pip install openpyxl")
    
    raise HTTPException(status_code=400, detail=f"Formato '{formato}' non supportato. Usa: json, docx, xlsx")


# ==========================================
# HELPER: Applica field_config da template a valori_configurazione
# ==========================================

def _apply_template_field_config(db, preventivo_id: int, template_id: int):
    """
    Legge field_config dal template e aggiorna i flag su valori_configurazione.
    Chiamare DOPO che i valori sono stati inseriti in valori_configurazione.
    """
    template = db.query(ProductTemplate).filter(ProductTemplate.id == template_id).first()
    if not template or not template.template_data:
        return 0

    try:
        td = json.loads(template.template_data)
    except json.JSONDecodeError:
        return 0

    field_config = td.get("field_config", {})
    if not field_config:
        return 0

    count = 0
    for codice_campo, cfg in field_config.items():
        # Costruisci UPDATE dinamico: solo flag esplicitamente presenti nel template
        updates = []
        params = {"pid": preventivo_id, "campo": codice_campo}

        if "readonly" in cfg:
            updates.append("is_readonly = :ro")
            params["ro"] = 1 if cfg["readonly"] else 0

        if "includi_preventivo" in cfg:
            updates.append("includi_preventivo = :inc")
            params["inc"] = 1 if cfg["includi_preventivo"] else 0

        if "mostra_default" in cfg:
            updates.append("mostra_default_preventivo = :md")
            params["md"] = 1 if cfg["mostra_default"] else 0

        if not updates:
            continue

        result = db.execute(text(f"""
            UPDATE valori_configurazione
            SET {', '.join(updates)}
            WHERE preventivo_id = :pid AND codice_campo = :campo
        """), params)

        if result.rowcount > 0:
            count += 1

    return count


# ==========================================
# PRODUCT TEMPLATES
# ==========================================
@app.get("/templates")
def get_templates(categoria: str = None, db: Session = Depends(get_db)):
    query = db.query(ProductTemplate).filter(ProductTemplate.attivo == True)
    if categoria:
        query = query.filter(ProductTemplate.categoria == categoria.upper())
    return _orm_to_dict(query.order_by(ProductTemplate.categoria, ProductTemplate.ordine).all())

@app.get("/templates/all")
def get_all_templates(db: Session = Depends(get_db)):
    return _orm_to_dict(db.query(ProductTemplate).order_by(ProductTemplate.categoria, ProductTemplate.ordine).all())

@app.get("/templates/categories/summary")
def get_categories_summary(db: Session = Depends(get_db)):
    templates = db.query(ProductTemplate).filter(ProductTemplate.attivo == True).order_by(
        ProductTemplate.categoria, ProductTemplate.ordine
    ).all()
    categories = {}
    for t in templates:
        cat = t.categoria
        if cat not in categories:
            categories[cat] = {"categoria": cat, "sottocategorie": []}
        categories[cat]["sottocategorie"].append({
            "id": t.id, "sottocategoria": t.sottocategoria,
            "nome_display": t.nome_display, "descrizione": t.descrizione,
            "icona": t.icona, "ordine": t.ordine
        })
    return list(categories.values())

@app.get("/templates/{template_id}")
def get_template(template_id: int, db: Session = Depends(get_db)):
    template = db.query(ProductTemplate).filter(ProductTemplate.id == template_id).first()
    if not template:
        raise HTTPException(status_code=404, detail="Template non trovato")
    return _orm_to_dict(template)

@app.post("/templates", status_code=201)
def create_template(data: ProductTemplateCreate, db: Session = Depends(get_db)):
    if data.template_data:
        try:
            json.loads(data.template_data)
        except json.JSONDecodeError:
            raise HTTPException(status_code=400, detail="template_data non ÃƒÂ¨ un JSON valido")
    
    template_dict = data.dict()
    template_dict['categoria'] = template_dict['categoria'].upper()
    template_dict['created_at'] = datetime.now()
    template_dict['updated_at'] = datetime.now()
    template = ProductTemplate(**template_dict)
    db.add(template)
    db.commit()
    db.refresh(template)
    return _orm_to_dict(template)

@app.put("/templates/{template_id}")
def update_template(template_id: int, data: ProductTemplateUpdate, db: Session = Depends(get_db)):
    template = db.query(ProductTemplate).filter(ProductTemplate.id == template_id).first()
    if not template:
        raise HTTPException(status_code=404, detail="Template non trovato")
    if data.template_data:
        try:
            json.loads(data.template_data)
        except json.JSONDecodeError:
            raise HTTPException(status_code=400, detail="template_data non ÃƒÂ¨ un JSON valido")
    
    update_dict = data.dict(exclude_unset=True)
    if 'categoria' in update_dict and update_dict['categoria']:
        update_dict['categoria'] = update_dict['categoria'].upper()
    update_dict['updated_at'] = datetime.now()
    for key, value in update_dict.items():
        setattr(template, key, value)
    db.commit()
    db.refresh(template)
    return _orm_to_dict(template)

@app.delete("/templates/{template_id}")
def delete_template(template_id: int, db: Session = Depends(get_db)):
    template = db.query(ProductTemplate).filter(ProductTemplate.id == template_id).first()
    if not template:
        raise HTTPException(status_code=404, detail="Template non trovato")
    db.delete(template)
    db.commit()
    return {"message": "Template eliminato"}


@app.get("/templates/{template_id}/field-config")
def get_template_field_config(template_id: int, db: Session = Depends(get_db)):
    """
    Restituisce la configurazione campi per un template.
    Per ogni campo: readonly, includi_preventivo, mostra_default.
    """
    template = db.query(ProductTemplate).filter(ProductTemplate.id == template_id).first()
    if not template:
        raise HTTPException(status_code=404, detail="Template non trovato")

    field_config = {}
    if template.template_data:
        try:
            td = json.loads(template.template_data)
            field_config = td.get("field_config", {})
        except json.JSONDecodeError:
            pass

    try:
        campi = db.execute(text("""
            SELECT codice, etichetta, sezione, tipo, valore_default
            FROM campi_configuratore WHERE attivo = 1
            ORDER BY sezione, ordine
        """)).fetchall()
    except Exception:
        campi = []

    result = []
    for codice, etichetta, sezione, tipo, valore_default in campi:
        cfg = field_config.get(codice, {})
        result.append({
            "codice": codice,
            "etichetta": etichetta,
            "sezione": sezione,
            "tipo": tipo,
            "valore_default": valore_default,
            "readonly": cfg.get("readonly", False),
            "includi_preventivo": cfg.get("includi_preventivo", True),
            "mostra_default": cfg.get("mostra_default", False),
            "cliente_id": u.cliente_id,
        })

    return {"template_id": template_id, "fields": result, "raw_config": field_config}


@app.put("/templates/{template_id}/field-config")
def update_template_field_config(template_id: int, data: dict, db: Session = Depends(get_db)):
    """
    Aggiorna la configurazione campi per un template.
    Body: {"field_config": {"campo": {"readonly": true, "includi_preventivo": true, "mostra_default": false}, ...}}
    """
    template = db.query(ProductTemplate).filter(ProductTemplate.id == template_id).first()
    if not template:
        raise HTTPException(status_code=404, detail="Template non trovato")

    new_field_config = data.get("field_config", {})

    existing_td = {}
    if template.template_data:
        try:
            existing_td = json.loads(template.template_data)
        except json.JSONDecodeError:
            existing_td = {}

    existing_td["field_config"] = new_field_config
    template.template_data = json.dumps(existing_td, ensure_ascii=False)
    db.commit()

    return {"status": "ok", "fields_configured": len(new_field_config)}


# ==========================================
# RULE ENGINE ENDPOINT
# ==========================================
@app.post("/preventivi/{preventivo_id}/evaluate-rules")
def evaluate_rules_endpoint(preventivo_id: int, db: Session = Depends(get_db)):
    return safe_evaluate_rules(preventivo_id, db)

# ==========================================
# RULE ENGINE — TEST / DIAGNOSTICA
# ==========================================
@app.post("/preventivi/{preventivo_id}/test-rules")
def test_rules_endpoint(preventivo_id: int, db: Session = Depends(get_db)):
    """
    Testa le regole per un preventivo SENZA modificare il DB.
    
    Restituisce un report dettagliato con:
    - Context iniziale (cosa vede il rule engine)
    - Per ogni regola lookup: condizioni valutate, valori scritti nel context
    - Per ogni regola materiale: condizioni valutate, materiali che verrebbero aggiunti
    - Context _calc.* dopo i lookup
    - Summary con conteggi e warning
    """
    from rule_engine import RuleEngine
    
    preventivo = db.query(Preventivo).filter(Preventivo.id == preventivo_id).first()
    if not preventivo:
        raise HTTPException(404, "Preventivo non trovato")
    
    engine = RuleEngine(db)
    report = engine.test_rules(preventivo)
    return report


@app.get("/preventivi/{preventivo_id}/rule-context")
def get_rule_context(preventivo_id: int, db: Session = Depends(get_db)):
    """
    Restituisce il context completo che il rule engine vede per un preventivo.
    Utile per capire quali campi sono disponibili e con quali valori.
    """
    from rule_engine import RuleEngine
    
    preventivo = db.query(Preventivo).filter(Preventivo.id == preventivo_id).first()
    if not preventivo:
        raise HTTPException(404, "Preventivo non trovato")
    
    engine = RuleEngine(db)
    return engine.get_context_debug(preventivo)

# ==========================================
# REGOLE JSON - CRUD su file in ./rules/
# ==========================================
RULES_DIR = "./rules"

@app.get("/regole")
def get_regole(
    sort: str = Query("priority", regex="^(name|date|priority|id)$"),
    order: str = Query("asc", regex="^(asc|desc)$"),
):
    """Lista tutte le regole JSON con ordinamento."""
    rules = load_rules()
    
    reverse = order == "desc"
    
    if sort == "priority":
        rules.sort(key=lambda r: r.get("priority", r.get("priorita", 99)), reverse=reverse)
    elif sort == "name":
        rules.sort(key=lambda r: (r.get("name") or r.get("nome") or r.get("id", "")).lower(), reverse=reverse)
    elif sort == "date":
        def date_key(r):
            d = r.get("_imported_at", "")
            if not d:
                rule_id = r.get("id", "")
                fpath = os.path.join("./rules", f"rule_{rule_id}.json")
                if os.path.exists(fpath):
                    return datetime.fromtimestamp(os.path.getmtime(fpath)).isoformat()
                return ""
            return d
        rules.sort(key=date_key, reverse=reverse)
    elif sort == "id":
        rules.sort(key=lambda r: r.get("id", "").lower(), reverse=reverse)
    
    return rules

@app.get("/regole/{rule_id}")
def get_regola(rule_id: str):
    """Ottieni una regola per ID"""
    rules = load_rules()
    for r in rules:
        if r.get("id") == rule_id:
            return r
    raise HTTPException(status_code=404, detail=f"Regola {rule_id} non trovata")

@app.post("/regole/check-value-usage")
def check_value_usage(data: dict, db: Session = Depends(get_db)):
    """
    Controlla se un valore di un gruppo dropdown e' usato in condizioni di regole.
    Body: { "gruppo": "en_81_20_anno", "valore": "2020" }
    """
    gruppo = data.get("gruppo", "")
    valore = str(data.get("valore", ""))
    
    if not gruppo or not valore:
        return {"affected_rules": [], "count": 0}
    
    # Trova i codici campo che usano questo gruppo dropdown
    try:
        result = db.execute(text(
            "SELECT codice FROM campi_configuratore WHERE gruppo_dropdown = :g"
        ), {"g": gruppo})
        codici = [r[0] for r in result.fetchall()]
    except:
        codici = []
    
    # Genera varianti nome campo (regole possono usare con o senza prefisso sezione)
    field_variants = set(codici)
    for codice in codici:
        if "." in codice:
            field_variants.add(codice.split(".", 1)[1])
        field_variants.add(gruppo.replace("_anno", "").replace("_tipo", ""))
    
    # Cerca nelle regole
    rules = load_rules()
    affected = []
    for rule in rules:
        for cond in rule.get("conditions", []):
            if cond.get("field") in field_variants and str(cond.get("value")) == valore:
                affected.append({
                    "rule_id": rule.get("id"),
                    "rule_name": rule.get("name", rule.get("id")),
                    "enabled": rule.get("enabled", True),
                    "field": cond.get("field"),
                    "value": str(cond.get("value")),
                })
                break
    
    return {"affected_rules": affected, "count": len(affected), "field_variants": list(field_variants)}

@app.put("/regole/cascade-update-value")
def cascade_update_value(data: dict, db: Session = Depends(get_db)):
    """
    Aggiorna un valore nelle condizioni di tutte le regole impattate.
    Body: { "gruppo": "en_81_20_anno", "old_value": "2020", "new_value": "2020_rev" }
    """
    gruppo = data.get("gruppo", "")
    old_value = str(data.get("old_value", ""))
    new_value = str(data.get("new_value", ""))
    
    if not old_value or not new_value:
        raise HTTPException(status_code=400, detail="old_value e new_value obbligatori")
    
    # Stessa logica di field_variants
    try:
        result = db.execute(text(
            "SELECT codice FROM campi_configuratore WHERE gruppo_dropdown = :g"
        ), {"g": gruppo})
        codici = [r[0] for r in result.fetchall()]
    except:
        codici = []
    
    field_variants = set(codici)
    for codice in codici:
        if "." in codice:
            field_variants.add(codice.split(".", 1)[1])
        field_variants.add(gruppo.replace("_anno", "").replace("_tipo", ""))
    
    # Aggiorna
    rules = load_rules()
    updated_rules = []
    for rule in rules:
        modified = False
        for cond in rule.get("conditions", []):
            if cond.get("field") in field_variants and str(cond.get("value")) == old_value:
                cond["value"] = new_value
                modified = True
        if modified:
            rule_id = rule.get("id")
            filepath = os.path.join(RULES_DIR, f"{rule_id}.json")
            try:
                with open(filepath, "w", encoding="utf-8") as f:
                    json.dump(rule, f, indent=2, ensure_ascii=False)
                updated_rules.append(rule_id)
            except Exception as e:
                print(f"Errore aggiornamento regola {rule_id}: {e}")
    
    return {"status": "ok", "updated_rules": updated_rules, "count": len(updated_rules)}

@app.post("/regole")
def create_regola(data: dict):
    """Crea una nuova regola JSON"""
    rule_id = data.get("id")
    if not rule_id:
        raise HTTPException(status_code=400, detail="Campo 'id' obbligatorio")
    
    os.makedirs(RULES_DIR, exist_ok=True)
    filepath = os.path.join(RULES_DIR, f"{rule_id}.json")
    if os.path.exists(filepath):
        raise HTTPException(status_code=409, detail=f"Regola {rule_id} esiste gia")
    
    # Assicura campi obbligatori
    data.setdefault("enabled", True)
    data.setdefault("priority", 50)
    data.setdefault("conditions", [])
    data.setdefault("materials", [])
    data.setdefault("version", "1.0")
    
    with open(filepath, "w", encoding="utf-8") as f:
        json.dump(data, f, indent=2, ensure_ascii=False)
    
    return {"status": "created", "id": rule_id}

@app.put("/regole/{rule_id}")
def update_regola(rule_id: str, data: dict):
    """Aggiorna una regola JSON esistente"""
    os.makedirs(RULES_DIR, exist_ok=True)
    
    # Cerca il file: prima match diretto, poi per ID interno
    filepath = os.path.join(RULES_DIR, f"{rule_id}.json")
    if not os.path.exists(filepath):
        for filename in os.listdir(RULES_DIR):
            if filename.endswith(".json"):
                fpath = os.path.join(RULES_DIR, filename)
                try:
                    with open(fpath, "r", encoding="utf-8") as f:
                        rule = json.load(f)
                    if rule.get("id") == rule_id:
                        filepath = fpath
                        break
                except Exception:
                    continue
    
    data["id"] = rule_id  # Forza ID consistente
    data.setdefault("enabled", True)
    data.setdefault("priority", 50)
    data.setdefault("conditions", [])
    data.setdefault("materials", [])
    
    with open(filepath, "w", encoding="utf-8") as f:
        json.dump(data, f, indent=2, ensure_ascii=False)
    
    return {"status": "updated", "id": rule_id}

@app.delete("/regole/{rule_id}")
def delete_regola(rule_id: str):
    """Elimina una regola JSON cercando per ID interno"""
    # Prima prova match diretto col nome file
    filepath = os.path.join(RULES_DIR, f"{rule_id}.json")
    if os.path.exists(filepath):
        os.remove(filepath)
        return {"status": "deleted", "id": rule_id}
    
    # Altrimenti cerca in tutti i file JSON per campo "id"
    if os.path.exists(RULES_DIR):
        for filename in os.listdir(RULES_DIR):
            if filename.endswith(".json"):
                fpath = os.path.join(RULES_DIR, filename)
                try:
                    with open(fpath, "r", encoding="utf-8") as f:
                        rule = json.load(f)
                    if rule.get("id") == rule_id:
                        os.remove(fpath)
                        return {"status": "deleted", "id": rule_id, "file": filename}
                except Exception:
                    continue
    
    raise HTTPException(status_code=404, detail=f"Regola {rule_id} non trovata")

# ==========================================
# PIPELINE BUILDER - API
# ==========================================

@app.post("/pipeline/simulate")
def simulate_pipeline_endpoint(data: dict, db: Session = Depends(get_db)):
    """
    Simula l'esecuzione di una pipeline senza modificare il DB.
    Body: { "pipeline_rule": {...}, "preventivo_id": 1 }
    """
    from rule_engine import RuleEngine
    
    pipeline_rule = data.get("pipeline_rule", {})
    preventivo_id = data.get("preventivo_id")
    
    if preventivo_id:
        preventivo = db.query(Preventivo).filter(Preventivo.id == preventivo_id).first()
    else:
        preventivo = db.query(Preventivo).order_by(Preventivo.updated_at.desc()).first()
    
    if not preventivo:
        raise HTTPException(404, "Nessun preventivo disponibile per la simulazione")
    
    engine = RuleEngine(db)
    result = engine.simulate_pipeline(pipeline_rule, preventivo)
    return result


@app.get("/pipeline/list")
def list_pipelines_endpoint():
    """Lista tutte le regole pipeline (con pipeline_steps)."""
    rules = load_rules()
    pipelines = [r for r in rules if r.get("pipeline_steps")]
    return pipelines

@app.post("/pipeline/crea-campi-da-tabella")
def crea_campi_da_tabella(data: dict, db: Session = Depends(get_db)):
    """
    Generico: data una data table, crea sezione + checkbox nel configuratore.
    
    Body: {
        "tabella": "utilizzatori_trasformatore",
        "colonna_chiave": "componente",        // diventa codice campo
        "colonna_label": "label",              // diventa etichetta (opz, fallback=chiave)
        "sezione_codice": "selezione_trasformatore",
        "sezione_etichetta": "Selez. Trasformatore",
        "sezione_icona": "Zap",                // opzionale, default "CircleDot"
        "campi_extra": [                        // opzionale, campi non-checkbox da aggiungere
            {"codice": "power_factor", "etichetta": "Power Factor", "tipo": "testo", "default": "0.8"}
        ]
    }
    """
    from excel_data_loader import ExcelDataLoader
    
    tabella_nome = data.get("tabella", "")
    col_chiave = data.get("colonna_chiave", "")
    col_label = data.get("colonna_label", "")
    sez_codice = data.get("sezione_codice", "")
    sez_etichetta = data.get("sezione_etichetta", sez_codice.replace("_", " ").title())
    sez_icona = data.get("sezione_icona", "CircleDot")
    campi_extra = data.get("campi_extra", [])
    
    if not tabella_nome or not col_chiave or not sez_codice:
        raise HTTPException(400, "Servono: tabella, colonna_chiave, sezione_codice")
    
    # 1. Carica la data table
    loader = ExcelDataLoader()
    table = loader.load_table(tabella_nome)
    if not table:
        raise HTTPException(404, f"Tabella '{tabella_nome}' non trovata in ./data/")
    
    records = table.get("records", [])
    if not records:
        raise HTTPException(400, f"Tabella '{tabella_nome}' vuota")
    
    # Verifica colonna chiave
    if col_chiave not in records[0]:
        cols_disp = list(records[0].keys())
        raise HTTPException(400, f"Colonna '{col_chiave}' non trovata. Disponibili: {cols_disp}")
    
    # 2. Crea sezione (se non esiste)
    try:
        existing = db.execute(
            text("SELECT id FROM sezioni_configuratore WHERE codice = :c"),
            {"c": sez_codice}
        ).fetchone()
        
        if existing:
            sezione_id = existing[0]
        else:
            # Trova ordine max + 1
            max_ord = db.execute(
                text("SELECT COALESCE(MAX(ordine), 0) FROM sezioni_configuratore")
            ).scalar()
            db.execute(
                text("""INSERT INTO sezioni_configuratore (codice, etichetta, icona, ordine, attivo)
                        VALUES (:c, :e, :i, :o, 1)"""),
                {"c": sez_codice, "e": sez_etichetta, "i": sez_icona, "o": (max_ord or 0) + 1}
            )
            db.commit()
            sezione_id = db.execute(
                text("SELECT id FROM sezioni_configuratore WHERE codice = :c"),
                {"c": sez_codice}
            ).scalar()
    except Exception as e:
        raise HTTPException(500, f"Errore creazione sezione: {e}")
    
    # 3. Raggruppa records per chiave (un componente può avere più righe = più uscite)
    from collections import OrderedDict
    componenti: OrderedDict = OrderedDict()  # codice_safe → {label, righe: [{col: val}]}
    
    for rec in records:
        codice_campo = str(rec.get(col_chiave, "")).strip()
        if not codice_campo:
            continue
        
        codice_safe = codice_campo.lower().strip()
        codice_safe = codice_safe.replace(" ", "_").replace("-", "_")
        codice_safe = "".join(c for c in codice_safe if c.isalnum() or c == "_")
        
        label = str(rec.get(col_label, codice_campo)) if col_label else codice_campo
        
        if codice_safe not in componenti:
            componenti[codice_safe] = {"label": label, "righe": []}
        
        # Raccogli le colonne extra di questa riga
        extra = {k: v for k, v in rec.items() if k not in (col_chiave, col_label) and v is not None}
        componenti[codice_safe]["righe"].append(extra)
    
    # 3b. Crea un campo checkbox per ogni componente UNICO
    creati = 0
    skipped = 0
    
    for i, (codice_safe, info) in enumerate(componenti.items()):
        # Verifica se già esiste
        exists = db.execute(
            text("SELECT id FROM campi_configuratore WHERE codice = :c AND sezione = :s"),
            {"c": codice_safe, "s": sez_codice}
        ).fetchone()
        
        if exists:
            skipped += 1
            continue
        
        # Note: riassumi tutte le righe del componente
        righe = info["righe"]
        if len(righe) == 1:
            note = ", ".join(f"{k}={v}" for k, v in righe[0].items())
        else:
            # Più uscite → "3 uscite: 75V 150W, 15V 150W, 18V 150W"
            uscite_desc = []
            for r in righe:
                vals = [f"{v}" for v in r.values()]
                uscite_desc.append("/".join(vals))
            note = f"{len(righe)} uscite: " + " | ".join(uscite_desc)
        note = note[:200]
        
        db.execute(
            text("""INSERT INTO campi_configuratore 
                    (codice, etichetta, sezione, tipo, ordine, attivo, obbligatorio,
                     gruppo_dropdown, valore_default)
                    VALUES (:codice, :etichetta, :sezione, 'checkbox', :ordine, 1, 0,
                            '', 'false')"""),
            {
                "codice": codice_safe,
                "etichetta": info["label"],
                "sezione": sez_codice,
                "ordine": (i + 1) * 10,
            }
        )
        creati += 1
    
    # 4. Campi extra (es. power_factor)
    extra_creati = 0
    for campo in campi_extra:
        cod = campo.get("codice", "")
        if not cod:
            continue
        exists = db.execute(
            text("SELECT id FROM campi_configuratore WHERE codice = :c AND sezione = :s"),
            {"c": cod, "s": sez_codice}
        ).fetchone()
        if exists:
            continue
        
        db.execute(
            text("""INSERT INTO campi_configuratore 
                    (codice, etichetta, sezione, tipo, ordine, attivo, obbligatorio,
                     gruppo_dropdown, valore_default)
                    VALUES (:codice, :etichetta, :sezione, :tipo, :ordine, 1, 0,
                            '', :default)"""),
            {
                "codice": cod,
                "etichetta": campo.get("etichetta", cod),
                "sezione": sez_codice,
                "tipo": campo.get("tipo", "testo"),
                "ordine": 9000 + extra_creati * 10,
                "default": campo.get("default", ""),
            }
        )
        extra_creati += 1
    
    db.commit()
    
    # Notifica frontend
    return {
        "status": "ok",
        "sezione_codice": sez_codice,
        "sezione_id": sezione_id,
        "campi_creati": creati,
        "campi_skipped": skipped,
        "campi_extra_creati": extra_creati,
        "componenti_unici": len(componenti),
        "totale_record_tabella": len(records),
    }
@app.get("/regole-campi-disponibili")
def get_campi_disponibili(db: Session = Depends(get_db)):
    """
    Lista tutti i campi utilizzabili nelle condizioni delle regole, con tipo e opzioni.
    
    Approccio DYNAMIC-FIRST:
    1. Legge TUTTO da campi_configuratore (fonte di verità per campi sezione)
    2. Aggiunge campi template/preventivo (meta-campi, non in campi_configuratore)
    3. Aggiunge fallback per campi ORM legacy NON presenti in campi_configuratore
    """
    campi = []
    codici_gia_aggiunti = set()

    # Pre-carica mappa sezione codice → etichetta per nomi leggibili
    sez_labels = {}
    try:
        sez_rows = db.execute(text("SELECT codice, etichetta FROM sezioni_configuratore")).fetchall()
        sez_labels = {r[0]: r[1] for r in sez_rows}
    except Exception:
        pass

    def sez_display(codice_sezione):
        return sez_labels.get(codice_sezione, codice_sezione)

    def add_campo(field, source, label, tipo="testo", options=None):
        if field in codici_gia_aggiunti:
            return  # Evita duplicati
        codici_gia_aggiunti.add(field)
        entry = {"field": field, "source": source, "label": label, "type": tipo}
        if options:
            entry["options"] = options
        campi.append(entry)

    # =========================================================
    # STEP 1: Campi da campi_configuratore (FONTE DI VERITÀ)
    # =========================================================
    try:
        result = db.execute(text(
            "SELECT codice, etichetta, sezione, tipo, gruppo_dropdown "
            "FROM campi_configuratore WHERE attivo=1 ORDER BY sezione, ordine"
        ))
        for row in result.fetchall():
            codice, etichetta, sezione, tipo, gruppo = row[0], row[1], row[2], row[3], row[4]
            options = None
            if tipo == "dropdown" and gruppo:
                options = _load_dropdown_options(db, gruppo)
            add_campo(codice, sez_display(sezione), etichetta, tipo or "testo", options)
    except Exception:
        pass

    # =========================================================
    # STEP 2: Meta-campi TEMPLATE (non sono in campi_configuratore)
    # =========================================================
    try:
        cat_result = db.execute(text("SELECT DISTINCT categoria FROM product_templates WHERE attivo=1 ORDER BY categoria"))
        template_categorie = [r[0] for r in cat_result.fetchall() if r[0]]
    except Exception:
        template_categorie = ["RISE", "HOME"]
    try:
        sub_result = db.execute(text("SELECT DISTINCT sotto_categoria FROM product_templates WHERE attivo=1 ORDER BY sotto_categoria"))
        template_sottocategorie = [r[0] for r in sub_result.fetchall() if r[0]]
    except Exception:
        template_sottocategorie = []
    try:
        nome_result = db.execute(text("SELECT DISTINCT nome_display FROM product_templates WHERE attivo=1 ORDER BY nome_display"))
        template_nomi = [r[0] for r in nome_result.fetchall() if r[0]]
    except Exception:
        template_nomi = []
    try:
        id_result = db.execute(text("SELECT id, nome_display FROM product_templates WHERE attivo=1 ORDER BY id"))
        template_ids = [{"value": str(r[0]), "label": f"{r[0]} - {r[1]}"} for r in id_result.fetchall()]
    except Exception:
        template_ids = []

    add_campo("template_categoria", "Prodotto / Template", "Template Categoria", "dropdown", template_categorie)
    add_campo("template_sottocategoria", "Prodotto / Template", "Template Sottocategoria", "dropdown", template_sottocategorie)
    add_campo("template_nome", "Prodotto / Template", "Template Nome Display", "dropdown", template_nomi)
    add_campo("template_id", "Prodotto / Template", "Template ID", "dropdown", template_ids)
    # =========================================================
    # STEP 3: Meta-campi PREVENTIVO (non sono in campi_configuratore)
    # =========================================================
    add_campo("preventivo_tipo", "Preventivo", "Tipo Preventivo", "dropdown", ["COMPLETO", "RICAMBI"])
    add_campo("preventivo_categoria", "Preventivo", "Categoria Preventivo", "dropdown", template_categorie)
    # =========================================================
    # STEP 4: Fallback campi ORM legacy
    #   Solo per campi NON già definiti in campi_configuratore.
    #   Questi servono se il sistema ha tabelle ORM legacy
    #   non ancora migrate a campi_configuratore.
    # =========================================================
    
    # Dati principali
    legacy_dp = {
        "tipo_impianto": ("dropdown", "tipo_impianto"),
        "nuovo_impianto": ("booleano", None),
        "numero_fermate": ("numero", None),
        "numero_servizi": ("numero", None),
        "velocita": ("numero", None),
        "corsa": ("numero", None),
        "con_locale_macchina": ("booleano", None),
        "posizione_locale_macchina": ("dropdown", "posizione_locale_macchina"),
        "tipo_trazione": ("dropdown", "tipo_trazione"),
        "forza_motrice": ("dropdown", "forza_motrice"),
        "luce": ("numero", None),
        "tensione_manovra": ("dropdown", "tensione_manovra"),
        "tensione_freno": ("dropdown", "tensione_freno"),
    }
    for campo, (tipo, gruppo) in legacy_dp.items():
        if campo not in codici_gia_aggiunti:
            opts = _load_dropdown_options(db, gruppo) if gruppo else None
            add_campo(campo, sez_display("dati_principali"), campo.replace("_", " ").title(), tipo, opts)

    # Normative
    legacy_norm = {
        "en_81_1": ("dropdown", "en_81_1_anno", [{"value": "1998", "label": "1998"}, {"value": "2010", "label": "2010"}]),
        "en_81_20": ("dropdown", "en_81_20_anno", [{"value": "2014", "label": "2014"}, {"value": "2020", "label": "2020"}]),
        "en_81_21": ("dropdown", "en_81_21_anno", [{"value": "2009", "label": "2009"}, {"value": "2018", "label": "2018"}]),
        "en_81_28": ("booleano", None, None),
        "en_81_70": ("booleano", None, None),
        "en_81_72": ("booleano", None, None),
        "en_81_73": ("booleano", None, None),
        "a3_95_16": ("booleano", None, None),
        "dm236_legge13": ("booleano", None, None),
        "emendamento_a3": ("booleano", None, None),
        "uni_10411_1": ("booleano", None, None),
    }
    for campo, (tipo, gruppo, fallback_opts) in legacy_norm.items():
        if campo not in codici_gia_aggiunti:
            opts = None
            if gruppo:
                opts = _load_dropdown_options(db, gruppo) or fallback_opts
            add_campo(campo, sez_display("normative"), campo.replace("_", " ").upper(), tipo, opts)

    # Argano
    legacy_argano = {
        "trazione": ("dropdown", "trazione"),
        "potenza_motore_kw": ("numero", None),
        "corrente_nom_motore_amp": ("numero", None),
        "tipo_vvvf": ("dropdown", "tipo_vvvf"),
        "vvvf_nel_vano": ("booleano", None),
        "freno_tensione": ("numero", None),
        "ventilazione_forzata": ("booleano", None),
        "tipo_teleruttore": ("dropdown", "tipo_teleruttore"),
    }
    for campo, (tipo, gruppo) in legacy_argano.items():
        if campo not in codici_gia_aggiunti:
            opts = _load_dropdown_options(db, gruppo) if gruppo else None
            add_campo(campo, sez_display("argano"), campo.replace("_", " ").title(), tipo, opts)

    return campi


def _load_dropdown_options(db, gruppo: str):
    """Carica opzioni dropdown dal DB per un dato gruppo"""
    try:
        result = db.execute(text(
            "SELECT valore, etichetta FROM opzioni_dropdown "
            "WHERE gruppo = :g AND attivo = 1 ORDER BY ordine"
        ), {"g": gruppo})
        rows = result.fetchall()
        if rows:
            return [{"value": r[0], "label": r[1] or r[0]} for r in rows]
    except Exception:
        pass
    return None

# ==========================================
# CAMPI <-> REGOLE: USAGE MAP + RINOMINA
# ==========================================

def _scan_rules_fields() -> dict:
    """
    Scansiona tutti i file JSON in rules/ e restituisce:
      { "nome_campo": [ { rule_id, rule_name, file, enabled, usage, detail }, ... ] }
    """
    field_map = {}
    rules_dir = "./rules"
    if not os.path.exists(rules_dir):
        return field_map

    for filename in os.listdir(rules_dir):
        if not filename.endswith(".json"):
            continue
        try:
            filepath = os.path.join(rules_dir, filename)
            with open(filepath, "r", encoding="utf-8") as f:
                rule = json.load(f)

            rule_id = rule.get("id", filename)
            rule_name = rule.get("name", rule_id)
            enabled = rule.get("enabled", True)

            def _extract_conditions(conds):
                fields = []
                if isinstance(conds, list):
                    for c in conds:
                        if "field" in c:
                            fields.append((c["field"], c.get("operator", "?"), c.get("value", "?")))
                elif isinstance(conds, dict):
                    if "field" in conds:
                        fields.append((conds["field"], conds.get("operator", "?"), conds.get("value", "?")))
                    if "conditions" in conds:
                        fields.extend(_extract_conditions(conds["conditions"]))
                return fields

            for field, op, val in _extract_conditions(rule.get("conditions", [])):
                entry = {
                    "rule_id": rule_id, "rule_name": rule_name,
                    "file": filename, "enabled": enabled,
                    "usage": "condition",
                    "detail": f"{op} {val}",
                }
                field_map.setdefault(field, []).append(entry)

            for mat in rule.get("materials", []):
                for val in [mat.get("codice", ""), mat.get("descrizione", "")]:
                    for m in re.findall(r"\{\{(\w[\w.]*)\}\}", str(val)):
                        entry = {
                            "rule_id": rule_id, "rule_name": rule_name,
                            "file": filename, "enabled": enabled,
                            "usage": "placeholder",
                            "detail": f"in {mat.get('codice', '?')}",
                        }
                        field_map.setdefault(m, []).append(entry)

        except Exception:
            continue

    return field_map


@app.get("/campi-regole-usage")
def get_campi_regole_usage(db: Session = Depends(get_db)):
    """
    Per ogni campo restituisce info + regole che lo usano + preventivi con valori.
    Include anche campi che appaiono solo nelle regole (non nel DB).
    """
    # 1. Campi dal DB
    try:
        result = db.execute(text("""
            SELECT id, codice, etichetta, sezione, tipo, attivo, usabile_regole
            FROM campi_configuratore ORDER BY sezione, ordine
        """))
        campi = []
        for r in result.fetchall():
            campi.append({
                "id": r[0], "codice": r[1], "etichetta": r[2],
                "sezione": r[3], "tipo": r[4], "attivo": bool(r[5]),
                "usabile_regole": bool(r[6]) if r[6] is not None else True,
            })
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

    # 2. Mappa regole
    field_rules = _scan_rules_fields()

    # 3. Conteggio valori
    try:
        result = db.execute(text("""
            SELECT codice_campo, COUNT(DISTINCT preventivo_id) as n
            FROM valori_configurazione
            WHERE valore IS NOT NULL AND valore != ''
            GROUP BY codice_campo
        """))
        usage_count = {r[0]: r[1] for r in result.fetchall()}
    except Exception:
        usage_count = {}

    # 4. Assembla
    output = []
    for campo in campi:
        codice = campo["codice"]
        regole = list(field_rules.get(codice, []))
        if "." in codice:
            senza = codice.split(".", 1)[1]
            regole += field_rules.get(senza, [])
        # Deduplica
        seen = set()
        uniche = []
        for r in regole:
            key = (r["rule_id"], r["usage"])
            if key not in seen:
                seen.add(key)
                uniche.append(r)

        output.append({
            **campo,
            "regole": uniche,
            "num_regole": len(uniche),
            "num_preventivi": usage_count.get(codice, 0),
            "solo_regole": False,
        })

    # 5. Campi che esistono solo nelle regole (non in campi_configuratore)
    codici_db = {c["codice"] for c in campi}
    codici_senza = set()
    for c in campi:
        if "." in c["codice"]:
            codici_senza.add(c["codice"].split(".", 1)[1])

    for field, regole in field_rules.items():
        if field not in codici_db and field not in codici_senza:
            output.append({
                "id": None, "codice": field,
                "etichetta": field.replace("_", " ").title(),
                "sezione": field.split(".")[0] if "." in field else "???",
                "tipo": "???", "attivo": True, "usabile_regole": True,
                "regole": regole,
                "num_regole": len(regole),
                "num_preventivi": usage_count.get(field, 0),
                "solo_regole": True,
            })

    return output


@app.post("/campi-rinomina/preview")
def preview_rinomina(data: dict, db: Session = Depends(get_db)):
    """Anteprima impatto rinomina senza applicare."""
    vecchio = data.get("vecchio_codice", "").strip()
    nuovo = data.get("nuovo_codice", "").strip()
    if not vecchio or not nuovo:
        raise HTTPException(status_code=400, detail="vecchio_codice e nuovo_codice obbligatori")

    preview = {"vecchio": vecchio, "nuovo": nuovo, "impatto": []}

    # DB campi
    try:
        r = db.execute(text("SELECT COUNT(*) FROM campi_configuratore WHERE codice=:v"), {"v": vecchio}).fetchone()
        if r[0]:
            preview["impatto"].append({"area": "campi_configuratore", "desc": f"{r[0]} record", "tipo": "db"})
    except Exception:
        pass

    # Valori preventivi
    try:
        r = db.execute(text("""
            SELECT COUNT(*), COUNT(DISTINCT preventivo_id)
            FROM valori_configurazione WHERE codice_campo=:v
        """), {"v": vecchio}).fetchone()
        if r[0]:
            preview["impatto"].append({"area": "valori_configurazione", "desc": f"{r[0]} valori in {r[1]} preventivi", "tipo": "db"})
    except Exception:
        pass

    # Opzioni dropdown
    try:
        r = db.execute(text("SELECT COUNT(*) FROM opzioni_dropdown WHERE gruppo=:v"), {"v": vecchio}).fetchone()
        if r[0]:
            preview["impatto"].append({"area": "opzioni_dropdown", "desc": f"{r[0]} opzioni nel gruppo", "tipo": "db"})
    except Exception:
        pass

    # Regole JSON
    field_rules = _scan_rules_fields()
    vecchio_senza = vecchio.split(".", 1)[1] if "." in vecchio else vecchio
    files = set()
    for f in [vecchio, vecchio_senza]:
        for entry in field_rules.get(f, []):
            files.add(entry["file"])
    if files:
        preview["impatto"].append({"area": "regole_json", "desc": f"{len(files)} file regola", "tipo": "file", "files": sorted(files)})

    return preview


@app.post("/campi-rinomina")
def rinomina_campo(data: dict, db: Session = Depends(get_db)):
    """
    Rinomina campo su tutti i punti:
      1. campi_configuratore.codice
      2. valori_configurazione.codice_campo + sezione
      3. opzioni_dropdown.gruppo
      4. File JSON regole (conditions + placeholders)
    """
    vecchio = data.get("vecchio_codice", "").strip()
    nuovo = data.get("nuovo_codice", "").strip()
    if not vecchio or not nuovo:
        raise HTTPException(status_code=400, detail="vecchio_codice e nuovo_codice obbligatori")
    if vecchio == nuovo:
        raise HTTPException(status_code=400, detail="Codici identici")

    report = {"vecchio": vecchio, "nuovo": nuovo,
              "campi_configuratore": 0, "valori_configurazione": 0,
              "opzioni_dropdown": 0, "regole_aggiornate": [], "errori": []}

    # 1. campi_configuratore
    try:
        r = db.execute(text("UPDATE campi_configuratore SET codice=:n WHERE codice=:v"), {"v": vecchio, "n": nuovo})
        report["campi_configuratore"] = r.rowcount
        if "." in nuovo:
            nuova_sez = nuovo.split(".", 1)[0]
            db.execute(text("UPDATE campi_configuratore SET sezione=:s WHERE codice=:c"), {"s": nuova_sez, "c": nuovo})
    except Exception as e:
        report["errori"].append(f"campi_configuratore: {e}")

    # 2. valori_configurazione
    try:
        r = db.execute(text("UPDATE valori_configurazione SET codice_campo=:n WHERE codice_campo=:v"), {"v": vecchio, "n": nuovo})
        report["valori_configurazione"] = r.rowcount
        if "." in nuovo:
            nuova_sez = nuovo.split(".", 1)[0]
            db.execute(text("UPDATE valori_configurazione SET sezione=:s WHERE codice_campo=:c"), {"s": nuova_sez, "c": nuovo})
    except Exception as e:
        report["errori"].append(f"valori_configurazione: {e}")

    # 3. opzioni_dropdown
    try:
        r = db.execute(text("UPDATE opzioni_dropdown SET gruppo=:n WHERE gruppo=:v"), {"v": vecchio, "n": nuovo})
        report["opzioni_dropdown"] = r.rowcount
    except Exception as e:
        report["errori"].append(f"opzioni_dropdown: {e}")

    # 4. Regole JSON
    rules_dir = "./rules"
    if os.path.exists(rules_dir):
        vecchio_senza = vecchio.split(".", 1)[1] if "." in vecchio else vecchio
        nuovo_senza = nuovo.split(".", 1)[1] if "." in nuovo else nuovo

        for filename in os.listdir(rules_dir):
            if not filename.endswith(".json"):
                continue
            filepath = os.path.join(rules_dir, filename)
            try:
                with open(filepath, "r", encoding="utf-8") as f:
                    content = f.read()
                if vecchio not in content and vecchio_senza not in content:
                    continue

                rule = json.loads(content)
                mod = False

                def _upd(conds):
                    nonlocal mod
                    if isinstance(conds, list):
                        for c in conds:
                            if c.get("field") == vecchio:
                                c["field"] = nuovo; mod = True
                            elif c.get("field") == vecchio_senza:
                                c["field"] = nuovo_senza; mod = True
                    elif isinstance(conds, dict):
                        if conds.get("field") == vecchio:
                            conds["field"] = nuovo; mod = True
                        elif conds.get("field") == vecchio_senza:
                            conds["field"] = nuovo_senza; mod = True
                        if "conditions" in conds:
                            _upd(conds["conditions"])

                _upd(rule.get("conditions", []))

                for mat in rule.get("materials", []):
                    for key in ["codice", "descrizione"]:
                        if key in mat and isinstance(mat[key], str):
                            for old_f, new_f in [(vecchio, nuovo), (vecchio_senza, nuovo_senza)]:
                                old_ph = "{{" + old_f + "}}"
                                new_ph = "{{" + new_f + "}}"
                                if old_ph in mat[key]:
                                    mat[key] = mat[key].replace(old_ph, new_ph)
                                    mod = True

                if mod:
                    with open(filepath, "w", encoding="utf-8") as f:
                        json.dump(rule, f, indent=2, ensure_ascii=False)
                    report["regole_aggiornate"].append(filename)

            except Exception as e:
                report["errori"].append(f"{filename}: {e}")

    # 5. Template JSON — template_data keys e field_config keys
    try:
        templates = db.execute(text("SELECT id, template_data FROM product_templates WHERE template_data IS NOT NULL")).fetchall()
        report["templates_aggiornati"] = 0
        vecchio_sez = vecchio.split(".", 1)[0] if "." in vecchio else ""
        vecchio_campo_only = vecchio.split(".", 1)[1] if "." in vecchio else vecchio
        nuovo_sez = nuovo.split(".", 1)[0] if "." in nuovo else ""
        nuovo_campo_only = nuovo.split(".", 1)[1] if "." in nuovo else nuovo

        for tmpl_id, tmpl_json in templates:
            if not tmpl_json:
                continue
            tmpl_data = json.loads(tmpl_json)
            changed = False

            # Rinomina campo nelle sezioni dati (es: templateData.argano.trazione → templateData.argano.nuovo_nome)
            if vecchio_sez and vecchio_sez in tmpl_data:
                sez_data = tmpl_data[vecchio_sez]
                if isinstance(sez_data, dict) and vecchio_campo_only in sez_data:
                    sez_data[nuovo_campo_only] = sez_data.pop(vecchio_campo_only)
                    # Se la sezione è cambiata, sposta
                    if nuovo_sez != vecchio_sez:
                        if nuovo_sez not in tmpl_data:
                            tmpl_data[nuovo_sez] = {}
                        tmpl_data[nuovo_sez][nuovo_campo_only] = sez_data.pop(nuovo_campo_only, tmpl_data[nuovo_sez].get(nuovo_campo_only))
                    changed = True

            # Rinomina chiave in field_config
            if "field_config" in tmpl_data:
                fc = tmpl_data["field_config"]
                if vecchio in fc:
                    fc[nuovo] = fc.pop(vecchio)
                    changed = True

            if changed:
                db.execute(text("UPDATE product_templates SET template_data=:data WHERE id=:id"),
                    {"data": json.dumps(tmpl_data), "id": tmpl_id})
                report["templates_aggiornati"] += 1
    except Exception as e:
        report["errori"].append(f"templates: {e}")

    db.commit()
    return report

# ==========================================
# ARTICOLI
# ==========================================
@app.get("/articoli")
def get_articoli(
    search: str = None,
    limit: int = 100,
    offset: int = 0,
    db: Session = Depends(get_db)
):
    """Lista articoli con ricerca opzionale"""
    query = db.query(Articolo)
    if search:
        search_term = f"%{search}%"
        query = query.filter(
            or_(
                Articolo.codice.ilike(search_term),
                Articolo.descrizione.ilike(search_term)
            )
        )
    articoli = query.order_by(Articolo.codice).offset(offset).limit(limit).all()
    return [{
        "id": a.id, "codice": a.codice, "descrizione": a.descrizione,
        "descrizione_estesa": a.descrizione_estesa,
        "tipo_articolo": a.tipo_articolo, "categoria_id": a.categoria_id,
        "costo_fisso": a.costo_fisso,
        "costo_variabile_1": a.costo_variabile_1, "unita_misura_var_1": a.unita_misura_var_1,
        "descrizione_var_1": a.descrizione_var_1,
        "costo_variabile_2": a.costo_variabile_2, "unita_misura_var_2": a.unita_misura_var_2,
        "descrizione_var_2": a.descrizione_var_2,
        "costo_variabile_3": a.costo_variabile_3, "unita_misura_var_3": a.unita_misura_var_3,
        "descrizione_var_3": a.descrizione_var_3,
        "costo_variabile_4": a.costo_variabile_4, "unita_misura_var_4": a.unita_misura_var_4,
        "descrizione_var_4": a.descrizione_var_4,
        "ricarico_percentuale": a.ricarico_percentuale,
        "unita_misura": a.unita_misura,
        "fornitore": a.fornitore, "codice_fornitore": a.codice_fornitore,
        "is_active": a.is_active
    } for a in articoli]

@app.get("/articoli/search")
def search_articoli(q: str = "", limit: int = 10, db: Session = Depends(get_db)):
    """Ricerca articoli per codice o descrizione"""
    if not q:
        return []
    search_term = f"%{q}%"
    articoli = db.query(Articolo).filter(
        or_(
            Articolo.codice.ilike(search_term),
            Articolo.descrizione.ilike(search_term)
        )
    ).limit(limit).all()
    return [{
        "id": a.id, "codice": a.codice, "descrizione": a.descrizione,
        "tipo_articolo": a.tipo_articolo, "costo_fisso": a.costo_fisso,
        "costo_variabile_1": a.costo_variabile_1, "unita_misura_var_1": a.unita_misura_var_1,
        "descrizione_var_1": a.descrizione_var_1,
        "costo_variabile_2": a.costo_variabile_2, "unita_misura_var_2": a.unita_misura_var_2,
        "descrizione_var_2": a.descrizione_var_2,
        "costo_variabile_3": a.costo_variabile_3, "unita_misura_var_3": a.unita_misura_var_3,
        "descrizione_var_3": a.descrizione_var_3,
        "costo_variabile_4": a.costo_variabile_4, "unita_misura_var_4": a.unita_misura_var_4,
        "descrizione_var_4": a.descrizione_var_4,
        "ricarico_percentuale": a.ricarico_percentuale,
        "unita_misura": a.unita_misura
    } for a in articoli]

@app.post("/articoli")
def create_articolo(data: dict, db: Session = Depends(get_db)):
    articolo = Articolo()
    for key in [
        "codice", "descrizione", "descrizione_estesa", "tipo_articolo", "categoria_id",
        "costo_fisso", "costo_variabile_1", "unita_misura_var_1", "descrizione_var_1",
        "costo_variabile_2", "unita_misura_var_2", "descrizione_var_2",
        "costo_variabile_3", "unita_misura_var_3", "descrizione_var_3",
        "costo_variabile_4", "unita_misura_var_4", "descrizione_var_4",
        "ricarico_percentuale", "unita_misura", "fornitore", "codice_fornitore",
        "giacenza", "scorta_minima", "lead_time_giorni", "manodopera_giorni", "is_active"
    ]:
        if key in data:
            setattr(articolo, key, data[key])
    db.add(articolo)
    db.commit()
    db.refresh(articolo)
    return {"id": articolo.id, "status": "ok"}

@app.put("/articoli/{articolo_id}")
def update_articolo(articolo_id: int, data: dict, db: Session = Depends(get_db)):
    articolo = db.query(Articolo).filter(Articolo.id == articolo_id).first()
    if not articolo:
        raise HTTPException(status_code=404, detail="Articolo non trovato")
    for key in [
        "codice", "descrizione", "descrizione_estesa", "tipo_articolo", "categoria_id",
        "costo_fisso", "costo_variabile_1", "unita_misura_var_1", "descrizione_var_1",
        "costo_variabile_2", "unita_misura_var_2", "descrizione_var_2",
        "costo_variabile_3", "unita_misura_var_3", "descrizione_var_3",
        "costo_variabile_4", "unita_misura_var_4", "descrizione_var_4",
        "ricarico_percentuale", "unita_misura", "fornitore", "codice_fornitore",
        "giacenza", "scorta_minima", "lead_time_giorni", "manodopera_giorni", "is_active"
    ]:
        if key in data:
            setattr(articolo, key, data[key])
    articolo.updated_at = datetime.now()
    db.commit()
    return {"id": articolo.id, "status": "ok"}

@app.delete("/articoli/{articolo_id}")
def delete_articolo(articolo_id: int, db: Session = Depends(get_db)):
    articolo = db.query(Articolo).filter(Articolo.id == articolo_id).first()
    if not articolo:
        raise HTTPException(status_code=404, detail="Articolo non trovato")
    db.delete(articolo)
    db.commit()
    return {"status": "ok"}

@app.post("/articoli/calcola-prezzo")
def calcola_prezzo_articolo(data: dict, db: Session = Depends(get_db)):
    """Calcola prezzo di un articolo in base ai parametri variabili"""
    costo_fisso = data.get("costo_fisso", 0) or 0
    costo_totale = costo_fisso
    
    for i in range(1, 5):
        costo_var = data.get(f"costo_variabile_{i}", 0) or 0
        parametro = data.get(f"parametro_{i}", 0) or 0
        costo_totale += costo_var * parametro
    
    ricarico = data.get("ricarico_percentuale", 0) or 0
    prezzo_listino = costo_totale * (1 + ricarico / 100) if ricarico else costo_totale
    
    sconto = data.get("sconto_cliente", 0) or 0
    prezzo_cliente = prezzo_listino * (1 - sconto / 100) if sconto else prezzo_listino
    
    return {
        "costo_base": costo_totale,
        "prezzo_listino": round(prezzo_listino, 4),
        "prezzo_cliente": round(prezzo_cliente, 4)
    }

# --- BOM: Distinta Base ---

def _detect_bom_schema(db: Session) -> str:
    """
    Rileva se bom_struttura usa lo schema vecchio (padre_codice/figlio_codice TEXT)
    o quello nuovo (articolo_padre_id/articolo_figlio_id INT).
    """
    try:
        conn = db.get_bind().raw_connection()
        cursor = conn.cursor()
        cursor.execute("PRAGMA table_info(bom_struttura)")
        columns = {row[1] for row in cursor.fetchall()}
        if 'articolo_padre_id' in columns and 'articolo_figlio_id' in columns:
            return 'id'
        return 'codice'
    except Exception:
        return 'id'


@app.get("/bom/counts")
def get_bom_counts(db: Session = Depends(get_db)):
    """Conta i figli per ogni articolo padre"""
    schema = _detect_bom_schema(db)
    conn = db.get_bind().raw_connection()
    cursor = conn.cursor()

    if schema == 'id':
        rows = db.execute(text(
            "SELECT articolo_padre_id, COUNT(*) as cnt FROM bom_struttura GROUP BY articolo_padre_id"
        )).fetchall()
        return {row[0]: row[1] for row in rows}
    else:
        # Schema vecchio: padre_codice â†’ mappa a ID articolo dalla tabella articoli
        cursor.execute(
            "SELECT bs.padre_codice, COUNT(*) as cnt "
            "FROM bom_struttura bs GROUP BY bs.padre_codice"
        )
        counts_by_codice = {row[0]: row[1] for row in cursor.fetchall()}
        if not counts_by_codice:
            return {}

        # Mappa codice â†’ articoli.id
        result = {}
        placeholders = ','.join(['?'] * len(counts_by_codice))
        cursor.execute(
            f"SELECT id, codice FROM articoli WHERE codice IN ({placeholders})",
            list(counts_by_codice.keys())
        )
        for row in cursor.fetchall():
            art_id, codice = row[0], row[1]
            if codice in counts_by_codice:
                result[art_id] = counts_by_codice[codice]
        return result


@app.get("/bom/{articolo_padre_id}")
def get_bom(articolo_padre_id: int, db: Session = Depends(get_db)):
    """Lista componenti di un articolo padre"""
    schema = _detect_bom_schema(db)

    if schema == 'id':
        # â”€â”€ Schema nuovo: SQLAlchemy con articolo_padre_id/articolo_figlio_id â”€â”€
        righe = db.query(BomStruttura).filter(
            BomStruttura.articolo_padre_id == articolo_padre_id
        ).order_by(BomStruttura.ordine, BomStruttura.id).all()

        result = []
        for r in righe:
            figlio = db.query(Articolo).filter(Articolo.id == r.articolo_figlio_id).first()
            result.append({
                "id": r.id,
                "articolo_padre_id": r.articolo_padre_id,
                "articolo_figlio_id": r.articolo_figlio_id,
                "articolo_figlio": {
                    "id": figlio.id, "codice": figlio.codice,
                    "descrizione": figlio.descrizione,
                    "tipo_articolo": figlio.tipo_articolo,
                    "costo_fisso": figlio.costo_fisso,
                    "unita_misura": figlio.unita_misura,
                    "fornitore": getattr(figlio, 'fornitore', None),
                    "is_active": figlio.is_active,
                } if figlio else None,
                "quantita": r.quantita,
                "formula_quantita": r.formula_quantita,
                "unita_misura": r.unita_misura,
                "condizione_esistenza": r.condizione_esistenza,
                "note": r.note,
                "ordine": r.ordine,
            })
        return result

    else:
        # â”€â”€ Schema vecchio: padre_codice/figlio_codice (TEXT) â”€â”€
        # 1. Trova il codice dell'articolo padre dalla tabella articoli
        articolo_padre = db.query(Articolo).filter(Articolo.id == articolo_padre_id).first()
        if not articolo_padre:
            return []

        conn = db.get_bind().raw_connection()
        cursor = conn.cursor()

        # 2. Query bom_struttura con padre_codice
        cursor.execute(
            "SELECT bs.id, bs.padre_codice, bs.figlio_codice, bs.quantita, "
            "bs.formula_quantita, bs.condizione_esistenza, "
            "COALESCE(bs.note, '') as note, "
            "COALESCE(bs.posizione, 0) as ordine "
            "FROM bom_struttura bs "
            "WHERE bs.padre_codice = ? "
            "ORDER BY ordine, bs.id",
            (articolo_padre.codice,)
        )
        rows = cursor.fetchall()

        result = []
        for r in rows:
            bs_id, padre_cod, figlio_cod, qta, formula, condizione, note, ordine = r

            # 3. Cerca figlio prima in 'articoli' (SQLAlchemy), poi in 'articoli_bom' (raw)
            figlio = db.query(Articolo).filter(Articolo.codice == figlio_cod).first()
            figlio_data = None

            if figlio:
                figlio_data = {
                    "id": figlio.id,
                    "codice": figlio.codice,
                    "descrizione": figlio.descrizione,
                    "tipo_articolo": figlio.tipo_articolo,
                    "costo_fisso": figlio.costo_fisso,
                    "unita_misura": figlio.unita_misura,
                    "fornitore": getattr(figlio, 'fornitore', None),
                    "is_active": figlio.is_active,
                }
            else:
                # Fallback: cerca in articoli_bom (tabella creata da migrate_prototipo)
                cursor.execute(
                    "SELECT codice, descrizione, tipo, categoria, costo_fisso, "
                    "unita_misura FROM articoli_bom WHERE codice = ?",
                    (figlio_cod,)
                )
                ab = cursor.fetchone()
                if ab:
                    figlio_data = {
                        "id": 0,
                        "codice": ab[0],
                        "descrizione": ab[1],
                        "tipo_articolo": ab[2] or "ACQUISTO",
                        "costo_fisso": ab[4] or 0,
                        "unita_misura": ab[5] or "PZ",
                        "fornitore": None,
                        "is_active": True,
                    }

            result.append({
                "id": bs_id,
                "articolo_padre_id": articolo_padre_id,
                "articolo_figlio_id": figlio.id if figlio else 0,
                "articolo_figlio": figlio_data,
                "quantita": qta or 1,
                "formula_quantita": formula,
                "unita_misura": "PZ",
                "condizione_esistenza": condizione,
                "note": note,
                "ordine": ordine or 0,
            })

        return result


@app.post("/bom")
def create_bom_riga(data: dict, db: Session = Depends(get_db)):
    """Aggiunge un componente alla distinta"""
    padre_id = data.get("articolo_padre_id")
    figlio_id = data.get("articolo_figlio_id")

    if not padre_id or not figlio_id:
        raise HTTPException(400, "articolo_padre_id e articolo_figlio_id obbligatori")
    if padre_id == figlio_id:
        raise HTTPException(400, "Un articolo non puÃ² contenere sÃ© stesso")

    schema = _detect_bom_schema(db)

    if schema == 'id':
        exists = db.query(BomStruttura).filter(
            BomStruttura.articolo_padre_id == padre_id,
            BomStruttura.articolo_figlio_id == figlio_id,
        ).first()
        if exists:
            raise HTTPException(400, "Questo componente Ã¨ giÃ  nella distinta")

        def is_ancestor(potential_ancestor_id, target_id, visited=None):
            if visited is None:
                visited = set()
            if potential_ancestor_id in visited:
                return False
            visited.add(potential_ancestor_id)
            parents = db.query(BomStruttura.articolo_padre_id).filter(
                BomStruttura.articolo_figlio_id == target_id
            ).all()
            for (pid,) in parents:
                if pid == potential_ancestor_id:
                    return True
                if is_ancestor(potential_ancestor_id, pid, visited):
                    return True
            return False

        if is_ancestor(figlio_id, padre_id):
            raise HTTPException(400, "Riferimento circolare")

        max_ord = db.query(BomStruttura).filter(
            BomStruttura.articolo_padre_id == padre_id
        ).count()

        riga = BomStruttura(
            articolo_padre_id=padre_id,
            articolo_figlio_id=figlio_id,
            quantita=data.get("quantita", 1),
            formula_quantita=data.get("formula_quantita"),
            unita_misura=data.get("unita_misura", "PZ"),
            condizione_esistenza=data.get("condizione_esistenza"),
            note=data.get("note"),
            ordine=max_ord + 1,
        )
        db.add(riga)
        db.commit()
        db.refresh(riga)
        return {"id": riga.id, "status": "ok"}

    else:
        # Schema vecchio: padre_codice/figlio_codice
        padre = db.query(Articolo).filter(Articolo.id == padre_id).first()
        figlio = db.query(Articolo).filter(Articolo.id == figlio_id).first()
        if not padre or not figlio:
            raise HTTPException(400, "Articolo padre o figlio non trovato")

        conn = db.get_bind().raw_connection()
        cursor = conn.cursor()

        cursor.execute(
            "SELECT COUNT(*) FROM bom_struttura WHERE padre_codice=? AND figlio_codice=?",
            (padre.codice, figlio.codice)
        )
        if cursor.fetchone()[0] > 0:
            raise HTTPException(400, "Questo componente Ã¨ giÃ  nella distinta")

        cursor.execute(
            "SELECT COALESCE(MAX(posizione), 0) FROM bom_struttura WHERE padre_codice=?",
            (padre.codice,)
        )
        next_pos = (cursor.fetchone()[0] or 0) + 1

        cursor.execute(
            "INSERT INTO bom_struttura (padre_codice, figlio_codice, quantita, "
            "formula_quantita, condizione_esistenza, note, posizione, obbligatorio) "
            "VALUES (?, ?, ?, ?, ?, ?, ?, 1)",
            (padre.codice, figlio.codice, data.get("quantita", 1),
             data.get("formula_quantita"), data.get("condizione_esistenza"),
             data.get("note"), next_pos)
        )
        conn.commit()
        return {"id": cursor.lastrowid, "status": "ok"}


@app.put("/bom/{riga_id}")
def update_bom_riga(riga_id: int, data: dict, db: Session = Depends(get_db)):
    """Aggiorna una riga della distinta"""
    schema = _detect_bom_schema(db)

    if schema == 'id':
        riga = db.query(BomStruttura).filter(BomStruttura.id == riga_id).first()
        if not riga:
            raise HTTPException(404, "Riga BOM non trovata")
        for key in ["quantita", "formula_quantita", "unita_misura", "condizione_esistenza", "note", "ordine"]:
            if key in data:
                setattr(riga, key, data[key])
        db.commit()
    else:
        conn = db.get_bind().raw_connection()
        cursor = conn.cursor()
        updates = []
        values = []
        for key in ["quantita", "formula_quantita", "condizione_esistenza", "note"]:
            if key in data:
                updates.append(f"{key} = ?")
                values.append(data[key])
        if "ordine" in data:
            updates.append("posizione = ?")
            values.append(data["ordine"])
        if not updates:
            return {"id": riga_id, "status": "ok"}
        values.append(riga_id)
        cursor.execute(f"UPDATE bom_struttura SET {', '.join(updates)} WHERE id = ?", values)
        if cursor.rowcount == 0:
            raise HTTPException(404, "Riga BOM non trovata")
        conn.commit()

    return {"id": riga_id, "status": "ok"}


@app.delete("/bom/{riga_id}")
def delete_bom_riga(riga_id: int, db: Session = Depends(get_db)):
    """Rimuove una riga dalla distinta"""
    schema = _detect_bom_schema(db)

    if schema == 'id':
        riga = db.query(BomStruttura).filter(BomStruttura.id == riga_id).first()
        if not riga:
            raise HTTPException(404, "Riga BOM non trovata")
        db.delete(riga)
        db.commit()
    else:
        conn = db.get_bind().raw_connection()
        cursor = conn.cursor()
        cursor.execute("DELETE FROM bom_struttura WHERE id = ?", (riga_id,))
        if cursor.rowcount == 0:
            raise HTTPException(404, "Riga BOM non trovata")
        conn.commit()

    return {"status": "ok"}


@app.post("/bom/duplica")
def duplica_bom(data: dict, db: Session = Depends(get_db)):
    """Duplica l'intera BOM da un articolo sorgente a un articolo destinazione"""
    source_id = data.get("source_articolo_id")
    target_id = data.get("target_articolo_id")

    if not source_id or not target_id:
        raise HTTPException(400, "source_articolo_id e target_articolo_id obbligatori")
    if source_id == target_id:
        raise HTTPException(400, "Sorgente e destinazione devono essere diversi")

    schema = _detect_bom_schema(db)

    if schema == 'id':
        db.query(BomStruttura).filter(BomStruttura.articolo_padre_id == target_id).delete()
        righe_source = db.query(BomStruttura).filter(
            BomStruttura.articolo_padre_id == source_id
        ).order_by(BomStruttura.ordine).all()
        for r in righe_source:
            nuova = BomStruttura(
                articolo_padre_id=target_id,
                articolo_figlio_id=r.articolo_figlio_id,
                quantita=r.quantita,
                formula_quantita=r.formula_quantita,
                unita_misura=r.unita_misura,
                condizione_esistenza=r.condizione_esistenza,
                note=r.note,
                ordine=r.ordine,
            )
            db.add(nuova)
        db.commit()
        return {"status": "ok", "righe_duplicate": len(righe_source)}

    else:
        source = db.query(Articolo).filter(Articolo.id == source_id).first()
        target = db.query(Articolo).filter(Articolo.id == target_id).first()
        if not source or not target:
            raise HTTPException(400, "Articoli non trovati")

        conn = db.get_bind().raw_connection()
        cursor = conn.cursor()
        cursor.execute("DELETE FROM bom_struttura WHERE padre_codice = ?", (target.codice,))
        cursor.execute(
            "SELECT figlio_codice, quantita, formula_quantita, condizione_esistenza, "
            "note, posizione, obbligatorio FROM bom_struttura WHERE padre_codice = ? ORDER BY posizione",
            (source.codice,)
        )
        righe = cursor.fetchall()
        for r in righe:
            cursor.execute(
                "INSERT INTO bom_struttura (padre_codice, figlio_codice, quantita, "
                "formula_quantita, condizione_esistenza, note, posizione, obbligatorio) "
                "VALUES (?, ?, ?, ?, ?, ?, ?, ?)",
                (target.codice, r[0], r[1], r[2], r[3], r[4], r[5], r[6])
            )
        conn.commit()
        return {"status": "ok", "righe_duplicate": len(righe)}

# ==========================================
# CATEGORIE ARTICOLI
# ==========================================
@app.get("/categorie-articoli")
def get_categorie_articoli(db: Session = Depends(get_db)):
    categorie = db.query(CategoriaArticoli).order_by(CategoriaArticoli.ordine).all()
    return [{
        "id": c.id, "codice": c.codice, "nome": c.nome,
        "descrizione": c.descrizione, "ordine": c.ordine, "is_active": c.is_active
    } for c in categorie]

# ==========================================
# CLIENTI
# ==========================================
@app.get("/clienti")
def get_clienti(db: Session = Depends(get_db)):
    clienti = db.query(Cliente).filter(Cliente.is_active == True).order_by(Cliente.ragione_sociale).all()
    return [{
        "id": c.id, "codice": c.codice, "ragione_sociale": c.ragione_sociale,
        "partita_iva": c.partita_iva, "codice_fiscale": c.codice_fiscale,
        "indirizzo": c.indirizzo, "cap": c.cap, "citta": c.citta,
        "provincia": c.provincia, "nazione": c.nazione,
        "telefono": c.telefono, "email": c.email, "pec": c.pec,
        "sconto_globale": c.sconto_globale, "sconto_produzione": c.sconto_produzione,
        "sconto_acquisto": c.sconto_acquisto, "aliquota_iva": c.aliquota_iva,
        "pagamento_default": c.pagamento_default, "imballo_default": c.imballo_default,
        "reso_fco_default": c.reso_fco_default, "trasporto_default": c.trasporto_default,
        "destinazione_default": c.destinazione_default,
        "riferimento_cliente_default": c.riferimento_cliente_default,
        "listino": c.listino, "note": c.note, "is_active": c.is_active
    } for c in clienti]

@app.get("/clienti/search")
def search_clienti(q: str = "", db: Session = Depends(get_db)):
    if not q:
        return []
    search_term = f"%{q}%"
    clienti = db.query(Cliente).filter(
        or_(
            Cliente.codice.ilike(search_term),
            Cliente.ragione_sociale.ilike(search_term),
            Cliente.partita_iva.ilike(search_term)
        )
    ).limit(20).all()
    return [{
        "id": c.id, "codice": c.codice, "ragione_sociale": c.ragione_sociale,
        "partita_iva": c.partita_iva, "citta": c.citta, "provincia": c.provincia,
        "sconto_globale": c.sconto_globale, "sconto_produzione": c.sconto_produzione,
        "sconto_acquisto": c.sconto_acquisto,
        "pagamento_default": c.pagamento_default, "imballo_default": c.imballo_default,
        "reso_fco_default": c.reso_fco_default, "trasporto_default": c.trasporto_default,
        "destinazione_default": c.destinazione_default,
        "riferimento_cliente_default": c.riferimento_cliente_default
    } for c in clienti]

@app.get("/clienti/{cliente_id}")
def get_cliente(cliente_id: int, db: Session = Depends(get_db)):
    cliente = db.query(Cliente).filter(Cliente.id == cliente_id).first()
    if not cliente:
        raise HTTPException(status_code=404, detail="Cliente non trovato")
    return {
        "id": cliente.id, "codice": cliente.codice, "ragione_sociale": cliente.ragione_sociale,
        "partita_iva": cliente.partita_iva, "codice_fiscale": cliente.codice_fiscale,
        "indirizzo": cliente.indirizzo, "cap": cliente.cap, "citta": cliente.citta,
        "provincia": cliente.provincia, "nazione": cliente.nazione,
        "telefono": cliente.telefono, "email": cliente.email, "pec": cliente.pec,
        "sconto_globale": cliente.sconto_globale, "sconto_produzione": cliente.sconto_produzione,
        "sconto_acquisto": cliente.sconto_acquisto, "aliquota_iva": cliente.aliquota_iva,
        "pagamento_default": cliente.pagamento_default, "imballo_default": cliente.imballo_default,
        "reso_fco_default": cliente.reso_fco_default, "trasporto_default": cliente.trasporto_default,
        "destinazione_default": cliente.destinazione_default,
        "riferimento_cliente_default": cliente.riferimento_cliente_default,
        "listino": cliente.listino, "note": cliente.note, "is_active": cliente.is_active
    }

@app.post("/clienti")
def create_cliente(data: dict, db: Session = Depends(get_db)):
    cliente = Cliente()
    for key in [
        "codice", "ragione_sociale", "partita_iva", "codice_fiscale",
        "indirizzo", "cap", "citta", "provincia", "nazione",
        "telefono", "email", "pec",
        "sconto_globale", "sconto_produzione", "sconto_acquisto", "aliquota_iva",
        "pagamento_default", "imballo_default", "reso_fco_default",
        "trasporto_default", "destinazione_default", "riferimento_cliente_default",
        "listino", "note", "is_active"
    ]:
        if key in data:
            setattr(cliente, key, data[key])
    db.add(cliente)
    db.commit()
    db.refresh(cliente)
    return {"id": cliente.id, "status": "ok"}

@app.put("/clienti/{cliente_id}")
def update_cliente(cliente_id: int, data: dict, db: Session = Depends(get_db)):
    cliente = db.query(Cliente).filter(Cliente.id == cliente_id).first()
    if not cliente:
        raise HTTPException(status_code=404, detail="Cliente non trovato")
    for key in [
        "codice", "ragione_sociale", "partita_iva", "codice_fiscale",
        "indirizzo", "cap", "citta", "provincia", "nazione",
        "telefono", "email", "pec",
        "sconto_globale", "sconto_produzione", "sconto_acquisto", "aliquota_iva",
        "pagamento_default", "imballo_default", "reso_fco_default",
        "trasporto_default", "destinazione_default", "riferimento_cliente_default",
        "listino", "note", "is_active"
    ]:
        if key in data:
            setattr(cliente, key, data[key])
    cliente.updated_at = datetime.now()
    db.commit()
    return {"id": cliente.id, "status": "ok"}

# ==========================================
# UTENTI
# ==========================================
@app.get("/utenti")
def get_utenti(db: Session = Depends(get_db)):
    utenti = db.query(Utente).order_by(Utente.username).all()
    result = []
    for u in utenti:
        gruppo_nome = None
        if u.gruppo_id:
            g = db.query(GruppoUtenti).filter(GruppoUtenti.id == u.gruppo_id).first()
            gruppo_nome = g.nome if g else None
        ruolo_nome = None
        ruolo_codice = None
        if u.ruolo_id:
            r = db.query(Ruolo).filter(Ruolo.id == u.ruolo_id).first()
            if r:
                ruolo_nome = r.nome
                ruolo_codice = r.codice
        result.append({
            "id": u.id,
            "username": u.username,
            "nome": u.nome,
            "cognome": u.cognome,
            "email": u.email,
            "gruppo_id": u.gruppo_id,
            "gruppo_nome": gruppo_nome,
            "ruolo_id": u.ruolo_id,
            "ruolo_nome": ruolo_nome,
            "ruolo_codice": ruolo_codice,
            "is_admin": u.is_admin,
            "is_active": u.is_active,
            "created_at": str(u.created_at) if u.created_at else None,
            "last_login": str(u.last_login) if u.last_login else None,
        })
    return result

@app.post("/utenti")
def create_utente(
    username: str = Query(...),
    password: str = Query(None),
    nome: str = Query(None),
    cognome: str = Query(None),
    email: str = Query(None),
    gruppo_id: int = Query(None),
    ruolo_id: int = Query(None),
    cliente_id: int = Query(None),
    is_admin: bool = Query(False),
    is_active: bool = Query(True),
    db: Session = Depends(get_db)
    
):
    existing = db.query(Utente).filter(Utente.username == username).first()
    if existing:
        raise HTTPException(status_code=400, detail="Username già esistente")
    
    utente = Utente(
        username=username,
        password_hash=get_password_hash(password or "changeme"),
        nome=nome, cognome=cognome, email=email,
        gruppo_id=gruppo_id, ruolo_id=ruolo_id,
        cliente_id=cliente_id,
        is_admin=is_admin, is_active=is_active
    )
    db.add(utente)
    db.commit()
    db.refresh(utente)
    return {"id": utente.id, "status": "ok"}

@app.put("/utenti/{utente_id}")
def update_utente(
    utente_id: int,
    username: str = Query(None),
    password: str = Query(None),
    nome: str = Query(None),
    cognome: str = Query(None),
    email: str = Query(None),
    gruppo_id: int = Query(None),
    ruolo_id: int = Query(None),
    cliente_id: int = Query(None),
    is_admin: bool = Query(None),
    is_active: bool = Query(None),
    db: Session = Depends(get_db)
):
    utente = db.query(Utente).filter(Utente.id == utente_id).first()
    if not utente:
        raise HTTPException(status_code=404, detail="Utente non trovato")
    
    if username is not None:
        utente.username = username
    if password is not None:
        utente.password_hash = get_password_hash(password)
    if nome is not None:
        utente.nome = nome
    if cognome is not None:
        utente.cognome = cognome
    if email is not None:
        utente.email = email
    if gruppo_id is not None:
        utente.gruppo_id = gruppo_id
    if ruolo_id is not None:
        utente.ruolo_id = ruolo_id
    if cliente_id is not None:
        utente.cliente_id = cliente_id if cliente_id != 0 else None
    if is_admin is not None:
        utente.is_admin = is_admin
    if is_active is not None:
        utente.is_active = is_active
    
    db.commit()
    db.refresh(utente)
    return {"id": utente.id, "status": "ok"}

@app.delete("/utenti/{utente_id}")
def delete_utente(utente_id: int, db: Session = Depends(get_db)):
    utente = db.query(Utente).filter(Utente.id == utente_id).first()
    if not utente:
        raise HTTPException(status_code=404, detail="Utente non trovato")
    db.delete(utente)
    db.commit()
    return {"status": "ok"}

# ==========================================
# PARAMETRI SISTEMA
# ==========================================
@app.get("/parametri-sistema")
def get_parametri_sistema(db: Session = Depends(get_db)):
    parametri = db.query(ParametriSistema).order_by(ParametriSistema.gruppo, ParametriSistema.chiave).all()
    return [{
        "id": p.id, "chiave": p.chiave, "valore": p.valore,
        "descrizione": p.descrizione, "tipo_dato": p.tipo_dato,
        "gruppo": p.gruppo
    } for p in parametri]

@app.put("/parametri-sistema/{chiave}")
def update_parametro_sistema(chiave: str, data: dict, db: Session = Depends(get_db)):
    parametro = db.query(ParametriSistema).filter(ParametriSistema.chiave == chiave).first()
    if not parametro:
        raise HTTPException(status_code=404, detail="Parametro non trovato")
    if "valore" in data:
        parametro.valore = data["valore"]
    parametro.updated_at = datetime.now()
    db.commit()
    return {"status": "ok"}

# ==========================================
# OPZIONI DROPDOWN
# ==========================================
@app.get("/opzioni-dropdown/gruppi")
def get_gruppi_dropdown(db: Session = Depends(get_db)):
    try:
        result = db.execute(text("""
            SELECT gruppo,
                   COUNT(*) as totale,
                   SUM(CASE WHEN attivo = 1 THEN 1 ELSE 0 END) as attive
            FROM opzioni_dropdown
            GROUP BY gruppo
            ORDER BY gruppo
        """))
        return [{"gruppo": r[0], "totale": r[1], "attive": r[2]} for r in result.fetchall()]
    except:
        return []
    
@app.get("/opzioni-dropdown/gruppi-sezioni-map")
def get_gruppi_sezioni_map(db: Session = Depends(get_db)):
    """
    Restituisce { gruppo_dropdown: { sezione_codice, sezione_etichetta } }
    costruita dai campi_configuratore + sezioni_configuratore.
    """
    try:
        result = db.execute(text("""
            SELECT DISTINCT c.gruppo_dropdown, c.sezione, 
                   COALESCE(s.etichetta, c.sezione) as sezione_etichetta
            FROM campi_configuratore c
            LEFT JOIN sezioni_configuratore s ON s.codice = c.sezione
            WHERE c.gruppo_dropdown IS NOT NULL 
              AND c.gruppo_dropdown != ''
        """))
        mapping = {}
        for row in result.fetchall():
            mapping[row[0]] = {
                "sezione_codice": row[1],
                "sezione_etichetta": row[2]
            }
        return mapping
    except Exception as e:
        print(f"Errore gruppi-sezioni-map: {e}")
        return {}

@app.get("/opzioni-dropdown/{gruppo}")
def get_opzioni_dropdown(gruppo: str, solo_attive: bool = True, db: Session = Depends(get_db)):
    try:
        q = "SELECT id, gruppo, valore, etichetta, ordine, attivo FROM opzioni_dropdown WHERE gruppo = :g"
        if solo_attive:
            q += " AND attivo = 1"
        q += " ORDER BY ordine"
        result = db.execute(text(q), {"g": gruppo})
        return [{
            "id": r[0], "gruppo": r[1], "valore": r[2],
            "etichetta": r[3],   # <-- era "label"
            "label": r[3],       # backward compat
            "ordine": r[4], "attivo": bool(r[5])
        } for r in result.fetchall()]
    except:
        return []


# SOSTITUIRE @app.post("/opzioni-dropdown") (riga ~2371-2379)
@app.post("/opzioni-dropdown")
def create_opzione_dropdown(data: dict, db: Session = Depends(get_db)):
    try:
        # Accetta sia "etichetta" che "label"
        etichetta = data.get("etichetta") or data.get("label") or data.get("valore", "")
        db.execute(text(
            "INSERT INTO opzioni_dropdown (gruppo, valore, etichetta, ordine, attivo) "
            "VALUES (:gruppo, :valore, :etichetta, :ordine, :attivo)"
        ), {
            "gruppo": data["gruppo"],
            "valore": data["valore"],
            "etichetta": etichetta,
            "ordine": data.get("ordine", 0),
            "attivo": 1 if data.get("attivo", True) else 0,
        })
        db.commit()
        return {"status": "ok"}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


# SOSTITUIRE @app.put("/opzioni-dropdown/{opzione_id}") (riga ~2381-2389)
@app.put("/opzioni-dropdown/{opzione_id}")
def update_opzione_dropdown(opzione_id: int, data: dict, db: Session = Depends(get_db)):
    """Aggiorna SOLO i campi presenti nel payload."""
    try:
        updates = []
        params = {"id": opzione_id}

        if "valore" in data and data["valore"] is not None:
            updates.append("valore = :valore")
            params["valore"] = data["valore"]

        # Accetta sia "etichetta" che "label"
        etichetta = data.get("etichetta", data.get("label"))
        if etichetta is not None:
            updates.append("etichetta = :etichetta")
            params["etichetta"] = etichetta

        if "ordine" in data:
            updates.append("ordine = :ordine")
            params["ordine"] = data["ordine"]

        if "attivo" in data:
            updates.append("attivo = :attivo")
            params["attivo"] = 1 if data["attivo"] else 0

        if not updates:
            return {"status": "noop"}

        sql = f"UPDATE opzioni_dropdown SET {', '.join(updates)} WHERE id = :id"
        db.execute(text(sql), params)
        db.commit()
        return {"status": "ok"}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.delete("/opzioni-dropdown/{opzione_id}")
def delete_opzione_dropdown(opzione_id: int, db: Session = Depends(get_db)):
    try:
        db.execute(text("DELETE FROM opzioni_dropdown WHERE id=:id"), {"id": opzione_id})
        db.commit()
        return {"status": "ok"}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/opzioni-dropdown/{gruppo}/riordina")
def riordina_opzioni(gruppo: str, data: dict, db: Session = Depends(get_db)):
    try:
        for i, oid in enumerate(data.get("ids", [])):
            db.execute(text("UPDATE opzioni_dropdown SET ordine=:o WHERE id=:id"), {"o": i, "id": oid})
        db.commit()
        return {"status": "ok"}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

# ==========================================
# DOCUMENT TEMPLATES — CRUD + Export
# ==========================================

@app.get("/document-templates/available-fields")
def get_available_fields(db: Session = Depends(get_db)):
    """Restituisce tutti i campi disponibili (statici + dinamici dal DB)"""
    return get_available_fields_from_db(db)


@app.get("/document-templates/default-config/{tipo}")
def get_default_config(tipo: str, db: Session = Depends(get_db)):
    """Restituisce una configurazione template di default per tipo (preventivo/ordine)"""
    if tipo not in ("preventivo", "ordine"):
        raise HTTPException(status_code=400, detail="Tipo deve essere 'preventivo' o 'ordine'")
    return get_default_template_config_from_db(db, tipo)


@app.get("/document-templates")
def list_document_templates(tipo: str = None, db: Session = Depends(get_db)):
    """Lista tutti i document templates, opzionalmente filtrati per tipo"""
    query = db.query(DocumentTemplate)
    if tipo:
        query = query.filter(DocumentTemplate.tipo == tipo)
    templates = query.order_by(DocumentTemplate.tipo, DocumentTemplate.nome).all()
    return [
        {
            "id": t.id,
            "nome": t.nome,
            "tipo": t.tipo,
            "descrizione": t.descrizione,
            "attivo": t.attivo,
            "is_default": t.is_default,
            "has_logo": t.logo_data is not None,
            "logo_filename": t.logo_filename,
            "config": t.config,
            "created_at": str(t.created_at) if t.created_at else None,
            "updated_at": str(t.updated_at) if t.updated_at else None,
        }
        for t in templates
    ]


@app.get("/document-templates/{template_id}")
def get_document_template(template_id: int, db: Session = Depends(get_db)):
    """Recupera un singolo document template con config completa"""
    t = db.query(DocumentTemplate).filter(DocumentTemplate.id == template_id).first()
    if not t:
        raise HTTPException(status_code=404, detail="Template non trovato")
    result = {
        "id": t.id,
        "nome": t.nome,
        "tipo": t.tipo,
        "descrizione": t.descrizione,
        "attivo": t.attivo,
        "is_default": t.is_default,
        "has_logo": t.logo_data is not None,
        "logo_filename": t.logo_filename,
        "logo_mime": t.logo_mime,
        "config": t.config,
        "created_at": str(t.created_at) if t.created_at else None,
        "updated_at": str(t.updated_at) if t.updated_at else None,
    }
    if t.logo_data:
        result["logo_base64"] = base64.b64encode(t.logo_data).decode("utf-8")
    return result


@app.post("/document-templates")
def create_document_template(
    nome: str = Form(...),
    tipo: str = Form(...),
    descrizione: str = Form(""),
    config_json: str = Form(...),
    is_default: bool = Form(False),
    logo: UploadFile = File(None),
    db: Session = Depends(get_db)
):
    """Crea un nuovo document template"""
    if tipo not in ("preventivo", "ordine"):
        raise HTTPException(status_code=400, detail="Tipo deve essere 'preventivo' o 'ordine'")

    try:
        config = json.loads(config_json)
    except json.JSONDecodeError:
        raise HTTPException(status_code=400, detail="config_json non e' un JSON valido")

    if is_default:
        db.query(DocumentTemplate).filter(
            DocumentTemplate.tipo == tipo,
            DocumentTemplate.is_default == True
        ).update({"is_default": False})

    template = DocumentTemplate(
        nome=nome, tipo=tipo, descrizione=descrizione,
        config=config, is_default=is_default, attivo=True,
    )

    if logo:
        logo_bytes = logo.file.read()
        template.logo_data = logo_bytes
        template.logo_filename = logo.filename
        template.logo_mime = logo.content_type

    db.add(template)
    db.commit()
    db.refresh(template)
    return {"id": template.id, "nome": template.nome, "message": "Template creato"}


@app.put("/document-templates/{template_id}")
def update_document_template(
    template_id: int,
    nome: str = Form(None),
    tipo: str = Form(None),
    descrizione: str = Form(None),
    config_json: str = Form(None),
    attivo: bool = Form(None),
    is_default: bool = Form(None),
    remove_logo: bool = Form(False),
    logo: UploadFile = File(None),
    db: Session = Depends(get_db)
):
    """Aggiorna un document template esistente"""
    template = db.query(DocumentTemplate).filter(DocumentTemplate.id == template_id).first()
    if not template:
        raise HTTPException(status_code=404, detail="Template non trovato")

    if nome is not None: template.nome = nome
    if tipo is not None: template.tipo = tipo
    if descrizione is not None: template.descrizione = descrizione
    if attivo is not None: template.attivo = attivo
    if config_json is not None:
        try:
            template.config = json.loads(config_json)
        except json.JSONDecodeError:
            raise HTTPException(status_code=400, detail="config_json non e' un JSON valido")

    if is_default is not None:
        if is_default:
            db.query(DocumentTemplate).filter(
                DocumentTemplate.tipo == (tipo or template.tipo),
                DocumentTemplate.is_default == True,
                DocumentTemplate.id != template_id
            ).update({"is_default": False})
        template.is_default = is_default

    if remove_logo:
        template.logo_data = None
        template.logo_filename = None
        template.logo_mime = None
    elif logo:
        logo_bytes = logo.file.read()
        template.logo_data = logo_bytes
        template.logo_filename = logo.filename
        template.logo_mime = logo.content_type

    db.commit()
    return {"id": template.id, "message": "Template aggiornato"}


@app.delete("/document-templates/{template_id}")
def delete_document_template(template_id: int, db: Session = Depends(get_db)):
    """Elimina un document template"""
    template = db.query(DocumentTemplate).filter(DocumentTemplate.id == template_id).first()
    if not template:
        raise HTTPException(status_code=404, detail="Template non trovato")
    db.delete(template)
    db.commit()
    return {"message": "Template eliminato"}

# ============================================================================
# SEZIONE A: CRUD sezioni_preventivo (NUOVI endpoint)
# ============================================================================

@app.get("/sezioni-preventivo")
def get_sezioni_preventivo(db: Session = Depends(get_db)):
    """Lista sezioni del documento preventivo, ordinate."""
    try:
        result = db.execute(text("""
            SELECT id, codice, titolo, ordine, tipo, visibile, mostra_titolo, nota_default, stile
            FROM sezioni_preventivo
            ORDER BY ordine
        """))
        return [
            {"id": r[0], "codice": r[1], "titolo": r[2], "ordine": r[3],
             "tipo": r[4], "visibile": bool(r[5]), "mostra_titolo": bool(r[6]),
             "nota_default": r[7], "stile": r[8]}
            for r in result.fetchall()
        ]
    except Exception as e:
        return []


@app.post("/sezioni-preventivo")
def create_sezione_preventivo(data: dict, db: Session = Depends(get_db)):
    """Crea nuova sezione documento."""
    try:
        db.execute(text("""
            INSERT INTO sezioni_preventivo (codice, titolo, ordine, tipo, visibile, mostra_titolo, nota_default, stile)
            VALUES (:codice, :titolo, :ordine, :tipo, :vis, :mt, :nota, :stile)
        """), {
            "codice": data["codice"], "titolo": data["titolo"],
            "ordine": data.get("ordine", 0), "tipo": data.get("tipo", "tabella"),
            "vis": 1 if data.get("visibile", True) else 0,
            "mt": 1 if data.get("mostra_titolo", True) else 0,
            "nota": data.get("nota_default"), "stile": data.get("stile"),
        })
        db.commit()
        return {"status": "ok"}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))
@app.put("/sezioni-preventivo/riordina")
def riordina_sezioni_preventivo(data: dict, db: Session = Depends(get_db)):
    """Riordina sezioni. Body: {"ordine": [{"id": 1, "ordine": 10}, ...]}"""
    try:
        for item in data.get("ordine", []):
            db.execute(text("UPDATE sezioni_preventivo SET ordine=:o WHERE id=:id"),
                       {"o": item["ordine"], "id": item["id"]})
        db.commit()
        return {"status": "ok"}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))
    

@app.put("/sezioni-preventivo/{sezione_id}")
def update_sezione_preventivo(sezione_id: int, data: dict, db: Session = Depends(get_db)):
    """Aggiorna sezione documento."""
    try:
        field_map = {
            "codice": "codice", "titolo": "titolo", "ordine": "ordine",
            "tipo": "tipo", "visibile": "visibile", "mostra_titolo": "mostra_titolo",
            "nota_default": "nota_default", "stile": "stile",
        }
        fields, params = [], {"id": sezione_id}
        for key, col in field_map.items():
            if key in data:
                val = data[key]
                if key in ("visibile", "mostra_titolo"):
                    val = 1 if val else 0
                fields.append(f"{col}=:{col}")
                params[col] = val
        if not fields:
            return {"status": "noop"}
        fields.append("updated_at=CURRENT_TIMESTAMP")
        db.execute(text(f"UPDATE sezioni_preventivo SET {', '.join(fields)} WHERE id=:id"), params)
        db.commit()
        return {"status": "ok"}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.delete("/sezioni-preventivo/{sezione_id}")
def delete_sezione_preventivo(sezione_id: int, db: Session = Depends(get_db)):
    """Elimina sezione documento."""
    try:
        db.execute(text("DELETE FROM sezioni_preventivo WHERE id=:id"), {"id": sezione_id})
        db.commit()
        return {"status": "ok"}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))




# ==========================================
# SEZIONI CONFIGURATORE
# ==========================================

def _ensure_sezioni_table(db):
    """Crea tabella sezioni_configuratore se non esiste"""
    try:
        db.execute(text('''
            CREATE TABLE IF NOT EXISTS sezioni_configuratore (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                codice TEXT NOT NULL UNIQUE,
                etichetta TEXT NOT NULL,
                descrizione TEXT,
                icona TEXT DEFAULT 'Settings',
                ordine INTEGER DEFAULT 0,
                attivo INTEGER DEFAULT 1,
                product_template_id INTEGER,
                product_template_ids TEXT,
                created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                updated_at TIMESTAMP
            )
        '''))
        db.commit()
    except:
        pass
    # Migrazione: aggiungi colonna se manca (tabella gia' esistente)
    try:
        db.execute(text("SELECT product_template_ids FROM sezioni_configuratore LIMIT 1"))
    except:
        try:
            db.execute(text("ALTER TABLE sezioni_configuratore ADD COLUMN product_template_ids TEXT"))
            db.commit()
        except:
            pass


@app.get("/sezioni-configuratore")
def get_sezioni_configuratore_list(db: Session = Depends(get_db)):
    """Lista tutte le sezioni con info prodotti associati e conteggio campi"""
    _ensure_sezioni_table(db)
    try:
        import json as _json
        result = db.execute(text("""
            SELECT s.id, s.codice, s.etichetta, s.descrizione, s.icona,
                   s.ordine, s.attivo, s.product_template_id, s.product_template_ids,
                   (SELECT COUNT(*) FROM campi_configuratore c WHERE c.sezione = s.codice) as num_campi,
                   (SELECT COUNT(*) FROM campi_configuratore c WHERE c.sezione = s.codice AND c.attivo = 1) as num_campi_attivi
            FROM sezioni_configuratore s
            ORDER BY s.ordine, s.etichetta
        """))
        rows = result.fetchall()
        
        # Carica tutti i product_templates per risolvere i nomi
        pt_result = db.execute(text("SELECT id, categoria, sottocategoria, nome_display FROM product_templates"))
        pt_map = {r[0]: {"categoria": r[1], "sottocategoria": r[2], "nome_display": r[3]} for r in pt_result.fetchall()}
        
        sezioni = []
        for r in rows:
            # Risolvi prodotti associati
            pt_ids = []
            if r[8]:  # product_template_ids (JSON array)
                try:
                    pt_ids = _json.loads(r[8])
                except:
                    pt_ids = []
            elif r[7]:  # fallback: vecchio product_template_id singolo
                pt_ids = [r[7]]
            
            prodotti = [pt_map[pid] for pid in pt_ids if pid in pt_map]
            
            sezioni.append({
                "id": r[0], "codice": r[1], "etichetta": r[2],
                "descrizione": r[3], "icona": r[4], "ordine": r[5],
                "attivo": bool(r[6]),
                "product_template_ids": pt_ids,
                "prodotti": prodotti,
                # Backward compat
                "product_template_id": pt_ids[0] if pt_ids else None,
                "prodotto": prodotti[0] if prodotti else None,
                "num_campi": r[9], "num_campi_attivi": r[10]
            })
        return sezioni
    except Exception as e:
        print(f"Errore get_sezioni: {e}")
        return []


@app.get("/sezioni-configuratore/campi-non-assegnati")
def get_campi_non_assegnati(db: Session = Depends(get_db)):
    """Campi la cui sezione non corrisponde a nessuna sezione registrata"""
    _ensure_sezioni_table(db)
    try:
        result = db.execute(text("""
            SELECT c.id, c.codice, c.etichetta, c.tipo, c.sezione, c.attivo
            FROM campi_configuratore c
            WHERE c.sezione NOT IN (SELECT codice FROM sezioni_configuratore)
               OR c.sezione IS NULL OR c.sezione = ''
            ORDER BY c.sezione, c.ordine
        """))
        return [{"id": r[0], "codice": r[1], "etichetta": r[2], "tipo": r[3], "sezione": r[4], "attivo": bool(r[5])} for r in result.fetchall()]
    except:
        return []


@app.get("/sezioni-configuratore/{sezione_id}")
def get_sezione_detail(sezione_id: int, db: Session = Depends(get_db)):
    """Dettaglio sezione con campi associati"""
    _ensure_sezioni_table(db)
    try:
        import json as _json
        result = db.execute(text("""
            SELECT s.id, s.codice, s.etichetta, s.descrizione, s.icona,
                   s.ordine, s.attivo, s.product_template_id, s.product_template_ids
            FROM sezioni_configuratore s
            WHERE s.id = :id
        """), {"id": sezione_id})
        r = result.fetchone()
        if not r:
            raise HTTPException(status_code=404, detail="Sezione non trovata")
        
        # Risolvi prodotti
        pt_ids = []
        if r[8]:
            try: pt_ids = _json.loads(r[8])
            except: pt_ids = []
        elif r[7]:
            pt_ids = [r[7]]
        
        prodotti = []
        if pt_ids:
            pt_result = db.execute(text("SELECT id, categoria, sottocategoria, nome_display FROM product_templates"))
            pt_map = {pr[0]: {"categoria": pr[1], "sottocategoria": pr[2], "nome_display": pr[3]} for pr in pt_result.fetchall()}
            prodotti = [pt_map[pid] for pid in pt_ids if pid in pt_map]
        
        campi_result = db.execute(text("""
            SELECT id, codice, etichetta, tipo, sezione, gruppo_dropdown, ordine, attivo,
                   obbligatorio, unita_misura, valore_default, visibile_form, usabile_regole
            FROM campi_configuratore WHERE sezione = :s ORDER BY ordine
        """), {"s": r[1]})
        campi = [{
            "id": c[0], "codice": c[1], "etichetta": c[2], "tipo": c[3],
            "sezione": c[4], "gruppo_dropdown": c[5], "ordine": c[6],
            "attivo": bool(c[7]), "obbligatorio": bool(c[8]),
            "unita_misura": c[9], "valore_default": c[10],
            "visibile_form": bool(c[11]) if c[11] is not None else True,
            "usabile_regole": bool(c[12]) if c[12] is not None else False
        } for c in campi_result.fetchall()]
        
        return {
            "id": r[0], "codice": r[1], "etichetta": r[2],
            "descrizione": r[3], "icona": r[4], "ordine": r[5],
            "attivo": bool(r[6]),
            "product_template_ids": pt_ids,
            "prodotti": prodotti,
            "product_template_id": pt_ids[0] if pt_ids else None,
            "prodotto": prodotti[0] if prodotti else None,
            "campi": campi
        }
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/sezioni-configuratore")
def create_sezione_configuratore(data: dict, db: Session = Depends(get_db)):
    """Crea una nuova sezione"""
    _ensure_sezioni_table(db)
    try:
        import json as _json
        existing = db.execute(text("SELECT id FROM sezioni_configuratore WHERE codice = :c"), {"c": data["codice"]}).fetchone()
        if existing:
            raise HTTPException(status_code=400, detail=f"Codice '{data['codice']}' gia' esistente")
        
        # Gestisci product_template_ids (array) o fallback a product_template_id (singolo)
        pt_ids_json = None
        if "product_template_ids" in data and data["product_template_ids"]:
            pt_ids_json = _json.dumps(data["product_template_ids"])
        elif "product_template_id" in data and data["product_template_id"]:
            pt_ids_json = _json.dumps([data["product_template_id"]])
        
        max_ord = db.execute(text("SELECT COALESCE(MAX(ordine), -1) FROM sezioni_configuratore")).fetchone()[0]
        db.execute(text("""
            INSERT INTO sezioni_configuratore (codice, etichetta, descrizione, icona, ordine, attivo, product_template_ids)
            VALUES (:codice, :etichetta, :descrizione, :icona, :ordine, :attivo, :pt_ids)
        """), {
            "codice": data["codice"],
            "etichetta": data.get("etichetta", data["codice"]),
            "descrizione": data.get("descrizione"),
            "icona": data.get("icona", "Settings"),
            "ordine": data.get("ordine", max_ord + 1),
            "attivo": 1 if data.get("attivo", True) else 0,
            "pt_ids": pt_ids_json
        })
        db.commit()
        new_id = db.execute(text("SELECT id FROM sezioni_configuratore WHERE codice = :c"), {"c": data["codice"]}).fetchone()[0]
        return {"status": "ok", "id": new_id}
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.put("/sezioni-configuratore/{sezione_id}")
def update_sezione_configuratore(sezione_id: int, data: dict, db: Session = Depends(get_db)):
    """Aggiorna una sezione"""
    _ensure_sezioni_table(db)
    try:
        import json as _json
        existing = db.execute(text("SELECT id, codice FROM sezioni_configuratore WHERE id = :id"), {"id": sezione_id}).fetchone()
        if not existing:
            raise HTTPException(status_code=404, detail="Sezione non trovata")
        old_codice = existing[1]
        
        field_map = {"codice": "codice", "etichetta": "etichetta", "descrizione": "descrizione",
                     "icona": "icona", "ordine": "ordine", "attivo": "attivo"}
        fields, params = [], {"id": sezione_id}
        for key, col in field_map.items():
            if key in data:
                val = data[key]
                if key == "attivo":
                    val = 1 if val else 0
                fields.append(f"{col}=:{col}")
                params[col] = val
        
        # Gestisci product_template_ids
        if "product_template_ids" in data:
            pt_ids = data["product_template_ids"]
            pt_ids_json = _json.dumps(pt_ids) if pt_ids else None
            fields.append("product_template_ids=:pt_ids")
            params["pt_ids"] = pt_ids_json
        
        if fields:
            fields.append("updated_at=CURRENT_TIMESTAMP")
            db.execute(text(f"UPDATE sezioni_configuratore SET {','.join(fields)} WHERE id=:id"), params)
            new_codice = data.get("codice")
            if new_codice and new_codice != old_codice:
                # ── CASCADE COMPLETO RINOMINA SEZIONE ──
                
                # 1. campi_configuratore.sezione
                db.execute(text("UPDATE campi_configuratore SET sezione=:new WHERE sezione=:old"),
                    {"new": new_codice, "old": old_codice})
                
                # 2. campi_configuratore.codice (prefisso sezione: "argano.trazione" → "motore.trazione")
                db.execute(text("""
                    UPDATE campi_configuratore 
                    SET codice = :new || SUBSTR(codice, LENGTH(:old) + 1)
                    WHERE codice LIKE :prefix
                """), {"new": new_codice, "old": old_codice, "prefix": old_codice + ".%"})
                
                # 3. valori_configurazione.sezione
                db.execute(text("UPDATE valori_configurazione SET sezione=:new WHERE sezione=:old"),
                    {"new": new_codice, "old": old_codice})
                
                # 4. valori_configurazione.codice_campo (prefisso)
                db.execute(text("""
                    UPDATE valori_configurazione
                    SET codice_campo = :new || SUBSTR(codice_campo, LENGTH(:old) + 1)
                    WHERE codice_campo LIKE :prefix
                """), {"new": new_codice, "old": old_codice, "prefix": old_codice + ".%"})
                
                # 5. template_data JSON keys nei template (rinomina chiave sezione e field_config keys)
                try:
                    templates = db.execute(text("SELECT id, template_data FROM product_templates WHERE template_data IS NOT NULL")).fetchall()
                    for tmpl_id, tmpl_json in templates:
                        if not tmpl_json:
                            continue
                        tmpl_data = _json.loads(tmpl_json)
                        changed = False
                        
                        # Rinomina chiave sezione (es: "argano" → "motore")
                        if old_codice in tmpl_data:
                            tmpl_data[new_codice] = tmpl_data.pop(old_codice)
                            changed = True
                        
                        # Rinomina chiavi in field_config (es: "argano.trazione" → "motore.trazione")
                        if "field_config" in tmpl_data:
                            fc = tmpl_data["field_config"]
                            new_fc = {}
                            for k, v in fc.items():
                                if k.startswith(old_codice + "."):
                                    new_fc[new_codice + k[len(old_codice):]] = v
                                    changed = True
                                else:
                                    new_fc[k] = v
                            if changed:
                                tmpl_data["field_config"] = new_fc
                        
                        if changed:
                            db.execute(text("UPDATE product_templates SET template_data=:data WHERE id=:id"),
                                {"data": _json.dumps(tmpl_data), "id": tmpl_id})
                except Exception as e_tmpl:
                    print(f"Warning: cascade template_data failed: {e_tmpl}")
                
                # 6. Rule JSON files (rinomina riferimenti campo nei file .json)
                try:
                    import glob
                    rules_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)), "rules")
                    if os.path.exists(rules_dir):
                        for filepath in glob.glob(os.path.join(rules_dir, "*.json")):
                            with open(filepath, "r", encoding="utf-8") as rf:
                                rule_text = rf.read()
                            if f'"{old_codice}.' in rule_text:
                                rule_text = rule_text.replace(f'"{old_codice}.', f'"{new_codice}.')
                                with open(filepath, "w", encoding="utf-8") as wf:
                                    wf.write(rule_text)
                except Exception as e_rules:
                    print(f"Warning: cascade rule files failed: {e_rules}")

            db.commit()
        return {"status": "ok"}
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.delete("/sezioni-configuratore/{sezione_id}")
def delete_sezione_configuratore(sezione_id: int, db: Session = Depends(get_db)):
    """Elimina una sezione"""
    _ensure_sezioni_table(db)
    try:
        existing = db.execute(text("SELECT id FROM sezioni_configuratore WHERE id = :id"), {"id": sezione_id}).fetchone()
        if not existing:
            raise HTTPException(status_code=404, detail="Sezione non trovata")
        db.execute(text("DELETE FROM sezioni_configuratore WHERE id=:id"), {"id": sezione_id})
        db.commit()
        return {"status": "ok"}
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/sezioni-configuratore/riordina")
def riordina_sezioni(data: dict, db: Session = Depends(get_db)):
    """Riordina le sezioni"""
    _ensure_sezioni_table(db)
    try:
        for i, sid in enumerate(data.get("ids", [])):
            db.execute(text("UPDATE sezioni_configuratore SET ordine=:o WHERE id=:id"), {"o": i, "id": sid})
        db.commit()
        return {"status": "ok"}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/sezioni-configuratore/{sezione_id}/associa-campi")
def associa_campi_a_sezione(sezione_id: int, data: dict, db: Session = Depends(get_db)):
    """Associa/sposta campi a una sezione"""
    _ensure_sezioni_table(db)
    try:
        sezione = db.execute(text("SELECT codice FROM sezioni_configuratore WHERE id = :id"), {"id": sezione_id}).fetchone()
        if not sezione:
            raise HTTPException(status_code=404, detail="Sezione non trovata")
        codice_sezione = sezione[0]
        for cid in data.get("campo_ids", []):
            db.execute(text("UPDATE campi_configuratore SET sezione=:s WHERE id=:id"), {"s": codice_sezione, "id": cid})
        db.commit()
        return {"status": "ok", "campi_associati": len(data.get("campo_ids", []))}
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/sezioni-configuratore/importa-esistenti")
def importa_sezioni_esistenti(db: Session = Depends(get_db)):
    """Importa sezioni esistenti come stringhe in campi_configuratore"""
    _ensure_sezioni_table(db)
    try:
        result = db.execute(text("SELECT DISTINCT sezione FROM campi_configuratore WHERE sezione IS NOT NULL AND sezione != ''"))
        sezioni_esistenti = [r[0] for r in result.fetchall()]
        importate = 0
        for i, codice in enumerate(sorted(sezioni_esistenti)):
            exists = db.execute(text("SELECT id FROM sezioni_configuratore WHERE codice = :c"), {"c": codice}).fetchone()
            if not exists:
                etichetta = codice.replace("_", " ").title()
                db.execute(text("INSERT INTO sezioni_configuratore (codice, etichetta, ordine, attivo) VALUES (:codice, :etichetta, :ordine, 1)"),
                          {"codice": codice, "etichetta": etichetta, "ordine": i * 10})
                importate += 1
        db.commit()
        return {"status": "ok", "importate": importate, "totale_trovate": len(sezioni_esistenti)}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


# ==========================================
# CAMPI CONFIGURATORE
# ==========================================
@app.get("/campi-configuratore/sezioni")
def get_sezioni_configuratore(db: Session = Depends(get_db)):
    try:
        result = db.execute(text("""
            SELECT sezione,
                   COUNT(*) as totale,
                   SUM(CASE WHEN attivo = 1 THEN 1 ELSE 0 END) as attivi
            FROM campi_configuratore
            GROUP BY sezione
            ORDER BY sezione
        """))
        return [{"codice": r[0], "etichetta": r[0], "totale": r[1], "attivi": r[2]} for r in result.fetchall()]
    except:
        return []

@app.get("/campi-configuratore/schema.json")
def get_schema_campi(db: Session = Depends(get_db)):
    try:
        result = db.execute(text("SELECT codice, etichetta, tipo, sezione FROM campi_configuratore WHERE attivo=1 ORDER BY sezione, ordine"))
        return {"campi": [{"codice": r[0], "label": r[1], "tipo": r[2], "sezione": r[3]} for r in result.fetchall()]}
    except:
        return {"campi": []}

@app.get("/campi-configuratore/{sezione}")
def get_campi_sezione(sezione: str, solo_attivi: bool = True, include_opzioni: bool = False, template_id: int = None, preventivo_id: int = None, db: Session = Depends(get_db)):
    try:
        # Se preventivo_id fornito, ricava template_id automaticamente
        effective_template_id = template_id
        if not effective_template_id and preventivo_id:
            prev_row = db.execute(text("SELECT template_id FROM preventivi WHERE id = :pid"),
                                  {"pid": preventivo_id}).fetchone()
            if prev_row and prev_row[0]:
                effective_template_id = prev_row[0]

        q = """SELECT id, codice, etichetta, tipo, sezione, gruppo_dropdown, ordine, attivo, obbligatorio,
                      unita_misura, valore_min, valore_max, valore_default, descrizione, visibile_form, usabile_regole,
                      product_template_ids
               FROM campi_configuratore WHERE sezione=:s"""
        if solo_attivi:
            q += " AND attivo=1"
        q += " ORDER BY ordine"
        result = db.execute(text(q), {"s": sezione})
        campi = []
        for r in result.fetchall():
            pt_ids_raw = r[16]
            pt_ids = None
            if pt_ids_raw:
                try:
                    pt_ids = json.loads(pt_ids_raw)
                except (json.JSONDecodeError, TypeError):
                    pt_ids = None

            # Filtro per template_id: se campo ha restrizioni e template non è nella lista, salta
            if effective_template_id and pt_ids and effective_template_id not in pt_ids:
                continue

            campi.append({
                "id": r[0], "codice": r[1], "label": r[2], "tipo": r[3], "sezione": r[4],
                "gruppo_opzioni": r[5], "ordine": r[6], "attivo": bool(r[7]), "obbligatorio": bool(r[8]),
                "unita_misura": r[9], "valore_min": r[10], "valore_max": r[11],
                "valore_default": r[12], "descrizione": r[13],
                "visibile_form": bool(r[14]) if r[14] is not None else True,
                "usabile_regole": bool(r[15]) if r[15] is not None else False,
                "product_template_ids": pt_ids,
            })
        return campi
    except:
        return []

@app.post("/campi-configuratore")
def create_campo(data: dict, db: Session = Depends(get_db)):
    try:
        db.execute(text("""INSERT INTO campi_configuratore
            (codice, etichetta, tipo, sezione, gruppo_dropdown, ordine, attivo, obbligatorio,
             unita_misura, valore_min, valore_max, valore_default, descrizione, visibile_form, usabile_regole)
            VALUES (:codice, :etichetta, :tipo, :sezione, :gruppo, :ordine, 1, :obb,
                    :unita, :vmin, :vmax, :vdef, :desc, :vform, :uregole)"""),
            {"codice": data["codice"], "etichetta": data.get("label", data.get("etichetta", "")),
             "tipo": data.get("tipo", "text"), "sezione": data["sezione"],
             "gruppo": data.get("gruppo_opzioni", data.get("gruppo_dropdown")),
             "ordine": data.get("ordine", 0), "obb": 1 if data.get("obbligatorio") else 0,
             "unita": data.get("unita_misura"), "vmin": data.get("valore_min"),
             "vmax": data.get("valore_max"), "vdef": data.get("valore_default"),
             "desc": data.get("descrizione"), "vform": 1 if data.get("visibile_form", True) else 0,
             "uregole": 1 if data.get("usabile_regole") else 0})
        db.commit()
        return {"status": "ok"}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.put("/campi-configuratore/{campo_id}")
def update_campo(campo_id: int, data: dict, db: Session = Depends(get_db)):
    try:
        # Mappa nomi frontend -> nomi DB
        field_map = {
            "codice": "codice", "label": "etichetta", "etichetta": "etichetta",
            "tipo": "tipo", "sezione": "sezione",
            "gruppo_opzioni": "gruppo_dropdown", "gruppo_dropdown": "gruppo_dropdown",
            "ordine": "ordine", "attivo": "attivo", "obbligatorio": "obbligatorio",
            "unita_misura": "unita_misura", "valore_min": "valore_min",
            "valore_max": "valore_max", "valore_default": "valore_default",
            "descrizione": "descrizione", "visibile_form": "visibile_form",
            "usabile_regole": "usabile_regole",
            # Campi preventivo/PDF
            "includi_preventivo": "includi_preventivo",
            "mostra_default_preventivo": "mostra_default_preventivo",
            "mostra_default": "mostra_default_preventivo",
            "sezione_preventivo": "sezione_preventivo",
            "ordine_preventivo": "ordine_preventivo",
            "etichetta_preventivo": "etichetta_preventivo",
        }
        fields, params = [], {"id": campo_id}
        for frontend_key, db_col in field_map.items():
            if frontend_key in data:
                fields.append(f"{db_col}=:{db_col}")
                params[db_col] = data[frontend_key]
        # product_template_ids va serializzato come JSON
        if "product_template_ids" in data:
            pt_ids = data["product_template_ids"]
            fields.append("product_template_ids=:pt_ids")
            params["pt_ids"] = json.dumps(pt_ids) if pt_ids else None
        if fields:
            # Deduplica (label e etichetta mappano entrambi a etichetta)
            unique_fields = list(dict.fromkeys(fields))
            db.execute(text(f"UPDATE campi_configuratore SET {','.join(unique_fields)} WHERE id=:id"), params)
            db.commit()
        return {"status": "ok"}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.delete("/campi-configuratore/{campo_id}")
def delete_campo(campo_id: int, db: Session = Depends(get_db)):
    try:
        db.execute(text("DELETE FROM campi_configuratore WHERE id=:id"), {"id": campo_id})
        db.commit()
        return {"status": "ok"}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

# ============================================================================
# SEZIONE B: GET/PUT valori_configurazione MODIFICATI
#            (SOSTITUISCONO gli endpoint esistenti ~linea 2619 e ~linea 2637)
# ============================================================================

@app.get("/preventivi/{preventivo_id}/configurazione/{sezione}")
def get_valori_sezione(preventivo_id: int, sezione: str, db: Session = Depends(get_db)):
    """
    Legge i valori di una sezione per un preventivo.
    AUTO-POPULATE: popola i default MANCANTI da campi_configuratore,
    anche se alcuni valori esistono già (es. da template).
    """
    try:
        # 1. Leggi valori esistenti
        result = db.execute(text("""
            SELECT codice_campo, valore, is_default, COALESCE(is_readonly, 0)
            FROM valori_configurazione
            WHERE preventivo_id = :pid AND sezione = :sez
        """), {"pid": preventivo_id, "sez": sezione})
        rows = result.fetchall()
        
        valori = {r[0]: r[1] for r in rows}
        defaults_info = {r[0]: bool(r[2]) for r in rows}
        readonly_info = {r[0]: bool(r[3]) for r in rows}
        
        # 2. Popola default MANCANTI da campi_configuratore (filtrati per prodotto)
        # Recupera template_id del preventivo per filtrare campi
        prev_row = db.execute(text("SELECT template_id FROM preventivi WHERE id = :pid"),
                              {"pid": preventivo_id}).fetchone()
        prev_template_id = prev_row[0] if prev_row and prev_row[0] else None

        campi = db.execute(text("""
            SELECT codice, valore_default, tipo, product_template_ids
            FROM campi_configuratore
            WHERE sezione = :sez AND attivo = 1
            ORDER BY ordine
        """), {"sez": sezione}).fetchall()
        
        # Filtra campi per prodotto
        campi_filtrati = []
        for c in campi:
            pt_ids_raw = c[3]
            if pt_ids_raw and prev_template_id:
                try:
                    pt_ids = json.loads(pt_ids_raw)
                    if prev_template_id not in pt_ids:
                        continue
                except (json.JSONDecodeError, TypeError):
                    pass
            campi_filtrati.append(c)

        nuovi = 0
        for codice, valore_default, tipo, _ in campi_filtrati:
            if codice not in valori and valore_default is not None and str(valore_default).strip() != "":
                val = _resolve_default_expr(str(valore_default), preventivo_id, db, context=valori)
                db.execute(text("""
                    INSERT INTO valori_configurazione
                        (preventivo_id, sezione, codice_campo, valore, is_default)
                    VALUES (:pid, :sez, :campo, :val, 1)
                """), {"pid": preventivo_id, "sez": sezione, "campo": codice, "val": val})
                valori[codice] = val
                defaults_info[codice] = True
                readonly_info[codice] = False
                nuovi += 1
        
        if nuovi:
            db.commit()

            # Applica field_config dal template se esiste
            prev = db.execute(text(
                "SELECT template_id FROM preventivi WHERE id = :pid"
            ), {"pid": preventivo_id}).fetchone()
            if prev and prev[0]:
                _apply_template_field_config(db, preventivo_id, prev[0])
                db.commit()
                # Rileggi readonly aggiornati
                for codice in list(readonly_info.keys()):
                    row = db.execute(text("""
                        SELECT COALESCE(is_readonly, 0) FROM valori_configurazione
                        WHERE preventivo_id = :pid AND codice_campo = :campo
                    """), {"pid": preventivo_id, "campo": codice}).fetchone()
                    if row:
                        readonly_info[codice] = bool(row[0])

        return {
            "preventivo_id": preventivo_id,
            "sezione": sezione,
            "valori": valori,
            "is_default": defaults_info,
            "is_readonly": readonly_info,
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.put("/preventivi/{preventivo_id}/configurazione/{sezione}")
def save_valori_sezione(preventivo_id: int, sezione: str, data: dict, db: Session = Depends(get_db)):
    """
    Salva valori sezione (upsert) con tracking is_default.
    I campi con is_readonly=1 vengono SALTATI (non modificabili).
    """
    try:
        valori = data.get("valori", {})

        # Carica valori default dalla definizione campi (per confronto)
        defaults_result = db.execute(text("""
            SELECT codice, valore_default
            FROM campi_configuratore
            WHERE sezione = :sez AND attivo = 1
        """), {"sez": sezione})
        campo_defaults = {r[0]: r[1] for r in defaults_result.fetchall()}

        # Carica stato attuale (valore, is_default, is_readonly)
        existing_result = db.execute(text("""
            SELECT codice_campo, valore, is_default, COALESCE(is_readonly, 0)
            FROM valori_configurazione
            WHERE preventivo_id = :pid AND sezione = :sez
        """), {"pid": preventivo_id, "sez": sezione})
        existing = {r[0]: {"valore": r[1], "is_default": r[2], "is_readonly": r[3]}
                    for r in existing_result.fetchall()}

        skipped_readonly = []

        for codice_campo, valore in valori.items():
            # Salta campi readonly
            if codice_campo in existing and existing[codice_campo]["is_readonly"]:
                skipped_readonly.append(codice_campo)
                continue

            valore_str = str(valore) if valore is not None else None

            if codice_campo in existing:
                old = existing[codice_campo]
                if old["is_default"] == 0:
                    new_is_default = 0
                elif valore_str != old["valore"]:
                    new_is_default = 0
                else:
                    new_is_default = 1

                db.execute(text("""
                    UPDATE valori_configurazione
                    SET valore = :val, is_default = :isd, updated_at = CURRENT_TIMESTAMP
                    WHERE preventivo_id = :pid AND sezione = :sez AND codice_campo = :campo
                """), {"val": valore_str, "isd": new_is_default,
                       "pid": preventivo_id, "sez": sezione, "campo": codice_campo})
            else:
                default_val = campo_defaults.get(codice_campo)
                is_def = 1 if (default_val is not None and valore_str == str(default_val)) else 0

                db.execute(text("""
                    INSERT INTO valori_configurazione
                        (preventivo_id, sezione, codice_campo, valore, is_default)
                    VALUES (:pid, :sez, :campo, :val, :isd)
                """), {"pid": preventivo_id, "sez": sezione, "campo": codice_campo,
                       "val": valore_str, "isd": is_def})

        db.commit()

        # Trigger rule engine dopo salvataggio
        result = safe_evaluate_rules(preventivo_id, db)
        touch_preventivo(preventivo_id, db)

        # Rileggi is_default aggiornati per feedback al frontend
        updated_defaults = db.execute(text("""
            SELECT codice_campo, COALESCE(is_default, 1)
            FROM valori_configurazione
            WHERE preventivo_id = :pid AND sezione = :sez
        """), {"pid": preventivo_id, "sez": sezione})
        is_default_map = {r[0]: bool(r[1]) for r in updated_defaults.fetchall()}

        response = {
            "status": "ok",
            "sezione": sezione,
            "campi_salvati": len(valori) - len(skipped_readonly),
            "rules_result": result,
            "is_default": is_default_map,
        }
        if skipped_readonly:
            response["campi_readonly_saltati"] = skipped_readonly

        return response
    except Exception as e:
        db.rollback()
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/preventivi/{preventivo_id}/configurazione")
def get_tutti_valori_configurazione(preventivo_id: int, db: Session = Depends(get_db)):
    """Leggi tutti i valori di tutte le sezioni dinamiche per un preventivo"""
    try:
        result = db.execute(
            text("""
                SELECT sezione, codice_campo, valore 
                FROM valori_configurazione 
                WHERE preventivo_id = :pid
                ORDER BY sezione, codice_campo
            """),
            {"pid": preventivo_id}
        )
        
        sezioni = {}
        for row in result.fetchall():
            sezione, campo, valore = row[0], row[1], row[2]
            if sezione not in sezioni:
                sezioni[sezione] = {}
            sezioni[sezione][campo] = valore
        
        return {"preventivo_id": preventivo_id, "sezioni": sezioni}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.delete("/preventivi/{preventivo_id}/configurazione/{sezione}")
def delete_valori_sezione(preventivo_id: int, sezione: str, db: Session = Depends(get_db)):
    """Elimina tutti i valori di una sezione per un preventivo"""
    try:
        db.execute(
            text("DELETE FROM valori_configurazione WHERE preventivo_id = :pid AND sezione = :sez"),
            {"pid": preventivo_id, "sez": sezione}
        )
        db.commit()
        return {"status": "ok", "sezione": sezione}
    except Exception as e:
        db.rollback()
        raise HTTPException(status_code=500, detail=str(e))

# ============================================================================
# SEZIONE C: DATI PREVENTIVO STRUTTURATI PER PDF
# ============================================================================
# SEZIONE C: DATI DOCUMENTO PREVENTIVO
# ============================================================================

@app.get("/preventivi/{preventivo_id}/dati-documento")
def get_dati_documento_preventivo(preventivo_id: int, db: Session = Depends(get_db)):
    """
    Restituisce tutti i dati del preventivo strutturati per sezione del documento.
    Usato dal generatore PDF/DOCX per assemblare il documento.
    
    Logica flag (per-preventivo > globale):
    - Se valori_configurazione ha includi_preventivo/mostra_default_preventivo -> usa quelli
    - Altrimenti fallback su campi_configuratore (globale)
    """
    try:
        # 1. Carica sezioni preventivo (visibili, ordinate)
        sez_result = db.execute(text("""
            SELECT codice, titolo, ordine, tipo, mostra_titolo, nota_default, stile
            FROM sezioni_preventivo
            WHERE visibile = 1
            ORDER BY ordine
        """))
        sezioni_doc = [
            {"codice": r[0], "titolo": r[1], "ordine": r[2], "tipo": r[3],
             "mostra_titolo": bool(r[4]), "nota": r[5], "stile": r[6]}
            for r in sez_result.fetchall()
        ]

        # 2. Carica tutti i campi con mapping preventivo (globale)
        campi_result = db.execute(text("""
            SELECT codice, etichetta, tipo, sezione, unita_misura, valore_default,
                   COALESCE(includi_preventivo, 0) as includi_glob,
                   sezione_preventivo,
                   COALESCE(ordine_preventivo, 0) as ordine_prev,
                   etichetta_preventivo,
                   COALESCE(mostra_default_preventivo, 0) as mostra_def_glob
            FROM campi_configuratore
            WHERE attivo = 1
            ORDER BY sezione_preventivo, ordine_preventivo
        """))
        campi_global = {}
        for r in campi_result.fetchall():
            campi_global[r[0]] = {
                "codice": r[0], "etichetta": r[1], "tipo_campo": r[2],
                "sezione_config": r[3], "unita_misura": r[4],
                "valore_default": r[5],
                "includi_preventivo_glob": bool(r[6]),
                "sezione_preventivo": r[7],
                "ordine_preventivo": r[8],
                "etichetta_preventivo": r[9],
                "mostra_default_glob": bool(r[10]),
            }

        # 3. Carica tutti i valori per questo preventivo (con flag per-preventivo)
        val_result = db.execute(text("""
            SELECT sezione, codice_campo, valore,
                   COALESCE(is_default, 1) as is_default,
                   includi_preventivo,
                   mostra_default_preventivo
            FROM valori_configurazione
            WHERE preventivo_id = :pid
        """), {"pid": preventivo_id})
        valori_all = {}
        for sez, campo, valore, is_def, incl_prev, mostra_def in val_result.fetchall():
            valori_all[campo] = {
                "valore": valore,
                "is_default": bool(is_def),
                "sezione_config": sez,
                "includi_preventivo_prev": incl_prev,
                "mostra_default_prev": mostra_def,
            }

        # 4. Assembla il documento
        sezioni_output = []
        valori_standard = []

        for sez_doc in sezioni_doc:
            cod_sez = sez_doc["codice"]

            if cod_sez == "valori_standard":
                continue

            if cod_sez == "materiali":
                sezioni_output.append({**sez_doc, "campi": [], "_tipo_speciale": "materiali"})
                continue

            # Trova campi assegnati a questa sezione preventivo
            campi_sez = [
                v for v in campi_global.values()
                if v["sezione_preventivo"] == cod_sez
            ]
            campi_sez.sort(key=lambda x: x["ordine_preventivo"])

            campi_output = []
            for campo_def in campi_sez:
                codice = campo_def["codice"]
                val_info = valori_all.get(codice, {})
                valore = val_info.get("valore")
                is_def = val_info.get("is_default", True)

                if valore is None and campo_def["valore_default"]:
                    valore = campo_def["valore_default"]
                    is_def = True

                if valore is None or str(valore).strip() == "":
                    continue

                # Determina includi_preventivo: per-preventivo > globale
                incl_prev = val_info.get("includi_preventivo_prev")
                if incl_prev is not None:
                    includi = bool(incl_prev)
                else:
                    includi = campo_def["includi_preventivo_glob"]

                if not includi:
                    continue

                # Determina mostra_default: per-preventivo > globale
                mostra_def_prev = val_info.get("mostra_default_prev")
                if mostra_def_prev is not None:
                    mostra_default = bool(mostra_def_prev)
                else:
                    mostra_default = campo_def["mostra_default_glob"]

                campo_out = {
                    "codice": codice,
                    "etichetta": campo_def["etichetta_preventivo"] or campo_def["etichetta"],
                    "valore": valore,
                    "unita_misura": campo_def["unita_misura"],
                    "is_default": is_def,
                    "tipo_campo": campo_def["tipo_campo"],
                }

                if is_def and mostra_default:
                    valori_standard.append({
                        **campo_out,
                        "sezione_config": campo_def["sezione_config"],
                    })

                campi_output.append(campo_out)

            if campi_output:
                sezioni_output.append({**sez_doc, "campi": campi_output})

        # 5. Aggiungi sezione "valori_standard" se ci sono
        sez_standard = next((s for s in sezioni_doc if s["codice"] == "valori_standard"), None)
        if sez_standard and valori_standard:
            sezioni_output.append({**sez_standard, "campi": valori_standard})

        return {
            "preventivo_id": preventivo_id,
            "sezioni": sezioni_output,
            "valori_standard": valori_standard,
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


# ============================================================================
# SEZIONE D: INIZIALIZZA DEFAULTS PER UN PREVENTIVO
# ============================================================================

@app.post("/preventivi/{preventivo_id}/inizializza-defaults")
def inizializza_defaults_preventivo(preventivo_id: int, db: Session = Depends(get_db)):
    """
    Popola TUTTE le sezioni con i valori di default da campi_configuratore.
    Poi applica field_config dal template se associato.
    """
    try:
        campi = db.execute(text("""
            SELECT codice, sezione, valore_default, tipo
            FROM campi_configuratore
            WHERE attivo = 1 AND valore_default IS NOT NULL AND valore_default != ''
        """)).fetchall()

        existing = db.execute(text("""
            SELECT codice_campo FROM valori_configurazione WHERE preventivo_id = :pid
        """), {"pid": preventivo_id}).fetchall()
        existing_set = {r[0] for r in existing}

        count = 0
        sezioni_touched = set()

        for codice, sezione, valore_default, tipo in campi:
            if codice in existing_set:
                continue
            db.execute(text("""
                INSERT INTO valori_configurazione
                    (preventivo_id, sezione, codice_campo, valore, is_default)
                VALUES (:pid, :sez, :campo, :val, 1)
            """), {"pid": preventivo_id, "sez": sezione,
                   "campo": codice, "val": str(valore_default)})
            count += 1
            sezioni_touched.add(sezione)

        if count:
            db.commit()

        # Applica field_config dal template
        field_config_count = 0
        prev = db.execute(text(
            "SELECT template_id FROM preventivi WHERE id = :pid"
        ), {"pid": preventivo_id}).fetchone()
        if prev and prev[0]:
            field_config_count = _apply_template_field_config(db, preventivo_id, prev[0])
            if field_config_count:
                db.commit()

        return {
            "status": "ok",
            "campi_inizializzati": count,
            "sezioni": sorted(sezioni_touched),
            "field_config_applicati": field_config_count,
        }
    except Exception as e:
        db.rollback()
        raise HTTPException(status_code=500, detail=str(e))

# ==========================================
# RULE ENGINE - DEBUG CONTEXT
# ==========================================

@app.get("/preventivi/{preventivo_id}/rule-context")
def get_rule_context(preventivo_id: int, db: Session = Depends(get_db)):
    """Debug: mostra il contesto completo che il rule engine usa per valutare."""
    preventivo = db.query(Preventivo).filter(Preventivo.id == preventivo_id).first()
    if not preventivo:
        raise HTTPException(status_code=404, detail="Preventivo non trovato")
    
    context = build_config_context(preventivo_id, db)
    rules = load_rules()
    
    # Mostra anche quali regole matcherebbero
    matching = []
    for rule in rules:
        if not rule.get("enabled", True):
            continue
        conditions = rule.get("conditions", [])
        all_met = all(evaluate_condition(c, context) for c in conditions) if conditions else False
        matching.append({
            "id": rule.get("id"),
            "name": rule.get("name"),
            "match": all_met,
            "conditions_detail": [
                {
                    "field": c.get("field"),
                    "operator": c.get("operator"),
                    "expected": c.get("value"),
                    "actual": context.get(c.get("field")),
                    "result": evaluate_condition(c, context)
                }
                for c in conditions
            ]
        })
    
    return {
        "preventivo_id": preventivo_id,
        "context": context,
        "context_keys_count": len(context),
        "rules_total": len(rules),
        "rules_matching": sum(1 for r in matching if r["match"]),
        "rules_detail": matching
    }


# ============================================================
# ORDINI E BOM - ENDPOINT
# ============================================================

@app.post("/preventivi/{preventivo_id}/conferma")
def api_conferma_preventivo(preventivo_id: int, body: dict = None, db: Session = Depends(get_db)):
    """
    Conferma preventivo e genera/aggiorna ordine.
    
    Body opzionale:
    - action: "new" (crea nuovo ordine, default), "update" (aggiorna ordine esistente), "check" (solo controlla)
    - ordine_id: ID ordine da aggiornare (se action=update)
    """
    action = (body or {}).get("action", "new")
    target_ordine_id = (body or {}).get("ordine_id")

    preventivo = db.query(Preventivo).filter(Preventivo.id == preventivo_id).first()
    if not preventivo:
        raise HTTPException(status_code=404, detail="Preventivo non trovato")

    try:
        conn = db.get_bind().raw_connection()
        cursor = conn.cursor()

        # Ensure ordini table has revision columns
        _ensure_ordini_revisione_columns(db)

        # Check existing orders
        try:
            cursor.execute(
                "SELECT id, numero_ordine, stato, numero_revisione_origine, bom_esplosa, bom_numero_revisione "
                "FROM ordini WHERE preventivo_id = ? ORDER BY created_at DESC",
                (preventivo_id,)
            )
            existing_ordini = []
            for r in cursor.fetchall():
                existing_ordini.append({
                    "id": r[0], "numero_ordine": r[1], "stato": r[2],
                    "numero_revisione_origine": r[3], "bom_esplosa": bool(r[4]),
                    "bom_numero_revisione": r[5]
                })
        except Exception:
            cursor.execute(
                "SELECT id, numero_ordine, stato, bom_esplosa FROM ordini WHERE preventivo_id = ? ORDER BY created_at DESC",
                (preventivo_id,)
            )
            existing_ordini = [{"id": r[0], "numero_ordine": r[1], "stato": r[2], "bom_esplosa": bool(r[3])} for r in cursor.fetchall()]

        # Se action=check, restituisci solo info sugli ordini esistenti
        if action == "check":
            rev_corrente = getattr(preventivo, 'revisione_corrente', 0) or 0
            return {
                "existing_ordini": existing_ordini,
                "revisione_corrente": rev_corrente,
                "preventivo_id": preventivo_id,
            }

        # Auto-snapshot prima della conferma
        try:
            snap_result = _crea_snapshot_preventivo(preventivo_id, db, motivo="Auto-snapshot pre-conferma")
            rev_id = snap_result["id"] if snap_result else None
            num_rev = snap_result["numero_revisione"] if snap_result else 0
        except Exception as snap_err:
            print(f"[WARN] Snapshot pre-conferma fallito: {snap_err}")
            rev_id = None
            num_rev = getattr(preventivo, 'revisione_corrente', 0) or 0

        # Lead time
        try:
            lead_time = calcola_lead_time(preventivo_id, db)
        except Exception:
            lead_time = 15

        # Calcola totali
        materiali = db.query(Materiale).filter(Materiale.preventivo_id == preventivo_id).all()
        totale = _prev_totale(preventivo) or sum(m.prezzo_totale or 0 for m in materiali)
        config_snap = json.dumps({
            "preventivo": _prev_numero(preventivo),
            "revisione": num_rev,
            "materiali": [{"codice": m.codice, "descrizione": m.descrizione,
                           "quantita": m.quantita, "prezzo_totale": m.prezzo_totale} for m in materiali]
        }, ensure_ascii=False)

        if action == "update" and target_ordine_id:
            # AGGIORNA ordine esistente
            cursor.execute(
                "UPDATE ordini SET stato='confermato', tipo_impianto=?, configurazione_json=?, "
                "totale_materiali=?, totale_netto=?, lead_time_giorni=?, "
                "data_consegna_prevista=?, revisione_id=?, numero_revisione_origine=?, "
                "updated_at=datetime('now') WHERE id=?",
                (_prev_tipo(preventivo), config_snap, totale,
                 _prev_netto(preventivo) or totale, lead_time,
                 (datetime.now() + timedelta(days=lead_time)).isoformat(),
                 rev_id, num_rev, target_ordine_id)
            )
            conn.commit()

            # Aggiorna preventivo
            _prev_set(preventivo, "confermato", 'stato', 'status')
            db.commit()

            cursor.execute("SELECT numero_ordine FROM ordini WHERE id=?", (target_ordine_id,))
            numero_ordine = cursor.fetchone()[0]

            return {
                "status": "aggiornato", "preventivo_id": preventivo_id,
                "ordine_id": target_ordine_id, "numero_ordine": numero_ordine,
                "revisione_origine": num_rev,
                "lead_time_giorni": lead_time,
                "data_consegna_prevista": (datetime.now() + timedelta(days=lead_time)).isoformat(),
                "totale": totale
            }

        # action == "new" — CREA nuovo ordine
        if _prev_stato(preventivo) not in ("draft", "bozza", "inviato", "confermato", None):
            raise HTTPException(status_code=400, detail=f"Stato attuale: {_prev_stato(preventivo)}")

        _prev_set(preventivo, "confermato", 'stato', 'status')
        db.commit()

        anno = datetime.now().year
        
        # Auto-crea/aggiorna tabella ordini
        
        cursor.execute("""
            CREATE TABLE IF NOT EXISTS ordini (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                numero_ordine TEXT UNIQUE NOT NULL,
                preventivo_id INTEGER,
                cliente_id INTEGER DEFAULT NULL,
                stato TEXT DEFAULT 'confermato',
                tipo_impianto TEXT,
                configurazione_json TEXT,
                totale_materiali REAL DEFAULT 0,
                totale_netto REAL DEFAULT 0,
                lead_time_giorni INTEGER DEFAULT 15,
                data_consegna_prevista TEXT,
                bom_esplosa INTEGER DEFAULT 0,
                data_esplosione_bom TEXT,
                revisione_id INTEGER,
                numero_revisione_origine INTEGER,
                bom_revisione_id INTEGER,
                bom_numero_revisione INTEGER,
                bom_esplosa_at TEXT,
                created_by TEXT,
                created_at TEXT DEFAULT (datetime('now')),
                updated_at TEXT DEFAULT (datetime('now'))
            )
        """)
        conn.commit()
        _ensure_ordini_revisione_columns(db)

        cursor.execute("SELECT MAX(CAST(SUBSTR(numero_ordine, 10) AS INTEGER)) FROM ordini WHERE numero_ordine LIKE ?", (f"ORD-{anno}-%",))
        last = cursor.fetchone()[0] or 0
        numero_ordine = f"ORD-{anno}-{last + 1:04d}"

        cursor.execute(
            "INSERT INTO ordini (numero_ordine, preventivo_id, cliente_id, stato, tipo_impianto, "
            "configurazione_json, totale_materiali, totale_netto, lead_time_giorni, "
            "data_consegna_prevista, revisione_id, numero_revisione_origine, created_by) "
            "VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?)",
            (numero_ordine, preventivo_id, getattr(preventivo, 'cliente_id', 0) or 0, "confermato",
             _prev_tipo(preventivo), config_snap, totale,
             _prev_netto(preventivo) or totale, lead_time,
             (datetime.now() + timedelta(days=lead_time)).isoformat(),
             rev_id, num_rev, "admin")
        )
        ordine_id = cursor.lastrowid
        conn.commit()

        # Salva ordine_id nel preventivo
        try:
            cursor.execute(
                "UPDATE preventivi SET ordine_id = ?, data_conferma = datetime('now'), "
                "lead_time_giorni = ? WHERE id = ?",
                (ordine_id, lead_time, preventivo_id)
            )
            conn.commit()
        except Exception as e:
            print(f"[WARN] ordine_id backlink failed: {e}")

        # Registra nel storico stati
        try:
            cursor.execute("""
                CREATE TABLE IF NOT EXISTS ordini_storico_stato (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    ordine_id INTEGER NOT NULL,
                    stato_precedente TEXT,
                    stato_nuovo TEXT NOT NULL,
                    motivo TEXT,
                    utente TEXT,
                    created_at TEXT DEFAULT (datetime('now')),
                    FOREIGN KEY (ordine_id) REFERENCES ordini(id)
                )
            """)
            cursor.execute(
                "INSERT INTO ordini_storico_stato (ordine_id, stato_precedente, stato_nuovo, motivo, utente) "
                "VALUES (?, NULL, 'confermato', 'Conferma preventivo', ?)",
                (ordine_id, 'admin')
            )
            conn.commit()
        except Exception as e:
            print(f"[WARN] Storico stato ordine: {e}")

        return {
            "status": "confermato", "preventivo_id": preventivo_id,
            "ordine_id": ordine_id, "numero_ordine": numero_ordine,
            "revisione_origine": num_rev,
            "lead_time_giorni": lead_time,
            "data_consegna_prevista": (datetime.now() + timedelta(days=lead_time)).isoformat(),
            "totale": totale
        }
    except HTTPException:
        raise
    except Exception as e:
        import traceback
        traceback.print_exc()
        db.rollback()
        raise HTTPException(status_code=500, detail=f"Errore conferma: {str(e)}")

@app.get("/ordini/search")
def search_ordini(
    q: str = None,
    stato: str = None,
    cliente_id: int = None,
    data_da: str = None,
    data_a: str = None,
    limit: int = 100,
    db: Session = Depends(get_db)
):
    print(f">>> SEARCH ORDINI q={q}")
    try:
        conn = db.get_bind().raw_connection()
        cursor = conn.cursor()

        sql = """
            SELECT o.*,
                   c.ragione_sociale AS cliente_ragione_sociale,
                   c.codice AS cliente_codice,
                   (SELECT COUNT(*) FROM materiali m WHERE m.preventivo_id = o.preventivo_id) AS n_materiali
            FROM ordini o
            LEFT JOIN clienti c ON c.id = o.cliente_id
            WHERE 1=1
        """
        params = []

        if q and q.strip():
            pattern = f"%{q.strip()}%"
            sql += " AND (o.numero_ordine LIKE ? OR c.ragione_sociale LIKE ? OR o.tipo_impianto LIKE ?)"
            params.extend([pattern, pattern, pattern])

        if stato:
            sql += " AND o.stato = ?"
            params.append(stato)

        if cliente_id:
            sql += " AND o.cliente_id = ?"
            params.append(cliente_id)

        if data_da:
            sql += " AND o.created_at >= ?"
            params.append(data_da)

        if data_a:
            sql += " AND o.created_at <= ?"
            params.append(data_a + "T23:59:59")

        sql += " ORDER BY o.created_at DESC LIMIT ?"
        params.append(limit)

        cursor.execute(sql, params)
        cols = [d[0] for d in cursor.description]
        rows = cursor.fetchall()

        return [dict(zip(cols, r)) for r in rows]
    except Exception as e:
        import traceback
        traceback.print_exc()
        raise HTTPException(500, str(e))


@app.get("/ordini")
def api_get_ordini(db: Session = Depends(get_db)):
    conn = db.get_bind().raw_connection()
    cursor = conn.cursor()
    cursor.execute(
        "SELECT o.id, o.numero_ordine, o.preventivo_id, o.stato, o.tipo_impianto, "
        "o.totale_materiali, o.lead_time_giorni, o.data_consegna_prevista, "
        "o.bom_esplosa, o.created_at, c.ragione_sociale "
        "FROM ordini o LEFT JOIN clienti c ON c.id = o.cliente_id "
        "ORDER BY o.created_at DESC"
    )
    return [{"id": r[0], "numero_ordine": r[1], "preventivo_id": r[2], "stato": r[3],
             "tipo_impianto": r[4], "totale_materiali": r[5], "lead_time_giorni": r[6],
             "data_consegna_prevista": r[7], "bom_esplosa": bool(r[8]),
             "created_at": r[9], "cliente": r[10]} for r in cursor.fetchall()]


@app.get("/ordini/{ordine_id}")
def api_get_ordine(ordine_id: int, db: Session = Depends(get_db)):
    conn = db.get_bind().raw_connection()
    cursor = conn.cursor()
    cursor.execute("SELECT * FROM ordini WHERE id = ?", (ordine_id,))
    r = cursor.fetchone()
    if not r:
        raise HTTPException(status_code=404, detail="Ordine non trovato")
    cols = [desc[0] for desc in cursor.description]
    return dict(zip(cols, r))

# --- START BLOCCO A ---

@app.put("/ordini/{ordine_id}")
def api_update_ordine(ordine_id: int, body: dict, db: Session = Depends(get_db)):
    """Aggiorna campi dell'ordine (note, condizioni, riferimento, lead_time, data_consegna)"""
    conn = db.get_bind().raw_connection()
    cursor = conn.cursor()

    cursor.execute("SELECT id FROM ordini WHERE id = ?", (ordine_id,))
    if not cursor.fetchone():
        raise HTTPException(404, "Ordine non trovato")

    # Campi aggiornabili
    campi_ammessi = {
        'note', 'note_interne', 'condizioni_pagamento', 'condizioni_consegna',
        'riferimento_cliente', 'lead_time_giorni', 'data_consegna_prevista',
    }

    updates = []
    values = []
    for campo, valore in body.items():
        if campo in campi_ammessi:
            updates.append(f"{campo} = ?")
            values.append(valore)

    if not updates:
        raise HTTPException(400, "Nessun campo valido da aggiornare")

    updates.append("updated_at = datetime('now')")
    values.append(ordine_id)

    cursor.execute(f"UPDATE ordini SET {', '.join(updates)} WHERE id = ?", values)
    conn.commit()

    # Ritorna ordine aggiornato
    cursor.execute("SELECT * FROM ordini WHERE id = ?", (ordine_id,))
    cols = [d[0] for d in cursor.description]
    return dict(zip(cols, cursor.fetchone()))


@app.get("/ordini/{ordine_id}/conferma-ordine/{formato}")
def api_conferma_ordine_doc(ordine_id: int, formato: str, request: Request, db: Session = Depends(get_db)):
    """
    Genera il documento "Conferma d'Ordine" basandosi sullo stesso template
    usato per l'export del preventivo, con sezioni aggiuntive per l'ordine.
    """
    from docx import Document as DocxDocument
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn

    conn = db.get_bind().raw_connection()
    cursor = conn.cursor()

    # ── 1. Carica ordine ──
    cursor.execute("SELECT * FROM ordini WHERE id = ?", (ordine_id,))
    r = cursor.fetchone()
    if not r:
        raise HTTPException(404, "Ordine non trovato")
    ord_cols = [d[0] for d in cursor.description]
    ordine_data = dict(zip(ord_cols, r))

    preventivo_id = ordine_data.get("preventivo_id")
    if not preventivo_id:
        raise HTTPException(400, "Ordine senza preventivo associato")

    # ── 2. Carica dati preventivo (stessa logica dell'export preventivo) ──
    preventivo = db.query(Preventivo).filter(Preventivo.id == preventivo_id).first()
    if not preventivo:
        raise HTTPException(404, "Preventivo non trovato")

    dati_commessa = db.query(DatiCommessa).filter(DatiCommessa.preventivo_id == preventivo_id).first()
    dati_principali = db.query(DatiPrincipali).filter(DatiPrincipali.preventivo_id == preventivo_id).first()
    normative_data = db.query(Normative).filter(Normative.preventivo_id == preventivo_id).first()
    materiali = db.query(Materiale).filter(Materiale.preventivo_id == preventivo_id).all()

    argano_data = None
    try:
        argano_data = db.query(Argano).filter(Argano.preventivo_id == preventivo_id).first()
    except Exception:
        pass

    cliente = None
    try:
        cid = getattr(preventivo, 'cliente_id', None)
        if cid:
            cliente = db.query(Cliente).filter(Cliente.id == cid).first()
    except Exception:
        pass

    # Icona prodotto
    product_icon_png = None
    product_name = ""
    product_category = ""
    try:
        tid = getattr(preventivo, 'template_id', None)
        if tid:
            prod_tpl = db.query(ProductTemplate).filter(ProductTemplate.id == tid).first()
            if prod_tpl:
                product_name = getattr(prod_tpl, 'nome_display', '') or ''
                product_category = getattr(prod_tpl, 'categoria', '') or ''
                icon_name = getattr(prod_tpl, 'icona', None)
                if icon_name:
                    from export_utils import load_product_icon_png
                    product_icon_png = load_product_icon_png(icon_name, size=160)
    except Exception as e:
        print(f"WARN: Icona prodotto conferma ordine: {e}")

    # ── 3. Genera DOCX base con il template preventivo ──
    doc_template = None
    doc_template_id = request.query_params.get("template_id")
    if doc_template_id:
        doc_template = db.query(DocumentTemplate).filter(
            DocumentTemplate.id == int(doc_template_id)
        ).first()
    else:
        doc_template = db.query(DocumentTemplate).filter(
            DocumentTemplate.tipo == "preventivo",
            DocumentTemplate.is_default == True,
            DocumentTemplate.attivo == True
        ).first()

    buf_base = None

    if doc_template and doc_template.config:
        try:
            available = get_available_fields_from_db(db)
            valori_din = load_valori_dinamici(db, preventivo_id)
            def_info = load_defaults_info(db, preventivo_id)
            buf_base = genera_docx_da_template(
                template_config=doc_template.config,
                preventivo=preventivo,
                dati_commessa=dati_commessa,
                dati_principali=dati_principali,
                normative=normative_data,
                argano=argano_data,
                materiali=materiali,
                cliente=cliente,
                logo_data=doc_template.logo_data,
                logo_mime=doc_template.logo_mime,
                valori_dinamici=valori_din,
                available_fields=available,
                defaults_info=def_info,
                product_icon_png=product_icon_png,
                product_name=product_name,
                product_category=product_category,
            )
        except Exception as e:
            print(f"WARN: Template conferma ordine fallito, uso fallback: {e}")

    if not buf_base:
        # Fallback: genera con v2
        try:
            dati_doc = get_dati_documento_preventivo(preventivo_id, db)
            preventivo_info = {
                "numero": _prev_numero(preventivo),
                "customer": getattr(preventivo, 'customer_name', '') or '',
                "status": _prev_stato(preventivo),
                "totale": safe_float(_prev_totale(preventivo)),
                "sconto": safe_float(getattr(preventivo, 'sconto_cliente', 0)) + safe_float(getattr(preventivo, 'sconto_extra_admin', 0)),
                "netto": safe_float(_prev_netto(preventivo)),
                "note": getattr(preventivo, 'note', '') or '',
                "revisione": getattr(preventivo, 'revisione_corrente', 0) or 0,
            }
            if cliente:
                preventivo_info["customer"] = getattr(cliente, 'ragione_sociale', preventivo_info["customer"]) or preventivo_info["customer"]
            from export_utils import genera_docx_preventivo_v2
            buf_base = genera_docx_preventivo_v2(preventivo_info, dati_doc, materiali,
                                                  product_icon_png=product_icon_png,
                                                  product_name=product_name,
                                                  product_category=product_category)
        except Exception as e:
            raise HTTPException(500, f"Errore generazione documento base: {e}")

    # ── 4. Apri il DOCX generato e appendi sezioni ordine ──
    buf_base.seek(0)
    doc = DocxDocument(buf_base)

    # Aggiungi page break prima delle sezioni ordine
    from docx.oxml.ns import qn as _qn
    doc.add_page_break()

    # --- TITOLO CONFERMA D'ORDINE ---
    numero_ordine = ordine_data.get('numero_ordine', '')
    h = doc.add_heading('CONFERMA D\'ORDINE', level=1)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)
        run.font.size = Pt(16)

    p_num = doc.add_paragraph()
    p_num.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_num = p_num.add_run(f"N° {numero_ordine}")
    r_num.font.size = Pt(14)
    r_num.font.bold = True

    p_data = doc.add_paragraph()
    p_data.alignment = WD_ALIGN_PARAGRAPH.CENTER
    from datetime import datetime as _dt
    r_data = p_data.add_run(f"Data: {_dt.now().strftime('%d/%m/%Y')}")
    r_data.font.size = Pt(10)
    r_data.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_paragraph()

    # --- DATI ORDINE ---
    doc.add_heading('Dati Ordine', level=2)

    t_ord = doc.add_table(rows=5, cols=4)
    t_ord.style = 'Table Grid'

    from export_utils import fmt_euro, safe_str

    totale_mat = ordine_data.get('totale_materiali', 0) or 0
    totale_netto = ordine_data.get('totale_netto', 0) or totale_mat
    lead_time = ordine_data.get('lead_time_giorni', 15) or 15
    data_consegna = ordine_data.get('data_consegna_prevista', '')
    if data_consegna and len(str(data_consegna)) >= 10:
        try:
            data_consegna = _dt.fromisoformat(str(data_consegna)[:19]).strftime("%d/%m/%Y")
        except Exception:
            data_consegna = str(data_consegna)[:10]
    data_creazione = ordine_data.get('created_at', '')
    if data_creazione and len(str(data_creazione)) >= 10:
        try:
            data_creazione = _dt.fromisoformat(str(data_creazione)[:19]).strftime("%d/%m/%Y")
        except Exception:
            data_creazione = str(data_creazione)[:10]

    rif_prev = f"PREV-{preventivo_id}"
    rev_corr = getattr(preventivo, 'revisione_corrente', 0) or 0
    if rev_corr:
        rif_prev += f" (REV.{rev_corr})"
    rif_cliente = ordine_data.get('riferimento_cliente', '') or ''

    info_rows = [
        ('N° Ordine', safe_str(numero_ordine), 'Data', safe_str(data_creazione)),
        ('Rif. Preventivo', safe_str(rif_prev), 'Rif. Cliente', safe_str(rif_cliente)),
        ('Tipo Impianto', safe_str(ordine_data.get('tipo_impianto', '')), 'Lead Time', f"{lead_time} gg lavorativi"),
        ('Consegna Prevista', safe_str(data_consegna), '', ''),
        ('Totale Ordine', fmt_euro(totale_netto), '', ''),
    ]

    for i, (l1, v1, l2, v2) in enumerate(info_rows):
        t_ord.cell(i, 0).text = l1
        t_ord.cell(i, 1).text = v1
        t_ord.cell(i, 2).text = l2
        t_ord.cell(i, 3).text = v2
        for j in [0, 2]:
            for run in t_ord.cell(i, j).paragraphs[0].runs:
                run.font.bold = True
                run.font.size = Pt(9)
        for j in [1, 3]:
            for run in t_ord.cell(i, j).paragraphs[0].runs:
                run.font.size = Pt(9)

    # Totale in evidenza
    for run in t_ord.cell(4, 1).paragraphs[0].runs:
        run.font.bold = True
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)

    # --- CONDIZIONI ---
    doc.add_paragraph()
    doc.add_heading('Condizioni', level=2)

    cond_pagamento = ordine_data.get('condizioni_pagamento', '') or ''
    cond_consegna = ordine_data.get('condizioni_consegna', '') or ''

    def _add_condizione(label, valore, default):
        p = doc.add_paragraph()
        r = p.add_run(f'{label}: ')
        r.font.bold = True
        r.font.size = Pt(10)
        r2 = p.add_run(valore if valore else default)
        r2.font.size = Pt(10)

    _add_condizione('Pagamento', cond_pagamento, 'Come da accordi')
    _add_condizione('Consegna', cond_consegna, f'Entro {lead_time} giorni lavorativi dalla data del presente ordine')
    _add_condizione('Validità', '', '30 giorni dalla data del presente documento')

    # Note
    note = ordine_data.get('note', '') or ''
    if note:
        doc.add_paragraph()
        doc.add_heading('Note', level=2)
        p = doc.add_paragraph(note)
        for run in p.runs:
            run.font.size = Pt(10)

    # --- FIRME ---
    doc.add_paragraph()
    doc.add_paragraph()

    # Linea separatrice
    p_line = doc.add_paragraph()
    r_line = p_line.add_run('_' * 80)
    r_line.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)

    # Nome azienda per la firma
    nome_azienda = ''
    try:
        cursor.execute("SELECT valore FROM parametri_sistema WHERE chiave = 'azienda_ragione_sociale'")
        row = cursor.fetchone()
        if row:
            nome_azienda = row[0]
    except Exception:
        pass
    if not nome_azienda:
        nome_azienda = 'Il Fornitore'

    t_firme = doc.add_table(rows=4, cols=2)
    t_firme.alignment = WD_TABLE_ALIGNMENT.CENTER

    t_firme.cell(0, 0).text = 'Per accettazione del Committente'
    t_firme.cell(0, 1).text = f'Per {nome_azienda}'
    for j in range(2):
        for run in t_firme.cell(0, j).paragraphs[0].runs:
            run.font.bold = True
            run.font.size = Pt(10)
        t_firme.cell(0, j).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    t_firme.cell(1, 0).text = ''
    t_firme.cell(1, 1).text = ''

    t_firme.cell(2, 0).text = '________________________________'
    t_firme.cell(2, 1).text = '________________________________'
    for j in range(2):
        t_firme.cell(2, j).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    t_firme.cell(3, 0).text = 'Firma e timbro'
    t_firme.cell(3, 1).text = 'Firma e timbro'
    for j in range(2):
        for run in t_firme.cell(3, j).paragraphs[0].runs:
            run.font.size = Pt(8)
            run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
        t_firme.cell(3, j).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Footer
    doc.add_paragraph()
    fp = doc.add_paragraph()
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = fp.add_run(f'Conferma d\'ordine generata il {_dt.now().strftime("%d/%m/%Y %H:%M")}')
    r.font.size = Pt(7)
    r.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)

    # ── 5. Salva e restituisci ──
    buf_out = io.BytesIO()
    doc.save(buf_out)
    buf_out.seek(0)

    num_safe = numero_ordine.replace('/', '_')
    return StreamingResponse(
        buf_out,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="Conferma_Ordine_{num_safe}.docx"'}
    )

@app.post("/ordini/{ordine_id}/esplodi-bom")

def api_esplodi_bom(ordine_id: int, body: dict = None, db: Session = Depends(get_db)):
    """
    Esplode BOM per un ordine.
    Body opzionale:
    - action: "esplodi" (default), "check" (controlla se gia esplosa)
    """
    action = (body or {}).get("action", "esplodi")
    conn = db.get_bind().raw_connection()
    cursor = conn.cursor()
    cursor.execute("SELECT id, preventivo_id FROM ordini WHERE id = ?", (ordine_id,))
    ordine = cursor.fetchone()
    if not ordine:
        raise HTTPException(status_code=404, detail="Ordine non trovato")

    preventivo_id = ordine[1]

    # Check se BOM gia esplosa
    if action == "check":
        try:
            cursor.execute(
                "SELECT bom_esplosa, bom_numero_revisione, bom_esplosa_at FROM ordini WHERE id=?",
                (ordine_id,)
            )
            r = cursor.fetchone()
            # Revisione corrente del preventivo
            cursor.execute("SELECT revisione_corrente FROM preventivi WHERE id=?", (preventivo_id,))
            pr = cursor.fetchone()
            rev_corrente = pr[0] if pr and pr[0] else 0
            return {
                "bom_esplosa": bool(r[0]) if r else False,
                "bom_numero_revisione": r[1] if r else None,
                "bom_esplosa_at": r[2] if r else None,
                "revisione_corrente": rev_corrente,
                "outdated": rev_corrente > (r[1] or 0) if r and r[0] else False,
            }
        except Exception:
            return {"bom_esplosa": False}

    preventivo = db.query(Preventivo).filter(Preventivo.id == preventivo_id).first()

    # Revisione corrente
    rev_corrente = getattr(preventivo, 'revisione_corrente', 0) or 0
    # Ultimo snapshot ID
    cursor.execute(
        "SELECT id FROM revisioni_preventivo WHERE preventivo_id=? ORDER BY id DESC LIMIT 1",
        (preventivo_id,)
    )
    last_snap = cursor.fetchone()
    rev_id = last_snap[0] if last_snap else None

    # Costruisci contesto usando build_config_context (stesso usato da regole)
    contesto = build_config_context(preventivo_id, db)
    if not contesto.get("numero_fermate") and not contesto.get("NUM_FERMATE"):
        contesto.update({"NUM_FERMATE": 6, "CORSA": 18, "TIPO_DISPLAY": "LCD",
                         "TIPO_PORTE": "Automatiche", "TENSIONE": "400V", "POTENZA_KW": 7.5})

    contesto_norm = {}
    for k, v in contesto.items():
        contesto_norm[k] = v
        contesto_norm[k.upper()] = v

    cursor.execute("DELETE FROM esplosi WHERE ordine_id = ?", (ordine_id,))

    materiali = db.query(Materiale).filter(Materiale.preventivo_id == preventivo_id).all()
    if not materiali:
        raise HTTPException(status_code=400, detail="Nessun materiale nel preventivo")

    tutti = []
    for mat in materiali:
        tutti.extend(esplodi_bom_ricorsiva(cursor, mat.codice, mat.quantita,
                                            contesto_norm, ordine_id, 0, mat.codice, []))

    aggregati = {}
    for esp in tutti:
        cod = esp["codice"]
        if cod in aggregati:
            aggregati[cod]["quantita"] += esp["quantita"]
            aggregati[cod]["costo_totale"] += esp["costo_totale"]
        else:
            aggregati[cod] = esp.copy()

    for esp in aggregati.values():
        cursor.execute(
            "INSERT INTO esplosi (ordine_id, codice, descrizione, tipo, categoria, "
            "quantita, unita_misura, costo_unitario, costo_totale, padre_codice, "
            "livello_esplosione, percorso, lead_time_giorni, "
            "parametro1_nome, parametro1_valore, parametro2_nome, parametro2_valore) "
            "VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)",
            (ordine_id, esp["codice"], esp.get("descrizione"), esp.get("tipo"), esp.get("categoria"),
             esp["quantita"], esp.get("unita_misura", "PZ"), esp.get("costo_unitario", 0),
             esp.get("costo_totale", 0), esp.get("padre_codice"), esp.get("livello_esplosione", 0),
             esp.get("percorso"), esp.get("lead_time_giorni", 0),
             esp.get("parametro1_nome"), esp.get("parametro1_valore"),
             esp.get("parametro2_nome"), esp.get("parametro2_valore"))
        )

    now_iso = datetime.now().isoformat()
    try:
        cursor.execute(
            "UPDATE ordini SET bom_esplosa=1, data_esplosione_bom=?, stato='in_produzione', "
            "bom_revisione_id=?, bom_numero_revisione=?, bom_esplosa_at=? WHERE id=?",
            (now_iso, rev_id, rev_corrente, now_iso, ordine_id)
        )
    except Exception:
        # Fallback se colonne nuove non esistono
        cursor.execute("UPDATE ordini SET bom_esplosa=1, data_esplosione_bom=?, stato='in_produzione' WHERE id=?",
                       (now_iso, ordine_id))
    conn.commit()

    costo_tot = sum(e.get("costo_totale", 0) for e in aggregati.values())
    max_lt = max((e.get("lead_time_giorni", 0) for e in aggregati.values()), default=0)

    return {
        "ordine_id": ordine_id,
        "revisione_bom": rev_corrente,
        "componenti_master": len(materiali),
        "componenti_esplosi": len(tutti),
        "componenti_aggregati": len(aggregati),
        "costo_totale_componenti": round(costo_tot, 2),
        "lead_time_max_giorni": max_lt,
        "dettaglio": [{"codice": e["codice"], "descrizione": e.get("descrizione"),
                       "tipo": e.get("tipo"), "quantita": e["quantita"],
                       "unita": e.get("unita_misura"), "costo_totale": round(e.get("costo_totale", 0), 2)}
                      for e in sorted(aggregati.values(), key=lambda x: x.get("categoria", ""))]
    }


@app.get("/ordini/{ordine_id}/esplosi")
def api_get_esplosi(ordine_id: int, tipo: str = None, db: Session = Depends(get_db)):
    conn = db.get_bind().raw_connection()
    cursor = conn.cursor()
    q = "SELECT * FROM esplosi WHERE ordine_id = ?"
    params = [ordine_id]
    if tipo:
        q += " AND tipo = ?"
        params.append(tipo)
    q += " ORDER BY categoria, codice"
    cursor.execute(q, params)
    cols = [d[0] for d in cursor.description]
    esplosi = [dict(zip(cols, r)) for r in cursor.fetchall()]
    totale = sum(e.get("costo_totale", 0) or 0 for e in esplosi)
    return {"ordine_id": ordine_id, "totale_componenti": len(esplosi),
            "costo_totale": round(totale, 2), "esplosi": esplosi}


@app.get("/ordini/{ordine_id}/lista-acquisti")
def api_lista_acquisti(ordine_id: int, db: Session = Depends(get_db)):
    conn = db.get_bind().raw_connection()
    cursor = conn.cursor()
    cursor.execute(
        "SELECT e.codice, e.descrizione, e.quantita, e.unita_misura, e.costo_unitario, e.costo_totale, "
        "ab.fornitore, ab.codice_fornitore, ab.lead_time_acquisto "
        "FROM esplosi e LEFT JOIN articoli_bom ab ON ab.codice = e.codice "
        "WHERE e.ordine_id = ? AND e.tipo = 'ACQUISTO' ORDER BY ab.fornitore, e.codice",
        (ordine_id,)
    )
    fornitori = {}
    for r in cursor.fetchall():
        f = r[6] or "Non assegnato"
        if f not in fornitori:
            fornitori[f] = {"articoli": [], "totale": 0}
        fornitori[f]["articoli"].append({"codice": r[0], "descrizione": r[1], "quantita": r[2],
            "unita": r[3], "costo_unitario": r[4], "costo_totale": r[5],
            "codice_fornitore": r[7], "lead_time_giorni": r[8]})
        fornitori[f]["totale"] += r[5] or 0
    return {
        "ordine_id": ordine_id, "num_fornitori": len(fornitori),
        "costo_totale_acquisti": round(sum(f["totale"] for f in fornitori.values()), 2),
        "fornitori": {n: {"num_articoli": len(d["articoli"]), "totale": round(d["totale"], 2),
                          "articoli": d["articoli"]} for n, d in fornitori.items()}
    }


@app.get("/preventivi/{preventivo_id}/lead-time")
def api_lead_time(preventivo_id: int, db: Session = Depends(get_db)):
    lt = calcola_lead_time(preventivo_id, db)
    return {"preventivo_id": preventivo_id, "lead_time_giorni": lt}

# ==========================================
# STARTUP
# ==========================================
@app.on_event("startup")
def startup_event():
    """Inizializzazione all'avvio"""
    db = SessionLocal()
    try:
        create_default_admin(db)
        _ensure_revisioni_table(db)
        _ensure_revisione_corrente_column(db)
        _ensure_ordini_revisione_columns(db)
        # Migrazione: product_template_ids su campi_configuratore
        try:
            db.execute(text("SELECT product_template_ids FROM campi_configuratore LIMIT 1"))
        except Exception:
            try:
                db.execute(text("ALTER TABLE campi_configuratore ADD COLUMN product_template_ids TEXT"))
                db.commit()
                print("[STARTUP] Aggiunta colonna product_template_ids a campi_configuratore")
            except Exception:
                db.rollback()
    except Exception as e:
        print(f"Errore startup: {e}")
    finally:
        db.close()

# ============================================================
# REVISIONI PREVENTIVO
# ============================================================

def _ensure_revisioni_table(db):
    """Crea tabella revisioni_preventivo se non esiste"""
    conn = db.get_bind().raw_connection()
    cursor = conn.cursor()
    cursor.execute("""
        CREATE TABLE IF NOT EXISTS revisioni_preventivo (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            preventivo_id INTEGER NOT NULL,
            numero_revisione INTEGER NOT NULL,
            motivo TEXT,
            snapshot_configurazione TEXT,
            snapshot_materiali TEXT,
            snapshot_totali TEXT,
            created_by TEXT,
            created_at TEXT DEFAULT (datetime('now')),
            FOREIGN KEY (preventivo_id) REFERENCES preventivi(id)
        )
    """)
    conn.commit()

def _ensure_revisione_corrente_column(db):
    """Aggiunge colonna revisione_corrente a preventivi se non esiste"""
    conn = db.get_bind().raw_connection()
    cursor = conn.cursor()
    try:
        cursor.execute("SELECT revisione_corrente FROM preventivi LIMIT 1")
    except Exception:
        cursor.execute("ALTER TABLE preventivi ADD COLUMN revisione_corrente INTEGER DEFAULT 0")
        conn.commit()
        print("[MIGRATION] Aggiunta colonna revisione_corrente a preventivi")

def _ensure_ordini_revisione_columns(db):
    """Aggiunge colonne di tracciamento revisione alla tabella ordini"""
    conn = db.get_bind().raw_connection()
    cursor = conn.cursor()
    try:
        cursor.execute("SELECT name FROM sqlite_master WHERE type='table' AND name='ordini'")
        if not cursor.fetchone():
            return  # Tabella ordini non esiste ancora
    except Exception:
        return

    new_cols = {
        "revisione_id": "INTEGER",
        "numero_revisione_origine": "INTEGER",
        "bom_revisione_id": "INTEGER",
        "bom_numero_revisione": "INTEGER",
        "bom_esplosa_at": "TEXT",
    }
    for col_name, col_type in new_cols.items():
        try:
            cursor.execute(f"SELECT {col_name} FROM ordini LIMIT 1")
        except Exception:
            try:
                cursor.execute(f"ALTER TABLE ordini ADD COLUMN {col_name} {col_type}")
                conn.commit()
                print(f"[MIGRATION] Aggiunta colonna {col_name} a ordini")
            except Exception as e:
                print(f"[WARN] Migration {col_name}: {e}")

def _crea_snapshot_preventivo(preventivo_id: int, db, motivo: str = None, created_by: str = "admin"):
    """Crea uno snapshot/revisione del preventivo corrente"""
    _ensure_revisioni_table(db)

    conn = db.get_bind().raw_connection()
    cursor = conn.cursor()

    # Prossimo numero revisione
    cursor.execute(
        "SELECT COALESCE(MAX(numero_revisione), 0) + 1 FROM revisioni_preventivo WHERE preventivo_id = ?",
        (preventivo_id,)
    )
    num_rev = cursor.fetchone()[0]

    # Snapshot configurazione
    cursor.execute("SELECT * FROM preventivi WHERE id = ?", (preventivo_id,))
    prev_row = cursor.fetchone()
    if not prev_row:
        return None
    prev_cols = [d[0] for d in cursor.description]
    prev_dict = dict(zip(prev_cols, prev_row))

    # Snapshot configurazione da sezioni — PRIMA di usare config_snap
    config_snap = {}

    # Metadata: info preventivo + cliente per ricostruire contesto
    try:
        cliente_info = {}
        cid = prev_dict.get("cliente_id")
        if cid:
            cliente_row = cursor.execute("SELECT id, codice, ragione_sociale FROM clienti WHERE id = ?", (cid,)).fetchone()
            if cliente_row:
                cliente_info = {"id": cliente_row[0], "codice": cliente_row[1], "ragione_sociale": cliente_row[2]}
        config_snap["_metadata"] = {
            "stato": prev_dict.get("status") or prev_dict.get("stato") or "draft",
            "template_id": prev_dict.get("template_id"),
            "numero_preventivo": prev_dict.get("numero_preventivo"),
            "cliente": cliente_info,
        }
    except Exception:
        pass

    for table in ['dati_commessa', 'dati_principali', 'normative', 'disposizione_vano', 'porte', 'argano']:
        try:
            cursor.execute(f"SELECT * FROM {table} WHERE preventivo_id = ?", (preventivo_id,))
            row = cursor.fetchone()
            if row:
                cols = [d[0] for d in cursor.description]
                config_snap[table] = dict(zip(cols, row))
        except Exception:
            pass

    # Valori configurazione (tabella chiave/valore) — ORA config_snap esiste
    try:
        cursor.execute(
            "SELECT codice_campo, valore, sezione FROM valori_configurazione WHERE preventivo_id = ?",
            (preventivo_id,)
        )
        valori_config = {}
        for row in cursor.fetchall():
            valori_config[row[0]] = {"valore": row[1], "sezione": row[2]}
        if valori_config:
            config_snap["valori_configurazione"] = valori_config
    except Exception:
        pass

    # Carica configurazione JSON dalle sezioni dinamiche
    try:
        cursor.execute(
            "SELECT sezione, valori FROM configurazione_preventivo WHERE preventivo_id = ?",
            (preventivo_id,)
        )
        for row in cursor.fetchall():
            config_snap[f"config_{row[0]}"] = row[1]
    except Exception:
        pass

    # Snapshot materiali
    materiali = db.query(Materiale).filter(Materiale.preventivo_id == preventivo_id).all()
    mat_list = []
    for m in materiali:
        mat_dict = {}
        for col in m.__table__.columns:
            val = getattr(m, col.name, None)
            if hasattr(val, 'isoformat'):
                val = val.isoformat()
            mat_dict[col.name] = val
        mat_list.append(mat_dict)

    # Snapshot totali
    totali = {
        "total_price": prev_dict.get("total_price", 0),
        "sconto_cliente": prev_dict.get("sconto_cliente", 0),
        "sconto_extra_admin": prev_dict.get("sconto_extra_admin", 0),
        "total_price_finale": prev_dict.get("total_price_finale", 0),
    }

    cursor.execute(
        "INSERT INTO revisioni_preventivo "
        "(preventivo_id, numero_revisione, motivo, snapshot_configurazione, snapshot_materiali, snapshot_totali, created_by) "
        "VALUES (?, ?, ?, ?, ?, ?, ?)",
        (
            preventivo_id, num_rev, motivo,
            json.dumps(config_snap, ensure_ascii=False, default=str),
            json.dumps(mat_list, ensure_ascii=False, default=str),
            json.dumps(totali, ensure_ascii=False, default=str),
            created_by
        )
    )
    rev_id = cursor.lastrowid

    # Aggiorna revisione_corrente sul preventivo
    cursor.execute(
        "UPDATE preventivi SET revisione_corrente = ?, updated_at = datetime('now') WHERE id = ?",
        (num_rev, preventivo_id)
    )
    conn.commit()

    return {
        "id": rev_id,
        "preventivo_id": preventivo_id,
        "numero_revisione": num_rev,
        "motivo": motivo,
        "created_at": datetime.now().isoformat()
    }


@app.get("/preventivi/{preventivo_id}/revisioni")
def get_revisioni(preventivo_id: int, db: Session = Depends(get_db)):
    """Lista revisioni di un preventivo"""
    _ensure_revisioni_table(db)
    conn = db.get_bind().raw_connection()
    cursor = conn.cursor()
    cursor.execute(
        "SELECT id, preventivo_id, numero_revisione, motivo, created_by, created_at "
        "FROM revisioni_preventivo WHERE preventivo_id = ? ORDER BY numero_revisione DESC",
        (preventivo_id,)
    )
    rows = cursor.fetchall()
    cols = [d[0] for d in cursor.description]
    return [dict(zip(cols, r)) for r in rows]


@app.get("/preventivi/{preventivo_id}/revisioni/{revisione_id}")
def get_revisione_dettaglio(preventivo_id: int, revisione_id: int, db: Session = Depends(get_db)):
    """Dettaglio di una revisione con snapshot completo"""
    _ensure_revisioni_table(db)
    conn = db.get_bind().raw_connection()
    cursor = conn.cursor()
    cursor.execute(
        "SELECT * FROM revisioni_preventivo WHERE id = ? AND preventivo_id = ?",
        (revisione_id, preventivo_id)
    )
    row = cursor.fetchone()
    if not row:
        raise HTTPException(status_code=404, detail="Revisione non trovata")
    cols = [d[0] for d in cursor.description]
    result = dict(zip(cols, row))
    # Parse JSON fields
    for field in ['snapshot_configurazione', 'snapshot_materiali', 'snapshot_totali']:
        if result.get(field):
            try:
                result[field] = json.loads(result[field])
            except Exception:
                pass
    return result


@app.get("/preventivi/{preventivo_id}/revisioni/{revisione_id}/diff")
def diff_revisione(preventivo_id: int, revisione_id: int, db: Session = Depends(get_db)):
    """Confronta una revisione con lo stato corrente del preventivo.
    Restituisce materiali aggiunti/rimossi/modificati e variazioni totali."""
    _ensure_revisioni_table(db)
    conn = db.get_bind().raw_connection()
    cursor = conn.cursor()

    # Carica snapshot revisione
    cursor.execute(
        "SELECT numero_revisione, snapshot_materiali, snapshot_totali, snapshot_configurazione "
        "FROM revisioni_preventivo WHERE id = ? AND preventivo_id = ?",
        (revisione_id, preventivo_id)
    )
    row = cursor.fetchone()
    if not row:
        raise HTTPException(404, "Revisione non trovata")

    num_rev = row[0]
    rev_materiali = json.loads(row[1]) if row[1] else []
    rev_totali = json.loads(row[2]) if row[2] else {}
    rev_config = json.loads(row[3]) if row[3] else {}

    # Stato corrente materiali
    materiali_correnti = db.query(Materiale).filter(Materiale.preventivo_id == preventivo_id).all()
    curr_mat_map = {}
    for m in materiali_correnti:
        curr_mat_map[m.codice] = {
            "codice": m.codice, "descrizione": m.descrizione,
            "quantita": m.quantita, "prezzo_unitario": m.prezzo_unitario,
            "prezzo_totale": m.prezzo_totale, "categoria": m.categoria,
        }

    rev_mat_map = {}
    for m in rev_materiali:
        cod = m.get("codice", "")
        rev_mat_map[cod] = m

    # Diff materiali
    materiali_aggiunti = []
    materiali_rimossi = []
    materiali_modificati = []

    for cod, curr in curr_mat_map.items():
        if cod not in rev_mat_map:
            materiali_aggiunti.append(curr)
        else:
            rev = rev_mat_map[cod]
            changes = {}
            for field in ["quantita", "prezzo_unitario", "prezzo_totale"]:
                cv = curr.get(field, 0) or 0
                rv = rev.get(field, 0) or 0
                if abs(float(cv) - float(rv)) > 0.001:
                    changes[field] = {"rev": rv, "corrente": cv}
            if changes:
                materiali_modificati.append({"codice": cod, "descrizione": curr.get("descrizione", ""), "changes": changes})

    for cod, rev in rev_mat_map.items():
        if cod not in curr_mat_map:
            materiali_rimossi.append({
                "codice": cod, "descrizione": rev.get("descrizione", ""),
                "quantita": rev.get("quantita", 0), "prezzo_totale": rev.get("prezzo_totale", 0),
            })

    # Diff totali
    prev_row = cursor.execute("SELECT * FROM preventivi WHERE id = ?", (preventivo_id,)).fetchone()
    prev_cols = [d[0] for d in cursor.description]
    prev_dict = dict(zip(prev_cols, prev_row)) if prev_row else {}

    totali_diff = {}
    for k in ["total_price", "sconto_cliente", "sconto_extra_admin", "total_price_finale"]:
        cv = float(prev_dict.get(k, 0) or 0)
        rv = float(rev_totali.get(k, 0) or 0)
        if abs(cv - rv) > 0.01:
            totali_diff[k] = {"rev": rv, "corrente": cv, "delta": round(cv - rv, 2)}

    # Diff configurazione (valori_configurazione — campi dinamici)
    config_changes = []
    rev_valori = rev_config.get("valori_configurazione", {})
    if rev_valori:
        curr_valori_rows = db.execute(text(
            "SELECT codice_campo, valore, sezione FROM valori_configurazione WHERE preventivo_id = :pid"
        ), {"pid": preventivo_id}).fetchall()
        curr_valori = {r[0]: {"valore": r[1], "sezione": r[2]} for r in curr_valori_rows}

        all_keys = set(list(curr_valori.keys()) + list(rev_valori.keys()))
        for key in sorted(all_keys):
            cv = curr_valori.get(key, {})
            rv = rev_valori.get(key, {})
            c_val = cv.get("valore", "") if isinstance(cv, dict) else str(cv)
            r_val = rv.get("valore", "") if isinstance(rv, dict) else str(rv)
            sez = cv.get("sezione", "") if isinstance(cv, dict) else (rv.get("sezione", "") if isinstance(rv, dict) else "")
            if str(c_val or "") != str(r_val or ""):
                config_changes.append({
                    "campo": key, "sezione": sez,
                    "rev": r_val or "(vuoto)", "corrente": c_val or "(vuoto)",
                })

    # Diff tabelle ORM fisse (dati_commessa, dati_principali, normative, disposizione_vano, porte, argano)
    SKIP_FIELDS = {'id', 'preventivo_id', 'created_at', 'updated_at'}
    for table_name in ['dati_commessa', 'dati_principali', 'normative', 'disposizione_vano', 'porte', 'argano']:
        rev_table_data = rev_config.get(table_name)
        if not rev_table_data or not isinstance(rev_table_data, dict):
            continue
        try:
            curr_row = cursor.execute(f"SELECT * FROM {table_name} WHERE preventivo_id = ?", (preventivo_id,)).fetchone()
            if not curr_row:
                continue
            curr_cols = [d[0] for d in cursor.description]
            curr_data = dict(zip(curr_cols, curr_row))

            for col in curr_cols:
                if col in SKIP_FIELDS:
                    continue
                c_val = str(curr_data.get(col) or "")
                r_val = str(rev_table_data.get(col) or "")
                if c_val != r_val:
                    config_changes.append({
                        "campo": col, "sezione": table_name,
                        "rev": r_val or "(vuoto)", "corrente": c_val or "(vuoto)",
                    })
        except Exception:
            pass

    return {
        "revisione_id": revisione_id,
        "numero_revisione": num_rev,
        "materiali_aggiunti": materiali_aggiunti,
        "materiali_rimossi": materiali_rimossi,
        "materiali_modificati": materiali_modificati,
        "totali_diff": totali_diff,
        "config_changes": config_changes,
        "summary": {
            "n_aggiunti": len(materiali_aggiunti),
            "n_rimossi": len(materiali_rimossi),
            "n_modificati": len(materiali_modificati),
            "n_config_changes": len(config_changes),
            "has_totali_changes": len(totali_diff) > 0,
        }
    }


@app.post("/preventivi/{preventivo_id}/revisioni")
def crea_revisione(preventivo_id: int, body: dict = None, db: Session = Depends(get_db)):
    """Crea manualmente una revisione/snapshot del preventivo"""
    motivo = (body or {}).get("motivo", "Snapshot manuale")
    result = _crea_snapshot_preventivo(preventivo_id, db, motivo=motivo)
    if not result:
        raise HTTPException(status_code=404, detail="Preventivo non trovato")
    return result


@app.post("/preventivi/{preventivo_id}/auto-snapshot")
def auto_snapshot(preventivo_id: int, db: Session = Depends(get_db)):
    """Crea snapshot solo se il preventivo è dirty (modificato dopo ultimo snapshot)"""
    _ensure_revisioni_table(db)
    conn = db.get_bind().raw_connection()
    cursor = conn.cursor()

    cursor.execute("SELECT updated_at FROM preventivi WHERE id = ?", (preventivo_id,))
    row = cursor.fetchone()
    if not row or not row[0]:
        return {"created": False, "reason": "preventivo non trovato o senza updated_at"}

    prev_updated = row[0]

    cursor.execute(
        "SELECT created_at FROM revisioni_preventivo WHERE preventivo_id = ? ORDER BY id DESC LIMIT 1",
        (preventivo_id,)
    )
    row = cursor.fetchone()
    if row and str(prev_updated) <= str(row[0]):
        return {"created": False, "reason": "nessuna modifica dall'ultimo snapshot"}

    result = _crea_snapshot_preventivo(preventivo_id, db, motivo="Auto-save all'uscita")
    return {"created": True, "revisione": result}


@app.post("/preventivi/{preventivo_id}/revisioni/{revisione_id}/ripristina")
def ripristina_revisione(preventivo_id: int, revisione_id: int, db: Session = Depends(get_db)):
    """Ripristina un preventivo da una revisione precedente.
    Il numero revisione viene SEMPRE incrementato (mai decrementato).
    Flow: snapshot stato attuale → ripristino dati → nuovo snapshot stato ripristinato.
    """
    _ensure_revisioni_table(db)

    try:
        # 1. Auto-snapshot stato corrente (salva "prima del ripristino")
        try:
            _crea_snapshot_preventivo(preventivo_id, db, motivo="Auto-snapshot prima di ripristino")
        except Exception as e:
            print(f"[WARN] Auto-snapshot fallito: {e}")

        conn = db.get_bind().raw_connection()
        cursor = conn.cursor()

        # 2. Carica la revisione target
        cursor.execute(
            "SELECT numero_revisione, snapshot_configurazione, snapshot_materiali, snapshot_totali "
            "FROM revisioni_preventivo WHERE id = ? AND preventivo_id = ?",
            (revisione_id, preventivo_id)
        )
        row = cursor.fetchone()
        if not row:
            raise HTTPException(status_code=404, detail="Revisione non trovata")

        num_rev_ripristinata = row[0]
        snap_config = json.loads(row[1]) if row[1] else {}
        snap_materiali = json.loads(row[2]) if row[2] else []
        snap_totali = json.loads(row[3]) if row[3] else {}

        # 3. Ripristina totali (SENZA toccare revisione_corrente — verrà aggiornata dallo snapshot finale)
        totali_updates = ["updated_at = datetime('now')"]
        totali_params = []
        if snap_totali:
            for k, v in snap_totali.items():
                totali_updates.append(f"{k} = ?")
                totali_params.append(v)
        totali_params.append(preventivo_id)
        cursor.execute(
            f"UPDATE preventivi SET {', '.join(totali_updates)} WHERE id = ?",
            totali_params
        )

        # 4. Ripristina materiali
        cursor.execute("DELETE FROM materiali WHERE preventivo_id = ?", (preventivo_id,))
        for mat in snap_materiali:
            mat.pop('id', None)
            mat.pop('created_at', None)
            mat['preventivo_id'] = preventivo_id
            cols = list(mat.keys())
            vals = list(mat.values())
            placeholders = ', '.join(['?'] * len(cols))
            try:
                cursor.execute(
                    f"INSERT INTO materiali ({', '.join(cols)}) VALUES ({placeholders})",
                    vals
                )
            except Exception as e:
                print(f"[WARN] Skip materiale restore: {e}")

        # 5. Ripristina sezioni configurazione
        for table_name, data in snap_config.items():
            if table_name == "valori_configurazione":
                try:
                    cursor.execute(
                        "DELETE FROM valori_configurazione WHERE preventivo_id = ?",
                        (preventivo_id,)
                    )
                    if isinstance(data, dict):
                        for campo, info in data.items():
                            valore = info.get("valore", "") if isinstance(info, dict) else str(info)
                            sezione = info.get("sezione", "") if isinstance(info, dict) else ""
                            cursor.execute(
                                "INSERT INTO valori_configurazione "
                                "(preventivo_id, sezione, codice_campo, valore) "
                                "VALUES (?, ?, ?, ?)",
                                (preventivo_id, sezione, campo, valore)
                            )
                except Exception as e:
                    print(f"[WARN] Skip valori_configurazione restore: {e}")

            elif table_name.startswith('config_'):
                sezione = table_name[7:]
                try:
                    cursor.execute(
                        "UPDATE configurazione_preventivo SET valori = ?, updated_at = datetime('now') "
                        "WHERE preventivo_id = ? AND sezione = ?",
                        (data if isinstance(data, str) else json.dumps(data), preventivo_id, sezione)
                    )
                except Exception:
                    pass
            elif table_name == "_metadata":
                pass  # metadata non va ripristinata, è solo info storica
            elif isinstance(data, dict):
                try:
                    data_copy = {k: v for k, v in data.items() if k not in ('id', 'preventivo_id')}
                    if data_copy:
                        updates = [f"{k} = ?" for k in data_copy.keys()]
                        params = list(data_copy.values()) + [preventivo_id]
                        cursor.execute(
                            f"UPDATE {table_name} SET {', '.join(updates)} WHERE preventivo_id = ?",
                            params
                        )
                except Exception as e:
                    print(f"[WARN] Skip config restore {table_name}: {e}")

        conn.commit()

        # 6. Riesegui regole dopo ripristino
        try:
            db.expire_all()  # forza ORM a rileggere i dati dopo raw SQL
            safe_evaluate_rules(preventivo_id, db)
        except Exception:
            pass

        # 7. Snapshot dello stato ripristinato → crea NUOVA revisione (numero sempre crescente)
        snap_result = _crea_snapshot_preventivo(
            preventivo_id, db,
            motivo=f"Ripristino da Rev #{num_rev_ripristinata}"
        )
        new_rev = snap_result["numero_revisione"] if snap_result else "?"

        return {
            "success": True,
            "message": f"Preventivo ripristinato dalla revisione #{num_rev_ripristinata} → nuova Rev #{new_rev}",
            "preventivo_id": preventivo_id,
            "revisione_ripristinata": num_rev_ripristinata,
            "revisione_corrente": new_rev
        }
    except HTTPException:
        raise
    except Exception as e:
        import traceback
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=f"Errore ripristino: {str(e)}")
    
@app.get("/preventivi/{preventivo_id}/dirty")
def check_dirty(preventivo_id: int, db: Session = Depends(get_db)):
    """Controlla se il preventivo ha modifiche non snapshot-ate."""
    _ensure_revisioni_table(db)
    conn = db.get_bind().raw_connection()
    cursor = conn.cursor()

    cursor.execute("SELECT updated_at FROM preventivi WHERE id = ?", (preventivo_id,))
    row = cursor.fetchone()
    if not row or not row[0]:
        return {"dirty": False, "preventivo_id": preventivo_id}

    prev_updated = row[0]

    cursor.execute(
        "SELECT created_at FROM revisioni_preventivo WHERE preventivo_id = ? ORDER BY id DESC LIMIT 1",
        (preventivo_id,)
    )
    row = cursor.fetchone()
    if not row:
        return {"dirty": True, "preventivo_id": preventivo_id}

    dirty = str(prev_updated) > str(row[0])
    return {"dirty": dirty, "preventivo_id": preventivo_id}


@app.get("/preventivi/{preventivo_id}/filiera")
def get_filiera(preventivo_id: int, db: Session = Depends(get_db)):
    """Restituisce la filiera compatta: preventivo -> ordini -> BOM con info revisioni"""
    _ensure_revisioni_table(db)
    conn = db.get_bind().raw_connection()
    cursor = conn.cursor()

    # Preventivo — prova entrambi i nomi colonna
    cursor.execute("SELECT * FROM preventivi WHERE id = ?", (preventivo_id,))
    prev = cursor.fetchone()
    if not prev:
        raise HTTPException(404, "Preventivo non trovato")
    prev_cols = [d[0] for d in cursor.description]
    prev_full = dict(zip(prev_cols, prev))
    
    prev_data = {
        "id": prev_full.get("id"),
        "revisione_corrente": prev_full.get("revisione_corrente", 0) or 0,
        "updated_at": prev_full.get("updated_at"),
        "status": prev_full.get("status") or prev_full.get("stato") or "draft",
    }

    # Ultima revisione
    cursor.execute(
        "SELECT numero_revisione, created_at FROM revisioni_preventivo WHERE preventivo_id = ? ORDER BY id DESC LIMIT 1",
        (preventivo_id,)
    )
    last_rev = cursor.fetchone()
    prev_data["ultima_revisione"] = last_rev[0] if last_rev else 0
    prev_data["ultima_revisione_data"] = last_rev[1] if last_rev else None

    # Ordini — con fallback robusto
    ordini = []
    try:
        cursor.execute("SELECT name FROM sqlite_master WHERE type='table' AND name='ordini'")
        if cursor.fetchone():
            # Prova prima con colonne nuove
            try:
                cursor.execute(
                    "SELECT id, numero_ordine, stato, revisione_id, numero_revisione_origine, "
                    "bom_esplosa, bom_numero_revisione, bom_esplosa_at, created_at, "
                    "totale_materiali, totale_netto, lead_time_giorni, data_consegna_prevista "
                    "FROM ordini WHERE preventivo_id = ? ORDER BY created_at DESC",
                    (preventivo_id,)
                )
                ordini_cols = [d[0] for d in cursor.description]
                ordini = [dict(zip(ordini_cols, r)) for r in cursor.fetchall()]
            except Exception:
                # Fallback senza colonne nuove
                cursor.execute(
                    "SELECT id, numero_ordine, stato, bom_esplosa, created_at, "
                    "totale_materiali, totale_netto, lead_time_giorni, data_consegna_prevista "
                    "FROM ordini WHERE preventivo_id = ? ORDER BY created_at DESC",
                    (preventivo_id,)
                )
                ordini_cols = [d[0] for d in cursor.description]
                ordini = [dict(zip(ordini_cols, r)) for r in cursor.fetchall()]
    except Exception as e:
        print(f"[WARN] filiera ordini: {e}")
        ordini_cols = [d[0] for d in cursor.description]
        ordini = [dict(zip(ordini_cols, r)) for r in cursor.fetchall()]
    except Exception:
        # Fallback se colonne nuove non esistono ancora
        cursor.execute(
            "SELECT id, numero_ordine, stato, bom_esplosa, created_at, "
            "totale_materiali, totale_netto, lead_time_giorni, data_consegna_prevista "
            "FROM ordini WHERE preventivo_id = ? ORDER BY created_at DESC",
            (preventivo_id,)
        )
        ordini_cols = [d[0] for d in cursor.description]
        ordini = [dict(zip(ordini_cols, r)) for r in cursor.fetchall()]

    # Per ogni ordine, calcola se è outdated
    rev_corrente = prev_data.get("revisione_corrente", 0) or 0
    for o in ordini:
        rev_origine = o.get("numero_revisione_origine") or 0
        o["outdated"] = rev_corrente > rev_origine if rev_origine > 0 else False
        o["revisioni_dietro"] = max(0, rev_corrente - rev_origine) if rev_origine > 0 else 0

        bom_rev = o.get("bom_numero_revisione") or 0
        o["bom_outdated"] = rev_corrente > bom_rev if bom_rev > 0 and o.get("bom_esplosa") else False
        o["bom_revisioni_dietro"] = max(0, rev_corrente - bom_rev) if bom_rev > 0 else 0

    return {
        "preventivo": prev_data,
        "ordini": ordini,
        "revisione_corrente": rev_corrente,
    }

# ============================================================================
# ENDPOINT: Parse Excel (usato dal wizard del Rule Designer)
# ============================================================================

@app.post("/api/v2/import/excel/parse")
async def parse_excel(
    file: UploadFile = File(...),
    sheet: str = None,
    header_row: int = 1,
    limit: int = 200,
):
    """
    Parsa un file Excel e restituisce dati strutturati.
    Usato dal wizard Excel del Rule Designer.
    
    - Se sheet non specificato: restituisce lista fogli + dati primo foglio
    - Se sheet specificato: restituisce colonne e righe di quel foglio
    """
    import tempfile, shutil
    try:
        import openpyxl
    except ImportError:
        raise HTTPException(500, "openpyxl non installato")

    suffix = os.path.splitext(file.filename)[1] or ".xlsx"
    with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
        shutil.copyfileobj(file.file, tmp)
        tmp_path = tmp.name

    try:
        wb = openpyxl.load_workbook(tmp_path, data_only=True, read_only=True)
        all_sheets = wb.sheetnames

        target_sheet = sheet or all_sheets[0]
        if target_sheet not in all_sheets:
            raise HTTPException(400, f"Foglio '{target_sheet}' non trovato. Disponibili: {all_sheets}")

        ws = wb[target_sheet]
        rows_raw = []
        for i, row in enumerate(ws.iter_rows(values_only=True)):
            if i >= header_row - 1 + limit + 1:  # header + limit righe dati
                break
            rows_raw.append(list(row))

        wb.close()

        if len(rows_raw) < header_row:
            return {"sheets": all_sheets, "columns": [], "rows": []}

        # Header dalla riga indicata
        headers = []
        for j, h in enumerate(rows_raw[header_row - 1]):
            if h is not None:
                headers.append(str(h).strip())
            else:
                headers.append(f"col_{j}")

        # Righe dati
        data_rows = []
        for row in rows_raw[header_row:]:
            d = {}
            for j, h in enumerate(headers):
                val = row[j] if j < len(row) else None
                if isinstance(val, float) and val == int(val):
                    val = int(val)
                d[h] = val
            data_rows.append(d)

        return {
            "sheets": all_sheets,
            "columns": headers,
            "rows": data_rows[:limit],
        }
    finally:
        os.unlink(tmp_path)



# ============================================================================
# ENDPOINT: Data Tables (gestione tabelle dati da Excel)
# ============================================================================

@app.post("/data-tables/save")
def save_data_table(data: dict):
    """
    Salva una data table JSON direttamente (usato dal wizard Excel).
    
    Il wizard nel Rule Designer manda il JSON già processato.
    Lo salviamo in ./data/{nome_tabella}.json.
    """
    import os
    meta = data.get("_meta", {})
    nome = meta.get("nome") or data.get("nome_tabella")
    if not nome:
        raise HTTPException(400, "Campo _meta.nome obbligatorio")

    # Sanitizza nome file
    nome_safe = "".join(c for c in nome if c.isalnum() or c == "_").lower()
    if not nome_safe:
        raise HTTPException(400, "Nome tabella non valido")

    data_dir = "./data"
    os.makedirs(data_dir, exist_ok=True)
    outpath = os.path.join(data_dir, f"{nome_safe}.json")

    with open(outpath, "w", encoding="utf-8") as f:
        json.dump(data, f, indent=2, ensure_ascii=False)

    return {"saved": nome_safe, "path": outpath}


@app.get("/data-tables")
def list_data_tables():
    """Lista tutte le data tables disponibili in ./data/"""
    loader = ExcelDataLoader()
    tables = loader.list_tables()
    return {"tables": tables, "data_dir": "./data"}


@app.get("/data-tables/{nome_tabella}")
def get_data_table(nome_tabella: str):
    """Restituisce il contenuto di una data table."""
    loader = ExcelDataLoader()
    table = loader.load_table(nome_tabella)
    if not table:
        raise HTTPException(404, f"Tabella '{nome_tabella}' non trovata")
    return table


@app.delete("/data-tables/{nome_tabella}")
def delete_data_table(nome_tabella: str):
    """Elimina una data table."""
    import os
    filepath = os.path.join("./data", f"{nome_tabella}.json")
    if not os.path.exists(filepath):
        raise HTTPException(404, f"Tabella '{nome_tabella}' non trovata")
    os.remove(filepath)
    return {"deleted": nome_tabella}


@app.post("/data-tables/upload")
async def upload_excel_data(file: UploadFile = File(...), overwrite: bool = True):
    """
    Carica un file Excel con foglio _MAPPA → genera data tables JSON.
    
    Il file Excel deve seguire la convenzione:
    - Foglio _MAPPA con colonne: foglio, tipo, nome_tabella, colonna_chiave, ...
    - Fogli dati con riga intestazioni + righe dati
    
    Returns: 
        Lista tabelle generate con eventuali errori/warning
    """
    import tempfile
    import shutil

    # Salva file temporaneo
    suffix = os.path.splitext(file.filename)[1] or ".xlsx"
    with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
        shutil.copyfileobj(file.file, tmp)
        tmp_path = tmp.name

    try:
        loader = ExcelDataLoader()
        result = loader.load_excel(tmp_path, overwrite=overwrite)
        
        # Salva anche il file Excel originale in ./data/excel_originals/
        originals_dir = os.path.join("./data", "excel_originals")
        os.makedirs(originals_dir, exist_ok=True)
        dest = os.path.join(originals_dir, file.filename)
        shutil.copy2(tmp_path, dest)
        result["original_saved"] = dest
        
        return result
    finally:
        os.unlink(tmp_path)


@app.post("/data-tables/validate")
async def validate_excel_data(file: UploadFile = File(...)):
    """
    Valida un file Excel senza generare tabelle.
    Utile per preview prima del caricamento.
    """
    import tempfile
    import shutil

    suffix = os.path.splitext(file.filename)[1] or ".xlsx"
    with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
        shutil.copyfileobj(file.file, tmp)
        tmp_path = tmp.name

    try:
        loader = ExcelDataLoader()
        return loader.validate_excel(tmp_path)
    finally:
        os.unlink(tmp_path)


@app.post("/data-tables/merge")
def merge_data_tables(data: dict):
    """
    Unisce più tabelle lookup_range partizionate in una sola.
    
    Body JSON:
    {
        "table_names": ["contattori_50hz", "contattori_60hz"],
        "merged_name": "contattori_oleo",
        "partition_field": "frequenza_tensione"
    }
    """
    table_names = data.get("table_names", [])
    merged_name = data.get("merged_name", "")
    partition_field = data.get("partition_field", "")

    if not table_names or not merged_name:
        raise HTTPException(400, "Servono table_names e merged_name")

    loader = ExcelDataLoader()
    result = loader.merge_partitioned_tables(table_names, merged_name, partition_field)
    
    if loader.errors:
        raise HTTPException(400, {"errors": loader.errors})
    
    return {"merged": merged_name, "partitions": list(result.get("partizioni", {}).keys())}

# ============================================================================
# IMPORT EXCEL V3 — WIZARD SENZA _MAPPA
# ============================================================================

@app.post("/import-excel/parse-sheet")
async def import_excel_parse_sheet(
    file: UploadFile = File(...),
    sheet: Optional[str] = Query(None),
    header_row: int = Query(1),
):
    """
    V3 Step 1: Parsa un foglio Excel, ritorna colonne + righe anteprima.
    Se sheet non specificato, usa il primo foglio.
    header_row = riga delle intestazioni (1-based).
    """
    import tempfile, shutil, openpyxl

    suffix = os.path.splitext(file.filename)[1] or ".xlsx"
    with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
        shutil.copyfileobj(file.file, tmp)
        tmp_path = tmp.name

    try:
        wb = openpyxl.load_workbook(tmp_path, data_only=True)
        sheets = wb.sheetnames
        selected = sheet if sheet and sheet in sheets else sheets[0]
        ws = wb[selected]

        # Leggi tutte le righe
        all_rows = []
        for row in ws.iter_rows(values_only=True):
            all_rows.append(list(row))

        # Headers dalla riga specificata (1-based)
        h_idx = max(0, header_row - 1)
        if h_idx >= len(all_rows):
            return {"sheets": sheets, "selected_sheet": selected, "header_row": header_row,
                    "columns": [], "rows": [], "total_rows": 0}

        raw_headers = all_rows[h_idx]
        columns = [str(h) if h is not None else f"Col_{i+1}" for i, h in enumerate(raw_headers)]

        # Righe dati (dopo header, max 50 per preview)
        data_rows = []
        for row in all_rows[h_idx + 1:]:
            if all(v is None for v in row):
                continue
            row_dict = {}
            for i, col in enumerate(columns):
                row_dict[col] = row[i] if i < len(row) else None
            data_rows.append(row_dict)

        wb.close()
        return {
            "sheets": sheets,
            "selected_sheet": selected,
            "header_row": header_row,
            "columns": columns,
            "rows": data_rows[:50],
            "total_rows": len(data_rows),
        }
    finally:
        os.unlink(tmp_path)


@app.post("/import-excel/analyze-v3")
async def import_excel_analyze_v3(
    file: UploadFile = File(...),
    config_json: str = Form(...),
    db: Session = Depends(get_db),
):
    """
    V3 Step 2→3: Analizza struttura, estrae valori distinti per colonne output,
    trova candidati campi configuratore per colonne chiave.
    """
    import tempfile, shutil, openpyxl

    config = json.loads(config_json)
    sheet = config.get("sheet")
    header_row = config.get("header_row", 1)
    key_columns = config.get("key_columns", [])
    output_columns = config.get("output_columns", [])

    suffix = os.path.splitext(file.filename)[1] or ".xlsx"
    with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
        shutil.copyfileobj(file.file, tmp)
        tmp_path = tmp.name

    try:
        wb = openpyxl.load_workbook(tmp_path, data_only=True)
        ws = wb[sheet] if sheet and sheet in wb.sheetnames else wb[wb.sheetnames[0]]

        all_rows = []
        for row in ws.iter_rows(values_only=True):
            all_rows.append(list(row))

        h_idx = max(0, header_row - 1)
        raw_headers = all_rows[h_idx] if h_idx < len(all_rows) else []
        columns = [str(h) if h is not None else f"Col_{i+1}" for i, h in enumerate(raw_headers)]

        data_rows = []
        for row in all_rows[h_idx + 1:]:
            if all(v is None for v in row):
                continue
            row_dict = {}
            for i, col in enumerate(columns):
                row_dict[col] = row[i] if i < len(row) else None
            data_rows.append(row_dict)

        # Valori distinti per colonne output
        distinct_values = {}
        for col in output_columns:
            vals = set()
            for row in data_rows:
                v = row.get(col)
                if v is not None:
                    vals.add(str(v))
            distinct_values[col] = sorted(vals)

        # Candidati campi configuratore per colonne chiave
        campi_db = []
        try:
            result = db.execute(text(
                "SELECT codice, etichetta, sezione, tipo FROM campi_configuratore WHERE attivo=1"
            ))
            campi_db = [{"codice": r[0], "etichetta": r[1], "sezione": r[2], "tipo": r[3]}
                        for r in result.fetchall()]
        except Exception:
            pass

        def find_candidates(keyword: str, limit: int = 5):
            """Trova i migliori campi candidati per una keyword."""
            if not keyword or not campi_db:
                return []
            kw_lower = keyword.lower().replace(" ", "_").replace(".", "_")
            kw_parts = [p for p in kw_lower.split("_") if len(p) > 1]

            scored = []
            for campo in campi_db:
                codice = (campo["codice"] or "").lower()
                etichetta = (campo["etichetta"] or "").lower()
                sezione = (campo["sezione"] or "").lower()
                full_ref = codice if codice.startswith(f"{sezione}.") else f"{sezione}.{codice}"

                score = 0
                if kw_lower == codice or kw_lower == codice.split(".")[-1]:
                    score = 100
                elif kw_lower in codice:
                    score = 80
                elif all(p in codice for p in kw_parts):
                    score = 70
                elif kw_lower in etichetta:
                    score = 60
                elif any(p in codice for p in kw_parts if len(p) > 2):
                    score = 40
                elif any(p in etichetta for p in kw_parts if len(p) > 2):
                    score = 30

                if score > 0:
                    scored.append({
                        "field": full_ref,
                        "codice": campo["codice"],
                        "etichetta": campo["etichetta"],
                        "sezione": sezione,
                        "tipo": campo["tipo"],
                        "score": score,
                    })

            scored.sort(key=lambda x: -x["score"])
            return scored[:limit]

        field_candidates = {}
        for col in key_columns:
            field_candidates[col] = find_candidates(col)

        # All fields (per compositi e mapping manuale)
        all_fields = []
        for campo in campi_db:
            sezione = (campo["sezione"] or "").lower()
            codice = campo["codice"] or ""
            full_ref = codice if codice.startswith(f"{sezione}.") else f"{sezione}.{codice}"
            all_fields.append({
                "field": full_ref,
                "codice": codice,
                "etichetta": campo["etichetta"],
                "sezione": sezione,
                "tipo": campo["tipo"],
            })

        wb.close()
        return {
            "success": True,
            "total_rows": len(data_rows),
            "distinct_values": distinct_values,
            "field_candidates": field_candidates,
            "all_fields": all_fields,
        }
    finally:
        os.unlink(tmp_path)


@app.post("/import-excel/genera-v3")
async def import_excel_genera_v3(
    file: UploadFile = File(...),
    config_json: str = Form(...),
):
    """
    V3 Step 3→4: Genera data table JSON (lookup_multi) + regola JSON.
    
    config_json contiene:
    - nome_tabella, sheet, header_row
    - key_columns, output_columns, key_configs (match type per colonna)
    - field_mappings: {col_excel: {type, field/fields, separator}}
    - value_mappings: {valore: {tipo, codice_articolo, ...}}
    - conditions
    """
    import tempfile, shutil, openpyxl
    from datetime import datetime

    config = json.loads(config_json)
    nome_tabella = config.get("nome_tabella", "imported_table")
    sheet_name = config.get("sheet")
    header_row = config.get("header_row", 1)
    key_columns = config.get("key_columns", [])
    output_columns = config.get("output_columns", [])
    key_configs = config.get("key_configs", {})
    field_mappings = config.get("field_mappings", {})
    value_mappings = config.get("value_mappings", {})
    conditions = config.get("conditions", [])

    suffix = os.path.splitext(file.filename)[1] or ".xlsx"
    with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
        shutil.copyfileobj(file.file, tmp)
        tmp_path = tmp.name

    import_timestamp = datetime.now().isoformat()

    try:
        wb = openpyxl.load_workbook(tmp_path, data_only=True)
        ws = wb[sheet_name] if sheet_name and sheet_name in wb.sheetnames else wb[wb.sheetnames[0]]

        all_rows = []
        for row in ws.iter_rows(values_only=True):
            all_rows.append(list(row))

        h_idx = max(0, header_row - 1)
        raw_headers = all_rows[h_idx] if h_idx < len(all_rows) else []
        columns = [str(h) if h is not None else f"Col_{i+1}" for i, h in enumerate(raw_headers)]

        data_rows = []
        for row in all_rows[h_idx + 1:]:
            if all(v is None for v in row):
                continue
            row_dict = {}
            for i, col in enumerate(columns):
                row_dict[col] = row[i] if i < len(row) else None
            data_rows.append(row_dict)

        wb.close()

        # Normalizza nome colonna per JSON
        def norm(s):
            import re as _re
            s = str(s).lower().strip()
            s = _re.sub(r'[^a-z0-9]+', '_', s)
            return s.strip('_')

        # Costruisci chiavi per data table
        chiavi = []
        for col in key_columns:
            match_type = key_configs.get(col, "exact")
            chiavi.append({
                "colonna": norm(col),
                "colonna_originale": col,
                "match": match_type,
            })

        # Costruisci records
        all_data = []
        for row in data_rows:
            record = {}
            # Chiavi
            for col in key_columns:
                val = row.get(col)
                if val is not None:
                    try:
                        record[norm(col)] = float(val)
                    except (ValueError, TypeError):
                        record[norm(col)] = str(val).strip()
                else:
                    record[norm(col)] = None

            # Output
            output = {}
            materiali_record = []
            for col in output_columns:
                val = row.get(col)
                if val is None:
                    continue
                val_str = str(val).strip()
                col_norm = norm(col)

                # Controlla se questo valore ha un mapping articolo
                vm = value_mappings.get(val_str, {})
                if vm.get("tipo") == "articolo":
                    output[col_norm] = val_str

                    # Supporta sia formato vecchio (codice_articolo singolo) che nuovo (articoli array)
                    vm_articoli = vm.get("articoli", [])
                    if not vm_articoli and vm.get("codice_articolo"):
                        # Backward compat: singolo articolo → array
                        vm_articoli = [{
                            "codice": vm["codice_articolo"],
                            "descrizione": vm.get("descrizione_articolo", ""),
                            "quantita": 1,
                        }]

                    for art in vm_articoli:
                        materiali_record.append({
                            "codice": art.get("codice", ""),
                            "descrizione": art.get("descrizione", ""),
                            "quantita": art.get("quantita", 1),
                            "categoria": art.get("categoria", "Materiale Automatico"),
                            "unita_misura": art.get("unita_misura", "pz"),
                            "from_output": col_norm,
                            "from_value": val_str,
                        })
                else:
                    # Parametro o non mappato
                    try:
                        output[col_norm] = float(val)
                    except (ValueError, TypeError):
                        output[col_norm] = val_str

            record["output"] = output
            if materiali_record:
                record["materiali"] = materiali_record
            all_data.append(record)

        # Salva data table
        data_dir = "./data"
        os.makedirs(data_dir, exist_ok=True)
        table_path = os.path.join(data_dir, f"{nome_tabella}.json")

        table_json = {
            "tipo": "lookup_multi",
            "chiavi": chiavi,
            "records": all_data,
            "_meta": {
                "nome": nome_tabella,
                "generato_il": import_timestamp,
                "file_origine": file.filename,
                "foglio": sheet_name,
                "righe": len(all_data),
                "colonne_chiave": key_columns,
                "colonne_output": output_columns,
            }
        }
        with open(table_path, "w", encoding="utf-8") as f:
            json.dump(table_json, f, indent=2, ensure_ascii=False)

        # Genera regola
        rules_dir = "./rules"
        os.makedirs(rules_dir, exist_ok=True)

        rule_id = f"LOOKUP_{nome_tabella.upper()}"
        has_todo = False

        # Costruisci input_fields
        input_fields = []
        warnings = []
        for col in key_columns:
            kc = norm(col)
            match_type = key_configs.get(col, "exact")
            fm = field_mappings.get(col, {})

            if fm.get("type") == "composite" and fm.get("fields"):
                comp_fields = fm["fields"]
                # Validazione: campi duplicati nel composite
                if len(comp_fields) != len(set(comp_fields)):
                    warnings.append(
                        f"Colonna '{col}': campo composito ha campi duplicati {comp_fields}. "
                        f"Ogni parte del composite deve mappare un campo diverso del configuratore."
                    )
                    has_todo = True
                # Validazione: campi vuoti
                if any(not f or f.startswith("TODO") for f in comp_fields):
                    has_todo = True
                input_fields.append({
                    "colonna_tabella": kc,
                    "match": match_type,
                    "type": "composite",
                    "fields": comp_fields,
                    "separator": fm.get("separator", "_"),
                })
            elif fm.get("field"):
                input_fields.append({
                    "colonna_tabella": kc,
                    "match": match_type,
                    "type": "single",
                    "field": fm["field"],
                })
            else:
                input_fields.append({
                    "colonna_tabella": kc,
                    "match": match_type,
                    "type": "single",
                    "field": f"TODO.{kc}",
                })
                has_todo = True

        rule = {
            "id": rule_id,
            "name": f"Lookup {nome_tabella.replace('_', ' ').title()}",
            "description": f"Lookup generata da {file.filename} — foglio {sheet_name}",
            "version": "2.0",
            "enabled": not has_todo,
            "priority": 10,
            "conditions": conditions,
            "actions": [{
                "action": "lookup_multi",
                "tabella": nome_tabella,
                "input_fields": input_fields,
                "output_prefix": f"_calc.{nome_tabella}.",
            }],
            "materials": [],
            "_source": "excel_import_v3",
            "_imported_at": import_timestamp,
            "_source_file": file.filename,
            "_field_mappings": field_mappings,
            "_value_mappings": value_mappings,
        }

        rule_path = os.path.join(rules_dir, f"rule_{rule_id}.json")
        with open(rule_path, "w", encoding="utf-8") as f:
            json.dump(rule, f, indent=2, ensure_ascii=False)

        return {
            "success": True,
            "data_table": nome_tabella,
            "data_table_path": table_path,
            "rule_id": rule_id,
            "rule_path": rule_path,
            "has_todo": has_todo,
            "rows_imported": len(all_data),
            "value_mappings_count": len([v for v in value_mappings.values() if isinstance(v, dict) and v.get("tipo") == "articolo"]),
            "materiali_count": sum(len(r.get("materiali", [])) for r in all_data),
            "warnings": warnings,
        }
    finally:
        os.unlink(tmp_path)

# ============================================================================
# IMPORT EXCEL CON _MAPPA
# ============================================================================

# ==========================================
# IMPORT EXCEL — PREVIEW + GENERA
# (usati dalla pagina "Importa da Excel")
# ==========================================
@app.post("/import-excel/preview")
async def import_excel_preview(file: UploadFile = File(...)):
    """
    Fase 2 Step 1: analizza l'Excel, valida _MAPPA, restituisce preview
    con anteprima dati, colonne, fasce calcolate.
    """
    import tempfile, shutil

    suffix = os.path.splitext(file.filename)[1] or ".xlsx"
    with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
        shutil.copyfileobj(file.file, tmp)
        tmp_path = tmp.name

    try:
        loader = ExcelDataLoader()
        validation = loader.validate_excel(tmp_path)

        if not validation.get("valid", False):
            return {
                "success": False,
                "filename": file.filename,
                "mappa": [],
                "errors": validation.get("errors", []),
                "warnings": validation.get("warnings", []),
            }

        # Arricchisci con anteprima dati e fasce calcolate
        import openpyxl
        wb = openpyxl.load_workbook(tmp_path, data_only=True)

        mappa_out = []
        for tab_info in validation.get("tabelle", []):
            for foglio_info in tab_info.get("fogli", []):
                foglio_nome = foglio_info.get("foglio", "")
                if not foglio_info.get("esiste"):
                    continue

                ws = wb[foglio_nome]
                rows_raw = []
                for row in ws.iter_rows(values_only=True):
                    if all(v is None for v in row):
                        continue
                    rows_raw.append(list(row))

                # Headers e dati
                headers_raw = [str(h) for h in rows_raw[0]] if rows_raw else []
                headers_norm = [ExcelDataLoader._normalize_key(h) for h in rows_raw[0]] if rows_raw else []

                col_chiave = tab_info.get("colonna_chiave", "").strip()

                # Separa colonne tecniche e articoli
                colonne_tecniche = []
                colonne_articoli = []
                for h_raw, h_norm in zip(headers_raw, headers_norm):
                    if h_norm.startswith("art_"):
                        colonne_articoli.append(h_raw)
                    else:
                        colonne_tecniche.append(h_raw)

                # Anteprima (prime 5 righe)
                anteprima = []
                for row in rows_raw[1:6]:
                    row_dict = {}
                    art_dict = {}
                    for j, h_raw in enumerate(headers_raw):
                        val = row[j] if j < len(row) else None
                        h_norm = headers_norm[j] if j < len(headers_norm) else ""
                        if h_norm.startswith("art_"):
                            if val is not None:
                                art_dict[h_raw] = val
                        else:
                            row_dict[h_raw] = val
                    if art_dict:
                        row_dict["_articoli"] = art_dict
                    anteprima.append(row_dict)

                # Fasce calcolate (per lookup_range)
                fasce = None
                col_chiave_norm = ExcelDataLoader._normalize_key(col_chiave) if col_chiave else ""
                if tab_info.get("tipo") == "lookup_range" and col_chiave_norm:
                    idx = None
                    for k, h in enumerate(headers_norm):
                        if h == col_chiave_norm:
                            idx = k
                            break
                    if idx is not None:
                        values = []
                        for row in rows_raw[1:]:
                            v = row[idx] if idx < len(row) else None
                            if v is not None:
                                try:
                                    values.append(float(v))
                                except (ValueError, TypeError):
                                    pass
                        values.sort()
                        if values:
                            fasce = []
                            for i, val in enumerate(values):
                                da = 0 if i == 0 else round((values[i-1] + val) / 2, 4)
                                a = round((val + values[i+1]) / 2, 4) if i + 1 < len(values) else None
                                fasce.append({"valore": val, "da": da, "a": a})

                entry = {
                    "foglio": foglio_nome,
                    "tipo": tab_info.get("tipo", ""),
                    "nome_tabella": tab_info.get("nome_tabella", ""),
                    "colonna_chiave": col_chiave,
                    "tipo_chiave": tab_info.get("tipo_chiave", ""),
                    "partizionato_per": tab_info.get("partizionato_per") or None,
                    "valore_partizione": foglio_info.get("valore_partizione") or None,
                    "righe": foglio_info.get("righe", 0),
                    "colonne_tecniche": colonne_tecniche,
                    "colonne_articoli": colonne_articoli,
                    "anteprima": anteprima,
                }
                if fasce:
                    entry["fasce_calcolate"] = fasce

                mappa_out.append(entry)

        wb.close()

        return {
            "success": True,
            "filename": file.filename,
            "mappa": mappa_out,
            "errors": validation.get("errors", []),
            "warnings": validation.get("warnings", []),
        }
    finally:
        os.unlink(tmp_path)


@app.post("/import-excel/genera")
async def import_excel_genera(file: UploadFile = File(...), db: Session = Depends(get_db)):
    """
    Fase 2 Step 2: genera data tables JSON + regole con matching intelligente.
    
    Supporta tutti i tipi di tabella:
    - lookup_range → regola lookup_table + bozza MAT_ per colonne ART
    - constants    → regola lookup_table (chiave testuale → valori)
    - catalog      → bozza catalog_match (selezione multi-criterio, disabilitata)
    
    Ogni regola generata include _source="excel_import" e _imported_at per
    tracciabilità nel Rule Builder.
    """
    import tempfile, shutil
    from datetime import datetime

    suffix = os.path.splitext(file.filename)[1] or ".xlsx"
    with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
        shutil.copyfileobj(file.file, tmp)
        tmp_path = tmp.name

    # Timestamp unico per questa sessione di import
    import_timestamp = datetime.now().isoformat()

    try:
        # 1. Genera data tables
        loader = ExcelDataLoader()
        result = loader.load_excel(tmp_path, overwrite=True)

        # Salva Excel originale
        originals_dir = os.path.join("./data", "excel_originals")
        os.makedirs(originals_dir, exist_ok=True)
        dest = os.path.join(originals_dir, file.filename)
        shutil.copy2(tmp_path, dest)
        result["original_saved"] = dest

        # 2. Per ogni tabella generata, genera regola con matching intelligente
        rules_generated = []
        rules_details = []

        if result.get("success"):
            # Carica campi configuratore per matching
            campi_db = []
            try:
                campi_result = db.execute(text(
                    "SELECT codice, etichetta, sezione, tipo FROM campi_configuratore WHERE attivo=1"
                ))
                campi_db = [{"codice": r[0], "etichetta": r[1], "sezione": r[2], "tipo": r[3]}
                            for r in campi_result.fetchall()]
            except Exception:
                pass

            def find_best_match(keyword: str, prefer_tipo: str = None) -> tuple:
                """Trova il miglior campo candidato. Restituisce (field_ref, score)."""
                if not keyword or not campi_db:
                    return ("", 0)
                kw_lower = keyword.lower().replace(" ", "_").replace(".", "_")
                kw_parts = [p for p in kw_lower.split("_") if len(p) > 1]

                best_score = 0
                best_field = ""
                for campo in campi_db:
                    codice = (campo["codice"] or "").lower()
                    etichetta = (campo["etichetta"] or "").lower()
                    sezione = (campo["sezione"] or "").lower()

                    if codice.startswith(f"{sezione}."):
                        full_ref = codice
                    else:
                        full_ref = f"{sezione}.{codice}"

                    score = 0
                    if kw_lower == codice or kw_lower == codice.split(".")[-1]:
                        score = 100
                    elif kw_lower in codice:
                        score = 80
                    elif all(p in codice for p in kw_parts):
                        score = 70
                    elif kw_lower in etichetta:
                        score = 60
                    elif any(p in codice for p in kw_parts if len(p) > 2):
                        score = 40
                    elif any(p in etichetta for p in kw_parts if len(p) > 2):
                        score = 30

                    if prefer_tipo and campo["tipo"] == prefer_tipo:
                        score += 5

                    if score > best_score:
                        best_score = score
                        best_field = full_ref

                return (best_field, best_score)

            def stamp_and_save(rule: dict, rule_id: str, has_todo: bool, is_draft: bool,
                               input_field: str, input_score: int,
                               partition_field: str = "", partition_score: int = 0):
                """Aggiunge metadati di import, salva su disco e registra nei risultati."""
                # Metadati tracciabilità
                rule["_source"] = "excel_import"
                rule["_imported_at"] = import_timestamp
                rule["_source_file"] = file.filename

                rule_path = os.path.join(rules_dir, f"rule_{rule_id}.json")
                with open(rule_path, "w", encoding="utf-8") as f:
                    json.dump(rule, f, indent=2, ensure_ascii=False)

                rules_generated.append(rule_id)
                rules_details.append({
                    "rule_id": rule_id,
                    "has_todo": has_todo,
                    "is_draft": is_draft,
                    "input_field": input_field if isinstance(input_field, str) else "",
                    "input_score": input_score,
                    "partition_field": partition_field if isinstance(partition_field, str) else "",
                    "partition_score": partition_score,
                })

            rules_dir = "./rules"
            os.makedirs(rules_dir, exist_ok=True)

            for nome_tabella in result.get("tables_generated", []):
                try:
                    table = loader.load_table(nome_tabella)
                    if not table:
                        continue

                    meta = table.get("_meta", {})
                    tipo = table.get("tipo", "")
                    parametro_lookup = table.get("parametro_lookup", "")
                    partizionato_per = table.get("partizionato_per", "")
                    colonna_chiave_orig = meta.get("colonna_chiave", parametro_lookup)

                    # =============================================
                    # TIPO: lookup_range (numerico con fasce)
                    # =============================================
                    if tipo == "lookup_range":
                        input_field, input_score = find_best_match(parametro_lookup, "numero")
                        if not input_field and colonna_chiave_orig:
                            input_field, input_score = find_best_match(colonna_chiave_orig, "numero")

                        partition_field = ""
                        partition_score = 0
                        if partizionato_per:
                            partition_field, partition_score = find_best_match(partizionato_per, "dropdown")

                        rule_id = f"LOOKUP_{nome_tabella.upper()}"
                        conditions = []
                        if input_field:
                            conditions.append({
                                "field": input_field,
                                "operator": "greater_than",
                                "value": 0
                            })

                        action = {
                            "action": "lookup_table",
                            "tabella": nome_tabella,
                            "input_field": input_field or f"TODO.{parametro_lookup}",
                            "output_prefix": f"_calc.{nome_tabella}.",
                        }
                        if partizionato_per:
                            action["partition_field"] = partition_field or f"TODO.{partizionato_per}"

                        has_todo = "TODO." in action.get("input_field", "") or "TODO." in action.get("partition_field", "")

                        rule = {
                            "id": rule_id,
                            "name": f"Lookup {nome_tabella.replace('_', ' ').title()}",
                            "description": f"Cerca nella tabella {nome_tabella} per {colonna_chiave_orig or parametro_lookup}",
                            "version": "1.0",
                            "enabled": True,
                            "priority": 10,
                            "conditions": conditions,
                            "actions": [action],
                            "materials": [],
                        }

                        stamp_and_save(rule, rule_id, has_todo=has_todo, is_draft=False,
                                       input_field=action.get("input_field", ""),
                                       input_score=input_score,
                                       partition_field=action.get("partition_field", ""),
                                       partition_score=partition_score)

                        # ===== GENERA BOZZA REGOLA MATERIALI (solo per lookup_range con ART) =====
                        art_columns = []
                        ranges_data = None
                        if "partizioni" in table:
                            first_part = list(table["partizioni"].values())[0]
                            if first_part:
                                ranges_data = first_part
                        elif "ranges" in table:
                            ranges_data = table["ranges"]
                        if ranges_data and len(ranges_data) > 0:
                            art_columns = list(ranges_data[0].get("articoli", {}).keys())

                        if art_columns:
                            mat_rule_id = f"MAT_{nome_tabella.upper()}"
                            output_prefix = f"_calc.{nome_tabella}."

                            materials_list = []
                            for art_col in art_columns:
                                ctx_key = f"{output_prefix}{art_col}"
                                desc_parts = art_col.replace("art_", "").replace("_", " ").strip()
                                materials_list.append({
                                    "codice": "{{" + ctx_key + "}}",
                                    "descrizione": f"{desc_parts.title()} ({nome_tabella})",
                                    "quantita": 1,
                                    "categoria": nome_tabella.replace("_", " ").title(),
                                    "unita_misura": "pz",
                                })

                            mat_rule = {
                                "id": mat_rule_id,
                                "name": f"Materiali {nome_tabella.replace('_', ' ').title()}",
                                "description": f"Aggiunge materiali basati sul lookup {nome_tabella}. BOZZA: verificare condizioni e scenari.",
                                "version": "1.0",
                                "enabled": False,
                                "priority": 50,
                                "conditions": [
                                    {
                                        "field": f"{output_prefix}{art_columns[0]}",
                                        "operator": "is_not_empty",
                                        "_hint": "Verifica che il lookup abbia prodotto risultati"
                                    }
                                ],
                                "materials": materials_list,
                                "_hints": {
                                    "nota": "BOZZA AUTOMATICA - Da personalizzare:",
                                    "suggerimenti": [
                                        "Questa regola è DISABILITATA. Abilitatela dopo averla verificata.",
                                        "Aggiungete condizioni per lo scenario (es: tipo_avviamento = 'diretto')",
                                        "Se servono regole separate per scenario, duplicate e filtrate",
                                        "Verificate che i codici articolo esistano nel sistema"
                                    ],
                                    "output_disponibili": list(
                                        (ranges_data[0].get("output", {}).keys()) if ranges_data else []
                                    ),
                                    "articoli_disponibili": art_columns,
                                }
                            }

                            stamp_and_save(mat_rule, mat_rule_id, has_todo=False, is_draft=True,
                                           input_field=f"{len(art_columns)} materiali",
                                           input_score=0)

                    # =============================================
                    # TIPO: constants (chiave testuale → valori)
                    # =============================================
                    elif tipo == "constants":
                        input_field, input_score = find_best_match(parametro_lookup, "dropdown")
                        if not input_field and colonna_chiave_orig:
                            input_field, input_score = find_best_match(colonna_chiave_orig, "dropdown")

                        rule_id = f"LOOKUP_{nome_tabella.upper()}"
                        conditions = []
                        if input_field:
                            conditions.append({
                                "field": input_field,
                                "operator": "is_not_empty",
                            })

                        action = {
                            "action": "lookup_table",
                            "tabella": nome_tabella,
                            "input_field": input_field or f"TODO.{parametro_lookup}",
                            "output_prefix": f"_calc.{nome_tabella}.",
                        }

                        has_todo = "TODO." in action.get("input_field", "")

                        chiavi = list(table.get("valori", {}).keys())
                        desc_chiavi = ", ".join(chiavi[:5])
                        if len(chiavi) > 5:
                            desc_chiavi += f" ... ({len(chiavi)} totali)"

                        rule = {
                            "id": rule_id,
                            "name": f"Lookup {nome_tabella.replace('_', ' ').title()}",
                            "description": f"Cerca costanti {nome_tabella} per {colonna_chiave_orig}. Chiavi: {desc_chiavi}",
                            "version": "1.0",
                            "enabled": True,
                            "priority": 10,
                            "conditions": conditions,
                            "actions": [action],
                            "materials": [],
                        }

                        stamp_and_save(rule, rule_id, has_todo=has_todo, is_draft=False,
                                       input_field=action.get("input_field", ""),
                                       input_score=input_score)

                    # =============================================
                    # TIPO: catalog (selezione multi-criterio)
                    # =============================================
                    elif tipo == "catalog":
                        colonne = table.get("colonne", [])
                        colonna_id = table.get("colonna_id", "codice")
                        records = table.get("records", [])

                        rule_id = f"CATALOG_{nome_tabella.upper()}"

                        output_mapping = {}
                        for col in colonne:
                            output_mapping[f"_calc.{nome_tabella}.{col}"] = col

                        action = {
                            "action": "catalog_match",
                            "tabella": nome_tabella,
                            "criteri_fissi": [
                                {
                                    "colonna": "TODO_COLONNA",
                                    "operatore": ">=",
                                    "valore": "TODO_VALORE",
                                    "_hint": f"Colonne disponibili: {', '.join(colonne)}"
                                }
                            ],
                            "ordinamento": {
                                "colonna": colonne[1] if len(colonne) > 1 else colonna_id,
                                "direzione": "ASC"
                            },
                            "output": output_mapping,
                        }

                        rule = {
                            "id": rule_id,
                            "name": f"Selezione {nome_tabella.replace('_', ' ').title()}",
                            "description": (
                                f"Cerca nel catalogo {nome_tabella} ({len(records)} record). "
                                f"BOZZA: completare criteri di selezione nel Rule Designer."
                            ),
                            "version": "1.0",
                            "enabled": False,
                            "priority": 20,
                            "conditions": [],
                            "actions": [action],
                            "materials": [],
                        }

                        stamp_and_save(rule, rule_id, has_todo=True, is_draft=True,
                                       input_field=f"{len(colonne)} colonne",
                                       input_score=0)

                    else:
                        result.setdefault("warnings", []).append(
                            f"Tipo '{tipo}' non supportato per generazione regola '{nome_tabella}'"
                        )
                        continue

                except Exception as e:
                    result.setdefault("warnings", []).append(
                        f"Errore generazione regola per '{nome_tabella}': {e}"
                    )

        result["rules_generated"] = rules_generated
        result["rules_details"] = rules_details

        return result
    finally:
        os.unlink(tmp_path)


@app.get("/import-excel/tables")
def list_imported_tables():
    """Lista le data tables generate da import Excel."""
    tables = []
    data_dir = "./data"
    if not os.path.exists(data_dir):
        return {"tables": []}

    for filename in sorted(os.listdir(data_dir)):
        if filename.endswith(".json"):
            filepath = os.path.join(data_dir, filename)
            try:
                with open(filepath, "r", encoding="utf-8") as f:
                    data = json.load(f)
                meta = data.get("_meta", {})
                tables.append({
                    "filename": filename,
                    "nome": meta.get("nome", filename.replace(".json", "")),
                    "tipo": data.get("tipo", "sconosciuto"),
                    "generato_il": meta.get("generato_il", ""),
                    "file_origine": meta.get("file_origine", ""),
                    "righe": meta.get("righe", ""),
                })
            except Exception:
                pass

    return {"tables": tables}

from fastapi.responses import FileResponse

@app.get("/import-excel/esempio/{filename}")
def download_esempio_excel(filename: str):
    """Scarica un file Excel di esempio."""
    import re
    # Sicurezza: solo .xlsx, no path traversal
    if not re.match(r'^[a-zA-Z0-9_\-]+\.xlsx$', filename):
        raise HTTPException(status_code=400, detail="Nome file non valido")
    filepath = os.path.join("examples", filename)
    if not os.path.exists(filepath):
        raise HTTPException(status_code=404, detail=f"File '{filename}' non trovato in examples/")
    return FileResponse(
        filepath,
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        filename=filename
    )

@app.post("/import-excel/parse-sheet")
async def import_excel_parse_sheet(
    file: UploadFile = File(...),
    sheet: str = Query(None, description="Nome foglio (default: primo)"),
    header_row: int = Query(1, description="Riga intestazioni (1-based)"),
):
    """
    Step 1: Parsa un foglio Excel e restituisce colonne + righe dati.
    """
    import tempfile, shutil, openpyxl

    suffix = os.path.splitext(file.filename)[1] or ".xlsx"
    with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
        shutil.copyfileobj(file.file, tmp)
        tmp_path = tmp.name

    try:
        wb = openpyxl.load_workbook(tmp_path, data_only=True, read_only=True)
        all_sheets = [s for s in wb.sheetnames if not s.startswith("_")]

        target = sheet or (all_sheets[0] if all_sheets else wb.sheetnames[0])
        if target not in wb.sheetnames:
            raise HTTPException(400, f"Foglio '{target}' non trovato. Disponibili: {wb.sheetnames}")

        ws = wb[target]
        rows_raw = []
        for i, row in enumerate(ws.iter_rows(values_only=True)):
            if i >= header_row - 1 + 500:
                break
            rows_raw.append(list(row))
        wb.close()

        if len(rows_raw) < header_row:
            return {"sheets": all_sheets, "columns": [], "rows": [], "total_rows": 0}

        headers = []
        for j, h in enumerate(rows_raw[header_row - 1]):
            if h is not None and str(h).strip():
                headers.append(str(h).strip())
            else:
                headers.append(f"col_{j+1}")

        data_rows = []
        for row in rows_raw[header_row:]:
            d = {}
            has_data = False
            for j, h in enumerate(headers):
                val = row[j] if j < len(row) else None
                if val is not None:
                    if isinstance(val, float) and val == int(val):
                        val = int(val)
                    d[h] = val
                    has_data = True
                else:
                    d[h] = None
            if has_data:
                data_rows.append(d)

        return {
            "sheets": all_sheets,
            "selected_sheet": target,
            "header_row": header_row,
            "columns": headers,
            "rows": data_rows[:50],
            "total_rows": len(data_rows),
        }
    finally:
        os.unlink(tmp_path)


@app.post("/import-excel/analyze-v3")
async def import_excel_analyze_v3(
    file: UploadFile = File(...),
    config_json: str = Form(...),
    db: Session = Depends(get_db),
):
    """
    Step 2→3: Analizza struttura + estrae valori distinti + candidati campi.
    
    config_json:
    {
        "sheet": "Foglio1",
        "header_row": 3,
        "key_columns": ["kW", "avviamento_freq"],
        "output_columns": ["contattore_km", "morsetti", "filo", "soft_starter"]
    }
    """
    import tempfile, shutil, openpyxl

    config = json.loads(config_json)
    sheet_name = config.get("sheet", "")
    header_row = config.get("header_row", 1)
    key_columns = config.get("key_columns", [])
    output_columns = config.get("output_columns", [])

    suffix = os.path.splitext(file.filename)[1] or ".xlsx"
    with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
        shutil.copyfileobj(file.file, tmp)
        tmp_path = tmp.name

    try:
        wb = openpyxl.load_workbook(tmp_path, data_only=True, read_only=True)
        if sheet_name not in wb.sheetnames:
            raise HTTPException(400, f"Foglio '{sheet_name}' non trovato")

        ws = wb[sheet_name]
        rows_raw = list(ws.iter_rows(values_only=True))
        wb.close()

        if len(rows_raw) < header_row:
            return {"success": False, "errors": ["Riga intestazioni fuori range"]}

        headers = []
        for j, h in enumerate(rows_raw[header_row - 1]):
            headers.append(str(h).strip() if h is not None and str(h).strip() else f"col_{j+1}")

        all_data = []
        for row in rows_raw[header_row:]:
            d = {}
            has_data = False
            for j, h in enumerate(headers):
                val = row[j] if j < len(row) else None
                if val is not None:
                    if isinstance(val, float) and val == int(val):
                        val = int(val)
                    d[h] = val
                    has_data = True
            if has_data:
                all_data.append(d)

        # Valori distinti per colonne output
        distinct_values = {}
        for col in output_columns:
            vals = set()
            for row in all_data:
                v = row.get(col)
                if v is not None and str(v).strip():
                    vals.add(str(v).strip())
            distinct_values[col] = sorted(vals)

        # Campi configuratore per matching
        campi_db = []
        try:
            campi_result = db.execute(text(
                "SELECT codice, etichetta, sezione, tipo, gruppo_dropdown "
                "FROM campi_configuratore WHERE attivo=1 ORDER BY sezione, ordine"
            ))
            campi_db = [
                {"codice": r[0], "etichetta": r[1], "sezione": r[2],
                 "tipo": r[3], "gruppo_dropdown": r[4]}
                for r in campi_result.fetchall()
            ]
        except Exception as e:
            print(f"Errore caricamento campi: {e}")

        def find_candidates(keyword, prefer_tipo=None):
            if not keyword or not campi_db:
                return []
            kw_lower = keyword.lower().replace(" ", "_").replace(".", "_")
            kw_parts = [p for p in kw_lower.split("_") if len(p) > 1]
            candidates = []
            for campo in campi_db:
                codice = (campo["codice"] or "").lower()
                etichetta = (campo["etichetta"] or "").lower()
                sezione = (campo["sezione"] or "").lower()
                full_ref = codice if codice.startswith(f"{sezione}.") else f"{sezione}.{codice}"
                score = 0
                if kw_lower == codice or kw_lower == codice.split(".")[-1]:
                    score = 100
                elif kw_lower in codice:
                    score = 80
                elif all(p in codice for p in kw_parts):
                    score = 70
                elif kw_lower in etichetta:
                    score = 60
                elif any(p in codice for p in kw_parts if len(p) > 2):
                    score = 40
                elif any(p in etichetta for p in kw_parts if len(p) > 2):
                    score = 30
                if prefer_tipo and campo["tipo"] == prefer_tipo:
                    score += 5
                if score > 0:
                    candidates.append({
                        "field": full_ref, "codice": campo["codice"],
                        "etichetta": campo["etichetta"], "sezione": sezione,
                        "tipo": campo["tipo"], "score": score,
                    })
            candidates.sort(key=lambda c: -c["score"])
            return candidates[:10]

        field_candidates = {}
        for col in key_columns:
            prefer = "numero" if key_columns.index(col) == 0 else "dropdown"
            field_candidates[col] = find_candidates(col, prefer_tipo=prefer)

        all_fields = []
        for c in campi_db:
            full_ref = c["codice"] if c["codice"].startswith(f"{c['sezione']}.") else f"{c['sezione']}.{c['codice']}"
            all_fields.append({
                "field": full_ref, "codice": c["codice"],
                "etichetta": c["etichetta"], "sezione": c["sezione"], "tipo": c["tipo"],
            })

        return {
            "success": True,
            "total_rows": len(all_data),
            "distinct_values": distinct_values,
            "field_candidates": field_candidates,
            "all_fields": all_fields,
        }
    finally:
        os.unlink(tmp_path)


@app.post("/import-excel/genera-v3")
async def import_excel_genera_v3(
    file: UploadFile = File(...),
    config_json: str = Form(...),
    db: Session = Depends(get_db),
):
    """
    Step finale: genera data table JSON + regola lookup.
    
    config_json:
    {
        "nome_tabella": "contattori_oleo",
        "sheet": "Dati",
        "header_row": 3,
        "key_columns": ["kW", "avviamento_freq"],
        "output_columns": ["contattore_km", "contattore_ks", "morsetti", "filo", "soft_starter"],
        "key_configs": {
            "kW": "lte",
            "avviamento_freq": "exact"
        },
        "field_mappings": {
            "kW": {"type": "single", "field": "argano.potenza_motore_kw"},
            "avviamento_freq": {
                "type": "composite",
                "fields": ["argano.tipo_avviamento", "tensioni.frequenza_rete"],
                "separator": "_"
            }
        },
        "value_mappings": {
            "D18": {"tipo": "articolo", "codice_articolo": "CONT-D18", "articolo_id": 123},
            "10":  {"tipo": "parametro"}
        },
        "conditions": []
    }
    """
    import tempfile, shutil, openpyxl
    from datetime import datetime

    config = json.loads(config_json)
    nome_tabella = config.get("nome_tabella", "imported_table")
    sheet_name = config.get("sheet", "")
    header_row = config.get("header_row", 1)
    key_columns = config.get("key_columns", [])
    output_columns = config.get("output_columns", [])
    key_configs = config.get("key_configs", {})          # {"kW": "lte", "avviamento_freq": "exact"}
    field_mappings = config.get("field_mappings", {})    # compositi supportati
    value_mappings = config.get("value_mappings", {})
    conditions = config.get("conditions", [])

    suffix = os.path.splitext(file.filename)[1] or ".xlsx"
    with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
        shutil.copyfileobj(file.file, tmp)
        tmp_path = tmp.name

    import_timestamp = datetime.now().isoformat()

    try:
        wb = openpyxl.load_workbook(tmp_path, data_only=True, read_only=True)
        if sheet_name not in wb.sheetnames:
            raise HTTPException(400, f"Foglio '{sheet_name}' non trovato")

        ws = wb[sheet_name]
        rows_raw = list(ws.iter_rows(values_only=True))
        wb.close()

        if len(rows_raw) < header_row:
            raise HTTPException(400, "Riga intestazioni fuori range")

        headers = []
        for j, h in enumerate(rows_raw[header_row - 1]):
            headers.append(str(h).strip() if h is not None and str(h).strip() else f"col_{j+1}")

        all_data = []
        for row in rows_raw[header_row:]:
            d = {}
            has_data = False
            for j, h in enumerate(headers):
                val = row[j] if j < len(row) else None
                if val is not None:
                    if isinstance(val, float) and val == int(val):
                        val = int(val)
                    d[h] = val
                    has_data = True
            if has_data:
                all_data.append(d)

        # Normalizza nome per JSON key
        import re as _re
        def norm(s):
            s = s.lower().strip()
            s = _re.sub(r'[^a-z0-9]+', '_', s)
            return s.strip('_')

        # =============================================
        # COSTRUZIONE DATA TABLE JSON
        # =============================================
        data_dir = "./data"
        os.makedirs(data_dir, exist_ok=True)

        # Struttura generica: array di record con chiavi + output + articoli
        # Il match type è salvato nei metadati, il rule engine lo usa a runtime
        records = []
        for row in all_data:
            record = {}

            # Chiavi
            for kc in key_columns:
                v = row.get(kc)
                if v is not None:
                    record[norm(kc)] = v if not isinstance(v, str) else str(v).strip()

            # Output
            output = {}
            articoli = {}
            for col in output_columns:
                v = row.get(col)
                if v is not None:
                    v_str = str(v).strip()
                    col_norm = norm(col)
                    output[col_norm] = v if not isinstance(v, str) else v_str
                    # Value mapping → articolo
                    if v_str in value_mappings:
                        vm = value_mappings[v_str]
                        if vm.get("tipo") == "articolo" and vm.get("codice_articolo"):
                            articoli[col_norm] = vm["codice_articolo"]

            record["output"] = output
            if articoli:
                record["articoli"] = articoli
            records.append(record)

        table_json = {
            "tipo": "lookup_multi",
            "chiavi": [
                {
                    "colonna": norm(kc),
                    "colonna_originale": kc,
                    "match": key_configs.get(kc, "exact"),
                }
                for kc in key_columns
            ],
            "records": records,
            "_meta": {
                "nome": nome_tabella,
                "generato_il": import_timestamp,
                "file_origine": file.filename,
                "foglio": sheet_name,
                "header_row": header_row,
                "righe": len(all_data),
                "colonne_chiave": key_columns,
                "colonne_output": output_columns,
                "key_configs": key_configs,
                "field_mappings": field_mappings,
                "value_mappings_count": len([v for v in value_mappings.values() if v.get("tipo") == "articolo"]),
            },
        }

        table_path = os.path.join(data_dir, f"{nome_tabella}.json")
        with open(table_path, "w", encoding="utf-8") as f:
            json.dump(table_json, f, indent=2, ensure_ascii=False)

        # Salva originale
        originals_dir = os.path.join(data_dir, "excel_originals")
        os.makedirs(originals_dir, exist_ok=True)
        shutil.copy2(tmp_path, os.path.join(originals_dir, file.filename))

        # =============================================
        # COSTRUZIONE REGOLA JSON
        # =============================================
        rules_dir = "./rules"
        os.makedirs(rules_dir, exist_ok=True)

        rule_id = f"LOOKUP_{nome_tabella.upper()}"

        # Costruisci input_fields dalla field_mappings
        # Ogni chiave Excel → uno o più campi configuratore
        input_fields = []
        has_todo = False

        for kc in key_columns:
            fm = field_mappings.get(kc, {})
            mt = key_configs.get(kc, "exact")

            if isinstance(fm, dict) and fm.get("type") == "composite":
                # Campo composto: N campi → concatenati con separatore
                comp_fields = fm.get("fields", [])
                separator = fm.get("separator", "_")
                if not comp_fields or any(not f for f in comp_fields):
                    has_todo = True
                input_fields.append({
                    "colonna_tabella": norm(kc),
                    "match": mt,
                    "type": "composite",
                    "fields": comp_fields,
                    "separator": separator,
                })
            elif isinstance(fm, dict) and fm.get("type") == "single":
                field_ref = fm.get("field", "")
                if not field_ref:
                    has_todo = True
                    field_ref = f"TODO.{norm(kc)}"
                input_fields.append({
                    "colonna_tabella": norm(kc),
                    "match": mt,
                    "type": "single",
                    "field": field_ref,
                })
            elif isinstance(fm, str) and fm:
                # Retrocompatibilità: stringa semplice
                input_fields.append({
                    "colonna_tabella": norm(kc),
                    "match": mt,
                    "type": "single",
                    "field": fm,
                })
            else:
                has_todo = True
                input_fields.append({
                    "colonna_tabella": norm(kc),
                    "match": mt,
                    "type": "single",
                    "field": f"TODO.{norm(kc)}",
                })

        rule = {
            "id": rule_id,
            "name": f"Lookup {nome_tabella.replace('_', ' ').title()}",
            "description": f"Lookup generata da {file.filename} — foglio {sheet_name}",
            "version": "2.0",
            "enabled": not has_todo,
            "priority": 10,
            "conditions": conditions,
            "actions": [{
                "action": "lookup_multi",
                "tabella": nome_tabella,
                "input_fields": input_fields,
                "output_prefix": f"_calc.{nome_tabella}.",
            }],
            "materials": [],
            "_source": "excel_import_v3",
            "_imported_at": import_timestamp,
            "_source_file": file.filename,
            "_field_mappings": field_mappings,
            "_value_mappings": value_mappings,
        }

        rule_path = os.path.join(rules_dir, f"rule_{rule_id}.json")
        with open(rule_path, "w", encoding="utf-8") as f:
            json.dump(rule, f, indent=2, ensure_ascii=False)

        return {
            "success": True,
            "data_table": nome_tabella,
            "data_table_path": table_path,
            "rule_id": rule_id,
            "rule_path": rule_path,
            "has_todo": has_todo,
            "rows_imported": len(all_data),
            "value_mappings_count": len([v for v in value_mappings.values() if v.get("tipo") == "articolo"]),
            "warnings": [],
        }
    finally:
        os.unlink(tmp_path)

# ==========================================
# GENERA BOZZA REGOLA DA DATA TABLE
# ==========================================
@app.post("/data-tables/{nome_tabella}/genera-regola")
def genera_regola_da_tabella(nome_tabella: str, body: dict = None, db: Session = Depends(get_db)):
    """
    Genera una bozza di regola lookup a partire da una data table.
    
    Fase 3.1 delle istruzioni: il sistema propone la struttura,
    l'utente conferma/corregge.
    
    Analizza la data table, cerca nel configuratore i campi candidati
    per input_field e partition_field, e restituisce:
    - la regola bozza JSON
    - i candidati per ogni parametro (l'utente sceglie)
    - info sulla struttura della tabella
    
    Body opzionale:
    {
        "input_field": "argano.potenza_motore_kw",   // se l'utente ha già scelto
        "partition_field": "tensioni.frequenza_rete", // se l'utente ha già scelto
        "conditions": [...],                          // condizioni custom
        "save": true                                  // salva subito la regola
    }
    """
    body = body or {}
    
    # 1. Carica data table
    loader = ExcelDataLoader()
    table = loader.load_table(nome_tabella)
    if not table:
        raise HTTPException(404, f"Data table '{nome_tabella}' non trovata in ./data/")
    
    meta = table.get("_meta", {})
    tipo = table.get("tipo", "")
    parametro_lookup = table.get("parametro_lookup", "")
    partizionato_per = table.get("partizionato_per", "")
    
    # Colonna chiave originale dall'Excel (più leggibile per matching)
    colonna_chiave_orig = meta.get("colonna_chiave", parametro_lookup)
    
    # 2. Raccogli info sulla struttura output
    output_columns = []
    articoli_columns = []
    if tipo == "lookup_range":
        # Prendi le colonne dal primo range della prima partizione
        ranges = None
        if "partizioni" in table:
            first_part = list(table["partizioni"].values())[0]
            if first_part:
                ranges = first_part
        elif "ranges" in table:
            ranges = table["ranges"]
        
        if ranges and len(ranges) > 0:
            output_columns = list(ranges[0].get("output", {}).keys())
            articoli_columns = list(ranges[0].get("articoli", {}).keys())
    elif tipo in ("constants", "costanti"):
        valori = table.get("valori", {})
        if valori:
            first_val = list(valori.values())[0]
            output_columns = list(first_val.keys()) if isinstance(first_val, dict) else []
    elif tipo in ("catalog", "catalogo"):
        output_columns = table.get("colonne", [])
        articoli_columns = table.get("colonne_articoli", [])
    
    # 3. Cerca campi candidati nel configuratore
    campi_db = []
    try:
        result = db.execute(text(
            "SELECT codice, etichetta, sezione, tipo FROM campi_configuratore WHERE attivo=1"
        ))
        campi_db = [{"codice": r[0], "etichetta": r[1], "sezione": r[2], "tipo": r[3]} 
                    for r in result.fetchall()]
    except Exception as e:
        print(f"Errore caricamento campi: {e}")
    
    # Funzione di matching: cerca keyword nel codice e nell'etichetta
    def find_candidates(keyword: str, prefer_tipo: str = None) -> list:
        """Cerca campi che contengono la keyword, ordinati per rilevanza."""
        if not keyword:
            return []
        
        # Genera varianti della keyword per matching fuzzy
        kw_lower = keyword.lower().replace(" ", "_").replace(".", "_")
        kw_parts = [p for p in kw_lower.split("_") if len(p) > 1]
        
        candidates = []
        for campo in campi_db:
            codice = (campo["codice"] or "").lower()
            etichetta = (campo["etichetta"] or "").lower()
            sezione = (campo["sezione"] or "").lower()
            # Se il codice contiene già la sezione come prefisso, non raddoppiarla
            if codice.startswith(f"{sezione}."):
                full_ref = codice
            else:
                full_ref = f"{sezione}.{codice}"
            
            score = 0
            
            # Match esatto sul codice (migliore)
            if kw_lower == codice:
                score = 100
            # Keyword contenuta nel codice
            elif kw_lower in codice:
                score = 80
            # Codice contiene tutte le parti della keyword
            elif all(p in codice for p in kw_parts):
                score = 70
            # Keyword contenuta nell'etichetta
            elif kw_lower in etichetta:
                score = 60
            # Almeno una parte significativa della keyword nel codice
            elif any(p in codice for p in kw_parts if len(p) > 2):
                score = 40
            # Parte nell'etichetta
            elif any(p in etichetta for p in kw_parts if len(p) > 2):
                score = 30
            
            # Bonus se il tipo campo corrisponde
            if prefer_tipo and campo["tipo"] == prefer_tipo:
                score += 5
            
            if score > 0:
                candidates.append({
                    "field": full_ref,
                    "codice": campo["codice"],
                    "etichetta": campo["etichetta"],
                    "sezione": sezione,
                    "score": score,
                })
        
        # Ordina per score decrescente
        candidates.sort(key=lambda c: -c["score"])
        return candidates[:10]  # max 10 candidati
    
    # Cerca candidati per input_field (colonna chiave)
    # Usa sia il nome normalizzato che quello originale dall'Excel
    input_candidates = find_candidates(parametro_lookup, prefer_tipo="numero")
    if colonna_chiave_orig and colonna_chiave_orig.lower() != parametro_lookup:
        extra = find_candidates(colonna_chiave_orig, prefer_tipo="numero")
        # Merge senza duplicati
        seen = {c["field"] for c in input_candidates}
        for c in extra:
            if c["field"] not in seen:
                input_candidates.append(c)
                seen.add(c["field"])
    
    # Cerca candidati per partition_field
    partition_candidates = find_candidates(partizionato_per, prefer_tipo="dropdown")
    
    # 4. Determina i campi scelti (da body o dal miglior candidato)
    chosen_input = body.get("input_field") or (
        input_candidates[0]["field"] if input_candidates else f"TODO.{parametro_lookup}"
    )
    chosen_partition = body.get("partition_field") or (
        partition_candidates[0]["field"] if partition_candidates else (
            f"TODO.{partizionato_per}" if partizionato_per else ""
        )
    )
    
    # 5. Costruisci la bozza della regola
    rule_id = f"LOOKUP_{nome_tabella.upper()}"
    
    # Conditions di default: input_field deve essere > 0 (per lookup numerici)
    conditions = body.get("conditions", [])
    if not conditions and tipo == "lookup_range":
        conditions = [
            {
                "field": chosen_input,
                "operator": "greater_than",
                "value": 0
            }
        ]
    
    # Azione lookup
    action = {
        "action": "lookup_table",
        "tabella": nome_tabella,
        "input_field": chosen_input,
        "output_prefix": f"_calc.{nome_tabella}.",
    }
    if partizionato_per:
        action["partition_field"] = chosen_partition
    
    rule = {
        "id": rule_id,
        "name": f"Lookup {nome_tabella.replace('_', ' ').title()}",
        "description": f"Cerca nella tabella {nome_tabella} per {colonna_chiave_orig or parametro_lookup}",
        "version": "1.0",
        "enabled": True,
        "priority": 10,
        "conditions": conditions,
        "actions": [action],
        "materials": [],
    }
    
    # 6. Se richiesto, salva la regola
    saved = False
    if body.get("save"):
        rules_dir = "./rules"
        os.makedirs(rules_dir, exist_ok=True)
        filepath = os.path.join(rules_dir, f"rule_{rule_id}.json")
        with open(filepath, "w", encoding="utf-8") as f:
            json.dump(rule, f, indent=2, ensure_ascii=False)
        saved = True
    
    # 7. Risposta con bozza + candidati + info struttura
    has_todo = "TODO." in chosen_input or "TODO." in chosen_partition
    
    return {
        "rule": rule,
        "saved": saved,
        "has_todo": has_todo,
        "table_info": {
            "nome": nome_tabella,
            "tipo": tipo,
            "parametro_lookup": parametro_lookup,
            "colonna_chiave_originale": colonna_chiave_orig,
            "partizionato_per": partizionato_per,
            "valori_partizione": meta.get("valori_partizione", []),
            "output_columns": output_columns,
            "articoli_columns": articoli_columns,
            "righe_totali": meta.get("righe_totali", 0),
        },
        "matching": {
            "input_field": {
                "chosen": chosen_input,
                "candidates": input_candidates,
            },
            "partition_field": {
                "chosen": chosen_partition,
                "candidates": partition_candidates,
            } if partizionato_per else None,
        },
    }

if __name__ == "__main__":
    import uvicorn
    print("Avvio Configuratore Elettroquadri API v0.11.0")
    print("Server in ascolto su http://0.0.0.0:8000")
    print("Documentazione API: http://0.0.0.0:8000/docs")
    uvicorn.run(app, host="0.0.0.0", port=8000)
