from fastapi import FastAPI, Depends, HTTPException, Query, Header
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import StreamingResponse
from sqlalchemy.orm import Session
from sqlalchemy import text, or_
from typing import List, Optional
from datetime import datetime, timedelta
import json
import os
import io

from database import engine, SessionLocal
from models import (
    Base, Preventivo, DatiCommessa, DatiPrincipali, Normative,
    DisposizioneVano, Porte, Materiale, ProductTemplate,
    Utente, GruppoUtenti, PermessoGruppo,
    Articolo, CategoriaArticoli, Cliente,
    RigaRicambio, Argano, ParametriSistema
)
from schemas import (
    PreventivoCreate, PreventivoUpdate, Preventivo as PreventivoSchema,
    DatiCommessaCreate, DatiCommessaUpdate, DatiCommessa as DatiCommessaSchema,
    DatiPrincipaliCreate, DatiPrincipaliUpdate, DatiPrincipali as DatiPrincipaliSchema,
    NormativeCreate, NormativeUpdate, Normative as NormativeSchema,
    DisposizioneVanoCreate, DisposizioneVanoUpdate, DisposizioneVano as DisposizioneVanoSchema,
    PorteCreate, PorteUpdate, Porte as PorteSchema,
    MaterialeCreate, MaterialeUpdate, Materiale as MaterialeSchema,
    ProductTemplateCreate, ProductTemplateUpdate, ProductTemplate as ProductTemplateSchema
)
from auth import (
    authenticate_user, create_access_token, get_password_hash,
    get_user_from_token, create_default_admin,
    ACCESS_TOKEN_EXPIRE_MINUTES
)

# Crea tabelle
Base.metadata.create_all(bind=engine)

# FastAPI app
app = FastAPI(title="Configuratore Elettroquadri API", version="0.11.0")

# ==========================================
# CORS CONFIGURATION
# ==========================================
app.add_middleware(
    CORSMiddleware,
    allow_origins=[
        "http://localhost:5173",
        "http://localhost:3000",
        "http://127.0.0.1:5173",
        "http://127.0.0.1:3000"
    ],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# ==========================================
# DEPENDENCY
# ==========================================
def get_db():
    db = SessionLocal()
    try:
        yield db
    finally:
        db.close()

# ==========================================
# RULE ENGINE - VERSIONE MIGLIORATA
# ==========================================
def load_rules():
    """Carica le regole dai file JSON nella directory rules/"""
    rules = []
    rules_dir = "./rules"
    
    if not os.path.exists(rules_dir):
        print("âš ï¸ Directory rules/ non trovata")
        return rules
    
    for filename in os.listdir(rules_dir):
        if filename.endswith(".json"):
            try:
                with open(os.path.join(rules_dir, filename), "r", encoding="utf-8") as f:
                    rule = json.load(f)
                    rules.append(rule)
                    print(f"âœ… Regola caricata: {filename}")
            except Exception as e:
                print(f"âŒ Errore caricamento {filename}: {e}")
    
    print(f"ðŸ“‹ Caricate {len(rules)} regole")
    return rules

def evaluate_condition(condition: dict, config_data: dict) -> bool:
    """Valuta una singola condizione con gestione robusta di NULL/None"""
    if isinstance(condition, str):
        return False
    
    field = condition.get("field")
    operator = condition.get("operator")
    expected_value = condition.get("value")
    
    if field not in config_data:
        return False
    
    config_value = config_data.get(field)
    
    if config_value is None:
        return False
    
    if isinstance(config_value, str) and config_value.strip() == "":
        return False
    
    result = False
    
    if operator == "equals":
        result = config_value == expected_value
    elif operator == "not_equals":
        result = config_value != expected_value
    elif operator == "contains":
        result = expected_value in str(config_value) if config_value else False
    elif operator == "greater_than":
        try:
            result = float(config_value) > float(expected_value)
        except (ValueError, TypeError):
            result = False
    elif operator == "less_than":
        try:
            result = float(config_value) < float(expected_value)
        except (ValueError, TypeError):
            result = False
    elif operator == "in":
        if isinstance(expected_value, list):
            result = config_value in expected_value
        else:
            result = False
    
    return result

def evaluate_rules(preventivo_id: int, db: Session):
    """Valuta tutte le regole per un preventivo e aggiunge/rimuove materiali"""
    db.expire_all()
    
    preventivo = db.query(Preventivo).filter(Preventivo.id == preventivo_id).first()
    if not preventivo:
        return {"error": "Preventivo non trovato"}
    
    # Raccogli tutti i dati di configurazione
    config_data = {}
    
    # 1. Dati dalle tabelle specifiche (legacy)
    if preventivo.dati_principali:
        dp = preventivo.dati_principali
        config_data.update({
            "tipo_impianto": dp.tipo_impianto,
            "nuovo_impianto": dp.nuovo_impianto,
            "numero_fermate": dp.numero_fermate,
            "numero_servizi": dp.numero_servizi,
            "velocita": dp.velocita,
            "corsa": dp.corsa,
            "con_locale_macchina": dp.con_locale_macchina,
            "posizione_locale_macchina": dp.posizione_locale_macchina,
            "tipo_trazione": dp.tipo_trazione,
            "forza_motrice": dp.forza_motrice,
            "luce": dp.luce,
            "tensione_manovra": dp.tensione_manovra,
            "tensione_freno": dp.tensione_freno,
        })
    
    if preventivo.normative:
        norm = preventivo.normative
        config_data.update({
            "en_81_1": norm.en_81_1,
            "en_81_20": norm.en_81_20,
            "en_81_21": norm.en_81_21,
            "en_81_28": norm.en_81_28,
            "en_81_70": norm.en_81_70,
            "en_81_72": norm.en_81_72,
            "en_81_73": norm.en_81_73,
            "a3_95_16": norm.a3_95_16,
            "dm236_legge13": norm.dm236_legge13,
            "emendamento_a3": norm.emendamento_a3,
            "uni_10411_1": norm.uni_10411_1,
        })
    
    # 2. Dati dalla configurazione JSON dinamica (sovrascrivono se duplicati)
    if preventivo.configurazione and isinstance(preventivo.configurazione, dict):
        for sezione, valori in preventivo.configurazione.items():
            if isinstance(valori, dict):
                config_data.update(valori)
    
    # Carica regole
    rules = load_rules()
    active_rules = set()
    materials_to_add = []
    
    for rule in rules:
        rule_id = rule.get("id", "unknown")
        conditions = rule.get("conditions", [])
        all_conditions_met = True
        
        for condition in conditions:
            if not evaluate_condition(condition, config_data):
                all_conditions_met = False
                break
        
        if all_conditions_met:
            active_rules.add(rule_id)
            for material in rule.get("materials", []):
                materials_to_add.append({
                    "rule_id": rule_id,
                    "codice": material.get("codice"),
                    "descrizione": material.get("descrizione"),
                    "quantita": material.get("quantita", 1),
                    "prezzo_unitario": material.get("prezzo_unitario", 0.0),
                    "categoria": material.get("categoria", "Materiale Automatico")
                })
    
    # Rimuovi materiali di regole non piu attive
    existing_materials = db.query(Materiale).filter(
        Materiale.preventivo_id == preventivo_id,
        Materiale.aggiunto_da_regola == True
    ).all()
    
    removed_count = 0
    for material in existing_materials:
        if material.regola_id not in active_rules:
            db.delete(material)
            removed_count += 1
    
    # Aggiungi nuovi materiali (evita duplicati)
    added_count = 0
    for mat_data in materials_to_add:
        existing = db.query(Materiale).filter(
            Materiale.preventivo_id == preventivo_id,
            Materiale.codice == mat_data["codice"],
            Materiale.regola_id == mat_data["rule_id"]
        ).first()
        
        if not existing:
            prezzo_totale = mat_data["quantita"] * mat_data["prezzo_unitario"]
            new_material = Materiale(
                preventivo_id=preventivo_id,
                codice=mat_data["codice"],
                descrizione=mat_data["descrizione"],
                quantita=mat_data["quantita"],
                prezzo_unitario=mat_data["prezzo_unitario"],
                prezzo_totale=prezzo_totale,
                categoria=mat_data["categoria"],
                aggiunto_da_regola=True,
                regola_id=mat_data["rule_id"]
            )
            db.add(new_material)
            added_count += 1
    
    db.commit()
    
    # Aggiorna totale preventivo
    all_materials = db.query(Materiale).filter(
        Materiale.preventivo_id == preventivo_id
    ).all()
    
    totale = sum(m.prezzo_totale for m in all_materials)
    preventivo.total_price = totale
    db.commit()
    
    return {
        "active_rules": list(active_rules),
        "materials_added": added_count,
        "materials_removed": removed_count,
        "total_materials": len(all_materials),
        "total_price": totale
    }

# ==========================================
# AUTH
# ==========================================
@app.post("/auth/login")
def login(data: dict, db: Session = Depends(get_db)):
    """Login utente, restituisce JWT token"""
    username = data.get("username", "")
    password = data.get("password", "")
    
    user = authenticate_user(db, username, password)
    if not user:
        raise HTTPException(status_code=401, detail="Credenziali non valide")
    
    # Aggiorna last_login
    user.last_login = datetime.now()
    db.commit()
    
    access_token = create_access_token(data={"sub": user.username})
    return {
        "access_token": access_token,
        "token_type": "bearer",
        "user": {
            "id": user.id,
            "username": user.username,
            "nome": user.nome,
            "cognome": user.cognome,
            "email": user.email,
            "is_admin": user.is_admin,
            "gruppo_id": user.gruppo_id
        }
    }

@app.get("/auth/me")
def get_current_user_info(
    authorization: Optional[str] = Header(None),
    db: Session = Depends(get_db)
):
    """Ottieni info utente corrente dal token"""
    token = None
    if authorization and authorization.startswith("Bearer "):
        token = authorization[7:]
    
    if not token:
        raise HTTPException(status_code=401, detail="Token non fornito")
    
    user = get_user_from_token(token, db)
    if not user:
        raise HTTPException(status_code=401, detail="Token non valido")
    
    return {
        "id": user.id,
        "username": user.username,
        "nome": user.nome,
        "cognome": user.cognome,
        "email": user.email,
        "is_admin": user.is_admin,
        "is_active": user.is_active,
        "gruppo_id": user.gruppo_id,
        "permessi": []  # TODO: caricare permessi dal gruppo
    }

# ==========================================
# PREVENTIVI
# ==========================================
@app.get("/preventivi", response_model=List[PreventivoSchema])
def get_preventivi(db: Session = Depends(get_db)):
    """Ottieni tutti i preventivi ordinati per data"""
    preventivi = db.query(Preventivo).order_by(Preventivo.created_at.desc()).all()
    return preventivi

@app.get("/preventivi/{preventivo_id}", response_model=PreventivoSchema)
def get_preventivo(preventivo_id: int, db: Session = Depends(get_db)):
    """Ottieni un preventivo specifico"""
    preventivo = db.query(Preventivo).filter(Preventivo.id == preventivo_id).first()
    if not preventivo:
        raise HTTPException(status_code=404, detail="Preventivo non trovato")
    return preventivo

@app.post("/preventivi", response_model=PreventivoSchema, status_code=201)
def create_preventivo(data: PreventivoCreate, db: Session = Depends(get_db)):
    """Crea nuovo preventivo con generazione automatica numero preventivo."""
    # Genera numero preventivo automatico se non fornito
    if not data.numero_preventivo:
        current_year = datetime.now().year
        last_preventivo = (
            db.query(Preventivo)
            .filter(Preventivo.numero_preventivo.like(f"{current_year}/%"))
            .order_by(Preventivo.id.desc())
            .first()
        )
        
        if last_preventivo and "/" in last_preventivo.numero_preventivo:
            try:
                last_number = int(last_preventivo.numero_preventivo.split("/")[1])
                new_number = last_number + 1
            except (ValueError, IndexError):
                new_number = 1
        else:
            new_number = 1
        
        numero_preventivo = f"{current_year}/{new_number:04d}"
    else:
        numero_preventivo = data.numero_preventivo
    
    # Carica template se specificato
    template_data_parsed = None
    template = None
    if data.template_id:
        template = db.query(ProductTemplate).filter(ProductTemplate.id == data.template_id).first()
        if template and template.template_data:
            try:
                template_data_parsed = json.loads(template.template_data)
            except json.JSONDecodeError:
                pass
    
    # Crea preventivo
    preventivo_dict = data.dict(exclude_unset=True)
    preventivo_dict['numero_preventivo'] = numero_preventivo
    template_id = preventivo_dict.pop('template_id', None)
    
    if template_id:
        preventivo_dict['template_id'] = template_id
        if template:
            preventivo_dict['categoria'] = template.categoria
    
    if 'created_at' not in preventivo_dict:
        preventivo_dict['created_at'] = datetime.now()
    if 'updated_at' not in preventivo_dict:
        preventivo_dict['updated_at'] = datetime.now()
    
    preventivo = Preventivo(**preventivo_dict)
    db.add(preventivo)
    db.commit()
    db.refresh(preventivo)
    
    # Crea record correlati (vuoti o pre-compilati da template)
    dp_data = template_data_parsed.get("dati_principali", {}) if template_data_parsed else {}
    norm_data = template_data_parsed.get("normative", {}) if template_data_parsed else {}
    dc_data = template_data_parsed.get("dati_commessa", {}) if template_data_parsed else {}
    dv_data = template_data_parsed.get("disposizione_vano", {}) if template_data_parsed else {}
    porte_data = template_data_parsed.get("porte", {}) if template_data_parsed else {}
    
    dati_commessa = DatiCommessa(preventivo_id=preventivo.id, **dc_data)
    dati_principali = DatiPrincipali(preventivo_id=preventivo.id, **dp_data)
    normative = Normative(preventivo_id=preventivo.id, **norm_data)
    disposizione_vano = DisposizioneVano(preventivo_id=preventivo.id, **dv_data)
    porte = Porte(preventivo_id=preventivo.id, **porte_data)
    
    db.add_all([dati_commessa, dati_principali, normative, disposizione_vano, porte])
    db.commit()
    
    # Aggiungi materiali da template se presenti
    if template_data_parsed and "materiali" in template_data_parsed:
        for mat in template_data_parsed["materiali"]:
            quantita = mat.get("quantita", 1)
            prezzo_unitario = mat.get("prezzo_unitario", 0.0)
            materiale = Materiale(
                preventivo_id=preventivo.id,
                codice=mat.get("codice", ""),
                descrizione=mat.get("descrizione", ""),
                quantita=quantita,
                prezzo_unitario=prezzo_unitario,
                prezzo_totale=quantita * prezzo_unitario,
                categoria=mat.get("categoria", "Template"),
                aggiunto_da_regola=False,
                note=f"Da template: {template.nome_display}" if template else None
            )
            db.add(materiale)
        
        db.commit()
        all_materials = db.query(Materiale).filter(Materiale.preventivo_id == preventivo.id).all()
        preventivo.total_price = sum(m.prezzo_totale for m in all_materials)
        db.commit()
        db.refresh(preventivo)
    
    return preventivo

@app.put("/preventivi/{preventivo_id}", response_model=PreventivoSchema)
def update_preventivo(preventivo_id: int, data: PreventivoUpdate, db: Session = Depends(get_db)):
    """Aggiorna preventivo"""
    preventivo = db.query(Preventivo).filter(Preventivo.id == preventivo_id).first()
    if not preventivo:
        raise HTTPException(status_code=404, detail="Preventivo non trovato")
    
    update_dict = data.dict(exclude_unset=True)
    update_dict['updated_at'] = datetime.now()
    
    for key, value in update_dict.items():
        setattr(preventivo, key, value)
    
    db.commit()
    db.refresh(preventivo)
    return preventivo

@app.delete("/preventivi/{preventivo_id}")
def delete_preventivo(preventivo_id: int, db: Session = Depends(get_db)):
    """Elimina preventivo"""
    preventivo = db.query(Preventivo).filter(Preventivo.id == preventivo_id).first()
    if not preventivo:
        raise HTTPException(status_code=404, detail="Preventivo non trovato")
    db.delete(preventivo)
    db.commit()
    return {"message": "Preventivo eliminato"}

# ==========================================
# DATI COMMESSA
# ==========================================
@app.get("/preventivi/{preventivo_id}/dati-commessa", response_model=DatiCommessaSchema)
def get_dati_commessa(preventivo_id: int, db: Session = Depends(get_db)):
    dati = db.query(DatiCommessa).filter(DatiCommessa.preventivo_id == preventivo_id).first()
    if not dati:
        dati = DatiCommessa(preventivo_id=preventivo_id)
        db.add(dati)
        db.commit()
        db.refresh(dati)
    return dati

@app.put("/preventivi/{preventivo_id}/dati-commessa", response_model=DatiCommessaSchema)
def update_dati_commessa(preventivo_id: int, data: DatiCommessaUpdate, db: Session = Depends(get_db)):
    dati = db.query(DatiCommessa).filter(DatiCommessa.preventivo_id == preventivo_id).first()
    if not dati:
        dati = DatiCommessa(preventivo_id=preventivo_id)
        db.add(dati)
    for key, value in data.dict(exclude_unset=True).items():
        setattr(dati, key, value)
    db.commit()
    db.refresh(dati)
    evaluate_rules(preventivo_id, db)
    return dati

# ==========================================
# DATI PRINCIPALI
# ==========================================
@app.get("/preventivi/{preventivo_id}/dati-principali", response_model=DatiPrincipaliSchema)
def get_dati_principali(preventivo_id: int, db: Session = Depends(get_db)):
    dati = db.query(DatiPrincipali).filter(DatiPrincipali.preventivo_id == preventivo_id).first()
    if not dati:
        dati = DatiPrincipali(preventivo_id=preventivo_id)
        db.add(dati)
        db.commit()
        db.refresh(dati)
    return dati

@app.put("/preventivi/{preventivo_id}/dati-principali", response_model=DatiPrincipaliSchema)
def update_dati_principali(preventivo_id: int, data: DatiPrincipaliUpdate, db: Session = Depends(get_db)):
    dati = db.query(DatiPrincipali).filter(DatiPrincipali.preventivo_id == preventivo_id).first()
    if not dati:
        dati = DatiPrincipali(preventivo_id=preventivo_id)
        db.add(dati)
    for key, value in data.dict(exclude_unset=True).items():
        setattr(dati, key, value)
    db.commit()
    db.refresh(dati)
    evaluate_rules(preventivo_id, db)
    return dati

# ==========================================
# NORMATIVE
# ==========================================
@app.get("/preventivi/{preventivo_id}/normative", response_model=NormativeSchema)
def get_normative(preventivo_id: int, db: Session = Depends(get_db)):
    normative = db.query(Normative).filter(Normative.preventivo_id == preventivo_id).first()
    if not normative:
        normative = Normative(preventivo_id=preventivo_id)
        db.add(normative)
        db.commit()
        db.refresh(normative)
    return normative

@app.put("/preventivi/{preventivo_id}/normative", response_model=NormativeSchema)
def update_normative(preventivo_id: int, data: NormativeUpdate, db: Session = Depends(get_db)):
    normative = db.query(Normative).filter(Normative.preventivo_id == preventivo_id).first()
    if not normative:
        normative = Normative(preventivo_id=preventivo_id)
        db.add(normative)
    for key, value in data.dict(exclude_unset=True).items():
        setattr(normative, key, value)
    db.commit()
    db.refresh(normative)
    evaluate_rules(preventivo_id, db)
    return normative

# ==========================================
# DISPOSIZIONE VANO
# ==========================================
@app.get("/preventivi/{preventivo_id}/disposizione-vano", response_model=DisposizioneVanoSchema)
def get_disposizione_vano(preventivo_id: int, db: Session = Depends(get_db)):
    disposizione = db.query(DisposizioneVano).filter(DisposizioneVano.preventivo_id == preventivo_id).first()
    if not disposizione:
        disposizione = DisposizioneVano(preventivo_id=preventivo_id)
        db.add(disposizione)
        db.commit()
        db.refresh(disposizione)
    return disposizione

@app.put("/preventivi/{preventivo_id}/disposizione-vano", response_model=DisposizioneVanoSchema)
def update_disposizione_vano(preventivo_id: int, data: DisposizioneVanoUpdate, db: Session = Depends(get_db)):
    disposizione = db.query(DisposizioneVano).filter(DisposizioneVano.preventivo_id == preventivo_id).first()
    if not disposizione:
        disposizione = DisposizioneVano(preventivo_id=preventivo_id)
        db.add(disposizione)
    
    data_dict = data.dict(exclude_unset=True)
    cleaned_data = {}
    for key, value in data_dict.items():
        if value == '' or value == '{}' or value == '[]':
            cleaned_data[key] = None
        else:
            cleaned_data[key] = value
    
    for key, value in cleaned_data.items():
        setattr(disposizione, key, value)
    
    db.commit()
    db.refresh(disposizione)
    return disposizione

@app.delete("/preventivi/{preventivo_id}/disposizione-vano")
def delete_disposizione_vano(preventivo_id: int, db: Session = Depends(get_db)):
    disposizione = db.query(DisposizioneVano).filter(DisposizioneVano.preventivo_id == preventivo_id).first()
    if disposizione:
        db.delete(disposizione)
        db.commit()
    return {"message": "Disposizione vano eliminata"}

# ==========================================
# PORTE
# ==========================================
@app.get("/preventivi/{preventivo_id}/porte", response_model=PorteSchema)
def get_porte(preventivo_id: int, db: Session = Depends(get_db)):
    porte = db.query(Porte).filter(Porte.preventivo_id == preventivo_id).first()
    if not porte:
        porte = Porte(preventivo_id=preventivo_id)
        db.add(porte)
        db.commit()
        db.refresh(porte)
    return porte

@app.put("/preventivi/{preventivo_id}/porte", response_model=PorteSchema)
def update_porte(preventivo_id: int, data: PorteUpdate, db: Session = Depends(get_db)):
    porte = db.query(Porte).filter(Porte.preventivo_id == preventivo_id).first()
    if not porte:
        porte = Porte(preventivo_id=preventivo_id)
        db.add(porte)
    for key, value in data.dict(exclude_unset=True).items():
        setattr(porte, key, value)
    db.commit()
    db.refresh(porte)
    return porte

# ==========================================
# ARGANO
# ==========================================
@app.get("/preventivi/{preventivo_id}/argano")
def get_argano(preventivo_id: int, db: Session = Depends(get_db)):
    """Ottieni dati argano, crea se non esiste"""
    argano = db.query(Argano).filter(Argano.preventivo_id == preventivo_id).first()
    if not argano:
        argano = Argano(preventivo_id=preventivo_id)
        db.add(argano)
        db.commit()
        db.refresh(argano)
    return {
        "id": argano.id,
        "preventivo_id": argano.preventivo_id,
        "trazione": argano.trazione,
        "potenza_motore_kw": argano.potenza_motore_kw,
        "corrente_nom_motore_amp": argano.corrente_nom_motore_amp,
        "tipo_vvvf": argano.tipo_vvvf,
        "vvvf_nel_vano": argano.vvvf_nel_vano,
        "freno_tensione": argano.freno_tensione,
        "ventilazione_forzata": argano.ventilazione_forzata,
        "tipo_teleruttore": argano.tipo_teleruttore
    }

@app.put("/preventivi/{preventivo_id}/argano")
def update_argano(preventivo_id: int, data: dict, db: Session = Depends(get_db)):
    """Aggiorna dati argano"""
    argano = db.query(Argano).filter(Argano.preventivo_id == preventivo_id).first()
    if not argano:
        argano = Argano(preventivo_id=preventivo_id)
        db.add(argano)
    
    for key in ["trazione", "potenza_motore_kw", "corrente_nom_motore_amp",
                "tipo_vvvf", "vvvf_nel_vano", "freno_tensione",
                "ventilazione_forzata", "tipo_teleruttore"]:
        if key in data:
            setattr(argano, key, data[key])
    
    db.commit()
    db.refresh(argano)
    return {
        "id": argano.id,
        "preventivo_id": argano.preventivo_id,
        "trazione": argano.trazione,
        "potenza_motore_kw": argano.potenza_motore_kw,
        "corrente_nom_motore_amp": argano.corrente_nom_motore_amp,
        "tipo_vvvf": argano.tipo_vvvf,
        "vvvf_nel_vano": argano.vvvf_nel_vano,
        "freno_tensione": argano.freno_tensione,
        "ventilazione_forzata": argano.ventilazione_forzata,
        "tipo_teleruttore": argano.tipo_teleruttore
    }

# ==========================================
# CONFIGURAZIONE DINAMICA (per sezioni da DB)
# ==========================================
@app.get("/preventivi/{preventivo_id}/configurazione")
def get_configurazione_completa(preventivo_id: int, db: Session = Depends(get_db)):
    """Ritorna tutta la configurazione JSON del preventivo"""
    preventivo = db.query(Preventivo).filter(Preventivo.id == preventivo_id).first()
    if not preventivo:
        raise HTTPException(status_code=404, detail="Preventivo non trovato")
    return preventivo.configurazione or {}


@app.get("/preventivi/{preventivo_id}/configurazione/{sezione}")
def get_configurazione_sezione(preventivo_id: int, sezione: str, db: Session = Depends(get_db)):
    """Ritorna i dati di una sezione specifica dalla configurazione JSON"""
    preventivo = db.query(Preventivo).filter(Preventivo.id == preventivo_id).first()
    if not preventivo:
        raise HTTPException(status_code=404, detail="Preventivo non trovato")
    config = preventivo.configurazione or {}
    return config.get(sezione, {})


@app.put("/preventivi/{preventivo_id}/configurazione/{sezione}")
def update_configurazione_sezione(preventivo_id: int, sezione: str, data: dict, db: Session = Depends(get_db)):
    """Salva i dati di una sezione nella configurazione JSON e rivaluta le regole"""
    preventivo = db.query(Preventivo).filter(Preventivo.id == preventivo_id).first()
    if not preventivo:
        raise HTTPException(status_code=404, detail="Preventivo non trovato")
    
    config = preventivo.configurazione or {}
    if not isinstance(config, dict):
        config = {}
    
    # Aggiorna solo la sezione specificata
    config[sezione] = data.get("valori", data)
    preventivo.configurazione = config
    
    # Forza il detect della modifica JSON (SQLAlchemy)
    from sqlalchemy.orm.attributes import flag_modified
    flag_modified(preventivo, "configurazione")
    
    db.commit()
    
    # Rivaluta regole
    try:
        result = evaluate_rules(preventivo_id, db)
    except Exception as e:
        print(f"Errore evaluate_rules: {e}")
        result = {}
    
    return {"status": "ok", "sezione": sezione, "rules_result": result}


# ==========================================
# MATERIALI
# ==========================================
@app.get("/preventivi/{preventivo_id}/materiali", response_model=List[MaterialeSchema])
def get_materiali(preventivo_id: int, db: Session = Depends(get_db)):
    return db.query(Materiale).filter(Materiale.preventivo_id == preventivo_id).all()

@app.post("/preventivi/{preventivo_id}/materiali", response_model=MaterialeSchema)
def create_materiale(preventivo_id: int, data: MaterialeCreate, db: Session = Depends(get_db)):
    prezzo_totale = data.quantita * data.prezzo_unitario
    materiale = Materiale(
        preventivo_id=preventivo_id,
        codice=data.codice, descrizione=data.descrizione,
        quantita=data.quantita, prezzo_unitario=data.prezzo_unitario,
        prezzo_totale=prezzo_totale, categoria=data.categoria,
        aggiunto_da_regola=data.aggiunto_da_regola,
        regola_id=data.regola_id, note=data.note
    )
    db.add(materiale)
    db.commit()
    db.refresh(materiale)
    
    preventivo = db.query(Preventivo).filter(Preventivo.id == preventivo_id).first()
    if preventivo:
        all_materials = db.query(Materiale).filter(Materiale.preventivo_id == preventivo_id).all()
        preventivo.total_price = sum(m.prezzo_totale for m in all_materials)
        db.commit()
    return materiale

@app.put("/preventivi/{preventivo_id}/materiali/{materiale_id}", response_model=MaterialeSchema)
def update_materiale(preventivo_id: int, materiale_id: int, data: MaterialeUpdate, db: Session = Depends(get_db)):
    materiale = db.query(Materiale).filter(Materiale.id == materiale_id, Materiale.preventivo_id == preventivo_id).first()
    if not materiale:
        raise HTTPException(status_code=404, detail="Materiale non trovato")
    
    update_data = data.dict(exclude_unset=True)
    for key, value in update_data.items():
        setattr(materiale, key, value)
    
    if 'quantita' in update_data or 'prezzo_unitario' in update_data:
        materiale.prezzo_totale = materiale.quantita * materiale.prezzo_unitario
    
    db.commit()
    db.refresh(materiale)
    
    preventivo = db.query(Preventivo).filter(Preventivo.id == preventivo_id).first()
    if preventivo:
        all_materials = db.query(Materiale).filter(Materiale.preventivo_id == preventivo_id).all()
        preventivo.total_price = sum(m.prezzo_totale for m in all_materials)
        db.commit()
    return materiale

@app.delete("/preventivi/{preventivo_id}/materiali/{materiale_id}")
def delete_materiale(preventivo_id: int, materiale_id: int, db: Session = Depends(get_db)):
    materiale = db.query(Materiale).filter(Materiale.id == materiale_id, Materiale.preventivo_id == preventivo_id).first()
    if not materiale:
        raise HTTPException(status_code=404, detail="Materiale non trovato")
    db.delete(materiale)
    db.commit()
    
    preventivo = db.query(Preventivo).filter(Preventivo.id == preventivo_id).first()
    if preventivo:
        all_materials = db.query(Materiale).filter(Materiale.preventivo_id == preventivo_id).all()
        preventivo.total_price = sum(m.prezzo_totale for m in all_materials)
        db.commit()
    return {"message": "Materiale eliminato"}

# ==========================================
# RIGHE RICAMBIO
# ==========================================
@app.get("/preventivi/{preventivo_id}/righe-ricambio")
def get_righe_ricambio(preventivo_id: int, db: Session = Depends(get_db)):
    righe = db.query(RigaRicambio).filter(
        RigaRicambio.preventivo_id == preventivo_id
    ).order_by(RigaRicambio.ordine).all()
    return [{
        "id": r.id, "preventivo_id": r.preventivo_id,
        "articolo_id": r.articolo_id, "codice": r.codice,
        "descrizione": r.descrizione, "tipo_articolo": r.tipo_articolo,
        "quantita": r.quantita,
        "parametro_1": r.parametro_1, "unita_param_1": r.unita_param_1,
        "desc_param_1": r.desc_param_1, "costo_var_1": r.costo_var_1,
        "parametro_2": r.parametro_2, "unita_param_2": r.unita_param_2,
        "desc_param_2": r.desc_param_2, "costo_var_2": r.costo_var_2,
        "parametro_3": r.parametro_3, "unita_param_3": r.unita_param_3,
        "desc_param_3": r.desc_param_3, "costo_var_3": r.costo_var_3,
        "parametro_4": r.parametro_4, "unita_param_4": r.unita_param_4,
        "desc_param_4": r.desc_param_4, "costo_var_4": r.costo_var_4,
        "costo_fisso": r.costo_fisso, "costo_base_unitario": r.costo_base_unitario,
        "ricarico_percentuale": r.ricarico_percentuale,
        "prezzo_listino_unitario": r.prezzo_listino_unitario,
        "sconto_cliente": r.sconto_cliente,
        "prezzo_cliente_unitario": r.prezzo_cliente_unitario,
        "prezzo_totale_listino": r.prezzo_totale_listino,
        "prezzo_totale_cliente": r.prezzo_totale_cliente,
        "ordine": r.ordine, "note": r.note
    } for r in righe]

@app.post("/preventivi/{preventivo_id}/righe-ricambio")
def create_riga_ricambio(preventivo_id: int, data: dict, db: Session = Depends(get_db)):
    riga = RigaRicambio(preventivo_id=preventivo_id)
    for key in [
        "articolo_id", "codice", "descrizione", "tipo_articolo", "quantita",
        "parametro_1", "unita_param_1", "desc_param_1", "costo_var_1",
        "parametro_2", "unita_param_2", "desc_param_2", "costo_var_2",
        "parametro_3", "unita_param_3", "desc_param_3", "costo_var_3",
        "parametro_4", "unita_param_4", "desc_param_4", "costo_var_4",
        "costo_fisso", "costo_base_unitario", "ricarico_percentuale",
        "prezzo_listino_unitario", "sconto_cliente", "prezzo_cliente_unitario",
        "prezzo_totale_listino", "prezzo_totale_cliente", "ordine", "note"
    ]:
        if key in data:
            setattr(riga, key, data[key])
    
    db.add(riga)
    db.commit()
    db.refresh(riga)
    return {"id": riga.id, "status": "ok"}

@app.put("/preventivi/{preventivo_id}/righe-ricambio/{riga_id}")
def update_riga_ricambio(preventivo_id: int, riga_id: int, data: dict, db: Session = Depends(get_db)):
    riga = db.query(RigaRicambio).filter(
        RigaRicambio.id == riga_id,
        RigaRicambio.preventivo_id == preventivo_id
    ).first()
    if not riga:
        raise HTTPException(status_code=404, detail="Riga ricambio non trovata")
    
    for key in [
        "articolo_id", "codice", "descrizione", "tipo_articolo", "quantita",
        "parametro_1", "unita_param_1", "desc_param_1", "costo_var_1",
        "parametro_2", "unita_param_2", "desc_param_2", "costo_var_2",
        "parametro_3", "unita_param_3", "desc_param_3", "costo_var_3",
        "parametro_4", "unita_param_4", "desc_param_4", "costo_var_4",
        "costo_fisso", "costo_base_unitario", "ricarico_percentuale",
        "prezzo_listino_unitario", "sconto_cliente", "prezzo_cliente_unitario",
        "prezzo_totale_listino", "prezzo_totale_cliente", "ordine", "note"
    ]:
        if key in data:
            setattr(riga, key, data[key])
    
    db.commit()
    db.refresh(riga)
    return {"id": riga.id, "status": "ok"}

@app.delete("/preventivi/{preventivo_id}/righe-ricambio/{riga_id}")
def delete_riga_ricambio(preventivo_id: int, riga_id: int, db: Session = Depends(get_db)):
    riga = db.query(RigaRicambio).filter(
        RigaRicambio.id == riga_id,
        RigaRicambio.preventivo_id == preventivo_id
    ).first()
    if not riga:
        raise HTTPException(status_code=404, detail="Riga ricambio non trovata")
    db.delete(riga)
    db.commit()
    return {"status": "ok"}

# ==========================================
# SCONTO ADMIN
# ==========================================
@app.put("/preventivi/{preventivo_id}/sconto-admin")
def update_sconto_admin(preventivo_id: int, data: dict, db: Session = Depends(get_db)):
    """Aggiorna sconto extra admin per un preventivo"""
    preventivo = db.query(Preventivo).filter(Preventivo.id == preventivo_id).first()
    if not preventivo:
        raise HTTPException(status_code=404, detail="Preventivo non trovato")
    
    if "sconto_extra_admin" in data:
        preventivo.sconto_extra_admin = data["sconto_extra_admin"]
    
    # Ricalcola prezzo finale
    sconto_totale = (preventivo.sconto_cliente or 0) + (preventivo.sconto_extra_admin or 0)
    if preventivo.total_price:
        preventivo.total_price_finale = preventivo.total_price * (1 - sconto_totale / 100)
    
    preventivo.updated_at = datetime.now()
    db.commit()
    db.refresh(preventivo)
    return {
        "sconto_extra_admin": preventivo.sconto_extra_admin,
        "total_price_finale": preventivo.total_price_finale
    }

# ==========================================
# EXPORT PREVENTIVO
# ==========================================
@app.get("/preventivi/{preventivo_id}/export/{formato}")
def export_preventivo(preventivo_id: int, formato: str, db: Session = Depends(get_db)):
    """Esporta preventivo in formato PDF o DOCX"""
    preventivo = db.query(Preventivo).filter(Preventivo.id == preventivo_id).first()
    if not preventivo:
        raise HTTPException(status_code=404, detail="Preventivo non trovato")
    
    # Raccogli tutti i dati
    dati_commessa = db.query(DatiCommessa).filter(DatiCommessa.preventivo_id == preventivo_id).first()
    dati_principali = db.query(DatiPrincipali).filter(DatiPrincipali.preventivo_id == preventivo_id).first()
    normative_data = db.query(Normative).filter(Normative.preventivo_id == preventivo_id).first()
    materiali = db.query(Materiale).filter(Materiale.preventivo_id == preventivo_id).all()
    
    if formato.lower() == "json":
        # Export JSON
        return {
            "preventivo": {
                "id": preventivo.id,
                "numero_preventivo": preventivo.numero_preventivo,
                "status": preventivo.status,
                "total_price": preventivo.total_price,
                "customer_name": preventivo.customer_name,
                "created_at": str(preventivo.created_at) if preventivo.created_at else None,
            },
            "dati_commessa": {k: v for k, v in (dati_commessa.__dict__.items() if dati_commessa else {}) if not k.startswith('_')},
            "dati_principali": {k: v for k, v in (dati_principali.__dict__.items() if dati_principali else {}) if not k.startswith('_')},
            "materiali": [{
                "codice": m.codice, "descrizione": m.descrizione,
                "quantita": m.quantita, "prezzo_unitario": m.prezzo_unitario,
                "prezzo_totale": m.prezzo_totale, "categoria": m.categoria
            } for m in materiali]
        }
    
    elif formato.lower() in ["pdf", "docx"]:
        try:
            from docx import Document
            from docx.shared import Inches, Pt
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            
            doc = Document()
            doc.add_heading(f'Preventivo {preventivo.numero_preventivo}', level=1)
            
            if preventivo.customer_name:
                doc.add_paragraph(f'Cliente: {preventivo.customer_name}')
            if dati_commessa and dati_commessa.data_offerta:
                doc.add_paragraph(f'Data: {dati_commessa.data_offerta}')
            
            doc.add_heading('Materiali', level=2)
            if materiali:
                table = doc.add_table(rows=1, cols=5)
                table.style = 'Table Grid'
                headers = table.rows[0].cells
                headers[0].text = 'Codice'
                headers[1].text = 'Descrizione'
                headers[2].text = 'QtÃ '
                headers[3].text = 'Prezzo Unit.'
                headers[4].text = 'Totale'
                
                for m in materiali:
                    row = table.add_row().cells
                    row[0].text = m.codice or ''
                    row[1].text = m.descrizione or ''
                    row[2].text = str(m.quantita)
                    row[3].text = f'â‚¬{m.prezzo_unitario:.2f}'
                    row[4].text = f'â‚¬{m.prezzo_totale:.2f}'
            
            doc.add_paragraph(f'\nTotale: â‚¬{preventivo.total_price:.2f}')
            
            buffer = io.BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            
            if formato.lower() == "docx":
                return StreamingResponse(
                    buffer,
                    media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    headers={"Content-Disposition": f"attachment; filename=preventivo_{preventivo.numero_preventivo.replace('/', '_')}.docx"}
                )
            else:
                # PDF: restituisce il docx per ora, il frontend gestisce la conversione
                return StreamingResponse(
                    buffer,
                    media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    headers={"Content-Disposition": f"attachment; filename=preventivo_{preventivo.numero_preventivo.replace('/', '_')}.docx"}
                )
        except ImportError:
            raise HTTPException(status_code=500, detail="python-docx non installato. Installa con: pip install python-docx")
    
    raise HTTPException(status_code=400, detail=f"Formato '{formato}' non supportato. Usa: json, pdf, docx")

# ==========================================
# PRODUCT TEMPLATES
# ==========================================
@app.get("/templates", response_model=List[ProductTemplateSchema])
def get_templates(categoria: str = None, db: Session = Depends(get_db)):
    query = db.query(ProductTemplate).filter(ProductTemplate.attivo == True)
    if categoria:
        query = query.filter(ProductTemplate.categoria == categoria.upper())
    return query.order_by(ProductTemplate.categoria, ProductTemplate.ordine).all()

@app.get("/templates/all", response_model=List[ProductTemplateSchema])
def get_all_templates(db: Session = Depends(get_db)):
    return db.query(ProductTemplate).order_by(ProductTemplate.categoria, ProductTemplate.ordine).all()

@app.get("/templates/categories/summary")
def get_categories_summary(db: Session = Depends(get_db)):
    templates = db.query(ProductTemplate).filter(ProductTemplate.attivo == True).order_by(
        ProductTemplate.categoria, ProductTemplate.ordine
    ).all()
    categories = {}
    for t in templates:
        cat = t.categoria
        if cat not in categories:
            categories[cat] = {"categoria": cat, "sottocategorie": []}
        categories[cat]["sottocategorie"].append({
            "id": t.id, "sottocategoria": t.sottocategoria,
            "nome_display": t.nome_display, "descrizione": t.descrizione,
            "icona": t.icona, "ordine": t.ordine
        })
    return list(categories.values())

@app.get("/templates/{template_id}", response_model=ProductTemplateSchema)
def get_template(template_id: int, db: Session = Depends(get_db)):
    template = db.query(ProductTemplate).filter(ProductTemplate.id == template_id).first()
    if not template:
        raise HTTPException(status_code=404, detail="Template non trovato")
    return template

@app.post("/templates", response_model=ProductTemplateSchema, status_code=201)
def create_template(data: ProductTemplateCreate, db: Session = Depends(get_db)):
    if data.template_data:
        try:
            json.loads(data.template_data)
        except json.JSONDecodeError:
            raise HTTPException(status_code=400, detail="template_data non Ã¨ un JSON valido")
    
    template_dict = data.dict()
    template_dict['categoria'] = template_dict['categoria'].upper()
    template_dict['created_at'] = datetime.now()
    template_dict['updated_at'] = datetime.now()
    template = ProductTemplate(**template_dict)
    db.add(template)
    db.commit()
    db.refresh(template)
    return template

@app.put("/templates/{template_id}", response_model=ProductTemplateSchema)
def update_template(template_id: int, data: ProductTemplateUpdate, db: Session = Depends(get_db)):
    template = db.query(ProductTemplate).filter(ProductTemplate.id == template_id).first()
    if not template:
        raise HTTPException(status_code=404, detail="Template non trovato")
    if data.template_data:
        try:
            json.loads(data.template_data)
        except json.JSONDecodeError:
            raise HTTPException(status_code=400, detail="template_data non Ã¨ un JSON valido")
    
    update_dict = data.dict(exclude_unset=True)
    if 'categoria' in update_dict and update_dict['categoria']:
        update_dict['categoria'] = update_dict['categoria'].upper()
    update_dict['updated_at'] = datetime.now()
    for key, value in update_dict.items():
        setattr(template, key, value)
    db.commit()
    db.refresh(template)
    return template

@app.delete("/templates/{template_id}")
def delete_template(template_id: int, db: Session = Depends(get_db)):
    template = db.query(ProductTemplate).filter(ProductTemplate.id == template_id).first()
    if not template:
        raise HTTPException(status_code=404, detail="Template non trovato")
    db.delete(template)
    db.commit()
    return {"message": "Template eliminato"}

# ==========================================
# RULE ENGINE ENDPOINT
# ==========================================
@app.post("/preventivi/{preventivo_id}/evaluate-rules")
def evaluate_rules_endpoint(preventivo_id: int, db: Session = Depends(get_db)):
    return evaluate_rules(preventivo_id, db)

# ==========================================
# ARTICOLI
# ==========================================
@app.get("/articoli")
def get_articoli(
    search: str = None,
    limit: int = 100,
    offset: int = 0,
    db: Session = Depends(get_db)
):
    """Lista articoli con ricerca opzionale"""
    query = db.query(Articolo)
    if search:
        search_term = f"%{search}%"
        query = query.filter(
            or_(
                Articolo.codice.ilike(search_term),
                Articolo.descrizione.ilike(search_term)
            )
        )
    articoli = query.order_by(Articolo.codice).offset(offset).limit(limit).all()
    return [{
        "id": a.id, "codice": a.codice, "descrizione": a.descrizione,
        "descrizione_estesa": a.descrizione_estesa,
        "tipo_articolo": a.tipo_articolo, "categoria_id": a.categoria_id,
        "costo_fisso": a.costo_fisso,
        "costo_variabile_1": a.costo_variabile_1, "unita_misura_var_1": a.unita_misura_var_1,
        "descrizione_var_1": a.descrizione_var_1,
        "costo_variabile_2": a.costo_variabile_2, "unita_misura_var_2": a.unita_misura_var_2,
        "descrizione_var_2": a.descrizione_var_2,
        "costo_variabile_3": a.costo_variabile_3, "unita_misura_var_3": a.unita_misura_var_3,
        "descrizione_var_3": a.descrizione_var_3,
        "costo_variabile_4": a.costo_variabile_4, "unita_misura_var_4": a.unita_misura_var_4,
        "descrizione_var_4": a.descrizione_var_4,
        "ricarico_percentuale": a.ricarico_percentuale,
        "unita_misura": a.unita_misura,
        "fornitore": a.fornitore, "codice_fornitore": a.codice_fornitore,
        "is_active": a.is_active
    } for a in articoli]

@app.get("/articoli/search")
def search_articoli(q: str = "", limit: int = 10, db: Session = Depends(get_db)):
    """Ricerca articoli per codice o descrizione"""
    if not q:
        return []
    search_term = f"%{q}%"
    articoli = db.query(Articolo).filter(
        or_(
            Articolo.codice.ilike(search_term),
            Articolo.descrizione.ilike(search_term)
        )
    ).limit(limit).all()
    return [{
        "id": a.id, "codice": a.codice, "descrizione": a.descrizione,
        "tipo_articolo": a.tipo_articolo, "costo_fisso": a.costo_fisso,
        "costo_variabile_1": a.costo_variabile_1, "unita_misura_var_1": a.unita_misura_var_1,
        "descrizione_var_1": a.descrizione_var_1,
        "costo_variabile_2": a.costo_variabile_2, "unita_misura_var_2": a.unita_misura_var_2,
        "descrizione_var_2": a.descrizione_var_2,
        "costo_variabile_3": a.costo_variabile_3, "unita_misura_var_3": a.unita_misura_var_3,
        "descrizione_var_3": a.descrizione_var_3,
        "costo_variabile_4": a.costo_variabile_4, "unita_misura_var_4": a.unita_misura_var_4,
        "descrizione_var_4": a.descrizione_var_4,
        "ricarico_percentuale": a.ricarico_percentuale,
        "unita_misura": a.unita_misura
    } for a in articoli]

@app.post("/articoli")
def create_articolo(data: dict, db: Session = Depends(get_db)):
    articolo = Articolo()
    for key in [
        "codice", "descrizione", "descrizione_estesa", "tipo_articolo", "categoria_id",
        "costo_fisso", "costo_variabile_1", "unita_misura_var_1", "descrizione_var_1",
        "costo_variabile_2", "unita_misura_var_2", "descrizione_var_2",
        "costo_variabile_3", "unita_misura_var_3", "descrizione_var_3",
        "costo_variabile_4", "unita_misura_var_4", "descrizione_var_4",
        "ricarico_percentuale", "unita_misura", "fornitore", "codice_fornitore",
        "giacenza", "scorta_minima", "lead_time_giorni", "manodopera_giorni", "is_active"
    ]:
        if key in data:
            setattr(articolo, key, data[key])
    db.add(articolo)
    db.commit()
    db.refresh(articolo)
    return {"id": articolo.id, "status": "ok"}

@app.put("/articoli/{articolo_id}")
def update_articolo(articolo_id: int, data: dict, db: Session = Depends(get_db)):
    articolo = db.query(Articolo).filter(Articolo.id == articolo_id).first()
    if not articolo:
        raise HTTPException(status_code=404, detail="Articolo non trovato")
    for key in [
        "codice", "descrizione", "descrizione_estesa", "tipo_articolo", "categoria_id",
        "costo_fisso", "costo_variabile_1", "unita_misura_var_1", "descrizione_var_1",
        "costo_variabile_2", "unita_misura_var_2", "descrizione_var_2",
        "costo_variabile_3", "unita_misura_var_3", "descrizione_var_3",
        "costo_variabile_4", "unita_misura_var_4", "descrizione_var_4",
        "ricarico_percentuale", "unita_misura", "fornitore", "codice_fornitore",
        "giacenza", "scorta_minima", "lead_time_giorni", "manodopera_giorni", "is_active"
    ]:
        if key in data:
            setattr(articolo, key, data[key])
    articolo.updated_at = datetime.now()
    db.commit()
    return {"id": articolo.id, "status": "ok"}

@app.delete("/articoli/{articolo_id}")
def delete_articolo(articolo_id: int, db: Session = Depends(get_db)):
    articolo = db.query(Articolo).filter(Articolo.id == articolo_id).first()
    if not articolo:
        raise HTTPException(status_code=404, detail="Articolo non trovato")
    db.delete(articolo)
    db.commit()
    return {"status": "ok"}

@app.post("/articoli/calcola-prezzo")
def calcola_prezzo_articolo(data: dict, db: Session = Depends(get_db)):
    """Calcola prezzo di un articolo in base ai parametri variabili"""
    costo_fisso = data.get("costo_fisso", 0) or 0
    costo_totale = costo_fisso
    
    for i in range(1, 5):
        costo_var = data.get(f"costo_variabile_{i}", 0) or 0
        parametro = data.get(f"parametro_{i}", 0) or 0
        costo_totale += costo_var * parametro
    
    ricarico = data.get("ricarico_percentuale", 0) or 0
    prezzo_listino = costo_totale * (1 + ricarico / 100) if ricarico else costo_totale
    
    sconto = data.get("sconto_cliente", 0) or 0
    prezzo_cliente = prezzo_listino * (1 - sconto / 100) if sconto else prezzo_listino
    
    return {
        "costo_base": costo_totale,
        "prezzo_listino": round(prezzo_listino, 4),
        "prezzo_cliente": round(prezzo_cliente, 4)
    }

# ==========================================
# CATEGORIE ARTICOLI
# ==========================================
@app.get("/categorie-articoli")
def get_categorie_articoli(db: Session = Depends(get_db)):
    categorie = db.query(CategoriaArticoli).order_by(CategoriaArticoli.ordine).all()
    return [{
        "id": c.id, "codice": c.codice, "nome": c.nome,
        "descrizione": c.descrizione, "ordine": c.ordine, "is_active": c.is_active
    } for c in categorie]

# ==========================================
# CLIENTI
# ==========================================
@app.get("/clienti")
def get_clienti(db: Session = Depends(get_db)):
    clienti = db.query(Cliente).filter(Cliente.is_active == True).order_by(Cliente.ragione_sociale).all()
    return [{
        "id": c.id, "codice": c.codice, "ragione_sociale": c.ragione_sociale,
        "partita_iva": c.partita_iva, "codice_fiscale": c.codice_fiscale,
        "indirizzo": c.indirizzo, "cap": c.cap, "citta": c.citta,
        "provincia": c.provincia, "nazione": c.nazione,
        "telefono": c.telefono, "email": c.email, "pec": c.pec,
        "sconto_globale": c.sconto_globale, "sconto_produzione": c.sconto_produzione,
        "sconto_acquisto": c.sconto_acquisto, "aliquota_iva": c.aliquota_iva,
        "pagamento_default": c.pagamento_default, "imballo_default": c.imballo_default,
        "reso_fco_default": c.reso_fco_default, "trasporto_default": c.trasporto_default,
        "destinazione_default": c.destinazione_default,
        "riferimento_cliente_default": c.riferimento_cliente_default,
        "listino": c.listino, "note": c.note, "is_active": c.is_active
    } for c in clienti]

@app.get("/clienti/search")
def search_clienti(q: str = "", db: Session = Depends(get_db)):
    if not q:
        return []
    search_term = f"%{q}%"
    clienti = db.query(Cliente).filter(
        or_(
            Cliente.codice.ilike(search_term),
            Cliente.ragione_sociale.ilike(search_term),
            Cliente.partita_iva.ilike(search_term)
        )
    ).limit(20).all()
    return [{
        "id": c.id, "codice": c.codice, "ragione_sociale": c.ragione_sociale,
        "partita_iva": c.partita_iva, "citta": c.citta, "provincia": c.provincia,
        "sconto_globale": c.sconto_globale, "sconto_produzione": c.sconto_produzione,
        "sconto_acquisto": c.sconto_acquisto,
        "pagamento_default": c.pagamento_default, "imballo_default": c.imballo_default,
        "reso_fco_default": c.reso_fco_default, "trasporto_default": c.trasporto_default,
        "destinazione_default": c.destinazione_default,
        "riferimento_cliente_default": c.riferimento_cliente_default
    } for c in clienti]

@app.get("/clienti/{cliente_id}")
def get_cliente(cliente_id: int, db: Session = Depends(get_db)):
    cliente = db.query(Cliente).filter(Cliente.id == cliente_id).first()
    if not cliente:
        raise HTTPException(status_code=404, detail="Cliente non trovato")
    return {
        "id": cliente.id, "codice": cliente.codice, "ragione_sociale": cliente.ragione_sociale,
        "partita_iva": cliente.partita_iva, "codice_fiscale": cliente.codice_fiscale,
        "indirizzo": cliente.indirizzo, "cap": cliente.cap, "citta": cliente.citta,
        "provincia": cliente.provincia, "nazione": cliente.nazione,
        "telefono": cliente.telefono, "email": cliente.email, "pec": cliente.pec,
        "sconto_globale": cliente.sconto_globale, "sconto_produzione": cliente.sconto_produzione,
        "sconto_acquisto": cliente.sconto_acquisto, "aliquota_iva": cliente.aliquota_iva,
        "pagamento_default": cliente.pagamento_default, "imballo_default": cliente.imballo_default,
        "reso_fco_default": cliente.reso_fco_default, "trasporto_default": cliente.trasporto_default,
        "destinazione_default": cliente.destinazione_default,
        "riferimento_cliente_default": cliente.riferimento_cliente_default,
        "listino": cliente.listino, "note": cliente.note, "is_active": cliente.is_active
    }

@app.post("/clienti")
def create_cliente(data: dict, db: Session = Depends(get_db)):
    cliente = Cliente()
    for key in [
        "codice", "ragione_sociale", "partita_iva", "codice_fiscale",
        "indirizzo", "cap", "citta", "provincia", "nazione",
        "telefono", "email", "pec",
        "sconto_globale", "sconto_produzione", "sconto_acquisto", "aliquota_iva",
        "pagamento_default", "imballo_default", "reso_fco_default",
        "trasporto_default", "destinazione_default", "riferimento_cliente_default",
        "listino", "note", "is_active"
    ]:
        if key in data:
            setattr(cliente, key, data[key])
    db.add(cliente)
    db.commit()
    db.refresh(cliente)
    return {"id": cliente.id, "status": "ok"}

@app.put("/clienti/{cliente_id}")
def update_cliente(cliente_id: int, data: dict, db: Session = Depends(get_db)):
    cliente = db.query(Cliente).filter(Cliente.id == cliente_id).first()
    if not cliente:
        raise HTTPException(status_code=404, detail="Cliente non trovato")
    for key in [
        "codice", "ragione_sociale", "partita_iva", "codice_fiscale",
        "indirizzo", "cap", "citta", "provincia", "nazione",
        "telefono", "email", "pec",
        "sconto_globale", "sconto_produzione", "sconto_acquisto", "aliquota_iva",
        "pagamento_default", "imballo_default", "reso_fco_default",
        "trasporto_default", "destinazione_default", "riferimento_cliente_default",
        "listino", "note", "is_active"
    ]:
        if key in data:
            setattr(cliente, key, data[key])
    cliente.updated_at = datetime.now()
    db.commit()
    return {"id": cliente.id, "status": "ok"}

# ==========================================
# UTENTI
# ==========================================
@app.get("/utenti")
def get_utenti(db: Session = Depends(get_db)):
    utenti = db.query(Utente).order_by(Utente.username).all()
    return [{
        "id": u.id, "username": u.username,
        "nome": u.nome, "cognome": u.cognome,
        "email": u.email, "gruppo_id": u.gruppo_id,
        "is_admin": u.is_admin, "is_active": u.is_active,
        "created_at": str(u.created_at) if u.created_at else None,
        "last_login": str(u.last_login) if u.last_login else None
    } for u in utenti]

@app.post("/utenti")
def create_utente(
    username: str = Query(...),
    password: str = Query(None),
    nome: str = Query(None),
    cognome: str = Query(None),
    email: str = Query(None),
    gruppo_id: int = Query(None),
    is_admin: bool = Query(False),
    is_active: bool = Query(True),
    db: Session = Depends(get_db)
):
    existing = db.query(Utente).filter(Utente.username == username).first()
    if existing:
        raise HTTPException(status_code=400, detail="Username giÃ  esistente")
    
    utente = Utente(
        username=username,
        password_hash=get_password_hash(password or "changeme"),
        nome=nome, cognome=cognome, email=email,
        gruppo_id=gruppo_id, is_admin=is_admin, is_active=is_active
    )
    db.add(utente)
    db.commit()
    db.refresh(utente)
    return {"id": utente.id, "status": "ok"}

@app.put("/utenti/{utente_id}")
def update_utente(
    utente_id: int,
    username: str = Query(None),
    password: str = Query(None),
    nome: str = Query(None),
    cognome: str = Query(None),
    email: str = Query(None),
    gruppo_id: int = Query(None),
    is_admin: bool = Query(None),
    is_active: bool = Query(None),
    db: Session = Depends(get_db)
):
    utente = db.query(Utente).filter(Utente.id == utente_id).first()
    if not utente:
        raise HTTPException(status_code=404, detail="Utente non trovato")
    
    if username is not None:
        utente.username = username
    if password is not None:
        utente.password_hash = get_password_hash(password)
    if nome is not None:
        utente.nome = nome
    if cognome is not None:
        utente.cognome = cognome
    if email is not None:
        utente.email = email
    if gruppo_id is not None:
        utente.gruppo_id = gruppo_id
    if is_admin is not None:
        utente.is_admin = is_admin
    if is_active is not None:
        utente.is_active = is_active
    
    db.commit()
    db.refresh(utente)
    return {"id": utente.id, "status": "ok"}

@app.delete("/utenti/{utente_id}")
def delete_utente(utente_id: int, db: Session = Depends(get_db)):
    utente = db.query(Utente).filter(Utente.id == utente_id).first()
    if not utente:
        raise HTTPException(status_code=404, detail="Utente non trovato")
    db.delete(utente)
    db.commit()
    return {"status": "ok"}

# ==========================================
# PARAMETRI SISTEMA
# ==========================================
@app.get("/parametri-sistema")
def get_parametri_sistema(db: Session = Depends(get_db)):
    parametri = db.query(ParametriSistema).order_by(ParametriSistema.gruppo, ParametriSistema.chiave).all()
    return [{
        "id": p.id, "chiave": p.chiave, "valore": p.valore,
        "descrizione": p.descrizione, "tipo_dato": p.tipo_dato,
        "gruppo": p.gruppo
    } for p in parametri]

@app.put("/parametri-sistema/{chiave}")
def update_parametro_sistema(chiave: str, data: dict, db: Session = Depends(get_db)):
    parametro = db.query(ParametriSistema).filter(ParametriSistema.chiave == chiave).first()
    if not parametro:
        raise HTTPException(status_code=404, detail="Parametro non trovato")
    if "valore" in data:
        parametro.valore = data["valore"]
    parametro.updated_at = datetime.now()
    db.commit()
    return {"status": "ok"}

# ==========================================
# OPZIONI DROPDOWN
# ==========================================
@app.get("/opzioni-dropdown/gruppi")
def get_gruppi_dropdown(db: Session = Depends(get_db)):
    try:
        result = db.execute(text("""
            SELECT gruppo,
                   COUNT(*) as totale,
                   SUM(CASE WHEN attivo = 1 THEN 1 ELSE 0 END) as attive
            FROM opzioni_dropdown
            GROUP BY gruppo
            ORDER BY gruppo
        """))
        return [{"gruppo": r[0], "totale": r[1], "attive": r[2]} for r in result.fetchall()]
    except:
        return []

@app.get("/opzioni-dropdown/{gruppo}")
def get_opzioni_dropdown(gruppo: str, solo_attive: bool = True, db: Session = Depends(get_db)):
    try:
        q = "SELECT id, gruppo, valore, etichetta, ordine, attivo FROM opzioni_dropdown WHERE gruppo = :g"
        if solo_attive:
            q += " AND attivo = 1"
        q += " ORDER BY ordine"
        result = db.execute(text(q), {"g": gruppo})
        return [{"id": r[0], "gruppo": r[1], "valore": r[2], "label": r[3], "ordine": r[4], "attivo": bool(r[5])} for r in result.fetchall()]
    except:
        return []

@app.post("/opzioni-dropdown")
def create_opzione_dropdown(data: dict, db: Session = Depends(get_db)):
    try:
        db.execute(text("INSERT INTO opzioni_dropdown (gruppo, valore, etichetta, ordine, attivo) VALUES (:gruppo, :valore, :etichetta, :ordine, 1)"),
            {"gruppo": data["gruppo"], "valore": data["valore"], "etichetta": data.get("label", data["valore"]), "ordine": data.get("ordine", 0)})
        db.commit()
        return {"status": "ok"}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.put("/opzioni-dropdown/{opzione_id}")
def update_opzione_dropdown(opzione_id: int, data: dict, db: Session = Depends(get_db)):
    try:
        db.execute(text("UPDATE opzioni_dropdown SET valore=:valore, etichetta=:etichetta, ordine=:ordine, attivo=:attivo WHERE id=:id"),
            {"valore": data.get("valore"), "etichetta": data.get("label"), "ordine": data.get("ordine", 0), "attivo": 1 if data.get("attivo", True) else 0, "id": opzione_id})
        db.commit()
        return {"status": "ok"}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.delete("/opzioni-dropdown/{opzione_id}")
def delete_opzione_dropdown(opzione_id: int, db: Session = Depends(get_db)):
    try:
        db.execute(text("DELETE FROM opzioni_dropdown WHERE id=:id"), {"id": opzione_id})
        db.commit()
        return {"status": "ok"}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/opzioni-dropdown/{gruppo}/riordina")
def riordina_opzioni(gruppo: str, data: dict, db: Session = Depends(get_db)):
    try:
        for i, oid in enumerate(data.get("ids", [])):
            db.execute(text("UPDATE opzioni_dropdown SET ordine=:o WHERE id=:id"), {"o": i, "id": oid})
        db.commit()
        return {"status": "ok"}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

# ==========================================
# SEZIONI CONFIGURATORE
# ==========================================

def _ensure_sezioni_table(db):
    """Crea tabella sezioni_configuratore se non esiste"""
    try:
        db.execute(text('''
            CREATE TABLE IF NOT EXISTS sezioni_configuratore (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                codice TEXT NOT NULL UNIQUE,
                etichetta TEXT NOT NULL,
                descrizione TEXT,
                icona TEXT DEFAULT 'Settings',
                ordine INTEGER DEFAULT 0,
                attivo INTEGER DEFAULT 1,
                product_template_id INTEGER,
                created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                updated_at TIMESTAMP
            )
        '''))
        db.commit()
    except:
        pass


@app.get("/sezioni-configuratore")
def get_sezioni_configuratore_list(db: Session = Depends(get_db)):
    """Lista tutte le sezioni con info prodotto associato e conteggio campi"""
    _ensure_sezioni_table(db)
    try:
        result = db.execute(text("""
            SELECT s.id, s.codice, s.etichetta, s.descrizione, s.icona,
                   s.ordine, s.attivo, s.product_template_id,
                   pt.categoria, pt.sottocategoria, pt.nome_display,
                   (SELECT COUNT(*) FROM campi_configuratore c WHERE c.sezione = s.codice) as num_campi,
                   (SELECT COUNT(*) FROM campi_configuratore c WHERE c.sezione = s.codice AND c.attivo = 1) as num_campi_attivi
            FROM sezioni_configuratore s
            LEFT JOIN product_templates pt ON s.product_template_id = pt.id
            ORDER BY s.ordine, s.etichetta
        """))
        return [{
            "id": r[0], "codice": r[1], "etichetta": r[2],
            "descrizione": r[3], "icona": r[4], "ordine": r[5],
            "attivo": bool(r[6]), "product_template_id": r[7],
            "prodotto": {"categoria": r[8], "sottocategoria": r[9], "nome_display": r[10]} if r[7] else None,
            "num_campi": r[11], "num_campi_attivi": r[12]
        } for r in result.fetchall()]
    except Exception as e:
        print(f"Errore get_sezioni: {e}")
        return []


@app.get("/sezioni-configuratore/campi-non-assegnati")
def get_campi_non_assegnati(db: Session = Depends(get_db)):
    """Campi la cui sezione non corrisponde a nessuna sezione registrata"""
    _ensure_sezioni_table(db)
    try:
        result = db.execute(text("""
            SELECT c.id, c.codice, c.etichetta, c.tipo, c.sezione, c.attivo
            FROM campi_configuratore c
            WHERE c.sezione NOT IN (SELECT codice FROM sezioni_configuratore)
               OR c.sezione IS NULL OR c.sezione = ''
            ORDER BY c.sezione, c.ordine
        """))
        return [{"id": r[0], "codice": r[1], "etichetta": r[2], "tipo": r[3], "sezione": r[4], "attivo": bool(r[5])} for r in result.fetchall()]
    except:
        return []


@app.get("/sezioni-configuratore/{sezione_id}")
def get_sezione_detail(sezione_id: int, db: Session = Depends(get_db)):
    """Dettaglio sezione con campi associati"""
    _ensure_sezioni_table(db)
    try:
        result = db.execute(text("""
            SELECT s.id, s.codice, s.etichetta, s.descrizione, s.icona,
                   s.ordine, s.attivo, s.product_template_id,
                   pt.categoria, pt.sottocategoria, pt.nome_display
            FROM sezioni_configuratore s
            LEFT JOIN product_templates pt ON s.product_template_id = pt.id
            WHERE s.id = :id
        """), {"id": sezione_id})
        r = result.fetchone()
        if not r:
            raise HTTPException(status_code=404, detail="Sezione non trovata")
        
        campi_result = db.execute(text("""
            SELECT id, codice, etichetta, tipo, sezione, gruppo_dropdown, ordine, attivo,
                   obbligatorio, unita_misura, valore_default, visibile_form, usabile_regole
            FROM campi_configuratore WHERE sezione = :s ORDER BY ordine
        """), {"s": r[1]})
        campi = [{
            "id": c[0], "codice": c[1], "etichetta": c[2], "tipo": c[3],
            "sezione": c[4], "gruppo_dropdown": c[5], "ordine": c[6],
            "attivo": bool(c[7]), "obbligatorio": bool(c[8]),
            "unita_misura": c[9], "valore_default": c[10],
            "visibile_form": bool(c[11]) if c[11] is not None else True,
            "usabile_regole": bool(c[12]) if c[12] is not None else False
        } for c in campi_result.fetchall()]
        
        return {
            "id": r[0], "codice": r[1], "etichetta": r[2],
            "descrizione": r[3], "icona": r[4], "ordine": r[5],
            "attivo": bool(r[6]), "product_template_id": r[7],
            "prodotto": {"categoria": r[8], "sottocategoria": r[9], "nome_display": r[10]} if r[7] else None,
            "campi": campi
        }
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/sezioni-configuratore")
def create_sezione_configuratore(data: dict, db: Session = Depends(get_db)):
    """Crea una nuova sezione"""
    _ensure_sezioni_table(db)
    try:
        existing = db.execute(text("SELECT id FROM sezioni_configuratore WHERE codice = :c"), {"c": data["codice"]}).fetchone()
        if existing:
            raise HTTPException(status_code=400, detail=f"Codice '{data['codice']}' già esistente")
        
        max_ord = db.execute(text("SELECT COALESCE(MAX(ordine), -1) FROM sezioni_configuratore")).fetchone()[0]
        db.execute(text("""
            INSERT INTO sezioni_configuratore (codice, etichetta, descrizione, icona, ordine, attivo, product_template_id)
            VALUES (:codice, :etichetta, :descrizione, :icona, :ordine, :attivo, :pt_id)
        """), {
            "codice": data["codice"],
            "etichetta": data.get("etichetta", data["codice"]),
            "descrizione": data.get("descrizione"),
            "icona": data.get("icona", "Settings"),
            "ordine": data.get("ordine", max_ord + 1),
            "attivo": 1 if data.get("attivo", True) else 0,
            "pt_id": data.get("product_template_id")
        })
        db.commit()
        new_id = db.execute(text("SELECT id FROM sezioni_configuratore WHERE codice = :c"), {"c": data["codice"]}).fetchone()[0]
        return {"status": "ok", "id": new_id}
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.put("/sezioni-configuratore/{sezione_id}")
def update_sezione_configuratore(sezione_id: int, data: dict, db: Session = Depends(get_db)):
    """Aggiorna una sezione"""
    _ensure_sezioni_table(db)
    try:
        existing = db.execute(text("SELECT id, codice FROM sezioni_configuratore WHERE id = :id"), {"id": sezione_id}).fetchone()
        if not existing:
            raise HTTPException(status_code=404, detail="Sezione non trovata")
        old_codice = existing[1]
        
        field_map = {"codice": "codice", "etichetta": "etichetta", "descrizione": "descrizione",
                     "icona": "icona", "ordine": "ordine", "attivo": "attivo", "product_template_id": "product_template_id"}
        fields, params = [], {"id": sezione_id}
        for key, col in field_map.items():
            if key in data:
                val = data[key]
                if key == "attivo":
                    val = 1 if val else 0
                fields.append(f"{col}=:{col}")
                params[col] = val
        
        if fields:
            fields.append("updated_at=CURRENT_TIMESTAMP")
            db.execute(text(f"UPDATE sezioni_configuratore SET {','.join(fields)} WHERE id=:id"), params)
            new_codice = data.get("codice")
            if new_codice and new_codice != old_codice:
                db.execute(text("UPDATE campi_configuratore SET sezione=:new WHERE sezione=:old"), {"new": new_codice, "old": old_codice})
            db.commit()
        return {"status": "ok"}
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.delete("/sezioni-configuratore/{sezione_id}")
def delete_sezione_configuratore(sezione_id: int, db: Session = Depends(get_db)):
    """Elimina una sezione"""
    _ensure_sezioni_table(db)
    try:
        existing = db.execute(text("SELECT id FROM sezioni_configuratore WHERE id = :id"), {"id": sezione_id}).fetchone()
        if not existing:
            raise HTTPException(status_code=404, detail="Sezione non trovata")
        db.execute(text("DELETE FROM sezioni_configuratore WHERE id=:id"), {"id": sezione_id})
        db.commit()
        return {"status": "ok"}
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/sezioni-configuratore/riordina")
def riordina_sezioni(data: dict, db: Session = Depends(get_db)):
    """Riordina le sezioni"""
    _ensure_sezioni_table(db)
    try:
        for i, sid in enumerate(data.get("ids", [])):
            db.execute(text("UPDATE sezioni_configuratore SET ordine=:o WHERE id=:id"), {"o": i, "id": sid})
        db.commit()
        return {"status": "ok"}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/sezioni-configuratore/{sezione_id}/associa-campi")
def associa_campi_a_sezione(sezione_id: int, data: dict, db: Session = Depends(get_db)):
    """Associa/sposta campi a una sezione"""
    _ensure_sezioni_table(db)
    try:
        sezione = db.execute(text("SELECT codice FROM sezioni_configuratore WHERE id = :id"), {"id": sezione_id}).fetchone()
        if not sezione:
            raise HTTPException(status_code=404, detail="Sezione non trovata")
        codice_sezione = sezione[0]
        for cid in data.get("campo_ids", []):
            db.execute(text("UPDATE campi_configuratore SET sezione=:s WHERE id=:id"), {"s": codice_sezione, "id": cid})
        db.commit()
        return {"status": "ok", "campi_associati": len(data.get("campo_ids", []))}
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/sezioni-configuratore/importa-esistenti")
def importa_sezioni_esistenti(db: Session = Depends(get_db)):
    """Importa sezioni esistenti come stringhe in campi_configuratore"""
    _ensure_sezioni_table(db)
    try:
        result = db.execute(text("SELECT DISTINCT sezione FROM campi_configuratore WHERE sezione IS NOT NULL AND sezione != ''"))
        sezioni_esistenti = [r[0] for r in result.fetchall()]
        importate = 0
        for i, codice in enumerate(sorted(sezioni_esistenti)):
            exists = db.execute(text("SELECT id FROM sezioni_configuratore WHERE codice = :c"), {"c": codice}).fetchone()
            if not exists:
                etichetta = codice.replace("_", " ").title()
                db.execute(text("INSERT INTO sezioni_configuratore (codice, etichetta, ordine, attivo) VALUES (:codice, :etichetta, :ordine, 1)"),
                          {"codice": codice, "etichetta": etichetta, "ordine": i * 10})
                importate += 1
        db.commit()
        return {"status": "ok", "importate": importate, "totale_trovate": len(sezioni_esistenti)}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


# ==========================================
# CAMPI CONFIGURATORE
# ==========================================
@app.get("/campi-configuratore/sezioni")
def get_sezioni_configuratore(db: Session = Depends(get_db)):
    try:
        result = db.execute(text("""
            SELECT sezione,
                   COUNT(*) as totale,
                   SUM(CASE WHEN attivo = 1 THEN 1 ELSE 0 END) as attivi
            FROM campi_configuratore
            GROUP BY sezione
            ORDER BY sezione
        """))
        return [{"codice": r[0], "etichetta": r[0], "totale": r[1], "attivi": r[2]} for r in result.fetchall()]
    except:
        return []

@app.get("/campi-configuratore/schema.json")
def get_schema_campi(db: Session = Depends(get_db)):
    try:
        result = db.execute(text("SELECT codice, etichetta, tipo, sezione FROM campi_configuratore WHERE attivo=1 ORDER BY sezione, ordine"))
        return {"campi": [{"codice": r[0], "label": r[1], "tipo": r[2], "sezione": r[3]} for r in result.fetchall()]}
    except:
        return {"campi": []}

@app.get("/campi-configuratore/{sezione}")
def get_campi_sezione(sezione: str, solo_attivi: bool = True, include_opzioni: bool = False, db: Session = Depends(get_db)):
    try:
        q = """SELECT id, codice, etichetta, tipo, sezione, gruppo_dropdown, ordine, attivo, obbligatorio,
                      unita_misura, valore_min, valore_max, valore_default, descrizione, visibile_form, usabile_regole
               FROM campi_configuratore WHERE sezione=:s"""
        if solo_attivi:
            q += " AND attivo=1"
        q += " ORDER BY ordine"
        result = db.execute(text(q), {"s": sezione})
        return [{"id": r[0], "codice": r[1], "label": r[2], "tipo": r[3], "sezione": r[4],
                 "gruppo_opzioni": r[5], "ordine": r[6], "attivo": bool(r[7]), "obbligatorio": bool(r[8]),
                 "unita_misura": r[9], "valore_min": r[10], "valore_max": r[11],
                 "valore_default": r[12], "descrizione": r[13],
                 "visibile_form": bool(r[14]) if r[14] is not None else True,
                 "usabile_regole": bool(r[15]) if r[15] is not None else False
                } for r in result.fetchall()]
    except:
        return []

@app.post("/campi-configuratore")
def create_campo(data: dict, db: Session = Depends(get_db)):
    try:
        db.execute(text("""INSERT INTO campi_configuratore
            (codice, etichetta, tipo, sezione, gruppo_dropdown, ordine, attivo, obbligatorio,
             unita_misura, valore_min, valore_max, valore_default, descrizione, visibile_form, usabile_regole)
            VALUES (:codice, :etichetta, :tipo, :sezione, :gruppo, :ordine, 1, :obb,
                    :unita, :vmin, :vmax, :vdef, :desc, :vform, :uregole)"""),
            {"codice": data["codice"], "etichetta": data.get("label", data.get("etichetta", "")),
             "tipo": data.get("tipo", "text"), "sezione": data["sezione"],
             "gruppo": data.get("gruppo_opzioni", data.get("gruppo_dropdown")),
             "ordine": data.get("ordine", 0), "obb": 1 if data.get("obbligatorio") else 0,
             "unita": data.get("unita_misura"), "vmin": data.get("valore_min"),
             "vmax": data.get("valore_max"), "vdef": data.get("valore_default"),
             "desc": data.get("descrizione"), "vform": 1 if data.get("visibile_form", True) else 0,
             "uregole": 1 if data.get("usabile_regole") else 0})
        db.commit()
        return {"status": "ok"}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.put("/campi-configuratore/{campo_id}")
def update_campo(campo_id: int, data: dict, db: Session = Depends(get_db)):
    try:
        # Mappa nomi frontend -> nomi DB
        field_map = {
            "codice": "codice", "label": "etichetta", "etichetta": "etichetta",
            "tipo": "tipo", "sezione": "sezione",
            "gruppo_opzioni": "gruppo_dropdown", "gruppo_dropdown": "gruppo_dropdown",
            "ordine": "ordine", "attivo": "attivo", "obbligatorio": "obbligatorio",
            "unita_misura": "unita_misura", "valore_min": "valore_min",
            "valore_max": "valore_max", "valore_default": "valore_default",
            "descrizione": "descrizione", "visibile_form": "visibile_form",
            "usabile_regole": "usabile_regole"
        }
        fields, params = [], {"id": campo_id}
        for frontend_key, db_col in field_map.items():
            if frontend_key in data:
                fields.append(f"{db_col}=:{db_col}")
                params[db_col] = data[frontend_key]
        if fields:
            # Deduplica (label e etichetta mappano entrambi a etichetta)
            unique_fields = list(dict.fromkeys(fields))
            db.execute(text(f"UPDATE campi_configuratore SET {','.join(unique_fields)} WHERE id=:id"), params)
            db.commit()
        return {"status": "ok"}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.delete("/campi-configuratore/{campo_id}")
def delete_campo(campo_id: int, db: Session = Depends(get_db)):
    try:
        db.execute(text("DELETE FROM campi_configuratore WHERE id=:id"), {"id": campo_id})
        db.commit()
        return {"status": "ok"}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

# ==========================================
# STARTUP
# ==========================================
@app.on_event("startup")
def startup_event():
    """Inizializzazione all'avvio"""
    db = SessionLocal()
    try:
        create_default_admin(db)
    except Exception as e:
        print(f"âš ï¸ Errore creazione admin: {e}")
    finally:
        db.close()

if __name__ == "__main__":
    import uvicorn
    print("ðŸš€ Avvio Configuratore Elettroquadri API v0.11.0")
    print("ðŸ“¡ Server in ascolto su http://0.0.0.0:8000")
    print("ðŸ“— Documentazione API: http://0.0.0.0:8000/docs")
    uvicorn.run(app, host="0.0.0.0", port=8000)
