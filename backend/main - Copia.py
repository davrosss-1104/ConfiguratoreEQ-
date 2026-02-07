from fastapi import FastAPI, Depends, HTTPException, UploadFile, File, Query
import shutil
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import StreamingResponse
from sqlalchemy.orm import Session
from sqlalchemy import or_, func
from typing import List, Optional
from pydantic import BaseModel
import json
import os
import io
import xml.etree.ElementTree as ET
from datetime import datetime

from database import engine, SessionLocal, Base
from models import (
    Preventivo, DatiCommessa, DatiPrincipali, Normative, 
    DisposizioneVano, Porte, Argano, Materiale, RigaRicambio,
    Cliente, Articolo, CategoriaArticolo, ParametriSistema,
    GruppoUtenti, PermessoGruppo, Utente, TemplatePreventivo,
    OpzioneDropdown, CampoConfiguratore
)
from schemas import (
    PreventivoCreate, PreventivoUpdate, Preventivo as PreventivoSchema,
    DatiCommessaCreate, DatiCommessaUpdate, DatiCommessa as DatiCommessaSchema,
    DatiPrincipaliCreate, DatiPrincipaliUpdate, DatiPrincipali as DatiPrincipaliSchema,
    NormativeCreate, NormativeUpdate, Normative as NormativeSchema,
    DisposizioneVanoCreate, DisposizioneVanoUpdate, DisposizioneVano as DisposizioneVanoSchema,
    PorteCreate, PorteUpdate, Porte as PorteSchema,
    ArganoCreate, ArganoUpdate, Argano as ArganoSchema,
    MaterialeCreate, MaterialeUpdate, Materiale as MaterialeSchema,
    RigaRicambioCreate, RigaRicambioUpdate, RigaRicambio as RigaRicambioSchema,
    ClienteCreate, ClienteUpdate, Cliente as ClienteSchema,
    ArticoloCreate, ArticoloUpdate, Articolo as ArticoloSchema,
    CategoriaArticolo as CategoriaArticoloSchema,
    ParametriSistema as ParametriSistemaSchema,
    CalcoloPrezzoOutput,
    TemplatePreventivoCreate, TemplatePreventivoUpdate, TemplatePreventivo as TemplatePreventivoSchema,
    OpzioneDropdownCreate, OpzioneDropdownUpdate, OpzioneDropdown as OpzioneDropdownSchema,
    CampoConfiguratoreCreate, CampoConfiguratoreUpdate, CampoConfiguratore as CampoConfiguratoreSchema
)

# Crea tabelle
Base.metadata.create_all(bind=engine)

app = FastAPI(title="Configuratore Elettroquadri API", version="2.0.0")

# CORS
app.add_middleware(
    CORSMiddleware,
    allow_origins=["http://localhost:5173", "http://localhost:3000", "http://localhost:3001", "http://127.0.0.1:5173", "http://127.0.0.1:3000"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

def get_db():
    db = SessionLocal()
    try:
        yield db
    finally:
        db.close()


# ==========================================
# AUTENTICAZIONE
# ==========================================
class LoginRequest(BaseModel):
    username: str
    password: str

class LoginResponse(BaseModel):
    access_token: str
    token_type: str = "bearer"
    user: dict

@app.post("/auth/login", response_model=LoginResponse)
def login(request: LoginRequest, db: Session = Depends(get_db)):
    """Login utente - restituisce token JWT"""
    try:
        # Cerca utente nel database
        utente = db.query(Utente).filter(Utente.username == request.username).first()
        
        if not utente:
            # Utente demo se non esiste nel DB
            if request.username == "admin" and request.password == "admin":
                return LoginResponse(
                    access_token="demo-token-admin-12345",
                    user={
                        "id": 1,
                        "username": "admin",
                        "nome": "Amministratore",
                        "email": "admin@elettroquadri.it",
                        "is_admin": True
                    }
                )
            if request.username == "utente" and request.password == "utente":
                return LoginResponse(
                    access_token="demo-token-utente-12345",
                    user={
                        "id": 2,
                        "username": "utente",
                        "nome": "Mario",
                        "cognome": "Rossi",
                        "email": "mario.rossi@elettroquadri.it",
                        "is_admin": False
                    }
                )
            raise HTTPException(status_code=401, detail="Credenziali non valide")
        
        # Verifica password (semplificata - in produzione usare hash)
        if utente.password_hash != request.password:
            raise HTTPException(status_code=401, detail="Password non valida")
        
        if not utente.is_active:
            raise HTTPException(status_code=401, detail="Utente disabilitato")
        
        # Genera token (semplificato)
        token = f"token-{utente.id}-{utente.username}-{datetime.now().timestamp()}"
        
        # Controlla se is_admin esiste
        is_admin = getattr(utente, 'is_admin', False)
        
        return LoginResponse(
            access_token=token,
            user={
                "id": utente.id,
                "username": utente.username,
                "nome": utente.nome,
                "cognome": utente.cognome,
                "email": utente.email,
                "is_admin": is_admin
            }
        )
    except HTTPException:
        raise
    except Exception as e:
        print(f"Errore login: {e}")
        raise HTTPException(status_code=500, detail=f"Errore interno: {str(e)}")

@app.get("/auth/me")
def get_current_user():
    """Restituisce utente corrente (placeholder)"""
    return {
        "id": 1,
        "username": "admin",
        "nome": "Amministratore",
        "is_admin": True
    }


# ==========================================
# GESTIONE UTENTI (Admin only)
# ==========================================
@app.get("/api/utenti")
def get_utenti(db: Session = Depends(get_db)):
    """Lista tutti gli utenti"""
    utenti = db.query(Utente).order_by(Utente.id).all()
    return [
        {
            "id": u.id,
            "username": u.username,
            "nome": u.nome,
            "cognome": u.cognome,
            "email": u.email,
            "is_admin": u.is_admin,
            "is_active": u.is_active,
            "created_at": u.created_at.isoformat() if u.created_at else None
        }
        for u in utenti
    ]

@app.post("/api/utenti")
def create_utente(
    username: str = Query(...),
    password: str = Query(...),
    nome: Optional[str] = Query(None),
    cognome: Optional[str] = Query(None),
    email: Optional[str] = Query(None),
    is_admin: str = Query("false"),
    db: Session = Depends(get_db)
):
    """Crea un nuovo utente"""
    # Verifica username unico
    existing = db.query(Utente).filter(Utente.username == username).first()
    if existing:
        raise HTTPException(status_code=400, detail="Username già esistente")
    
    utente = Utente(
        username=username,
        password_hash=password,  # In produzione: usare hash
        nome=nome,
        cognome=cognome,
        email=email,
        is_admin=is_admin.lower() == "true",
        is_active=True
    )
    db.add(utente)
    db.commit()
    db.refresh(utente)
    
    return {
        "id": utente.id,
        "username": utente.username,
        "nome": utente.nome,
        "cognome": utente.cognome,
        "is_admin": utente.is_admin,
        "message": "Utente creato con successo"
    }

@app.put("/api/utenti/{utente_id}")
def update_utente(
    utente_id: int,
    nome: Optional[str] = Query(None),
    cognome: Optional[str] = Query(None),
    email: Optional[str] = Query(None),
    password: Optional[str] = Query(None),
    is_admin: Optional[str] = Query(None),
    is_active: Optional[str] = Query(None),
    db: Session = Depends(get_db)
):
    """Modifica un utente"""
    utente = db.query(Utente).filter(Utente.id == utente_id).first()
    if not utente:
        raise HTTPException(status_code=404, detail="Utente non trovato")
    
    if nome is not None:
        utente.nome = nome
    if cognome is not None:
        utente.cognome = cognome
    if email is not None:
        utente.email = email
    if password is not None:
        utente.password_hash = password
    if is_admin is not None:
        utente.is_admin = is_admin.lower() == "true"
    if is_active is not None:
        utente.is_active = is_active.lower() == "true"
    
    db.commit()
    return {"message": "Utente aggiornato", "id": utente.id}

@app.delete("/api/utenti/{utente_id}")
def delete_utente(utente_id: int, db: Session = Depends(get_db)):
    """Elimina un utente"""
    utente = db.query(Utente).filter(Utente.id == utente_id).first()
    if not utente:
        raise HTTPException(status_code=404, detail="Utente non trovato")
    
    db.delete(utente)
    db.commit()
    return {"message": "Utente eliminato"}


# ==========================================
# OPZIONI DROPDOWN
# ==========================================
@app.get("/api/opzioni-dropdown", response_model=List[OpzioneDropdownSchema])
def get_all_opzioni(
    gruppo: Optional[str] = Query(None, description="Filtra per gruppo"),
    solo_attive: bool = Query(True, description="Solo opzioni attive"),
    db: Session = Depends(get_db)
):
    """Lista tutte le opzioni dropdown, opzionalmente filtrate per gruppo"""
    query = db.query(OpzioneDropdown)
    
    if gruppo:
        query = query.filter(OpzioneDropdown.gruppo == gruppo)
    
    if solo_attive:
        query = query.filter(OpzioneDropdown.attivo == True)
    
    return query.order_by(OpzioneDropdown.gruppo, OpzioneDropdown.ordine).all()


@app.get("/api/opzioni-dropdown/gruppi")
def get_gruppi_opzioni(db: Session = Depends(get_db)):
    """Lista tutti i gruppi disponibili con conteggio opzioni"""
    result = db.query(
        OpzioneDropdown.gruppo,
        func.count(OpzioneDropdown.id).label('count'),
        func.sum(func.cast(OpzioneDropdown.attivo, Integer)).label('attive')
    ).group_by(OpzioneDropdown.gruppo).all()
    
    return [{"gruppo": r[0], "totale": r[1], "attive": r[2]} for r in result]


@app.get("/api/opzioni-dropdown/{gruppo}", response_model=List[OpzioneDropdownSchema])
def get_opzioni_by_gruppo(
    gruppo: str, 
    solo_attive: bool = Query(True),
    db: Session = Depends(get_db)
):
    """Lista opzioni per un gruppo specifico (per i form)"""
    query = db.query(OpzioneDropdown).filter(OpzioneDropdown.gruppo == gruppo)
    
    if solo_attive:
        query = query.filter(OpzioneDropdown.attivo == True)
    
    return query.order_by(OpzioneDropdown.ordine).all()


@app.post("/api/opzioni-dropdown", response_model=OpzioneDropdownSchema, status_code=201)
def create_opzione(opzione: OpzioneDropdownCreate, db: Session = Depends(get_db)):
    """Crea nuova opzione dropdown"""
    # Verifica duplicato
    existing = db.query(OpzioneDropdown).filter(
        OpzioneDropdown.gruppo == opzione.gruppo,
        OpzioneDropdown.valore == opzione.valore
    ).first()
    if existing:
        raise HTTPException(status_code=400, detail="Opzione già esistente per questo gruppo")
    
    db_opzione = OpzioneDropdown(**opzione.model_dump())
    db.add(db_opzione)
    db.commit()
    db.refresh(db_opzione)
    return db_opzione


@app.put("/api/opzioni-dropdown/{opzione_id}", response_model=OpzioneDropdownSchema)
def update_opzione(opzione_id: int, opzione: OpzioneDropdownUpdate, db: Session = Depends(get_db)):
    """Aggiorna opzione dropdown"""
    db_opzione = db.query(OpzioneDropdown).filter(OpzioneDropdown.id == opzione_id).first()
    if not db_opzione:
        raise HTTPException(status_code=404, detail="Opzione non trovata")
    
    update_data = opzione.model_dump(exclude_unset=True)
    for key, value in update_data.items():
        setattr(db_opzione, key, value)
    
    db.commit()
    db.refresh(db_opzione)
    return db_opzione


@app.delete("/api/opzioni-dropdown/{opzione_id}")
def delete_opzione(opzione_id: int, db: Session = Depends(get_db)):
    """Elimina opzione dropdown"""
    db_opzione = db.query(OpzioneDropdown).filter(OpzioneDropdown.id == opzione_id).first()
    if not db_opzione:
        raise HTTPException(status_code=404, detail="Opzione non trovata")
    
    db.delete(db_opzione)
    db.commit()
    return {"message": "Opzione eliminata"}


@app.post("/api/opzioni-dropdown/batch")
def create_opzioni_batch(opzioni: List[OpzioneDropdownCreate], db: Session = Depends(get_db)):
    """Crea multiple opzioni in batch"""
    created = 0
    skipped = 0
    
    for opzione in opzioni:
        existing = db.query(OpzioneDropdown).filter(
            OpzioneDropdown.gruppo == opzione.gruppo,
            OpzioneDropdown.valore == opzione.valore
        ).first()
        
        if existing:
            skipped += 1
            continue
        
        db_opzione = OpzioneDropdown(**opzione.model_dump())
        db.add(db_opzione)
        created += 1
    
    db.commit()
    return {"created": created, "skipped": skipped}


@app.put("/api/opzioni-dropdown/{gruppo}/riordina")
def riordina_opzioni(gruppo: str, ordini: List[dict], db: Session = Depends(get_db)):
    """Riordina le opzioni di un gruppo. ordini = [{"id": 1, "ordine": 0}, ...]"""
    for item in ordini:
        db.query(OpzioneDropdown).filter(
            OpzioneDropdown.id == item["id"],
            OpzioneDropdown.gruppo == gruppo
        ).update({"ordine": item["ordine"]})
    
    db.commit()
    return {"message": f"Riordinato gruppo {gruppo}"}


@app.put("/api/opzioni-dropdown/{gruppo}/sposta-sezione")
def sposta_gruppo_sezione(gruppo: str, nuova_sezione: str = Query(...), db: Session = Depends(get_db)):
    """
    Sposta un gruppo di opzioni in una nuova sezione.
    Aggiorna automaticamente tutti i campi che usano quel gruppo_dropdown.
    """
    # Verifica che il gruppo esista
    opzioni_count = db.query(OpzioneDropdown).filter(OpzioneDropdown.gruppo == gruppo).count()
    if opzioni_count == 0:
        raise HTTPException(status_code=404, detail=f"Gruppo '{gruppo}' non trovato")
    
    # Trova tutti i campi che usano questo gruppo e aggiorna la loro sezione
    campi_aggiornati = db.query(CampoConfiguratore).filter(
        CampoConfiguratore.gruppo_dropdown == gruppo
    ).update({"sezione": nuova_sezione})
    
    db.commit()
    
    return {
        "message": f"Gruppo '{gruppo}' spostato in '{nuova_sezione}'",
        "campi_aggiornati": campi_aggiornati
    }


# Importa Integer per la query di conteggio
from sqlalchemy import Integer


# ==========================================
# CAMPI CONFIGURATORE (per Rule Designer)
# ==========================================
@app.get("/api/campi-configuratore")
def get_all_campi(
    sezione: Optional[str] = Query(None, description="Filtra per sezione"),
    tipo: Optional[str] = Query(None, description="Filtra per tipo (dropdown, numero, testo, booleano)"),
    solo_attivi: bool = Query(True, description="Solo campi attivi"),
    solo_regole: bool = Query(False, description="Solo campi usabili nelle regole"),
    include_opzioni: bool = Query(True, description="Include opzioni per dropdown"),
    db: Session = Depends(get_db)
):
    """
    Lista tutti i campi del configuratore.
    Usato dal Rule Designer per costruire l'interfaccia regole.
    """
    query = db.query(CampoConfiguratore)
    
    if sezione:
        query = query.filter(CampoConfiguratore.sezione == sezione)
    if tipo:
        query = query.filter(CampoConfiguratore.tipo == tipo)
    if solo_attivi:
        query = query.filter(CampoConfiguratore.attivo == True)
    if solo_regole:
        query = query.filter(CampoConfiguratore.usabile_regole == True)
    
    campi = query.order_by(CampoConfiguratore.sezione, CampoConfiguratore.ordine).all()
    
    # Se richiesto, aggiungi le opzioni per i dropdown
    if include_opzioni:
        result = []
        for campo in campi:
            campo_dict = {
                "id": campo.id,
                "codice": campo.codice,
                "etichetta": campo.etichetta,
                "tipo": campo.tipo,
                "sezione": campo.sezione,
                "gruppo_dropdown": campo.gruppo_dropdown,
                "unita_misura": campo.unita_misura,
                "valore_min": campo.valore_min,
                "valore_max": campo.valore_max,
                "valore_default": campo.valore_default,
                "descrizione": campo.descrizione,
                "obbligatorio": campo.obbligatorio,
                "ordine": campo.ordine,
                "attivo": campo.attivo,
                "visibile_form": campo.visibile_form,
                "usabile_regole": campo.usabile_regole,
                "opzioni": None
            }
            
            # Carica opzioni se è un dropdown
            if campo.tipo == "dropdown" and campo.gruppo_dropdown:
                opzioni = db.query(OpzioneDropdown).filter(
                    OpzioneDropdown.gruppo == campo.gruppo_dropdown,
                    OpzioneDropdown.attivo == True
                ).order_by(OpzioneDropdown.ordine).all()
                campo_dict["opzioni"] = [{"value": o.valore, "label": o.etichetta} for o in opzioni]
            
            result.append(campo_dict)
        return result
    
    return campi


@app.get("/api/campi-configuratore/sezioni")
def get_sezioni_campi(db: Session = Depends(get_db)):
    """Lista tutte le sezioni disponibili con conteggio campi"""
    result = db.query(
        CampoConfiguratore.sezione,
        func.count(CampoConfiguratore.id).label('totale'),
        func.sum(func.cast(CampoConfiguratore.attivo, Integer)).label('attivi')
    ).group_by(CampoConfiguratore.sezione).all()
    
    # Etichette sezioni
    SEZIONI_LABELS = {
        "dati_commessa": "Dati Commessa",
        "dati_principali": "Dati Principali",
        "tensioni": "Tensioni",
        "normative": "Normative",
        "argano": "Argano / Trazione",
        "porte_lato_a": "Porte Lato A",
        "operatore_a": "Operatore Porte A",
        "porte_lato_b": "Porte Lato B",
        "cabina": "Cabina",
        "vano": "Vano",
        "quadro": "Quadro Elettrico",
    }
    
    return [{
        "codice": r[0], 
        "etichetta": SEZIONI_LABELS.get(r[0], r[0]),
        "totale": r[1], 
        "attivi": r[2]
    } for r in result]


@app.get("/api/campi-configuratore/schema.json")
def export_schema_json(solo_attivi: bool = Query(False), db: Session = Depends(get_db)):
    """
    Esporta lo schema completo in formato JSON.
    Usato dal Rule Designer per import configurazione.
    NOTA: Questo endpoint DEVE essere prima di /{sezione} per non essere intercettato
    """
    query = db.query(CampoConfiguratore)
    if solo_attivi:
        query = query.filter(CampoConfiguratore.attivo == True)
    
    campi = query.order_by(CampoConfiguratore.sezione, CampoConfiguratore.ordine).all()
    
    # Raggruppa per sezione
    schema = {"version": "1.0", "sezioni": {}}
    
    for campo in campi:
        if campo.sezione not in schema["sezioni"]:
            schema["sezioni"][campo.sezione] = {"campi": []}
        
        campo_dict = {
            "codice": campo.codice,
            "etichetta": campo.etichetta,
            "tipo": campo.tipo,
            "obbligatorio": campo.obbligatorio,
            "attivo": campo.attivo,
        }
        
        if campo.tipo == "dropdown" and campo.gruppo_dropdown:
            opzioni = db.query(OpzioneDropdown).filter(
                OpzioneDropdown.gruppo == campo.gruppo_dropdown,
                OpzioneDropdown.attivo == True
            ).order_by(OpzioneDropdown.ordine).all()
            campo_dict["opzioni"] = [{"value": o.valore, "label": o.etichetta} for o in opzioni]
            campo_dict["gruppo_dropdown"] = campo.gruppo_dropdown
        
        if campo.tipo == "numero":
            if campo.unita_misura:
                campo_dict["unita"] = campo.unita_misura
            if campo.valore_min is not None:
                campo_dict["min"] = campo.valore_min
            if campo.valore_max is not None:
                campo_dict["max"] = campo.valore_max
        
        if campo.valore_default:
            campo_dict["default"] = campo.valore_default
        
        if campo.descrizione:
            campo_dict["descrizione"] = campo.descrizione
        
        schema["sezioni"][campo.sezione]["campi"].append(campo_dict)
    
    return schema


@app.get("/api/campi-configuratore/{sezione}")
def get_campi_by_sezione(
    sezione: str,
    solo_attivi: bool = Query(True),
    include_opzioni: bool = Query(True),
    db: Session = Depends(get_db)
):
    """Lista campi per una sezione specifica"""
    query = db.query(CampoConfiguratore).filter(CampoConfiguratore.sezione == sezione)
    
    if solo_attivi:
        query = query.filter(CampoConfiguratore.attivo == True)
    
    campi = query.order_by(CampoConfiguratore.ordine).all()
    
    if include_opzioni:
        result = []
        for campo in campi:
            campo_dict = {
                "id": campo.id,
                "codice": campo.codice,
                "etichetta": campo.etichetta,
                "tipo": campo.tipo,
                "sezione": campo.sezione,
                "gruppo_dropdown": campo.gruppo_dropdown,
                "unita_misura": campo.unita_misura,
                "valore_min": campo.valore_min,
                "valore_max": campo.valore_max,
                "valore_default": campo.valore_default,
                "descrizione": campo.descrizione,
                "obbligatorio": campo.obbligatorio,
                "ordine": campo.ordine,
                "attivo": campo.attivo,
                "visibile_form": campo.visibile_form,
                "usabile_regole": campo.usabile_regole,
                "opzioni": None
            }
            
            if campo.tipo == "dropdown" and campo.gruppo_dropdown:
                opzioni = db.query(OpzioneDropdown).filter(
                    OpzioneDropdown.gruppo == campo.gruppo_dropdown,
                    OpzioneDropdown.attivo == True
                ).order_by(OpzioneDropdown.ordine).all()
                campo_dict["opzioni"] = [{"value": o.valore, "label": o.etichetta} for o in opzioni]
            
            result.append(campo_dict)
        return result
    
    return campi


@app.post("/api/campi-configuratore", status_code=201)
def create_campo(campo: CampoConfiguratoreCreate, db: Session = Depends(get_db)):
    """Crea nuovo campo configuratore"""
    existing = db.query(CampoConfiguratore).filter(CampoConfiguratore.codice == campo.codice).first()
    if existing:
        raise HTTPException(status_code=400, detail="Campo con questo codice già esistente")
    
    db_campo = CampoConfiguratore(**campo.model_dump())
    db.add(db_campo)
    db.commit()
    db.refresh(db_campo)
    return db_campo


@app.put("/api/campi-configuratore/{campo_id}")
def update_campo(campo_id: int, campo: CampoConfiguratoreUpdate, db: Session = Depends(get_db)):
    """Aggiorna campo configuratore"""
    db_campo = db.query(CampoConfiguratore).filter(CampoConfiguratore.id == campo_id).first()
    if not db_campo:
        raise HTTPException(status_code=404, detail="Campo non trovato")
    
    update_data = campo.model_dump(exclude_unset=True)
    for key, value in update_data.items():
        setattr(db_campo, key, value)
    
    db.commit()
    db.refresh(db_campo)
    return db_campo


@app.delete("/api/campi-configuratore/{campo_id}")
def delete_campo(campo_id: int, db: Session = Depends(get_db)):
    """Elimina campo configuratore"""
    db_campo = db.query(CampoConfiguratore).filter(CampoConfiguratore.id == campo_id).first()
    if not db_campo:
        raise HTTPException(status_code=404, detail="Campo non trovato")
    
    db.delete(db_campo)
    db.commit()
    return {"message": "Campo eliminato"}


# ==========================================
# TEMPLATE PREVENTIVI
# ==========================================
@app.get("/api/templates", response_model=List[TemplatePreventivoSchema])
def get_templates(user_id: Optional[int] = Query(None), db: Session = Depends(get_db)):
    """
    Lista template disponibili per l'utente.
    Restituisce: template pubblici + template personali dell'utente
    """
    if user_id:
        # Template pubblici + template dell'utente
        templates = db.query(TemplatePreventivo).filter(
            or_(
                TemplatePreventivo.is_public == True,
                TemplatePreventivo.created_by == user_id
            )
        ).order_by(TemplatePreventivo.is_public.desc(), TemplatePreventivo.nome).all()
    else:
        # Solo template pubblici
        templates = db.query(TemplatePreventivo).filter(
            TemplatePreventivo.is_public == True
        ).order_by(TemplatePreventivo.nome).all()
    
    return templates

@app.post("/api/templates", response_model=TemplatePreventivoSchema)
def create_template(data: TemplatePreventivoCreate, user_id: int = Query(...), is_admin: str = Query("false"), db: Session = Depends(get_db)):
    """
    Crea un nuovo template.
    Solo admin può creare template pubblici.
    """
    is_admin_bool = str(is_admin).lower() == "true"
    
    # Solo admin può creare template pubblici
    if data.is_public and not is_admin_bool:
        raise HTTPException(status_code=403, detail="Solo gli amministratori possono creare template pubblici")
    
    template = TemplatePreventivo(
        nome=data.nome,
        descrizione=data.descrizione,
        is_public=data.is_public if is_admin_bool else False,
        dati_json=data.dati_json,
        created_by=user_id
    )
    
    db.add(template)
    db.commit()
    db.refresh(template)
    return template

@app.get("/api/templates/{template_id}", response_model=TemplatePreventivoSchema)
def get_template(template_id: int, db: Session = Depends(get_db)):
    """Dettaglio template"""
    template = db.query(TemplatePreventivo).filter(TemplatePreventivo.id == template_id).first()
    if not template:
        raise HTTPException(status_code=404, detail="Template non trovato")
    return template

@app.delete("/api/templates/{template_id}")
def delete_template(template_id: int, user_id: int = Query(...), is_admin: str = Query("false"), db: Session = Depends(get_db)):
    """
    Cancella un template.
    - Admin può cancellare qualsiasi template
    - Utente normale può cancellare solo i propri template
    """
    is_admin_bool = str(is_admin).lower() == "true"
    
    template = db.query(TemplatePreventivo).filter(TemplatePreventivo.id == template_id).first()
    if not template:
        raise HTTPException(status_code=404, detail="Template non trovato")
    
    # Verifica permessi
    if not is_admin_bool and template.created_by != user_id:
        raise HTTPException(status_code=403, detail="Non hai i permessi per cancellare questo template")
    
    db.delete(template)
    db.commit()
    return {"message": "Template cancellato", "id": template_id}

@app.post("/api/templates/{template_id}/create-preventivo")
def create_preventivo_from_template(
    template_id: int, 
    cliente_id: Optional[int] = Query(None),
    user_id: Optional[int] = Query(None),
    db: Session = Depends(get_db)
):
    """
    Crea un nuovo preventivo usando un template.
    Copia tutti i dati dal template nel nuovo preventivo.
    """
    template = db.query(TemplatePreventivo).filter(TemplatePreventivo.id == template_id).first()
    if not template:
        raise HTTPException(status_code=404, detail="Template non trovato")
    
    # Genera numero preventivo
    year = datetime.now().year
    last_preventivo = db.query(Preventivo).filter(
        Preventivo.numero_preventivo.like(f"{year}/%")
    ).order_by(Preventivo.id.desc()).first()
    
    if last_preventivo:
        try:
            last_num = int(last_preventivo.numero_preventivo.split('/')[1])
            new_num = last_num + 1
        except:
            new_num = 1
    else:
        new_num = 1
    
    numero_preventivo = f"{year}/{new_num:04d}"
    
    # Crea preventivo
    preventivo = Preventivo(
        numero_preventivo=numero_preventivo,
        tipo_preventivo="COMPLETO",
        cliente_id=cliente_id,
        created_by=user_id,  # Chi crea il preventivo
        status="draft"
    )
    db.add(preventivo)
    db.commit()
    db.refresh(preventivo)
    
    # Applica dati dal template se presenti
    if template.dati_json:
        try:
            dati = json.loads(template.dati_json)
            
            # Dati Principali
            if 'dati_principali' in dati:
                dp = DatiPrincipali(preventivo_id=preventivo.id, **dati['dati_principali'])
                db.add(dp)
            
            # Normative
            if 'normative' in dati:
                norm = Normative(preventivo_id=preventivo.id, **dati['normative'])
                db.add(norm)
            
            # Argano
            if 'argano' in dati:
                arg = Argano(preventivo_id=preventivo.id, **dati['argano'])
                db.add(arg)
            
            # Disposizione Vano
            if 'disposizione_vano' in dati:
                dv = DisposizioneVano(preventivo_id=preventivo.id, **dati['disposizione_vano'])
                db.add(dv)
            
            # Porte
            if 'porte' in dati:
                porte = Porte(preventivo_id=preventivo.id, **dati['porte'])
                db.add(porte)
            
            db.commit()
        except Exception as e:
            print(f"Errore applicazione template: {e}")
    
    return {
        "preventivo_id": preventivo.id,
        "numero_preventivo": preventivo.numero_preventivo,
        "template_usato": template.nome
    }

@app.post("/api/preventivi/{preventivo_id}/save-as-template")
def save_preventivo_as_template(
    preventivo_id: int,
    nome: str = Query(...),
    descrizione: Optional[str] = Query(None),
    user_id: int = Query(0),
    is_admin: str = Query("false"),
    is_public: str = Query("false"),
    db: Session = Depends(get_db)
):
    """
    Salva un preventivo esistente come template.
    """
    try:
        # Converti stringhe in boolean
        is_admin_bool = is_admin.lower() == "true"
        is_public_bool = is_public.lower() == "true"
        
        preventivo = db.query(Preventivo).filter(Preventivo.id == preventivo_id).first()
        if not preventivo:
            raise HTTPException(status_code=404, detail="Preventivo non trovato")
        
        if preventivo.tipo_preventivo != "COMPLETO":
            raise HTTPException(status_code=400, detail="Solo preventivi COMPLETO possono essere salvati come template")
        
        # Solo admin può creare template pubblici
        if is_public_bool and not is_admin_bool:
            is_public_bool = False
        
        # Raccogli dati da salvare nel template
        dati = {}
        
        # Dati Principali
        dp = db.query(DatiPrincipali).filter(DatiPrincipali.preventivo_id == preventivo_id).first()
        if dp:
            dati['dati_principali'] = {
                'tipo_impianto': dp.tipo_impianto,
                'nuovo_impianto': dp.nuovo_impianto,
                'numero_fermate': dp.numero_fermate,
                'numero_servizi': dp.numero_servizi,
                'velocita': float(dp.velocita) if dp.velocita else None,
                'corsa': float(dp.corsa) if dp.corsa else None,
                'forza_motrice': dp.forza_motrice,
                'luce': dp.luce,
                'tensione_manovra': dp.tensione_manovra,
                'tensione_freno': dp.tensione_freno,
            }
        
        # Normative (con campi corretti)
        norm = db.query(Normative).filter(Normative.preventivo_id == preventivo_id).first()
        if norm:
            dati['normative'] = {
                'en_81_1': norm.en_81_1,
                'en_81_20': norm.en_81_20,
                'en_81_21': norm.en_81_21,
                'en_81_28': norm.en_81_28,
                'en_81_70': norm.en_81_70,
                'en_81_72': norm.en_81_72,
                'en_81_73': norm.en_81_73,
                'a3_95_16': norm.a3_95_16,
                'dm236_legge13': norm.dm236_legge13,
                'emendamento_a3': norm.emendamento_a3,
                'uni_10411_1': norm.uni_10411_1,
            }
        
        # Argano
        arg = db.query(Argano).filter(Argano.preventivo_id == preventivo_id).first()
        if arg:
            dati['argano'] = {
                'trazione': arg.trazione,
                'potenza_motore_kw': float(arg.potenza_motore_kw) if arg.potenza_motore_kw else None,
                'corrente_nom_motore_amp': float(arg.corrente_nom_motore_amp) if arg.corrente_nom_motore_amp else None,
                'tipo_vvvf': arg.tipo_vvvf,
                'vvvf_nel_vano': arg.vvvf_nel_vano,
                'freno_tensione': arg.freno_tensione,
                'ventilazione_forzata': arg.ventilazione_forzata,
                'tipo_teleruttore': arg.tipo_teleruttore,
            }
        
        # Disposizione Vano
        dv = db.query(DisposizioneVano).filter(DisposizioneVano.preventivo_id == preventivo_id).first()
        if dv:
            dati['disposizione_vano'] = {
                'posizione_quadro_lato': dv.posizione_quadro_lato,
                'posizione_quadro_piano': dv.posizione_quadro_piano,
                'altezza_vano': float(dv.altezza_vano) if dv.altezza_vano else None,
                'piano_piu_alto': dv.piano_piu_alto,
                'piano_piu_basso': dv.piano_piu_basso,
                'note': dv.note,
            }
        
        # Porte
        porte = db.query(Porte).filter(Porte.preventivo_id == preventivo_id).first()
        if porte:
            dati['porte'] = {
                'tipo_porte_piano': porte.tipo_porte_piano,
                'tipo_porte_cabina': porte.tipo_porte_cabina,
                'numero_accessi': porte.numero_accessi,
                'tipo_operatore': porte.tipo_operatore,
                'marca_operatore': porte.marca_operatore,
                'stazionamento_porte': porte.stazionamento_porte,
                'tipo_apertura': porte.tipo_apertura,
                'distanza_minima_accessi': float(porte.distanza_minima_accessi) if porte.distanza_minima_accessi else None,
                'alimentazione_operatore': porte.alimentazione_operatore,
                'con_scheda': porte.con_scheda,
            }
        
        # Crea template
        template = TemplatePreventivo(
            nome=nome,
            descrizione=descrizione,
            is_public=is_public_bool,
            dati_json=json.dumps(dati),
            created_by=user_id if user_id > 0 else None
        )
        
        db.add(template)
        db.commit()
        db.refresh(template)
        
        return {
            "template_id": template.id,
            "nome": template.nome,
            "is_public": template.is_public,
            "message": "Template creato con successo"
        }
    except HTTPException:
        raise
    except Exception as e:
        print(f"Errore salvataggio template: {e}")
        raise HTTPException(status_code=500, detail=f"Errore salvataggio template: {str(e)}")


# ==========================================
# CALCOLO PREZZO A 3 LIVELLI (fino a 4 parametri)
# ==========================================
class ParametroInput(BaseModel):
    param_1: Optional[float] = None
    param_2: Optional[float] = None
    param_3: Optional[float] = None
    param_4: Optional[float] = None

def calcola_prezzo_articolo(
    db: Session,
    articolo: Articolo,
    quantita: float = 1,
    parametri: Optional[ParametroInput] = None,
    cliente_id: Optional[int] = None
) -> dict:
    """
    Calcola il prezzo con logica a 3 livelli e fino a 4 parametri variabili:
    
    STEP 1 - COSTO BASE:
    costo_base = costo_fisso + Σ(costo_variabile_i × parametro_i)
    
    Es: Quadro elettrico con:
    - costo_fisso = 100€
    - costo_var_1 = 5€/fermata, param_1 = 8 fermate → 40€
    - costo_var_2 = 2€/metro cavo, param_2 = 50 metri → 100€
    - costo_base = 100 + 40 + 100 = 240€
    
    STEP 2 - RICARICO (dipende da tipo articolo):
    prezzo_listino = costo_base × (1 + ricarico%)
    
    STEP 3 - SCONTO CLIENTE:
    prezzo_cliente = prezzo_listino × (1 - sconto_cliente%)
    """
    
    # Dati articolo
    costo_fisso = float(articolo.costo_fisso or 0)
    tipo_articolo = articolo.tipo_articolo or "PRODUZIONE"
    
    # Prepara parametri
    params = parametri or ParametroInput()
    
    # Estrai costi variabili e parametri
    costi_var = [
        (float(articolo.costo_variabile_1 or 0), params.param_1, articolo.unita_misura_var_1, articolo.descrizione_var_1),
        (float(articolo.costo_variabile_2 or 0), params.param_2, articolo.unita_misura_var_2, articolo.descrizione_var_2),
        (float(articolo.costo_variabile_3 or 0), params.param_3, articolo.unita_misura_var_3, articolo.descrizione_var_3),
        (float(articolo.costo_variabile_4 or 0), params.param_4, articolo.unita_misura_var_4, articolo.descrizione_var_4),
    ]
    
    # STEP 1: Calcola costo base
    costo_base = costo_fisso
    parametri_output = []
    dettaglio_parts = [f"Fisso €{costo_fisso:.2f}"]
    
    for i, (costo_var, param_val, um, desc) in enumerate(costi_var, 1):
        if costo_var > 0:
            # Se parametro non fornito, usa 1 come default
            p = param_val if param_val is not None else 1
            subtotale = costo_var * p
            costo_base += subtotale
            
            parametri_output.append({
                "valore": p,
                "costo_variabile": costo_var,
                "unita_misura": um,
                "descrizione": desc or f"Parametro {i}",
                "subtotale": round(subtotale, 4)
            })
            
            dettaglio_parts.append(f"+ (€{costo_var:.4f} × {p} {um or ''}) = €{subtotale:.2f}")
    
    # STEP 2: Ricarico
    ricarico = float(articolo.ricarico_percentuale or 0)
    if ricarico == 0:
        chiave_default = f"ricarico_{tipo_articolo.lower()}_default"
        param_ricarico = db.query(ParametriSistema).filter(
            ParametriSistema.chiave == chiave_default
        ).first()
        if param_ricarico:
            ricarico = float(param_ricarico.valore)
    
    prezzo_listino = costo_base * (1 + ricarico / 100)
    
    # STEP 3: Sconto Cliente
    sconto_cliente = 0.0
    if cliente_id:
        cliente = db.query(Cliente).filter(Cliente.id == cliente_id).first()
        if cliente:
            if tipo_articolo == "PRODUZIONE":
                sconto_cliente = float(cliente.sconto_produzione or 0)
            else:
                sconto_cliente = float(cliente.sconto_acquisto or 0)
    
    prezzo_cliente = prezzo_listino * (1 - sconto_cliente / 100)
    
    # Totali
    prezzo_totale_listino = prezzo_listino * quantita
    prezzo_totale_cliente = prezzo_cliente * quantita
    
    # Dettaglio
    dettaglio = " ".join(dettaglio_parts)
    dettaglio += f" | Ricarico {ricarico:.1f}% → €{prezzo_listino:.2f}"
    if sconto_cliente > 0:
        dettaglio += f" | Sconto {sconto_cliente:.1f}% → €{prezzo_cliente:.2f}"
    
    return {
        "costo_fisso": round(costo_fisso, 4),
        "parametri": parametri_output,
        "costo_base_unitario": round(costo_base, 4),
        "ricarico_percentuale": round(ricarico, 2),
        "prezzo_listino_unitario": round(prezzo_listino, 4),
        "sconto_cliente_percentuale": round(sconto_cliente, 2),
        "prezzo_cliente_unitario": round(prezzo_cliente, 4),
        "quantita": quantita,
        "prezzo_totale_listino": round(prezzo_totale_listino, 2),
        "prezzo_totale_cliente": round(prezzo_totale_cliente, 2),
        "tipo_articolo": tipo_articolo,
        "dettaglio_calcolo": dettaglio
    }


# ==========================================
# RULE ENGINE
# ==========================================
def load_rules():
    rules = []
    rules_dir = "./rules"
    if not os.path.exists(rules_dir):
        return rules
    for filename in os.listdir(rules_dir):
        if filename.endswith(".json"):
            try:
                with open(os.path.join(rules_dir, filename), "r", encoding="utf-8") as f:
                    rules.append(json.load(f))
            except Exception as e:
                print(f"Errore caricamento {filename}: {e}")
    return rules

def evaluate_condition(condition: dict, config_data: dict) -> bool:
    if isinstance(condition, str):
        return False
    field = condition.get("field")
    operator = condition.get("operator")
    expected_value = condition.get("value")
    if field not in config_data:
        return False
    actual_value = config_data.get(field)
    if actual_value is None or actual_value == "":
        return False
    
    if operator == "equals":
        return str(actual_value) == str(expected_value)
    elif operator == "not_equals":
        return str(actual_value) != str(expected_value)
    elif operator == "contains":
        return str(expected_value).lower() in str(actual_value).lower()
    elif operator == "greater_than":
        return float(actual_value) > float(expected_value)
    elif operator == "less_than":
        return float(actual_value) < float(expected_value)
    elif operator == "in":
        return str(actual_value) in expected_value
    return False

def evaluate_rules(preventivo_id: int, db: Session) -> dict:
    """
    Valuta le regole e aggiunge/rimuove materiali automaticamente.
    FIX: Recupera il prezzo dalla tabella Articolo invece che dalla regola JSON.
    """
    preventivo = db.query(Preventivo).filter(Preventivo.id == preventivo_id).first()
    if not preventivo:
        return {"error": "Preventivo non trovato"}
    
    # Raccogli dati configurazione
    dati_principali = db.query(DatiPrincipali).filter(DatiPrincipali.preventivo_id == preventivo_id).first()
    normative = db.query(Normative).filter(Normative.preventivo_id == preventivo_id).first()
    
    config_data = {}
    if dati_principali:
        for col in ["tipo_impianto", "nuovo_impianto", "numero_fermate", "tipo_trazione", "velocita", "corsa"]:
            config_data[col] = getattr(dati_principali, col, None)
    if normative:
        for col in ["en_81_1", "en_81_20", "en_81_21", "en_81_28", "en_81_70", "en_81_72", "en_81_73"]:
            config_data[col] = getattr(normative, col, None)
    
    # Carica e valuta regole
    rules = load_rules()
    active_rules = []
    materials_added = 0
    materials_removed = 0
    
    for rule in rules:
        rule_id = rule.get("id", "unknown")
        conditions = rule.get("conditions", [])
        
        # FIX: Se non ci sono condizioni, la regola NON deve scattare automaticamente
        # (a meno che non sia marcata esplicitamente come "default")
        if not conditions and not rule.get("is_default", False):
            print(f"⚠️ Regola {rule_id} senza condizioni - saltata")
            continue
        
        # Valuta tutte le condizioni (se vuota e is_default=True, scatta sempre)
        all_conditions_met = len(conditions) == 0 or all(evaluate_condition(c, config_data) for c in conditions)
        
        if all_conditions_met:
            active_rules.append(rule_id)
            
            # Aggiungi materiali della regola
            for mat in rule.get("materials", []):
                # IMPORTANTE: definisco codice_articolo QUI
                codice_articolo = mat.get("codice", "")
                
                # Controlla se già esiste
                existing = db.query(Materiale).filter(
                    Materiale.preventivo_id == preventivo_id,
                    Materiale.codice == codice_articolo,
                    Materiale.regola_id == rule_id
                ).first()
                
                if not existing:
                    # ============================================
                    # FIX: Cerca il prezzo nella tabella Articolo
                    # ============================================
                    articolo = db.query(Articolo).filter(
                        Articolo.codice == codice_articolo,
                        Articolo.is_active == True
                    ).first()
                    
                    if articolo:
                        # Calcola prezzo_listino dall'articolo (costo_fisso * (1 + ricarico%))
                        costo = float(articolo.costo_fisso or 0)
                        ricarico = float(articolo.ricarico_percentuale or 0)
                        prezzo_unitario = costo * (1 + ricarico / 100)
                        descrizione = articolo.descrizione
                        # Campi corretti per lead time
                        lead_time_giorni = int(articolo.lead_time_giorni or 0)
                        manodopera_giorni = int(articolo.manodopera_giorni or 0)
                        categoria = articolo.categoria.nome if hasattr(articolo, 'categoria') and articolo.categoria else mat.get("categoria")
                    else:
                        # Fallback: usa dati dalla regola (per articoli non ancora in anagrafica)
                        prezzo_unitario = float(mat.get("prezzo_unitario", 0))
                        descrizione = mat.get("descrizione", codice_articolo)
                        lead_time_giorni = 0
                        manodopera_giorni = 0
                        categoria = mat.get("categoria")
                        print(f"⚠️ Articolo {codice_articolo} non trovato in anagrafica, uso dati regola")
                    
                    quantita = mat.get("quantita", 1)
                    prezzo_totale = quantita * prezzo_unitario
                    
                    nuovo = Materiale(
                        preventivo_id=preventivo_id,
                        codice=codice_articolo,
                        descrizione=descrizione,
                        quantita=quantita,
                        prezzo_unitario=prezzo_unitario,
                        prezzo_totale=prezzo_totale,
                        categoria=categoria,
                        aggiunto_da_regola=True,
                        regola_id=rule_id,
                        lead_time_giorni=lead_time_giorni,
                        manodopera_giorni=manodopera_giorni
                    )
                    db.add(nuovo)
                    materials_added += 1
                    print(f"➕ Aggiunto: {codice_articolo} - {descrizione} (€{prezzo_totale:.2f})")
        else:
            # Rimuovi materiali di regola non più attiva
            obsolete = db.query(Materiale).filter(
                Materiale.preventivo_id == preventivo_id,
                Materiale.regola_id == rule_id,
                Materiale.aggiunto_da_regola == True
            ).all()
            for m in obsolete:
                db.delete(m)
                materials_removed += 1
                print(f"➖ Rimosso: {m.codice} (regola {rule_id} non più attiva)")
    
    db.commit()
    
    # Aggiorna totale preventivo
    all_mats = db.query(Materiale).filter(Materiale.preventivo_id == preventivo_id).all()
    preventivo.total_price = sum(float(m.prezzo_totale or 0) for m in all_mats)
    
    # FIX: Aggiorna anche total_price_finale applicando gli sconti
    # Recupera sconti cliente e admin
    sconto_cliente = 0.0
    if preventivo.cliente_id:
        cliente = db.query(Cliente).filter(Cliente.id == preventivo.cliente_id).first()
        if cliente:
            sconto_cliente = float(cliente.sconto_globale or 0)
    sconto_admin = float(preventivo.sconto_extra_admin or 0)
    
    # Applica sconti
    totale_dopo_sconto_cliente = preventivo.total_price * (1 - sconto_cliente / 100)
    preventivo.total_price_finale = totale_dopo_sconto_cliente * (1 - sconto_admin / 100)
    
    db.commit()
    
    return {
        "active_rules": active_rules,
        "materials_added": materials_added,
        "materials_removed": materials_removed,
        "total_materials": len(all_mats),
        "total_price": float(preventivo.total_price)
    }

# ==========================================
# PARAMETRI SISTEMA
# ==========================================
@app.get("/api/parametri-sistema", response_model=List[ParametriSistemaSchema])
def get_parametri_sistema(db: Session = Depends(get_db)):
    return db.query(ParametriSistema).all()


# ==========================================
# CLIENTI
# ==========================================
@app.get("/api/clienti/search", response_model=List[ClienteSchema])
def search_clienti(
    q: str = Query(..., min_length=2),
    limit: int = Query(20, le=100),
    db: Session = Depends(get_db)
):
    search_term = f"%{q}%"
    return db.query(Cliente).filter(
        Cliente.is_active == True,
        or_(
            Cliente.codice.ilike(search_term),
            Cliente.ragione_sociale.ilike(search_term),
            Cliente.partita_iva.ilike(search_term)
        )
    ).limit(limit).all()

@app.get("/api/clienti", response_model=List[ClienteSchema])
def get_clienti(skip: int = 0, limit: int = 100, db: Session = Depends(get_db)):
    return db.query(Cliente).filter(Cliente.is_active == True).offset(skip).limit(limit).all()

@app.get("/api/clienti/{cliente_id}", response_model=ClienteSchema)
def get_cliente(cliente_id: int, db: Session = Depends(get_db)):
    cliente = db.query(Cliente).filter(Cliente.id == cliente_id).first()
    if not cliente:
        raise HTTPException(status_code=404, detail="Cliente non trovato")
    return cliente

@app.post("/api/clienti", response_model=ClienteSchema, status_code=201)
def create_cliente(data: ClienteCreate, db: Session = Depends(get_db)):
    if db.query(Cliente).filter(Cliente.codice == data.codice).first():
        raise HTTPException(status_code=400, detail="Codice già esistente")
    cliente = Cliente(**data.dict())
    db.add(cliente)
    db.commit()
    db.refresh(cliente)
    return cliente

@app.put("/api/clienti/{cliente_id}", response_model=ClienteSchema)
def update_cliente(cliente_id: int, data: ClienteUpdate, db: Session = Depends(get_db)):
    cliente = db.query(Cliente).filter(Cliente.id == cliente_id).first()
    if not cliente:
        raise HTTPException(status_code=404, detail="Cliente non trovato")
    for key, value in data.dict(exclude_unset=True).items():
        setattr(cliente, key, value)
    db.commit()
    db.refresh(cliente)
    return cliente


# ==========================================
# CATEGORIE ARTICOLI
# ==========================================
@app.get("/api/categorie-articoli", response_model=List[CategoriaArticoloSchema])
def get_categorie_articoli(db: Session = Depends(get_db)):
    return db.query(CategoriaArticolo).filter(CategoriaArticolo.is_active == True).order_by(CategoriaArticolo.ordine).all()


# ==========================================
# ARTICOLI
# ==========================================
@app.get("/api/articoli/search", response_model=List[ArticoloSchema])
def search_articoli(
    q: str = Query(..., min_length=1),
    categoria_id: Optional[int] = None,
    tipo_articolo: Optional[str] = None,
    limit: int = Query(20, le=100),
    db: Session = Depends(get_db)
):
    search_term = f"%{q}%"
    query = db.query(Articolo).filter(
        Articolo.is_active == True,
        or_(Articolo.codice.ilike(search_term), Articolo.descrizione.ilike(search_term))
    )
    if categoria_id:
        query = query.filter(Articolo.categoria_id == categoria_id)
    if tipo_articolo:
        query = query.filter(Articolo.tipo_articolo == tipo_articolo)
    return query.limit(limit).all()

@app.get("/api/articoli", response_model=List[ArticoloSchema])
def get_articoli(
    skip: int = 0, 
    limit: int = 100, 
    search: Optional[str] = Query(None),
    db: Session = Depends(get_db)
):
    query = db.query(Articolo).filter(Articolo.is_active == True)
    
    if search and len(search) >= 2:
        search_pattern = f"%{search}%"
        query = query.filter(
            (Articolo.codice.ilike(search_pattern)) | 
            (Articolo.descrizione.ilike(search_pattern))
        )
    
    return query.offset(skip).limit(limit).all()

@app.get("/api/articoli/{articolo_id}", response_model=ArticoloSchema)
def get_articolo(articolo_id: int, db: Session = Depends(get_db)):
    articolo = db.query(Articolo).filter(Articolo.id == articolo_id).first()
    if not articolo:
        raise HTTPException(status_code=404, detail="Articolo non trovato")
    return articolo

@app.post("/api/articoli", response_model=ArticoloSchema, status_code=201)
def create_articolo(data: ArticoloCreate, db: Session = Depends(get_db)):
    if db.query(Articolo).filter(Articolo.codice == data.codice).first():
        raise HTTPException(status_code=400, detail="Codice già esistente")
    articolo = Articolo(**data.dict())
    db.add(articolo)
    db.commit()
    db.refresh(articolo)
    return articolo

@app.put("/api/articoli/{articolo_id}", response_model=ArticoloSchema)
def update_articolo(articolo_id: int, data: ArticoloUpdate, db: Session = Depends(get_db)):
    articolo = db.query(Articolo).filter(Articolo.id == articolo_id).first()
    if not articolo:
        raise HTTPException(status_code=404, detail="Articolo non trovato")
    for key, value in data.dict(exclude_unset=True).items():
        setattr(articolo, key, value)
    db.commit()
    db.refresh(articolo)
    return articolo

@app.delete("/api/articoli/{articolo_id}")
def delete_articolo(articolo_id: int, db: Session = Depends(get_db)):
    articolo = db.query(Articolo).filter(Articolo.id == articolo_id).first()
    if not articolo:
        raise HTTPException(status_code=404, detail="Articolo non trovato")
    articolo.is_active = False
    db.commit()
    return {"message": "Articolo eliminato"}


# ==========================================
# CALCOLA PREZZO ENDPOINT
# ==========================================
class CalcoloPrezzoRequest(BaseModel):
    articolo_id: int
    quantita: float = 1.0
    param_1: Optional[float] = None
    param_2: Optional[float] = None
    param_3: Optional[float] = None
    param_4: Optional[float] = None
    cliente_id: Optional[int] = None

@app.post("/api/articoli/calcola-prezzo")
def calcola_prezzo_endpoint(data: CalcoloPrezzoRequest, db: Session = Depends(get_db)):
    articolo = db.query(Articolo).filter(Articolo.id == data.articolo_id).first()
    if not articolo:
        raise HTTPException(status_code=404, detail="Articolo non trovato")
    
    parametri = ParametroInput(
        param_1=data.param_1,
        param_2=data.param_2,
        param_3=data.param_3,
        param_4=data.param_4
    )
    
    return calcola_prezzo_articolo(db, articolo, data.quantita, parametri, data.cliente_id)


# ==========================================
# PREVENTIVI
# ==========================================
@app.get("/api/preventivi", response_model=List[PreventivoSchema])
def get_preventivi(
    skip: int = 0, 
    limit: int = 100, 
    user_id: Optional[int] = Query(None),
    is_admin: str = Query("false"),
    db: Session = Depends(get_db)
):
    """
    Lista preventivi.
    - Admin (is_admin=true): vede tutti i preventivi
    - Utente normale: vede solo i propri preventivi (created_by = user_id)
    """
    query = db.query(Preventivo)
    
    # Se non è admin e ha un user_id, filtra per created_by
    is_admin_bool = is_admin.lower() == "true"
    if not is_admin_bool and user_id:
        query = query.filter(Preventivo.created_by == user_id)
    
    return query.order_by(Preventivo.created_at.desc()).offset(skip).limit(limit).all()

@app.get("/api/preventivi/{preventivo_id}", response_model=PreventivoSchema)
def get_preventivo(preventivo_id: int, db: Session = Depends(get_db)):
    preventivo = db.query(Preventivo).filter(Preventivo.id == preventivo_id).first()
    if not preventivo:
        raise HTTPException(status_code=404, detail="Preventivo non trovato")
    return preventivo

@app.post("/api/preventivi", response_model=PreventivoSchema, status_code=201)
def create_preventivo(data: PreventivoCreate, db: Session = Depends(get_db)):
    year = datetime.now().year
    count = db.query(Preventivo).filter(Preventivo.numero_preventivo.like(f"{year}/%")).count()
    numero = f"{year}/{str(count + 1).zfill(4)}"
    
    preventivo = Preventivo(
        numero_preventivo=numero,
        tipo_preventivo=data.tipo_preventivo,
        cliente_id=data.cliente_id,
        customer_name=data.customer_name,
        created_by=data.user_id,  # Salva chi ha creato il preventivo
        status="draft"
    )
    db.add(preventivo)
    db.commit()
    db.refresh(preventivo)
    
    # Se COMPLETO, crea tabelle correlate
    if data.tipo_preventivo == "COMPLETO":
        db.add(DatiCommessa(preventivo_id=preventivo.id))
        db.add(DatiPrincipali(preventivo_id=preventivo.id))
        db.add(Normative(preventivo_id=preventivo.id))
        db.add(DisposizioneVano(preventivo_id=preventivo.id))
        db.add(Porte(preventivo_id=preventivo.id))
        db.commit()
    
    return preventivo

@app.put("/api/preventivi/{preventivo_id}", response_model=PreventivoSchema)
def update_preventivo(preventivo_id: int, data: PreventivoUpdate, db: Session = Depends(get_db)):
    preventivo = db.query(Preventivo).filter(Preventivo.id == preventivo_id).first()
    if not preventivo:
        raise HTTPException(status_code=404, detail="Preventivo non trovato")
    for key, value in data.dict(exclude_unset=True).items():
        setattr(preventivo, key, value)
    preventivo.updated_at = datetime.utcnow()
    db.commit()
    db.refresh(preventivo)
    return preventivo

@app.delete("/api/preventivi/{preventivo_id}")
def delete_preventivo(preventivo_id: int, db: Session = Depends(get_db)):
    preventivo = db.query(Preventivo).filter(Preventivo.id == preventivo_id).first()
    if not preventivo:
        raise HTTPException(status_code=404, detail="Preventivo non trovato")
    db.delete(preventivo)
    db.commit()
    return {"message": "Preventivo eliminato"}


# ==========================================
# DATI COMMESSA
# ==========================================
@app.get("/api/preventivi/{preventivo_id}/dati-commessa")
def get_dati_commessa(preventivo_id: int, db: Session = Depends(get_db)):
    dati = db.query(DatiCommessa).filter(DatiCommessa.preventivo_id == preventivo_id).first()
    if not dati:
        # Crea record vuoto se non esiste
        dati = DatiCommessa(preventivo_id=preventivo_id)
        db.add(dati)
        db.commit()
        db.refresh(dati)
    
    # Ritorna con data convertita in stringa per JSON
    return {
        "id": dati.id,
        "preventivo_id": dati.preventivo_id,
        "numero_offerta": dati.numero_offerta,
        "data_offerta": dati.data_offerta,
        "riferimento_cliente": dati.riferimento_cliente,
        "quantita": dati.quantita,
        "data_consegna_richiesta": dati.data_consegna_richiesta.isoformat() if dati.data_consegna_richiesta else None,
        "imballo": dati.imballo,
        "reso_fco": dati.reso_fco,
        "pagamento": dati.pagamento,
        "trasporto": dati.trasporto,
        "destinazione": dati.destinazione,
    }

@app.put("/api/preventivi/{preventivo_id}/dati-commessa")
def update_dati_commessa(preventivo_id: int, data: DatiCommessaUpdate, db: Session = Depends(get_db)):
    from datetime import date as date_type
    
    dati = db.query(DatiCommessa).filter(DatiCommessa.preventivo_id == preventivo_id).first()
    if not dati:
        dati = DatiCommessa(preventivo_id=preventivo_id)
        db.add(dati)
    
    for key, value in data.dict(exclude_unset=True).items():
        # Gestione speciale per data_consegna_richiesta
        if key == 'data_consegna_richiesta' and value:
            try:
                if isinstance(value, str) and value:
                    value = date_type.fromisoformat(value)
                elif value == '':
                    value = None
            except (ValueError, TypeError):
                value = None
        setattr(dati, key, value)
    
    db.commit()
    db.refresh(dati)
    
    # ⚡ SINCRONIZZA NOME CLIENTE CON PREVENTIVO
    preventivo = db.query(Preventivo).filter(Preventivo.id == preventivo_id).first()
    if preventivo and hasattr(dati, 'cliente_id') and dati.cliente_id:
        cliente = db.query(Cliente).filter(Cliente.id == dati.cliente_id).first()
        if cliente:
            preventivo.customer_name = cliente.ragione_sociale
            preventivo.cliente_id = cliente.id
            db.commit()
    elif preventivo and hasattr(dati, 'cliente_id') and not dati.cliente_id:
        # Se cliente rimosso, pulisci anche il preventivo
        preventivo.customer_name = None
        preventivo.cliente_id = None
        db.commit()
    
    # Ritorna con data convertita in stringa per JSON
    result = {
        "id": dati.id,
        "preventivo_id": dati.preventivo_id,
        "numero_offerta": dati.numero_offerta,
        "data_offerta": dati.data_offerta,
        "riferimento_cliente": dati.riferimento_cliente,
        "quantita": dati.quantita,
        "data_consegna_richiesta": dati.data_consegna_richiesta.isoformat() if dati.data_consegna_richiesta else None,
        "imballo": dati.imballo,
        "reso_fco": dati.reso_fco,
        "pagamento": dati.pagamento,
        "trasporto": dati.trasporto,
        "destinazione": dati.destinazione,
    }
    return result


# ==========================================
# DATI PRINCIPALI
# ==========================================
@app.get("/api/preventivi/{preventivo_id}/dati-principali", response_model=DatiPrincipaliSchema)
def get_dati_principali(preventivo_id: int, db: Session = Depends(get_db)):
    dati = db.query(DatiPrincipali).filter(DatiPrincipali.preventivo_id == preventivo_id).first()
    if not dati:
        dati = DatiPrincipali(preventivo_id=preventivo_id)
        db.add(dati)
        db.commit()
        db.refresh(dati)
    return dati

@app.put("/api/preventivi/{preventivo_id}/dati-principali", response_model=DatiPrincipaliSchema)
def update_dati_principali(preventivo_id: int, data: DatiPrincipaliUpdate, db: Session = Depends(get_db)):
    dati = db.query(DatiPrincipali).filter(DatiPrincipali.preventivo_id == preventivo_id).first()
    if not dati:
        dati = DatiPrincipali(preventivo_id=preventivo_id)
        db.add(dati)
    for key, value in data.dict(exclude_unset=True).items():
        setattr(dati, key, value)
    db.commit()
    db.refresh(dati)
    
    # ⚡ Rule Engine - valuta regole dopo salvataggio
    evaluate_rules(preventivo_id, db)
    
    return dati


# ==========================================
# NORMATIVE
# ==========================================
@app.get("/api/preventivi/{preventivo_id}/normative", response_model=NormativeSchema)
def get_normative(preventivo_id: int, db: Session = Depends(get_db)):
    dati = db.query(Normative).filter(Normative.preventivo_id == preventivo_id).first()
    if not dati:
        # Crea record vuoto se non esiste
        dati = Normative(preventivo_id=preventivo_id)
        db.add(dati)
        db.commit()
        db.refresh(dati)
    return dati

@app.put("/api/preventivi/{preventivo_id}/normative", response_model=NormativeSchema)
def update_normative(preventivo_id: int, data: NormativeUpdate, db: Session = Depends(get_db)):
    dati = db.query(Normative).filter(Normative.preventivo_id == preventivo_id).first()
    if not dati:
        dati = Normative(preventivo_id=preventivo_id)
        db.add(dati)
    
    for key, value in data.dict(exclude_unset=True).items():
        setattr(dati, key, value)
    
    db.commit()
    db.refresh(dati)
    
    # ⚡ Rule Engine - valuta regole dopo salvataggio
    evaluate_rules(preventivo_id, db)
    
    return dati


# ==========================================
# DISPOSIZIONE VANO
# ==========================================
@app.get("/api/preventivi/{preventivo_id}/disposizione-vano", response_model=DisposizioneVanoSchema)
def get_disposizione_vano(preventivo_id: int, db: Session = Depends(get_db)):
    dati = db.query(DisposizioneVano).filter(DisposizioneVano.preventivo_id == preventivo_id).first()
    if not dati:
        dati = DisposizioneVano(preventivo_id=preventivo_id)
        db.add(dati)
        db.commit()
        db.refresh(dati)
    return dati

@app.put("/api/preventivi/{preventivo_id}/disposizione-vano", response_model=DisposizioneVanoSchema)
def update_disposizione_vano(preventivo_id: int, data: DisposizioneVanoUpdate, db: Session = Depends(get_db)):
    dati = db.query(DisposizioneVano).filter(DisposizioneVano.preventivo_id == preventivo_id).first()
    if not dati:
        dati = DisposizioneVano(preventivo_id=preventivo_id)
        db.add(dati)
    for key, value in data.dict(exclude_unset=True).items():
        setattr(dati, key, value)
    db.commit()
    db.refresh(dati)
    return dati


# ==========================================
# PORTE
# ==========================================
@app.get("/api/preventivi/{preventivo_id}/porte", response_model=PorteSchema)
def get_porte(preventivo_id: int, db: Session = Depends(get_db)):
    dati = db.query(Porte).filter(Porte.preventivo_id == preventivo_id).first()
    if not dati:
        dati = Porte(preventivo_id=preventivo_id)
        db.add(dati)
        db.commit()
        db.refresh(dati)
    return dati

@app.put("/api/preventivi/{preventivo_id}/porte", response_model=PorteSchema)
def update_porte(preventivo_id: int, data: PorteUpdate, db: Session = Depends(get_db)):
    dati = db.query(Porte).filter(Porte.preventivo_id == preventivo_id).first()
    if not dati:
        dati = Porte(preventivo_id=preventivo_id)
        db.add(dati)
    for key, value in data.dict(exclude_unset=True).items():
        setattr(dati, key, value)
    db.commit()
    db.refresh(dati)
    return dati


# ==========================================
# ARGANO
# ==========================================
@app.get("/api/preventivi/{preventivo_id}/argano", response_model=ArganoSchema)
def get_argano(preventivo_id: int, db: Session = Depends(get_db)):
    dati = db.query(Argano).filter(Argano.preventivo_id == preventivo_id).first()
    if not dati:
        dati = Argano(preventivo_id=preventivo_id)
        db.add(dati)
        db.commit()
        db.refresh(dati)
    return dati

@app.post("/api/preventivi/{preventivo_id}/argano")
def save_argano(preventivo_id: int, data: dict, db: Session = Depends(get_db)):
    """Salva dati argano e applica regole automatiche per materiali"""
    
    # Verifica preventivo esista
    preventivo = db.query(Preventivo).filter(Preventivo.id == preventivo_id).first()
    if not preventivo:
        raise HTTPException(status_code=404, detail="Preventivo non trovato")
    
    # Crea o aggiorna argano
    argano = db.query(Argano).filter(Argano.preventivo_id == preventivo_id).first()
    
    # Salva trazione precedente per confronto
    trazione_precedente = argano.trazione if argano else None
    
    if not argano:
        argano = Argano(preventivo_id=preventivo_id)
        db.add(argano)
    
    # Campi validi per argano
    valid_fields = ['trazione', 'potenza_motore_kw', 'corrente_nom_motore_amp', 
                    'tipo_vvvf', 'vvvf_nel_vano', 'freno_tensione', 
                    'ventilazione_forzata', 'tipo_teleruttore']
    
    for key, value in data.items():
        if key in valid_fields:
            # Converti stringhe vuote in None
            if value == '':
                value = None
            setattr(argano, key, value)
    
    db.commit()
    db.refresh(argano)
    
    # Prendi il valore trazione per le regole
    trazione = data.get('trazione', '')
    
    # ==========================================
    # RULE ENGINE - Gestione materiali automatici
    # ==========================================
    materials_added = 0
    materials_removed = 0
    
    # Se la trazione è cambiata, rimuovi i materiali della regola precedente
    if trazione_precedente and trazione_precedente != trazione:
        # Rimuovi materiali della regola Gearless MRL
        deleted_gearless = db.query(Materiale).filter(
            Materiale.preventivo_id == preventivo_id,
            Materiale.regola_id == "RULE_GEARLESS_MRL"
        ).delete()
        materials_removed += deleted_gearless
        
        # Rimuovi materiali della regola Geared
        deleted_geared = db.query(Materiale).filter(
            Materiale.preventivo_id == preventivo_id,
            Materiale.regola_id == "RULE_GEARED"
        ).delete()
        materials_removed += deleted_geared
        
        db.commit()
    
    # Regola: Gearless MRL aggiunge componenti specifici
    if trazione == "Gearless MRL":
        # Verifica se già aggiunti
        existing = db.query(Materiale).filter(
            Materiale.preventivo_id == preventivo_id,
            Materiale.regola_id == "RULE_GEARLESS_MRL"
        ).first()
        
        if not existing:
            # (codice, descrizione, quantità, prezzo, categoria, lead_time_giorni, manodopera_giorni)
            materiali_gearless = [
                ("QEL-GRL-001", "Quadro elettrico Gearless MRL - Configurazione standard", 1, 2850.00, "Quadri Elettrici", 20, 5),
                ("INV-GRL-001", "Inverter per motore Gearless - Alta efficienza", 1, 1650.00, "Quadri Elettrici", 12, 2),
                ("ENC-MRL-001", "Encoder assoluto per trazione MRL", 1, 420.00, "Componenti Trazione", 8, 1),
                ("PCB-GRL-001", "Scheda controllo principale Gearless", 1, 580.00, "Componenti Elettronici", 10, 3),
            ]
            
            for codice, desc, qty, prezzo, cat, lt, mo in materiali_gearless:
                mat = Materiale(
                    preventivo_id=preventivo_id,
                    codice=codice,
                    descrizione=desc,
                    quantita=qty,
                    prezzo_unitario=prezzo,
                    prezzo_totale=qty * prezzo,
                    categoria=cat,
                    aggiunto_da_regola=True,
                    regola_id="RULE_GEARLESS_MRL",
                    lead_time_giorni=lt,
                    manodopera_giorni=mo
                )
                db.add(mat)
                materials_added += 1
    
    # Regola: Geared aggiunge componenti diversi
    elif trazione == "Geared":
        existing = db.query(Materiale).filter(
            Materiale.preventivo_id == preventivo_id,
            Materiale.regola_id == "RULE_GEARED"
        ).first()
        
        if not existing:
            # (codice, descrizione, quantità, prezzo, categoria, lead_time_giorni, manodopera_giorni)
            materiali_geared = [
                ("QEL-GRD-001", "Quadro elettrico Geared - Configurazione standard", 1, 2200.00, "Quadri Elettrici", 18, 4),
                ("INV-GRD-001", "Inverter per motore Geared", 1, 1200.00, "Quadri Elettrici", 10, 2),
                ("ENC-GRD-001", "Encoder incrementale per Geared", 1, 280.00, "Componenti Trazione", 6, 1),
                ("RID-GRD-001", "Riduttore meccanico", 1, 950.00, "Componenti Trazione", 25, 3),
            ]
            
            for codice, desc, qty, prezzo, cat, lt, mo in materiali_geared:
                mat = Materiale(
                    preventivo_id=preventivo_id,
                    codice=codice,
                    descrizione=desc,
                    quantita=qty,
                    prezzo_unitario=prezzo,
                    prezzo_totale=qty * prezzo,
                    categoria=cat,
                    aggiunto_da_regola=True,
                    regola_id="RULE_GEARED",
                    lead_time_giorni=lt,
                    manodopera_giorni=mo
                )
                db.add(mat)
                materials_added += 1
    
    if materials_added > 0 or materials_removed > 0:
        db.commit()
        # Ricalcola totale preventivo
        totale = db.query(func.sum(Materiale.prezzo_totale)).filter(
            Materiale.preventivo_id == preventivo_id
        ).scalar() or 0
        preventivo.total_price = totale
        db.commit()
    
    message_parts = []
    if materials_removed > 0:
        message_parts.append(f"{materials_removed} materiali rimossi")
    if materials_added > 0:
        message_parts.append(f"{materials_added} materiali aggiunti")
    
    return {
        "argano": argano,
        "materials_added": materials_added,
        "materials_removed": materials_removed,
        "message": f"Argano salvato. {', '.join(message_parts)}." if message_parts else "Argano salvato."
    }

@app.put("/api/preventivi/{preventivo_id}/argano", response_model=ArganoSchema)
def update_argano(preventivo_id: int, data: ArganoUpdate, db: Session = Depends(get_db)):
    dati = db.query(Argano).filter(Argano.preventivo_id == preventivo_id).first()
    if not dati:
        dati = Argano(preventivo_id=preventivo_id)
        db.add(dati)
    for key, value in data.dict(exclude_unset=True).items():
        setattr(dati, key, value)
    db.commit()
    db.refresh(dati)
    return dati


# ==========================================
# MATERIALI
# ==========================================
@app.get("/api/preventivi/{preventivo_id}/materiali", response_model=List[MaterialeSchema])
def get_materiali(preventivo_id: int, db: Session = Depends(get_db)):
    return db.query(Materiale).filter(Materiale.preventivo_id == preventivo_id).all()

@app.post("/api/preventivi/{preventivo_id}/materiali", response_model=MaterialeSchema)
def create_materiale(preventivo_id: int, data: MaterialeCreate, db: Session = Depends(get_db)):
    prezzo_totale = float(data.quantita) * float(data.prezzo_unitario)
    materiale = Materiale(
        preventivo_id=preventivo_id,
        codice=data.codice,
        descrizione=data.descrizione,
        quantita=data.quantita,
        prezzo_unitario=data.prezzo_unitario,
        prezzo_totale=prezzo_totale,
        categoria=data.categoria,
        aggiunto_da_regola=data.aggiunto_da_regola,
        regola_id=data.regola_id,
        note=data.note,
        lead_time_giorni=data.lead_time_giorni,
        manodopera_giorni=data.manodopera_giorni
    )
    db.add(materiale)
    db.commit()
    db.refresh(materiale)
    _aggiorna_totale_preventivo(db, preventivo_id)
    return materiale

@app.put("/api/preventivi/{preventivo_id}/materiali/{materiale_id}", response_model=MaterialeSchema)
def update_materiale(preventivo_id: int, materiale_id: int, data: MaterialeUpdate, db: Session = Depends(get_db)):
    materiale = db.query(Materiale).filter(Materiale.id == materiale_id, Materiale.preventivo_id == preventivo_id).first()
    if not materiale:
        raise HTTPException(status_code=404, detail="Materiale non trovato")
    for key, value in data.dict(exclude_unset=True).items():
        setattr(materiale, key, value)
    if data.quantita is not None or data.prezzo_unitario is not None:
        materiale.prezzo_totale = float(materiale.quantita) * float(materiale.prezzo_unitario)
    db.commit()
    db.refresh(materiale)
    _aggiorna_totale_preventivo(db, preventivo_id)
    return materiale

@app.delete("/api/preventivi/{preventivo_id}/materiali/{materiale_id}")
def delete_materiale(preventivo_id: int, materiale_id: int, db: Session = Depends(get_db)):
    materiale = db.query(Materiale).filter(Materiale.id == materiale_id, Materiale.preventivo_id == preventivo_id).first()
    if not materiale:
        raise HTTPException(status_code=404, detail="Materiale non trovato")
    db.delete(materiale)
    db.commit()
    _aggiorna_totale_preventivo(db, preventivo_id)
    return {"message": "Materiale eliminato"}


# ==========================================
# RIGHE RICAMBIO
# ==========================================
@app.get("/api/preventivi/{preventivo_id}/righe-ricambio", response_model=List[RigaRicambioSchema])
def get_righe_ricambio(preventivo_id: int, db: Session = Depends(get_db)):
    return db.query(RigaRicambio).filter(RigaRicambio.preventivo_id == preventivo_id).order_by(RigaRicambio.ordine).all()

@app.post("/api/preventivi/{preventivo_id}/righe-ricambio", response_model=RigaRicambioSchema, status_code=201)
def create_riga_ricambio(preventivo_id: int, data: RigaRicambioCreate, db: Session = Depends(get_db)):
    preventivo = db.query(Preventivo).filter(Preventivo.id == preventivo_id).first()
    if not preventivo:
        raise HTTPException(status_code=404, detail="Preventivo non trovato")
    
    max_ordine = db.query(func.max(RigaRicambio.ordine)).filter(RigaRicambio.preventivo_id == preventivo_id).scalar() or 0
    
    riga = RigaRicambio(
        preventivo_id=preventivo_id,
        ordine=max_ordine + 1,
        articolo_id=data.articolo_id,
        codice=data.codice,
        descrizione=data.descrizione,
        tipo_articolo=data.tipo_articolo,
        quantita=data.quantita,
        # 4 parametri
        parametro_1=data.parametro_1,
        unita_param_1=data.unita_param_1,
        desc_param_1=data.desc_param_1,
        costo_var_1=data.costo_var_1,
        parametro_2=data.parametro_2,
        unita_param_2=data.unita_param_2,
        desc_param_2=data.desc_param_2,
        costo_var_2=data.costo_var_2,
        parametro_3=data.parametro_3,
        unita_param_3=data.unita_param_3,
        desc_param_3=data.desc_param_3,
        costo_var_3=data.costo_var_3,
        parametro_4=data.parametro_4,
        unita_param_4=data.unita_param_4,
        desc_param_4=data.desc_param_4,
        costo_var_4=data.costo_var_4,
        # Costi
        costo_fisso=data.costo_fisso,
        costo_base_unitario=data.costo_base_unitario,
        ricarico_percentuale=data.ricarico_percentuale,
        prezzo_listino_unitario=data.prezzo_listino_unitario,
        sconto_cliente=data.sconto_cliente,
        prezzo_cliente_unitario=data.prezzo_cliente_unitario,
        prezzo_totale_listino=data.prezzo_totale_listino,
        prezzo_totale_cliente=data.prezzo_totale_cliente,
        note=data.note
    )
    db.add(riga)
    db.commit()
    db.refresh(riga)
    _aggiorna_totale_preventivo_ricambio(db, preventivo_id)
    return riga

@app.put("/api/preventivi/{preventivo_id}/righe-ricambio/{riga_id}", response_model=RigaRicambioSchema)
def update_riga_ricambio(preventivo_id: int, riga_id: int, data: RigaRicambioUpdate, db: Session = Depends(get_db)):
    riga = db.query(RigaRicambio).filter(RigaRicambio.id == riga_id, RigaRicambio.preventivo_id == preventivo_id).first()
    if not riga:
        raise HTTPException(status_code=404, detail="Riga non trovata")
    for key, value in data.dict(exclude_unset=True).items():
        setattr(riga, key, value)
    db.commit()
    db.refresh(riga)
    _aggiorna_totale_preventivo_ricambio(db, preventivo_id)
    return riga

@app.delete("/api/preventivi/{preventivo_id}/righe-ricambio/{riga_id}")
def delete_riga_ricambio(preventivo_id: int, riga_id: int, db: Session = Depends(get_db)):
    riga = db.query(RigaRicambio).filter(RigaRicambio.id == riga_id, RigaRicambio.preventivo_id == preventivo_id).first()
    if not riga:
        raise HTTPException(status_code=404, detail="Riga non trovata")
    db.delete(riga)
    db.commit()
    _aggiorna_totale_preventivo_ricambio(db, preventivo_id)
    return {"message": "Riga eliminata"}


# ==========================================
# SCONTO EXTRA ADMIN
# ==========================================
class ScontoAdminRequest(BaseModel):
    sconto_extra_admin: float

@app.put("/api/preventivi/{preventivo_id}/sconto-admin")
def update_sconto_admin(preventivo_id: int, data: ScontoAdminRequest, db: Session = Depends(get_db)):
    """
    Aggiorna lo sconto extra admin sul preventivo.
    NOTA: Il frontend deve verificare che l'utente sia admin prima di mostrare questa opzione.
    """
    preventivo = db.query(Preventivo).filter(Preventivo.id == preventivo_id).first()
    if not preventivo:
        raise HTTPException(status_code=404, detail="Preventivo non trovato")
    
    # Valida sconto (0-100)
    if data.sconto_extra_admin < 0 or data.sconto_extra_admin > 100:
        raise HTTPException(status_code=400, detail="Sconto deve essere tra 0 e 100")
    
    preventivo.sconto_extra_admin = data.sconto_extra_admin
    
    # Ricalcola totale finale
    if preventivo.tipo_preventivo == "RICAMBIO":
        _aggiorna_totale_preventivo_ricambio(db, preventivo_id)
    else:
        _aggiorna_totale_preventivo(db, preventivo_id)
    
    db.refresh(preventivo)
    return {
        "sconto_extra_admin": float(preventivo.sconto_extra_admin),
        "total_price": float(preventivo.total_price),
        "total_price_finale": float(preventivo.total_price_finale)
    }


def _aggiorna_totale_preventivo(db: Session, preventivo_id: int):
    """Aggiorna totali preventivo COMPLETO (da materiali)"""
    preventivo = db.query(Preventivo).filter(Preventivo.id == preventivo_id).first()
    if preventivo:
        mats = db.query(Materiale).filter(Materiale.preventivo_id == preventivo_id).all()
        preventivo.total_price = sum(float(m.prezzo_totale or 0) for m in mats)
        # Applica sconto extra admin
        sconto_admin = float(preventivo.sconto_extra_admin or 0)
        preventivo.total_price_finale = float(preventivo.total_price) * (1 - sconto_admin / 100)
        db.commit()

def _aggiorna_totale_preventivo_ricambio(db: Session, preventivo_id: int):
    """
    Aggiorna totali preventivo RICAMBIO:
    - total_price = somma prezzi listino
    - sconto_cliente = media ponderata sconti cliente
    - total_price_finale = totale dopo sconto cliente e sconto admin
    """
    preventivo = db.query(Preventivo).filter(Preventivo.id == preventivo_id).first()
    if not preventivo:
        return
    
    righe = db.query(RigaRicambio).filter(RigaRicambio.preventivo_id == preventivo_id).all()
    
    # Somma totali
    totale_listino = sum(float(r.prezzo_totale_listino or 0) for r in righe)
    totale_cliente = sum(float(r.prezzo_totale_cliente or 0) for r in righe)
    
    preventivo.total_price = totale_listino
    
    # Calcola sconto cliente medio (per visualizzazione)
    if totale_listino > 0:
        sconto_medio = (1 - totale_cliente / totale_listino) * 100
        preventivo.sconto_cliente = round(sconto_medio, 2)
    
    # Applica sconto extra admin sul totale cliente
    sconto_admin = float(preventivo.sconto_extra_admin or 0)
    preventivo.total_price_finale = totale_cliente * (1 - sconto_admin / 100)
    
    db.commit()


# ==========================================
# RULE ENGINE ENDPOINT
# ==========================================
@app.post("/api/preventivi/{preventivo_id}/evaluate-rules")
def evaluate_rules_endpoint(preventivo_id: int, db: Session = Depends(get_db)):
    return evaluate_rules(preventivo_id, db)


# ==========================================
# ESPORTAZIONE PREVENTIVO
# ==========================================

def _get_preventivo_data(db: Session, preventivo_id: int):
    """Raccoglie tutti i dati del preventivo per l'export"""
    preventivo = db.query(Preventivo).filter(Preventivo.id == preventivo_id).first()
    if not preventivo:
        return None
    
    cliente = db.query(Cliente).filter(Cliente.id == preventivo.cliente_id).first() if preventivo.cliente_id else None
    dati_commessa = db.query(DatiCommessa).filter(DatiCommessa.preventivo_id == preventivo_id).first()
    dati_principali = db.query(DatiPrincipali).filter(DatiPrincipali.preventivo_id == preventivo_id).first()
    normative = db.query(Normative).filter(Normative.preventivo_id == preventivo_id).first()
    argano = db.query(Argano).filter(Argano.preventivo_id == preventivo_id).first()
    materiali = db.query(Materiale).filter(Materiale.preventivo_id == preventivo_id).all()
    righe_ricambio = db.query(RigaRicambio).filter(RigaRicambio.preventivo_id == preventivo_id).all()
    
    # Calcola totali al volo per sicurezza
    if preventivo.tipo_preventivo == "COMPLETO":
        totale_materiali = sum(float(m.prezzo_totale or 0) for m in materiali)
        if totale_materiali > 0:
            preventivo.total_price = totale_materiali
            sconto_cl = float(preventivo.sconto_cliente or 0)
            sconto_admin = float(preventivo.sconto_extra_admin or 0)
            totale_scontato = totale_materiali * (1 - sconto_cl / 100)
            preventivo.total_price_finale = totale_scontato * (1 - sconto_admin / 100)
    elif preventivo.tipo_preventivo == "RICAMBIO":
        totale_listino = sum(float(r.prezzo_totale_listino or 0) for r in righe_ricambio)
        totale_cliente = sum(float(r.prezzo_totale_cliente or 0) for r in righe_ricambio)
        if totale_listino > 0:
            preventivo.total_price = totale_listino
            sconto_admin = float(preventivo.sconto_extra_admin or 0)
            preventivo.total_price_finale = totale_cliente * (1 - sconto_admin / 100)
    
    return {
        "preventivo": preventivo,
        "cliente": cliente,
        "dati_commessa": dati_commessa,
        "dati_principali": dati_principali,
        "normative": normative,
        "argano": argano,
        "materiali": materiali,
        "righe_ricambio": righe_ricambio,
    }


@app.get("/api/preventivi/{preventivo_id}/export/xml")
def export_xml(preventivo_id: int, db: Session = Depends(get_db)):
    """Esporta preventivo in formato XML con indentazione"""
    from xml.dom import minidom
    
    data = _get_preventivo_data(db, preventivo_id)
    if not data:
        raise HTTPException(status_code=404, detail="Preventivo non trovato")
    
    p = data["preventivo"]
    aliquota_iva = 22.0  # IVA italiana
    
    root = ET.Element("preventivo")
    root.set("id", str(p.id))
    root.set("numero", p.numero_preventivo or "")
    root.set("tipo", p.tipo_preventivo or "COMPLETO")
    root.set("data", p.created_at.strftime("%Y-%m-%d") if p.created_at else "")
    
    # Azienda
    azienda_el = ET.SubElement(root, "azienda")
    ET.SubElement(azienda_el, "nome").text = "ELETTROQUADRI S.R.L."
    ET.SubElement(azienda_el, "indirizzo").text = "via Puccini, 1 - 21050 Bisuschio (VA)"
    ET.SubElement(azienda_el, "website").text = "www.elettroquadri.net"
    ET.SubElement(azienda_el, "email").text = "sales@elettroquadri.net"
    
    # Cliente
    cliente_el = ET.SubElement(root, "cliente")
    if data["cliente"]:
        c = data["cliente"]
        ET.SubElement(cliente_el, "codice").text = c.codice or ""
        ET.SubElement(cliente_el, "ragione_sociale").text = c.ragione_sociale or ""
        ET.SubElement(cliente_el, "partita_iva").text = c.partita_iva or ""
        ET.SubElement(cliente_el, "indirizzo").text = c.indirizzo or ""
        ET.SubElement(cliente_el, "cap").text = c.cap or ""
        ET.SubElement(cliente_el, "citta").text = c.citta or ""
        ET.SubElement(cliente_el, "provincia").text = c.provincia or ""
    
    # Dati Commessa
    commessa_el = ET.SubElement(root, "dati_commessa")
    if data["dati_commessa"]:
        dc = data["dati_commessa"]
        ET.SubElement(commessa_el, "numero_offerta").text = dc.numero_offerta or ""
        ET.SubElement(commessa_el, "data_offerta").text = dc.data_offerta or ""
        ET.SubElement(commessa_el, "riferimento_cliente").text = dc.riferimento_cliente or ""
        ET.SubElement(commessa_el, "quantita").text = str(dc.quantita or 1)
        ET.SubElement(commessa_el, "data_consegna_richiesta").text = str(dc.data_consegna_richiesta) if dc.data_consegna_richiesta else ""
        ET.SubElement(commessa_el, "pagamento").text = dc.pagamento or ""
        ET.SubElement(commessa_el, "imballo").text = dc.imballo or ""
        ET.SubElement(commessa_el, "trasporto").text = dc.trasporto or ""
    
    # Dati Principali
    principali_el = ET.SubElement(root, "dati_principali")
    if data["dati_principali"]:
        dp = data["dati_principali"]
        ET.SubElement(principali_el, "tipo_impianto").text = dp.tipo_impianto or ""
        ET.SubElement(principali_el, "nuovo_impianto").text = str(dp.nuovo_impianto or False).lower()
        ET.SubElement(principali_el, "numero_fermate").text = str(dp.numero_fermate or 0)
        ET.SubElement(principali_el, "numero_servizi").text = str(dp.numero_servizi or 0)
        ET.SubElement(principali_el, "velocita").text = str(dp.velocita or 0)
        ET.SubElement(principali_el, "corsa").text = str(dp.corsa or 0)
    
    # Normative
    normative_el = ET.SubElement(root, "normative")
    if data["normative"]:
        n = data["normative"]
        ET.SubElement(normative_el, "direttiva").text = n.direttiva or ""
        ET.SubElement(normative_el, "en_81_20").text = str(n.en_81_20 or False).lower()
        ET.SubElement(normative_el, "en_81_20_anno").text = n.en_81_20_anno or ""
        ET.SubElement(normative_el, "en_81_50").text = str(n.en_81_50 or False).lower()
    
    # Argano
    argano_el = ET.SubElement(root, "argano")
    if data["argano"]:
        a = data["argano"]
        ET.SubElement(argano_el, "trazione").text = a.trazione or ""
        ET.SubElement(argano_el, "potenza_motore_kw").text = str(a.potenza_motore_kw or 0)
        ET.SubElement(argano_el, "tipo_vvvf").text = a.tipo_vvvf or ""
        ET.SubElement(argano_el, "vvvf_nel_vano").text = str(a.vvvf_nel_vano or False).lower()
    
    # Materiali
    materiali_el = ET.SubElement(root, "materiali")
    materiali_el.set("count", str(len(data["materiali"])))
    for m in data["materiali"]:
        mat_el = ET.SubElement(materiali_el, "materiale")
        ET.SubElement(mat_el, "codice").text = m.codice or ""
        ET.SubElement(mat_el, "descrizione").text = m.descrizione or ""
        ET.SubElement(mat_el, "categoria").text = m.categoria or ""
        ET.SubElement(mat_el, "quantita").text = str(float(m.quantita or 0))
        ET.SubElement(mat_el, "prezzo_unitario").text = f"{float(m.prezzo_unitario or 0):.2f}"
        ET.SubElement(mat_el, "prezzo_totale").text = f"{float(m.prezzo_totale or 0):.2f}"
        ET.SubElement(mat_el, "aggiunto_da_regola").text = str(m.aggiunto_da_regola or False).lower()
        if m.regola_id:
            ET.SubElement(mat_el, "regola_id").text = m.regola_id
    
    # Righe Ricambio
    if data["righe_ricambio"]:
        ricambi_el = ET.SubElement(root, "righe_ricambio")
        ricambi_el.set("count", str(len(data["righe_ricambio"])))
        for r in data["righe_ricambio"]:
            riga_el = ET.SubElement(ricambi_el, "riga")
            ET.SubElement(riga_el, "codice").text = r.codice or ""
            ET.SubElement(riga_el, "descrizione").text = r.descrizione or ""
            ET.SubElement(riga_el, "quantita").text = str(float(r.quantita or 0))
            ET.SubElement(riga_el, "prezzo_listino_unitario").text = f"{float(r.prezzo_listino_unitario or 0):.2f}"
            ET.SubElement(riga_el, "sconto_cliente").text = f"{float(r.sconto_cliente or 0):.2f}"
            ET.SubElement(riga_el, "prezzo_cliente_unitario").text = f"{float(r.prezzo_cliente_unitario or 0):.2f}"
            ET.SubElement(riga_el, "prezzo_totale_listino").text = f"{float(r.prezzo_totale_listino or 0):.2f}"
            ET.SubElement(riga_el, "prezzo_totale_cliente").text = f"{float(r.prezzo_totale_cliente or 0):.2f}"
    
    # Totali
    totale_imponibile = float(p.total_price_finale or p.total_price or 0)
    iva = totale_imponibile * aliquota_iva / 100
    totale_con_iva = totale_imponibile + iva
    
    totali_el = ET.SubElement(root, "totali")
    ET.SubElement(totali_el, "totale_listino").text = f"{float(p.total_price or 0):.2f}"
    ET.SubElement(totali_el, "sconto_cliente_percentuale").text = f"{float(p.sconto_cliente or 0):.2f}"
    ET.SubElement(totali_el, "sconto_admin_percentuale").text = f"{float(p.sconto_extra_admin or 0):.2f}"
    ET.SubElement(totali_el, "totale_imponibile").text = f"{totale_imponibile:.2f}"
    ET.SubElement(totali_el, "aliquota_iva").text = f"{aliquota_iva:.2f}"
    ET.SubElement(totali_el, "importo_iva").text = f"{iva:.2f}"
    ET.SubElement(totali_el, "totale_con_iva").text = f"{totale_con_iva:.2f}"
    
    # Genera XML formattato con indentazione
    xml_str = ET.tostring(root, encoding='unicode', method='xml')
    dom = minidom.parseString(xml_str)
    xml_pretty = dom.toprettyxml(indent="  ", encoding=None)
    # Rimuovi la prima riga (xml declaration di minidom) e aggiungi la nostra
    lines = xml_pretty.split('\n')[1:]  # Salta la prima riga
    xml_final = '<?xml version="1.0" encoding="UTF-8"?>\n' + '\n'.join(lines)
    
    return StreamingResponse(
        io.BytesIO(xml_final.encode('utf-8')),
        media_type="application/xml",
        headers={"Content-Disposition": f'attachment; filename="preventivo_{p.numero_preventivo.replace("/", "_")}.xml"'}
    )


@app.get("/api/preventivi/{preventivo_id}/export/xlsx")
def export_xlsx(preventivo_id: int, db: Session = Depends(get_db)):
    """Esporta preventivo in formato Excel"""
    try:
        from openpyxl import Workbook
        from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    except ImportError:
        raise HTTPException(status_code=500, detail="openpyxl non installato. Esegui: pip install openpyxl")
    
    data = _get_preventivo_data(db, preventivo_id)
    if not data:
        raise HTTPException(status_code=404, detail="Preventivo non trovato")
    
    p = data["preventivo"]
    aliquota_iva = 22.0
    wb = Workbook()
    ws = wb.active
    ws.title = "Preventivo"
    
    # Stili
    header_font = Font(bold=True, color="FFFFFF")
    header_fill = PatternFill("solid", fgColor="4A4A4A")  # Grigio scuro come logo
    title_font = Font(bold=True, size=14)
    money_format = '€ #,##0.00'
    percent_format = '0.00%'
    border = Border(
        left=Side(style='thin'), right=Side(style='thin'),
        top=Side(style='thin'), bottom=Side(style='thin')
    )
    
    # Intestazione azienda
    ws['A1'] = "ELETTROQUADRI S.R.L."
    ws['A1'].font = Font(bold=True, size=16)
    ws['A2'] = "via Puccini, 1 - 21050 Bisuschio (VA)"
    ws['A3'] = "www.elettroquadri.net | sales@elettroquadri.net"
    ws.merge_cells('A1:C1')
    ws.merge_cells('A2:C2')
    ws.merge_cells('A3:C3')
    
    # Info preventivo
    ws['E1'] = "Numero Offerta"
    ws['F1'] = p.numero_preventivo
    ws['F1'].font = Font(bold=True, size=12)
    ws['E2'] = "Data"
    ws['F2'] = p.created_at.strftime('%d/%m/%Y') if p.created_at else ''
    ws['E3'] = "Tipo"
    ws['F3'] = p.tipo_preventivo
    
    # Cliente
    row = 5
    ws[f'A{row}'] = "CLIENTE"
    ws[f'A{row}'].font = title_font
    row += 1
    if data["cliente"]:
        c = data["cliente"]
        ws[f'A{row}'] = "Ragione Sociale:"
        ws[f'B{row}'] = c.ragione_sociale
        row += 1
        ws[f'A{row}'] = "P.IVA:"
        ws[f'B{row}'] = c.partita_iva
        row += 1
        ws[f'A{row}'] = "Indirizzo:"
        ws[f'B{row}'] = f"{c.indirizzo}, {c.cap} {c.citta} ({c.provincia})"
        row += 1
    
    # Materiali o Righe Ricambio
    row += 2
    if p.tipo_preventivo == "COMPLETO" and data["materiali"]:
        ws[f'A{row}'] = "MATERIALI"
        ws[f'A{row}'].font = title_font
        row += 1
        
        headers = ["Codice", "Descrizione", "Categoria", "Qtà", "Prezzo Unit.", "Totale"]
        for col, h in enumerate(headers, 1):
            cell = ws.cell(row=row, column=col, value=h)
            cell.font = header_font
            cell.fill = header_fill
            cell.border = border
            cell.alignment = Alignment(horizontal='center')
        row += 1
        
        for m in data["materiali"]:
            ws.cell(row=row, column=1, value=m.codice).border = border
            ws.cell(row=row, column=2, value=m.descrizione).border = border
            ws.cell(row=row, column=3, value=m.categoria).border = border
            ws.cell(row=row, column=4, value=float(m.quantita or 0)).border = border
            cell_pu = ws.cell(row=row, column=5, value=float(m.prezzo_unitario or 0))
            cell_pu.number_format = money_format
            cell_pu.border = border
            cell_tot = ws.cell(row=row, column=6, value=float(m.prezzo_totale or 0))
            cell_tot.number_format = money_format
            cell_tot.border = border
            row += 1
    
    elif p.tipo_preventivo == "RICAMBIO" and data["righe_ricambio"]:
        ws[f'A{row}'] = "RICAMBI"
        ws[f'A{row}'].font = title_font
        row += 1
        
        headers = ["Codice", "Descrizione", "Qtà", "Prezzo Listino", "Sconto %", "Prezzo Cliente", "Totale"]
        for col, h in enumerate(headers, 1):
            cell = ws.cell(row=row, column=col, value=h)
            cell.font = header_font
            cell.fill = header_fill
            cell.border = border
        row += 1
        
        for r in data["righe_ricambio"]:
            ws.cell(row=row, column=1, value=r.codice).border = border
            ws.cell(row=row, column=2, value=r.descrizione).border = border
            ws.cell(row=row, column=3, value=float(r.quantita or 0)).border = border
            cell_pl = ws.cell(row=row, column=4, value=float(r.prezzo_listino_unitario or 0))
            cell_pl.number_format = money_format
            cell_pl.border = border
            ws.cell(row=row, column=5, value=float(r.sconto_cliente or 0)/100).number_format = percent_format
            ws.cell(row=row, column=5).border = border
            cell_pc = ws.cell(row=row, column=6, value=float(r.prezzo_cliente_unitario or 0))
            cell_pc.number_format = money_format
            cell_pc.border = border
            cell_tot = ws.cell(row=row, column=7, value=float(r.prezzo_totale_cliente or 0))
            cell_tot.number_format = money_format
            cell_tot.border = border
            row += 1
    
    # Totali con IVA
    row += 2
    ws[f'A{row}'] = "RIEPILOGO ECONOMICO"
    ws[f'A{row}'].font = title_font
    row += 1
    
    totale_listino = float(p.total_price or 0)
    sconto_cliente = float(p.sconto_cliente or 0)
    sconto_admin = float(p.sconto_extra_admin or 0)
    totale_imponibile = float(p.total_price_finale or totale_listino)
    iva = totale_imponibile * aliquota_iva / 100
    totale_con_iva = totale_imponibile + iva
    
    ws[f'D{row}'] = "Totale Listino:"
    ws[f'E{row}'] = totale_listino
    ws[f'E{row}'].number_format = money_format
    row += 1
    
    ws[f'D{row}'] = "Sconto Cliente:"
    ws[f'E{row}'] = f"{sconto_cliente:.2f}%"
    row += 1
    
    ws[f'D{row}'] = "Sconto Extra:"
    ws[f'E{row}'] = f"{sconto_admin:.2f}%"
    row += 1
    
    ws[f'D{row}'] = "Totale Imponibile:"
    ws[f'D{row}'].font = Font(bold=True)
    ws[f'E{row}'] = totale_imponibile
    ws[f'E{row}'].number_format = money_format
    ws[f'E{row}'].font = Font(bold=True)
    row += 1
    
    ws[f'D{row}'] = f"IVA ({aliquota_iva:.0f}%):"
    ws[f'E{row}'] = iva
    ws[f'E{row}'].number_format = money_format
    row += 1
    
    ws[f'D{row}'] = "TOTALE CON IVA:"
    ws[f'D{row}'].font = Font(bold=True, size=12)
    ws[f'E{row}'] = totale_con_iva
    ws[f'E{row}'].number_format = money_format
    ws[f'E{row}'].font = Font(bold=True, size=12)
    
    # Larghezza colonne
    ws.column_dimensions['A'].width = 18
    ws.column_dimensions['B'].width = 40
    ws.column_dimensions['C'].width = 15
    ws.column_dimensions['D'].width = 18
    ws.column_dimensions['E'].width = 15
    ws.column_dimensions['F'].width = 15
    ws.column_dimensions['G'].width = 15
    
    buffer = io.BytesIO()
    wb.save(buffer)
    buffer.seek(0)
    
    return StreamingResponse(
        buffer,
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers={"Content-Disposition": f'attachment; filename="preventivo_{p.numero_preventivo.replace("/", "_")}.xlsx"'}
    )


@app.get("/api/preventivi/{preventivo_id}/export/pdf")
def export_pdf(preventivo_id: int, db: Session = Depends(get_db)):
    """Esporta preventivo in formato PDF professionale"""
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib import colors
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, Image
        from reportlab.lib.units import cm, mm
        from reportlab.lib.enums import TA_LEFT, TA_RIGHT, TA_CENTER
    except ImportError as e:
        raise HTTPException(status_code=500, detail=f"reportlab non installato: {str(e)}. Esegui: pip install reportlab")
    
    data = _get_preventivo_data(db, preventivo_id)
    if not data:
        raise HTTPException(status_code=404, detail="Preventivo non trovato")
    
    p = data["preventivo"]
    aliquota_iva = 22.0
    
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(
        buffer, 
        pagesize=A4, 
        topMargin=1.5*cm, 
        bottomMargin=1.5*cm,
        leftMargin=1.5*cm,
        rightMargin=1.5*cm
    )
    
    # Colori aziendali
    GRIGIO_SCURO = colors.HexColor('#4A4A4A')
    GRIGIO_CHIARO = colors.HexColor('#F5F5F5')
    BLU_ACCENT = colors.HexColor('#2196F3')
    
    styles = getSampleStyleSheet()
    
    # Stili personalizzati
    style_azienda = ParagraphStyle('Azienda', parent=styles['Normal'], fontSize=10, textColor=GRIGIO_SCURO)
    style_titolo = ParagraphStyle('Titolo', parent=styles['Heading1'], fontSize=14, textColor=GRIGIO_SCURO, spaceAfter=10)
    style_sezione = ParagraphStyle('Sezione', parent=styles['Heading2'], fontSize=11, textColor=colors.white, 
                                    backColor=GRIGIO_SCURO, spaceBefore=15, spaceAfter=8, leftIndent=5, rightIndent=5)
    style_label = ParagraphStyle('Label', parent=styles['Normal'], fontSize=9, textColor=colors.grey)
    style_value = ParagraphStyle('Value', parent=styles['Normal'], fontSize=10, fontName='Helvetica-Bold')
    style_totale = ParagraphStyle('Totale', parent=styles['Normal'], fontSize=12, fontName='Helvetica-Bold', alignment=TA_RIGHT)
    
    story = []
    
    # Header con logo e info azienda
    logo_path = os.path.join(os.path.dirname(__file__), 'logo_elettroquadri.png')
    header_data = []
    
    # Colonna sinistra: Logo e info azienda
    azienda_info = """<b>ELETTROQUADRI S.R.L.</b><br/>
    via Puccini, 1 - 21050 Bisuschio (VA)<br/>
    www.elettroquadri.net<br/>
    sales@elettroquadri.net"""
    
    # Colonna destra: Info preventivo
    prev_info = f"""<b>Numero Offerta</b><br/>
    <font size="14"><b>{p.numero_preventivo}</b></font><br/><br/>
    <b>Data:</b> {p.created_at.strftime('%d/%m/%Y') if p.created_at else ''}<br/>
    <b>Tipo:</b> {p.tipo_preventivo}"""
    
    header_table = Table([
        [Paragraph(azienda_info, style_azienda), Paragraph(prev_info, style_azienda)]
    ], colWidths=[10*cm, 7*cm])
    header_table.setStyle(TableStyle([
        ('VALIGN', (0, 0), (-1, -1), 'TOP'),
        ('ALIGN', (1, 0), (1, 0), 'RIGHT'),
    ]))
    story.append(header_table)
    story.append(Spacer(1, 0.5*cm))
    
    # Linea separatrice
    line_table = Table([['']], colWidths=[17*cm])
    line_table.setStyle(TableStyle([
        ('LINEABOVE', (0, 0), (-1, 0), 2, GRIGIO_SCURO),
    ]))
    story.append(line_table)
    story.append(Spacer(1, 0.5*cm))
    
    # Cliente
    if data["cliente"]:
        c = data["cliente"]
        story.append(Paragraph("Cliente", style_sezione))
        cliente_data = [
            [Paragraph("<b>Ragione Sociale</b>", style_label), c.ragione_sociale or ""],
            [Paragraph("<b>P.IVA</b>", style_label), c.partita_iva or ""],
            [Paragraph("<b>Indirizzo</b>", style_label), f"{c.indirizzo or ''}, {c.cap or ''} {c.citta or ''} ({c.provincia or ''})"],
        ]
        t = Table(cliente_data, colWidths=[4*cm, 13*cm])
        t.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, -1), GRIGIO_CHIARO),
            ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
            ('TOPPADDING', (0, 0), (-1, -1), 6),
        ]))
        story.append(t)
    
    # Dati Commessa (se presenti)
    if data["dati_commessa"]:
        dc = data["dati_commessa"]
        story.append(Paragraph("Dati Commessa", style_sezione))
        commessa_data = [
            ["Riferimento Cliente", dc.riferimento_cliente or "-", "Consegna Richiesta", str(dc.data_consegna_richiesta) if dc.data_consegna_richiesta else "-"],
            ["Pagamento", dc.pagamento or "-", "Trasporto", dc.trasporto or "-"],
        ]
        t = Table(commessa_data, colWidths=[4*cm, 4.5*cm, 4*cm, 4.5*cm])
        t.setStyle(TableStyle([
            ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
            ('FONTNAME', (2, 0), (2, -1), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, -1), 9),
            ('BACKGROUND', (0, 0), (-1, -1), GRIGIO_CHIARO),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 5),
            ('TOPPADDING', (0, 0), (-1, -1), 5),
        ]))
        story.append(t)
    
    # Materiali (per COMPLETO)
    if p.tipo_preventivo == "COMPLETO" and data["materiali"]:
        story.append(Paragraph("Materiali", style_sezione))
        mat_data = [["Codice", "Descrizione", "Qtà", "Prezzo Unit.", "Totale"]]
        for m in data["materiali"]:
            mat_data.append([
                m.codice or "",
                (m.descrizione or "")[:45],
                str(int(float(m.quantita or 0))) if float(m.quantita or 0) == int(float(m.quantita or 0)) else f"{float(m.quantita or 0):.2f}",
                f"€ {float(m.prezzo_unitario or 0):,.2f}",
                f"€ {float(m.prezzo_totale or 0):,.2f}",
            ])
        t = Table(mat_data, colWidths=[2.5*cm, 8*cm, 1.5*cm, 2.5*cm, 2.5*cm])
        t.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), GRIGIO_SCURO),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, -1), 8),
            ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
            ('ALIGN', (2, 0), (-1, -1), 'RIGHT'),
            ('BACKGROUND', (0, 1), (-1, -1), colors.white),
            ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, GRIGIO_CHIARO]),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 4),
            ('TOPPADDING', (0, 0), (-1, -1), 4),
        ]))
        story.append(t)
    
    # Ricambi (per RICAMBIO)
    elif p.tipo_preventivo == "RICAMBIO" and data["righe_ricambio"]:
        story.append(Paragraph("Ricambi", style_sezione))
        ric_data = [["Codice", "Descrizione", "Qtà", "Listino", "Sc.%", "Totale"]]
        for r in data["righe_ricambio"]:
            ric_data.append([
                r.codice or "",
                (r.descrizione or "")[:35],
                str(int(float(r.quantita or 0))),
                f"€ {float(r.prezzo_listino_unitario or 0):,.2f}",
                f"{float(r.sconto_cliente or 0):.0f}%",
                f"€ {float(r.prezzo_totale_cliente or 0):,.2f}",
            ])
        t = Table(ric_data, colWidths=[2.5*cm, 6.5*cm, 1.2*cm, 2.3*cm, 1.5*cm, 2.5*cm])
        t.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), GRIGIO_SCURO),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, -1), 8),
            ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
            ('ALIGN', (2, 0), (-1, -1), 'RIGHT'),
            ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, GRIGIO_CHIARO]),
        ]))
        story.append(t)
    
    # Riepilogo Economico
    story.append(Spacer(1, 0.5*cm))
    story.append(Paragraph("Riepilogo Economico", style_sezione))
    
    totale_listino = float(p.total_price or 0)
    sconto_cliente = float(p.sconto_cliente or 0)
    sconto_admin = float(p.sconto_extra_admin or 0)
    totale_imponibile = float(p.total_price_finale or totale_listino)
    iva = totale_imponibile * aliquota_iva / 100
    totale_con_iva = totale_imponibile + iva
    
    totali_data = [
        ["Totale Listino", f"€ {totale_listino:,.2f}"],
        ["Sconto Cliente", f"{sconto_cliente:.2f} %"],
        ["Sconto Extra", f"{sconto_admin:.2f} %"],
        ["Totale Imponibile", f"€ {totale_imponibile:,.2f}"],
        [f"IVA ({aliquota_iva:.0f}%)", f"€ {iva:,.2f}"],
    ]
    
    t = Table(totali_data, colWidths=[12*cm, 5*cm])
    t.setStyle(TableStyle([
        ('ALIGN', (0, 0), (0, -1), 'RIGHT'),
        ('ALIGN', (1, 0), (1, -1), 'RIGHT'),
        ('FONTSIZE', (0, 0), (-1, -1), 10),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 4),
        ('TOPPADDING', (0, 0), (-1, -1), 4),
    ]))
    story.append(t)
    
    # Totale finale grande
    totale_finale_data = [["TOTALE CON IVA", f"€ {totale_con_iva:,.2f}"]]
    t = Table(totale_finale_data, colWidths=[12*cm, 5*cm])
    t.setStyle(TableStyle([
        ('ALIGN', (0, 0), (-1, -1), 'RIGHT'),
        ('FONTNAME', (0, 0), (-1, -1), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, -1), 14),
        ('BACKGROUND', (0, 0), (-1, -1), GRIGIO_SCURO),
        ('TEXTCOLOR', (0, 0), (-1, -1), colors.white),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
        ('TOPPADDING', (0, 0), (-1, -1), 8),
    ]))
    story.append(t)
    
    # Condizioni
    story.append(Spacer(1, 1*cm))
    condizioni = """<b>Condizioni Generali:</b> art. 1495 C.C. e art. 39 Conv. di Vienna (1980)"""
    story.append(Paragraph(condizioni, ParagraphStyle('Condizioni', parent=styles['Normal'], fontSize=8, textColor=colors.grey)))
    
    doc.build(story)
    buffer.seek(0)
    
    return StreamingResponse(
        buffer,
        media_type="application/pdf",
        headers={"Content-Disposition": f'attachment; filename="preventivo_{p.numero_preventivo.replace("/", "_")}.pdf"'}
    )


@app.get("/api/preventivi/{preventivo_id}/export/docx")
def export_docx(preventivo_id: int, db: Session = Depends(get_db)):
    """Esporta preventivo in formato Word professionale"""
    try:
        from docx import Document
        from docx.shared import Inches, Pt, Cm, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
    except ImportError:
        raise HTTPException(status_code=500, detail="python-docx non installato. Esegui: pip install python-docx")
    
    data = _get_preventivo_data(db, preventivo_id)
    if not data:
        raise HTTPException(status_code=404, detail="Preventivo non trovato")
    
    p = data["preventivo"]
    aliquota_iva = 22.0
    doc = Document()
    
    # Imposta margini
    sections = doc.sections
    for section in sections:
        section.left_margin = Cm(1.5)
        section.right_margin = Cm(1.5)
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
    
    # Header azienda
    header = doc.add_paragraph()
    run = header.add_run("ELETTROQUADRI S.R.L.")
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(74, 74, 74)
    
    info = doc.add_paragraph("via Puccini, 1 - 21050 Bisuschio (VA)\nwww.elettroquadri.net | sales@elettroquadri.net")
    info.style.font.size = Pt(10)
    
    doc.add_paragraph()
    
    # Info preventivo
    p_info = doc.add_paragraph()
    run = p_info.add_run(f"PREVENTIVO N. {p.numero_preventivo}")
    run.bold = True
    run.font.size = Pt(14)
    
    doc.add_paragraph(f"Data: {p.created_at.strftime('%d/%m/%Y') if p.created_at else ''} | Tipo: {p.tipo_preventivo}")
    doc.add_paragraph()
    
    # Cliente
    if data["cliente"]:
        c = data["cliente"]
        doc.add_heading("Cliente", level=2)
        doc.add_paragraph(f"Ragione Sociale: {c.ragione_sociale or ''}")
        doc.add_paragraph(f"P.IVA: {c.partita_iva or ''}")
        doc.add_paragraph(f"Indirizzo: {c.indirizzo or ''}, {c.cap or ''} {c.citta or ''} ({c.provincia or ''})")
    
    # Materiali o Ricambi
    if p.tipo_preventivo == "COMPLETO" and data["materiali"]:
        doc.add_heading("Materiali", level=2)
        table = doc.add_table(rows=1, cols=5)
        table.style = 'Table Grid'
        
        hdr_cells = table.rows[0].cells
        headers = ["Codice", "Descrizione", "Qtà", "Prezzo Unit.", "Totale"]
        for i, h in enumerate(headers):
            hdr_cells[i].text = h
            hdr_cells[i].paragraphs[0].runs[0].bold = True
        
        for m in data["materiali"]:
            row_cells = table.add_row().cells
            row_cells[0].text = m.codice or ""
            row_cells[1].text = (m.descrizione or "")[:50]
            row_cells[2].text = str(float(m.quantita or 0))
            row_cells[3].text = f"€ {float(m.prezzo_unitario or 0):,.2f}"
            row_cells[4].text = f"€ {float(m.prezzo_totale or 0):,.2f}"
    
    elif p.tipo_preventivo == "RICAMBIO" and data["righe_ricambio"]:
        doc.add_heading("Ricambi", level=2)
        table = doc.add_table(rows=1, cols=6)
        table.style = 'Table Grid'
        
        hdr_cells = table.rows[0].cells
        headers = ["Codice", "Descrizione", "Qtà", "Listino", "Sconto", "Totale"]
        for i, h in enumerate(headers):
            hdr_cells[i].text = h
            hdr_cells[i].paragraphs[0].runs[0].bold = True
        
        for r in data["righe_ricambio"]:
            row_cells = table.add_row().cells
            row_cells[0].text = r.codice or ""
            row_cells[1].text = (r.descrizione or "")[:40]
            row_cells[2].text = str(float(r.quantita or 0))
            row_cells[3].text = f"€ {float(r.prezzo_listino_unitario or 0):,.2f}"
            row_cells[4].text = f"{float(r.sconto_cliente or 0):.1f}%"
            row_cells[5].text = f"€ {float(r.prezzo_totale_cliente or 0):,.2f}"
    
    # Riepilogo Economico
    doc.add_paragraph()
    doc.add_heading("Riepilogo Economico", level=2)
    
    totale_listino = float(p.total_price or 0)
    sconto_cliente = float(p.sconto_cliente or 0)
    sconto_admin = float(p.sconto_extra_admin or 0)
    totale_imponibile = float(p.total_price_finale or totale_listino)
    iva = totale_imponibile * aliquota_iva / 100
    totale_con_iva = totale_imponibile + iva
    
    doc.add_paragraph(f"Totale Listino: € {totale_listino:,.2f}")
    doc.add_paragraph(f"Sconto Cliente: {sconto_cliente:.2f}%")
    doc.add_paragraph(f"Sconto Extra: {sconto_admin:.2f}%")
    doc.add_paragraph(f"Totale Imponibile: € {totale_imponibile:,.2f}")
    doc.add_paragraph(f"IVA ({aliquota_iva:.0f}%): € {iva:,.2f}")
    
    total_para = doc.add_paragraph()
    run = total_para.add_run(f"TOTALE CON IVA: € {totale_con_iva:,.2f}")
    run.bold = True
    run.font.size = Pt(14)
    
    # Condizioni
    doc.add_paragraph()
    cond = doc.add_paragraph("Condizioni Generali: art. 1495 C.C. e art. 39 Conv. di Vienna (1980)")
    cond.style.font.size = Pt(8)
    
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    return StreamingResponse(
        buffer,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="preventivo_{p.numero_preventivo.replace("/", "_")}.docx"'}
    )

# ==========================================
# RULE DESIGNER IMPORT
# ==========================================

def convert_operator(op: str) -> str:
    """Converte operatore Rule Designer → Configuratore"""
    mapping = {
        "==": "equals",
        "!=": "not_equals",
        ">": "greater_than",
        ">=": "greater_or_equal",
        "<": "less_than",
        "<=": "less_or_equal",
        "contains": "contains",
        "in": "in"
    }
    return mapping.get(op, "equals")

def convert_rule_designer_to_configurator(rd_rule: dict, db: Session = None) -> dict:
    """
    Converte una regola dal formato Rule Designer al formato Configuratore.
    """
    # Estrai il nome del campo (ultimo segmento del path)
    campo_path = rd_rule.get("condizione", {}).get("campo", "")
    field_name = campo_path.split(".")[-1] if "." in campo_path else campo_path
    
    # Converti operatore
    operatore = rd_rule.get("condizione", {}).get("operatore", "==")
    operator = convert_operator(operatore)
    
    # Valore condizione
    valore = rd_rule.get("condizione", {}).get("valore", "")
    
    # Costruisci condizione
    condition = {
        "field": field_name,
        "operator": operator,
        "value": str(valore),
        "description": f"Campo {campo_path} {operatore} {valore}"
    }
    
    # Costruisci materiali
    materials = []
    
    # NUOVO FORMATO: outputs è un array
    outputs_list = rd_rule.get("outputs", [])
    
    # VECCHIO FORMATO (retrocompatibilità): output singolo + quantita separata
    if not outputs_list and rd_rule.get("output"):
        old_output = rd_rule.get("output", {})
        old_quantita = rd_rule.get("quantita", {})
        outputs_list = [{
            **old_output,
            "quantita": old_quantita
        }]
    
    # Processa ogni output
    for output in outputs_list:
        codice_articolo = output.get("codice_articolo", "")
        
        if not codice_articolo:
            continue
            
        descrizione = f"Articolo {codice_articolo}"
        prezzo = 0.0
        
        # Cerca articolo nel DB
        if db:
            try:
                articolo = db.query(Articolo).filter(Articolo.codice == codice_articolo).first()
                if articolo:
                    descrizione = articolo.descrizione or descrizione
                    # Calcola prezzo listino: costo_fisso * (1 + ricarico%)
                    costo = float(articolo.costo_fisso or 0)
                    ricarico = float(articolo.ricarico_percentuale or 0)
                    prezzo = costo * (1 + ricarico / 100)
            except Exception:
                pass
        
        # Calcola quantità (ora è dentro l'output)
        quantita_config = output.get("quantita", {})
        if quantita_config.get("tipo") == "fisso":
            qta = quantita_config.get("valore", 1)
        elif quantita_config.get("tipo") == "campo":
            qta = f"{{${quantita_config.get('campo', 'quantita')}}}"  # placeholder
        elif quantita_config.get("tipo") == "formula":
            qta = f"={quantita_config.get('espressione', '1')}"  # formula
        else:
            qta = 1
        
        material = {
            "codice": codice_articolo,
            "descrizione": descrizione,
            "quantita": qta,
            "prezzo_unitario": prezzo,
            "categoria": output.get("famiglia", "Materiale Automatico"),
            "unita": output.get("unita", "pz"),
            "codice_fonte": output.get("codice_fonte", "")
        }
        materials.append(material)
    
    # Costruisci regola finale
    nome = rd_rule.get("nome", "rule_imported")
    
    converted = {
        "id": nome,
        "name": nome.replace("_", " ").title(),
        "description": f"Regola importata da Rule Designer - {campo_path}",
        "version": "1.0",
        "enabled": True,
        "priority": rd_rule.get("priorita", 50),
        "conditions": [condition],
        "materials": materials
    }
    
    return converted


@app.get("/api/rules")
def list_rules():
    """Lista tutte le regole nella directory rules/"""
    rules = []
    rules_dir = "./rules"
    
    if not os.path.exists(rules_dir):
        os.makedirs(rules_dir)
        return {"rules": [], "count": 0}
    
    for filename in os.listdir(rules_dir):
        if filename.endswith(".json"):
            try:
                with open(os.path.join(rules_dir, filename), "r", encoding="utf-8") as f:
                    rule = json.load(f)
                    rule["_filename"] = filename
                    rules.append(rule)
            except Exception as e:
                rules.append({"_filename": filename, "_error": str(e)})
    
    return {"rules": rules, "count": len(rules)}


@app.get("/api/rules/{rule_id}")
def get_rule(rule_id: str):
    """Ottieni una regola specifica"""
    rules_dir = "./rules"
    
    for filename in os.listdir(rules_dir):
        if filename.endswith(".json"):
            filepath = os.path.join(rules_dir, filename)
            with open(filepath, "r", encoding="utf-8") as f:
                rule = json.load(f)
                if rule.get("id") == rule_id:
                    return rule
    
    raise HTTPException(status_code=404, detail=f"Regola '{rule_id}' non trovata")


@app.post("/api/rules/import")
def import_rule_from_designer(data: dict, db: Session = Depends(get_db)):
    """
    Importa regole da Rule Designer e le converte nel formato Configuratore.
    
    Accetta:
    - Singola regola BOM (con "nome" e "condizione")
    - Array di regole
    - Export completo Rule Designer (con bom_rules)
    """
    import shutil
    from datetime import datetime
    
    rules_dir = "./rules"
    os.makedirs(rules_dir, exist_ok=True)
    
    imported = []
    errors = []
    
    # Determina il tipo di input
    rules_to_import = []
    
    # Export completo Rule Designer
    if "bom_rules" in data:
        bom_rules = data["bom_rules"]
        for categoria, regole in bom_rules.items():
            if isinstance(regole, list):
                rules_to_import.extend(regole)
    
    # Array di regole
    elif "rules" in data:
        rules_to_import = data["rules"]
    
    # Singola regola (ha "nome" e "condizione")
    elif "nome" in data and "condizione" in data:
        rules_to_import = [data]
    
    # Formato già compatibile (ha "id" e "conditions")
    elif "id" in data and "conditions" in data:
        rule_id = data["id"]
        filename = f"rule_{rule_id}.json"
        filepath = os.path.join(rules_dir, filename)
        
        with open(filepath, "w", encoding="utf-8") as f:
            json.dump(data, f, indent=2, ensure_ascii=False)
        
        return {
            "success": True,
            "message": "Regola salvata (formato già compatibile)",
            "imported": [rule_id],
            "errors": []
        }
    
    else:
        raise HTTPException(
            status_code=400, 
            detail="Formato non riconosciuto. Atteso: singola regola BOM, array di regole, o export completo."
        )
    
    # Converti e salva ogni regola
    for rd_rule in rules_to_import:
        try:
            if not rd_rule.get("nome"):
                errors.append({"rule": str(rd_rule)[:50], "error": "Campo 'nome' mancante"})
                continue
            
            # Converti
            converted = convert_rule_designer_to_configurator(rd_rule, db)
            
            # Salva file
            rule_id = converted["id"]
            # Evita doppio prefisso rule_rule_
            if rule_id.startswith("rule_"):
                filename = f"{rule_id}.json"
            else:
                filename = f"rule_{rule_id}.json"
            filepath = os.path.join(rules_dir, filename)
            
            # Backup se esiste già
            if os.path.exists(filepath):
                backup_path = filepath.replace(".json", f"_backup_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json")
                shutil.copy(filepath, backup_path)
            
            with open(filepath, "w", encoding="utf-8") as f:
                json.dump(converted, f, indent=2, ensure_ascii=False)
            
            imported.append({
                "id": rule_id,
                "filename": filename,
                "materials_count": len(converted.get("materials", []))
            })
            
        except Exception as e:
            errors.append({"rule": rd_rule.get("nome", "unknown"), "error": str(e)})
    
    return {
        "success": len(errors) == 0,
        "message": f"Importate {len(imported)} regole" + (f", {len(errors)} errori" if errors else ""),
        "imported": imported,
        "errors": errors
    }


@app.delete("/api/rules/{rule_id}")
def delete_rule(rule_id: str):
    """Elimina una regola"""
    import shutil
    from datetime import datetime
    
    rules_dir = "./rules"
    
    for filename in os.listdir(rules_dir):
        if filename.endswith(".json"):
            filepath = os.path.join(rules_dir, filename)
            with open(filepath, "r", encoding="utf-8") as f:
                rule = json.load(f)
                if rule.get("id") == rule_id:
                    backup_path = filepath.replace(".json", f"_deleted_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json")
                    shutil.move(filepath, backup_path)
                    return {"success": True, "message": f"Regola '{rule_id}' eliminata", "backup": backup_path}
    
    raise HTTPException(status_code=404, detail=f"Regola '{rule_id}' non trovata")


@app.post("/api/rules/convert-preview")
def preview_conversion(data: dict):
    """Anteprima conversione senza salvare."""
    if "nome" not in data or "condizione" not in data:
        raise HTTPException(status_code=400, detail="Formato Rule Designer non valido")
    
    converted = convert_rule_designer_to_configurator(data)
    
    return {
        "original": data,
        "converted": converted,
        "preview": True
    }


# ==========================================
# STARTUP
# ==========================================
if __name__ == "__main__":
    import uvicorn
    print("🚀 Configuratore Elettroquadri API v2.0.0")
    print("📡 Server: http://0.0.0.0:8000")
    print("📖 Docs: http://0.0.0.0:8000/docs")
    print("\n✨ Novità v2.0:")
    print("   ✅ Preventivi COMPLETO e RICAMBIO")
    print("   ✅ Anagrafica Clienti con sconti")
    print("   ✅ Anagrafica Articoli")
    print("   ✅ Calcolo prezzi a 3 livelli")
    print("   ✅ Righe ricambio")
    uvicorn.run(app, host="0.0.0.0", port=8000)
